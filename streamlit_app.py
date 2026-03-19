from __future__ import annotations

import io
import os
import re
import zipfile
from datetime import datetime
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from huggingface_hub import InferenceClient
except ImportError:
    InferenceClient = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None


APP_DIR = Path(__file__).resolve().parent
ENV_PATH = APP_DIR / ".env"

SYSTEM_PROMPT = (
    "You are Hugyoku, an educational assistant for study use. Help the user analyze "
    "activities, quizzes, assignments, uploaded files, documents, essays, and code "
    "issues. You may produce sample answers, draft essays, structured summaries, "
    "guided explanations, generated study documents, and debugging help for offline "
    "learning and review. Keep outputs clear and well-organized, and avoid assisting "
    "with live cheating or deceptive conduct."
)

TOOL_FOLDERS = {
    "quiz": "quiz_solver",
    "assignment": "assignment_solver",
    "essay": "essay_generator",
    "activity": "activity_generator",
    "document": "document_generator",
    "codefix": "code_error_fixer",
}

KNOWN_EXPORT_LABELS = {
    "heading suggestion",
    "title suggestion",
    "suggested heading",
    "essay body",
    "essay",
    "self-check tip",
    "quick self-check tip",
    "closing self-check tip",
    "task understanding",
    "response",
    "short explanation or rationale",
    "summary",
    "objective",
    "instructions",
    "activity proper",
    "answer guide",
    "document type",
    "purpose",
    "main content",
    "solution plan",
    "sample answer or draft",
    "notes to review",
    "issue summary",
    "root cause",
    "fixed version",
    "why it works",
    "next checks",
}

STATE_DEFAULTS = {
    "active_page": "dashboard",
    "saved_name": "",
    "saved_include_date": False,
    "main_folder_name": "hugyoku_exports",
    "output_include_name": True,
    "output_include_date": False,
    "essay_include_heading": True,
    "essay_include_tip": True,
    "profile_name_input": "",
    "profile_include_date_input": False,
    "profile_main_folder_input": "hugyoku_exports",
    "profile_output_include_name_input": True,
    "profile_output_include_date_input": False,
    "profile_essay_include_heading_input": True,
    "profile_essay_include_tip_input": True,
    "quiz_upload_name": "No file loaded yet",
    "quiz_source_text": "",
    "quiz_summary": "",
    "quiz_prompt": "",
    "quiz_response": "",
    "quiz_mode": "complete",
    "assignment_upload_name": "No assignment file loaded yet",
    "assignment_source_text": "",
    "assignment_summary": "",
    "assignment_prompt": "",
    "assignment_response": "",
    "assignment_mode": "guided",
    "essay_title": "",
    "essay_prompt": "",
    "essay_word_count": 500,
    "essay_tagalog": False,
    "essay_english": True,
    "essay_specific_name": "",
    "essay_response": "",
    "activity_title": "",
    "activity_type": "Worksheet",
    "activity_level": "",
    "activity_prompt": "",
    "activity_response": "",
    "document_title": "",
    "document_type": "Study Handout",
    "document_audience": "",
    "document_prompt": "",
    "document_response": "",
    "codefix_title": "",
    "codefix_language": "Python",
    "codefix_error": "",
    "codefix_source": "",
    "codefix_expectation": "",
    "codefix_response": "",
}

PAGE_DETAILS = {
    "dashboard": {
        "title": "Dashboard",
        "subtitle": "Save your identity, choose one main export folder name, and keep all Hugyoku outputs organized in web-friendly download packages.",
    },
    "academics": {
        "title": "Academics",
        "subtitle": "Choose a focused school workspace. Each tool keeps the same Hugyoku content flow but is adapted for browser-based use.",
    },
    "developer": {
        "title": "Developer",
        "subtitle": "A separate hub for code support so the academic tools stay clean and easier to scan.",
    },
    "quiz": {
        "title": "Quiz Solver",
        "subtitle": "Read quiz files, summarize the task, and generate guided response support in a dedicated web workspace.",
    },
    "assignment": {
        "title": "Assignment Solver",
        "subtitle": "Analyze assignment instructions, understand the task first, then generate a guided response or full draft.",
    },
    "essay": {
        "title": "Essay Generator",
        "subtitle": "Build essay drafts with the same Hugyoku output controls, optional export name, and Word export formatting.",
    },
    "activity": {
        "title": "Activity Generator",
        "subtitle": "Generate worksheets, reflections, drills, or school activities from a topic and instruction set.",
    },
    "document": {
        "title": "Document Generator",
        "subtitle": "Create structured school documents like handouts, reports, reviewers, and formal academic materials.",
    },
    "codefix": {
        "title": "Code Error Fixer",
        "subtitle": "Paste code and the error details, then generate a cleaner fix and explanation in a separate developer workspace.",
    },
}

THEME_CSS = """
<style>
:root {
  --bg: #0c1017;
  --sidebar: #121924;
  --panel: rgba(27, 35, 49, 0.88);
  --border: rgba(80, 96, 122, 0.65);
  --text: #f6f0e5;
  --muted: #a9b2c3;
  --accent-soft: #f3d79a;
}

.stApp {
  background: radial-gradient(circle at top, #182132 0%, #0c1017 55%);
  color: var(--text);
}

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #111824 0%, #0d1420 100%);
  border-right: 1px solid rgba(80, 96, 122, 0.38);
}

.block-container {
  padding-top: 1.3rem;
  padding-bottom: 2rem;
}

.h-shell,
.h-card,
.h-stat {
  background: var(--panel);
  border: 1px solid var(--border);
  border-radius: 20px;
  backdrop-filter: blur(12px);
}

.h-shell {
  padding: 1.35rem 1.5rem;
  margin-bottom: 1rem;
}

.h-card {
  padding: 1.1rem 1.15rem;
  margin-bottom: 1rem;
}

.h-stat {
  padding: 0.95rem 1rem;
  min-height: 100%;
}

.h-hero-title {
  font-size: clamp(2rem, 3vw, 3rem);
  line-height: 1.1;
  color: var(--text);
  font-weight: 800;
  margin-bottom: 0.35rem;
}

.h-title {
  font-size: 1.65rem;
  font-weight: 800;
  color: var(--text);
  margin-bottom: 0.35rem;
}

.h-subtitle,
.h-copy,
.h-muted,
.h-folder {
  color: var(--muted);
}

.h-kicker {
  font-size: 0.78rem;
  text-transform: uppercase;
  letter-spacing: 0.12em;
  color: var(--accent-soft);
  margin-bottom: 0.5rem;
}

.h-chip {
  display: inline-block;
  padding: 0.38rem 0.7rem;
  border-radius: 999px;
  font-size: 0.82rem;
  font-weight: 700;
  margin-bottom: 0.5rem;
}

.h-chip.ready {
  background: rgba(80, 181, 142, 0.2);
  color: #c9ffe7;
}

.h-chip.offline {
  background: rgba(227, 123, 116, 0.16);
  color: #ffd4cf;
}

.h-chip.waiting {
  background: rgba(228, 182, 94, 0.16);
  color: #ffe7bb;
}

.h-tag-row {
  display: flex;
  flex-wrap: wrap;
  gap: 0.5rem;
  margin-top: 0.9rem;
}

.h-tag {
  background: rgba(243, 215, 154, 0.1);
  border: 1px solid rgba(243, 215, 154, 0.2);
  color: var(--accent-soft);
  border-radius: 999px;
  padding: 0.3rem 0.65rem;
  font-size: 0.8rem;
}

.h-divider {
  height: 1px;
  background: rgba(124, 184, 255, 0.32);
  margin: 1rem 0 1.2rem;
}

.h-folder {
  font-family: Consolas, monospace;
  font-size: 0.85rem;
  white-space: pre-wrap;
}

code {
  color: #ffe9bd;
}
</style>
"""


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", name).strip(" ._")
    return cleaned or "document"


def today_string() -> str:
    now = datetime.now()
    return f"{now.strftime('%B')} {now.day}, {now.year}"


def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text))


def collapse_paragraph_lines(lines: list[str]) -> str:
    return " ".join(line.strip() for line in lines if line.strip()).strip()


def split_text_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n\s*\n", text.strip()) if block.strip()]


def strip_heading_markers(text: str) -> str:
    cleaned = re.sub(r"^#+\s*", "", text.strip())
    return cleaned.strip().strip(":").strip()


def heading_level_from_line(text: str) -> int:
    stripped = text.lstrip()
    if stripped.startswith("#"):
        return min(max(len(stripped) - len(stripped.lstrip("#")), 1), 4)
    return 2


def normalize_section_label(text: str) -> str:
    cleaned = strip_heading_markers(text)
    cleaned = re.sub(r"^\d+\.\s*", "", cleaned)
    return re.sub(r"\s+", " ", cleaned).lower().strip()


def parse_structured_blocks(text: str) -> list[dict[str, object]]:
    items: list[dict[str, object]] = []
    for block in split_text_blocks(text):
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        first = lines[0].strip()
        normalized_first = normalize_section_label(first)
        is_heading = first.startswith("#") or first.endswith(":") or normalized_first in KNOWN_EXPORT_LABELS
        if is_heading:
            items.append(
                {
                    "type": "section",
                    "heading": strip_heading_markers(first),
                    "level": heading_level_from_line(first),
                    "content": collapse_paragraph_lines(lines[1:]),
                }
            )
        else:
            items.append({"type": "paragraph", "content": collapse_paragraph_lines(lines)})
    return items


def add_body_paragraph(document: Document, text: str, indent_first_line: bool = False, italic: bool = False) -> None:
    if not text.strip():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(8)
    if indent_first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
    run = paragraph.add_run(text.strip())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = italic


def render_docx_bytes(
    title: str,
    body: str,
    category: str = "generic",
    metadata_lines: list[str] | None = None,
    output_options: dict[str, bool] | None = None,
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    metadata_lines = metadata_lines or []
    output_options = output_options or {}
    export_title = title.strip() or "Untitled Document"

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(export_title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(48, 87, 138)

    if metadata_lines:
        meta_paragraph = document.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_paragraph.paragraph_format.space_after = Pt(10)
        meta_run = meta_paragraph.add_run(" | ".join(metadata_lines))
        meta_run.italic = True
        meta_run.font.name = "Times New Roman"
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(88, 88, 88)

    parsed_blocks = parse_structured_blocks(body)
    essay_subtitle: str | None = None
    essay_tip: str | None = None
    body_items: list[dict[str, object]] = []

    for item in parsed_blocks:
        if item["type"] == "section":
            label = normalize_section_label(str(item["heading"]))
            content = str(item.get("content", "")).strip()
            if category == "essay":
                if label in {"heading suggestion", "title suggestion", "suggested heading"}:
                    if content:
                        essay_subtitle = content
                    continue
                if label in {"essay body", "essay"}:
                    if content:
                        body_items.append({"type": "paragraph", "content": content})
                    continue
                if label in {"self-check tip", "quick self-check tip", "closing self-check tip"}:
                    if content:
                        essay_tip = content
                    continue
            body_items.append(item)
        else:
            body_items.append(item)

    if category == "essay" and output_options.get("essay_include_heading", True) and essay_subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.paragraph_format.space_after = Pt(12)
        subtitle_run = subtitle_paragraph.add_run(essay_subtitle)
        subtitle_run.italic = True
        subtitle_run.font.name = "Times New Roman"
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    for item in body_items:
        if item["type"] == "section":
            heading_text = str(item["heading"]).strip().rstrip(":")
            if heading_text:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(heading_text)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(60, 76, 104)
            content = str(item.get("content", "")).strip()
            if content:
                add_body_paragraph(document, content, indent_first_line=(category == "essay"))
        else:
            add_body_paragraph(document, str(item["content"]), indent_first_line=(category == "essay"))

    if category == "essay" and output_options.get("essay_include_tip", True) and essay_tip:
        tip_heading = document.add_paragraph()
        tip_heading.paragraph_format.space_before = Pt(10)
        tip_heading.paragraph_format.space_after = Pt(2)
        tip_run = tip_heading.add_run("Self-Check Tip")
        tip_run.bold = True
        tip_run.font.name = "Times New Roman"
        tip_run.font.size = Pt(13)
        tip_run.font.color.rgb = RGBColor(60, 76, 104)
        add_body_paragraph(document, essay_tip, italic=True)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def extract_completion_text(message: object) -> str | None:
    content = getattr(message, "content", None)
    if isinstance(content, str):
        return content.strip() or None
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, str) and item.strip():
                parts.append(item.strip())
            elif isinstance(item, dict):
                text = item.get("text") or item.get("content")
                if isinstance(text, str) and text.strip():
                    parts.append(text.strip())
        joined = "\n\n".join(parts).strip()
        return joined or None
    return None


def read_secret(name: str) -> str:
    load_dotenv(ENV_PATH)
    try:
        secret_value = st.secrets.get(name)
    except Exception:
        secret_value = None
    if isinstance(secret_value, str) and secret_value.strip():
        return secret_value.strip()
    return os.getenv(name, "").strip()


def load_client() -> tuple[object | None, str | None, str | None]:
    token = read_secret("HF_TOKEN")
    model = read_secret("HF_MODEL")

    if InferenceClient is None:
        return None, None, "Install the packages in requirements.txt to enable AI features."
    if not token or not model:
        return None, None, "Add HF_TOKEN and HF_MODEL to Streamlit secrets or local .env to enable AI features."

    return InferenceClient(api_key=token), model, None


def generate_text(prompt: str, label: str, client: object | None = None, model: str | None = None) -> str:
    local_client = client
    local_model = model
    if local_client is None or local_model is None:
        local_client, local_model, error = load_client()
        if error:
            raise RuntimeError(error)

    try:
        completion = local_client.chat.completions.create(
            model=local_model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1400,
            temperature=0.45,
        )
    except Exception as exc:
        raise RuntimeError(f"Could not generate {label}: {exc}") from exc

    message = completion.choices[0].message
    result = extract_completion_text(message)
    if not result:
        raise RuntimeError("The AI model returned an empty response.")
    return result


def read_docx_text(file_obj: io.BytesIO) -> str:
    document = Document(file_obj)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs).strip()


def read_pdf_text(file_obj: io.BytesIO) -> str:
    if PdfReader is None:
        raise RuntimeError("PDF reading needs the pypdf package. Install requirements.txt first.")

    reader = PdfReader(file_obj)
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    joined = "\n\n".join(parts).strip()
    if not joined:
        raise RuntimeError("The PDF did not contain readable text.")
    return joined


def read_text_file(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


def read_uploaded_document(uploaded_file: object) -> str:
    name = getattr(uploaded_file, "name", "uploaded.txt")
    suffix = Path(name).suffix.lower()
    data = uploaded_file.getvalue()
    if suffix == ".docx":
        return read_docx_text(io.BytesIO(data))
    if suffix == ".pdf":
        return read_pdf_text(io.BytesIO(data))
    if suffix in {".txt", ".md"}:
        return read_text_file(data)
    raise RuntimeError("Unsupported file type. Use .docx, .pdf, .txt, or .md.")

def ensure_state() -> None:
    for key, value in STATE_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value


def go(page: str) -> None:
    st.session_state.active_page = page
    st.rerun()


def save_profile() -> None:
    st.session_state.saved_name = st.session_state.profile_name_input.strip()
    st.session_state.saved_include_date = bool(st.session_state.profile_include_date_input)
    st.session_state.main_folder_name = sanitize_filename(st.session_state.profile_main_folder_input.strip() or "hugyoku_exports")
    st.session_state.output_include_name = bool(st.session_state.profile_output_include_name_input)
    st.session_state.output_include_date = bool(st.session_state.profile_output_include_date_input)
    st.session_state.essay_include_heading = bool(st.session_state.profile_essay_include_heading_input)
    st.session_state.essay_include_tip = bool(st.session_state.profile_essay_include_tip_input)


def clear_profile() -> None:
    st.session_state.saved_name = ""
    st.session_state.saved_include_date = False
    st.session_state.main_folder_name = "hugyoku_exports"
    st.session_state.output_include_name = False
    st.session_state.output_include_date = False
    st.session_state.essay_include_heading = True
    st.session_state.essay_include_tip = True
    st.session_state.profile_name_input = ""
    st.session_state.profile_include_date_input = False
    st.session_state.profile_main_folder_input = "hugyoku_exports"
    st.session_state.profile_output_include_name_input = False
    st.session_state.profile_output_include_date_input = False
    st.session_state.profile_essay_include_heading_input = True
    st.session_state.profile_essay_include_tip_input = True


def identity_block() -> str:
    lines: list[str] = []
    if st.session_state.saved_name:
        lines.append(f"Name: {st.session_state.saved_name}")
    if st.session_state.saved_include_date:
        lines.append(f"Date: {today_string()}")
    return "\n".join(lines)


def current_output_settings() -> dict[str, bool]:
    return {
        "include_name": bool(st.session_state.output_include_name),
        "include_date": bool(st.session_state.output_include_date),
        "essay_include_heading": bool(st.session_state.essay_include_heading),
        "essay_include_tip": bool(st.session_state.essay_include_tip),
    }


def export_metadata_lines(category: str = "generic", name_override: str | None = None) -> list[str]:
    lines: list[str] = []
    selected_name = st.session_state.saved_name
    if category == "essay" and name_override and name_override.strip():
        selected_name = name_override.strip()
    if st.session_state.output_include_name and selected_name:
        lines.append(selected_name)
    if st.session_state.output_include_date:
        lines.append(today_string())
    return lines


def folder_path_lines() -> dict[str, str]:
    main_folder = sanitize_filename(st.session_state.main_folder_name or "hugyoku_exports")
    paths = {"main": main_folder}
    for category, folder in TOOL_FOLDERS.items():
        paths[category] = f"{main_folder}/{folder}"
    return paths


def build_export_package(
    title: str,
    body: str,
    default_name: str,
    category: str,
    name_override: str | None = None,
) -> tuple[bytes, str, str]:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = sanitize_filename(default_name)
    docx_filename = f"{stem}_{stamp}.docx"
    main_folder = sanitize_filename(st.session_state.main_folder_name or "hugyoku_exports")
    internal_folder = TOOL_FOLDERS.get(category, category)
    internal_path = f"{main_folder}/{internal_folder}/{docx_filename}"
    docx_bytes = render_docx_bytes(
        title,
        body.strip(),
        category=category,
        metadata_lines=export_metadata_lines(category=category, name_override=name_override),
        output_options=current_output_settings(),
    )

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(internal_path, docx_bytes)
    zip_name = f"{stem}_{stamp}.zip"
    return zip_buffer.getvalue(), zip_name, internal_path


def clear_quiz_workspace() -> None:
    st.session_state.quiz_upload_name = "No file loaded yet"
    st.session_state.quiz_source_text = ""
    st.session_state.quiz_summary = ""
    st.session_state.quiz_prompt = ""
    st.session_state.quiz_response = ""
    st.session_state.quiz_mode = "complete"


def clear_assignment_workspace() -> None:
    st.session_state.assignment_upload_name = "No assignment file loaded yet"
    st.session_state.assignment_source_text = ""
    st.session_state.assignment_summary = ""
    st.session_state.assignment_prompt = ""
    st.session_state.assignment_response = ""
    st.session_state.assignment_mode = "guided"


def clear_essay_form() -> None:
    st.session_state.essay_title = ""
    st.session_state.essay_prompt = ""
    st.session_state.essay_word_count = 500
    st.session_state.essay_tagalog = False
    st.session_state.essay_english = True
    st.session_state.essay_specific_name = ""


def clear_essay_result() -> None:
    st.session_state.essay_response = ""


def clear_activity_form() -> None:
    st.session_state.activity_title = ""
    st.session_state.activity_type = "Worksheet"
    st.session_state.activity_level = ""
    st.session_state.activity_prompt = ""


def clear_activity_result() -> None:
    st.session_state.activity_response = ""


def clear_document_form() -> None:
    st.session_state.document_title = ""
    st.session_state.document_type = "Study Handout"
    st.session_state.document_audience = ""
    st.session_state.document_prompt = ""


def clear_document_result() -> None:
    st.session_state.document_response = ""


def clear_codefix_form() -> None:
    st.session_state.codefix_title = ""
    st.session_state.codefix_language = "Python"
    st.session_state.codefix_error = ""
    st.session_state.codefix_source = ""
    st.session_state.codefix_expectation = ""


def clear_codefix_result() -> None:
    st.session_state.codefix_response = ""


def run_generation(prompt: str, label: str) -> str | None:
    client, model, error = load_client()
    if error:
        st.error(error)
        return None
    with st.spinner(f"Generating {label}..."):
        try:
            return generate_text(prompt, label, client=client, model=model)
        except Exception as exc:
            st.error(str(exc))
            return None


def essay_language() -> str | None:
    tagalog = bool(st.session_state.essay_tagalog)
    english = bool(st.session_state.essay_english)
    if tagalog and english:
        return "Taglish"
    if tagalog:
        return "Tagalog"
    if english:
        return "English"
    return None


def render_card_header(title: str, subtitle: str, kicker: str | None = None) -> None:
    kicker_html = f'<div class="h-kicker">{kicker}</div>' if kicker else ""
    st.markdown(
        f"""
        <div class="h-card">
          {kicker_html}
          <div class="h-title">{title}</div>
          <div class="h-subtitle">{subtitle}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_stat_card(title: str, body: str) -> None:
    st.markdown(
        f"""
        <div class="h-stat">
          <div class="h-kicker">{title}</div>
          <div class="h-folder">{body}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_header(ai_ready: bool, model_label: str, ai_message: str) -> None:
    title = f"{st.session_state.saved_name or 'Guest'}'s Premium Study Desk"
    subtitle = (
        "A cleaner dashboard for identity settings, grouped academic workspaces, "
        "and a separate developer hub for code-fixing support."
    )
    chip_class = "ready" if ai_ready else ("waiting" if "Add HF_TOKEN" in ai_message or "Streamlit secrets" in ai_message else "offline")

    left, right = st.columns([3.2, 1.2])
    with left:
        st.markdown(
            f"""
            <div class="h-shell">
              <div class="h-hero-title">{title}</div>
              <div class="h-subtitle">{subtitle}</div>
              <div class="h-tag-row">
                <span class="h-tag">Dashboard Identity</span>
                <span class="h-tag">Academic Solvers</span>
                <span class="h-tag">Document Drafting</span>
                <span class="h-tag">Code Fix Support</span>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with right:
        st.markdown(
            f"""
            <div class="h-shell">
              <div class="h-chip {chip_class}">{'AI ready for academics' if ai_ready else ai_message}</div>
              <div class="h-copy"><strong>{st.session_state.saved_name or 'No saved name yet'}</strong></div>
              <div class="h-muted">{'Date today enabled: ' + today_string() if st.session_state.saved_include_date else 'Date today disabled'}</div>
              <div class="h-muted" style="margin-top:0.4rem;">{model_label}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        if st.button("Refresh AI Status", key="refresh_ai_status", use_container_width=True):
            st.rerun()


def render_sidebar(ai_ready: bool, model_label: str, ai_message: str) -> None:
    st.sidebar.markdown("## Hugyoku")
    st.sidebar.caption("Premium study and utility desk for academic generators, solvers, and code-fixing support.")

    nav_items = [
        ("dashboard", "01  Dashboard"),
        ("academics", "02  Academics"),
        ("developer", "03  Developer"),
    ]
    academic_pages = {"academics", "quiz", "assignment", "essay", "activity", "document"}
    developer_pages = {"developer", "codefix"}

    for page, label in nav_items:
        active = (
            page == "dashboard" and st.session_state.active_page == "dashboard"
        ) or (
            page == "academics" and st.session_state.active_page in academic_pages
        ) or (
            page == "developer" and st.session_state.active_page in developer_pages
        )
        if st.sidebar.button(label, key=f"nav_{page}", type="primary" if active else "secondary", use_container_width=True):
            go(page)

    st.sidebar.markdown("---")
    st.sidebar.markdown("**Educational use only**")
    st.sidebar.caption("Use these flows for studying, drafting, and understanding tasks from your files. Review outputs before submitting anything.")
    st.sidebar.markdown("---")
    st.sidebar.caption("AI Status")
    st.sidebar.write("Ready" if ai_ready else ai_message)
    st.sidebar.caption(model_label)


def render_page_intro(page_key: str) -> None:
    details = PAGE_DETAILS[page_key]
    st.markdown(f"## {details['title']}")
    st.markdown(f"<div class='h-copy'>{details['subtitle']}</div>", unsafe_allow_html=True)
    st.markdown("<div class='h-divider'></div>", unsafe_allow_html=True)


def render_tool_hub_card(title: str, subtitle: str, folder_text: str, button_label: str, target: str, primary: bool = False) -> None:
    render_card_header(title, subtitle)
    st.markdown(f"<div class='h-folder'>{folder_text}</div>", unsafe_allow_html=True)
    if st.button(button_label, key=f"open_{target}", use_container_width=True, type="primary" if primary else "secondary"):
        go(target)


def render_back_button(target: str, label: str) -> None:
    if st.button(label, key=f"back_{target}_{st.session_state.active_page}"):
        go(target)


def render_download_button(title: str, body: str, default_name: str, category: str, clear_callback: object, name_override: str | None = None) -> None:
    package_bytes, zip_name, internal_path = build_export_package(title, body, default_name, category, name_override=name_override)
    st.caption(f"Package path: {internal_path}")
    st.download_button(
        "Download Export Package",
        data=package_bytes,
        file_name=zip_name,
        mime="application/zip",
        key=f"download_{category}_{sanitize_filename(default_name)}",
        use_container_width=True,
        on_click=clear_callback,
    )


def render_dashboard() -> None:
    render_page_intro("dashboard")

    col_a, col_b = st.columns([1.1, 0.9])
    with col_a:
        render_card_header(
            "Profile Section",
            "Enter your name, decide whether to stamp today's date, then set one main export folder name for all web download packages.",
            "Dashboard Identity",
        )
        st.text_input("Enter your name", key="profile_name_input")
        st.checkbox("Add date today", key="profile_include_date_input")
        st.text_input("Main export folder name", key="profile_main_folder_input", help="Used inside downloaded zip packages as the main parent folder.")
        st.markdown("#### Output Options")
        st.checkbox("Include saved name in export", key="profile_output_include_name_input")
        st.checkbox("Include date in export", key="profile_output_include_date_input")
        st.checkbox("Include essay heading suggestion", key="profile_essay_include_heading_input")
        st.checkbox("Include self-check tip", key="profile_essay_include_tip_input")
        btn_a, btn_b = st.columns(2)
        if btn_a.button("Save Profile", use_container_width=True):
            save_profile()
            st.success("Profile saved for this session.")
        if btn_b.button("Clear Saved Profile", use_container_width=True):
            clear_profile()
            st.success("Profile cleared.")

    with col_b:
        render_card_header(
            "Export Folder Routing",
            "Web exports are downloaded as zip packages that preserve the same main folder and subfolder structure from the desktop version.",
            "Download Structure",
        )
        paths = folder_path_lines()
        render_stat_card("Main folder", paths["main"])
        render_stat_card(
            "Academics suite",
            "\n".join(
                [
                    paths["quiz"],
                    paths["assignment"],
                    paths["essay"],
                    paths["activity"],
                    paths["document"],
                ]
            ),
        )
        render_stat_card("Developer suite", paths["codefix"])

    launch_a, launch_b = st.columns(2)
    with launch_a:
        render_card_header(
            "Quick Launch",
            "Open the academics hub for school tools or the developer hub for code fixing.",
            "Next Step",
        )
        if st.button("Open Academics Hub", key="open_academics_hub", use_container_width=True, type="primary"):
            go("academics")
        if st.button("Open Developer Hub", key="open_developer_hub", use_container_width=True):
            go("developer")
    with launch_b:
        render_card_header(
            "How This Flow Works",
            "The structure stays simple: save your profile, open a hub, choose a tool, then download the finished package.",
            "Guide",
        )
        st.markdown(
            """
            1. Dashboard stores your active profile for this session.
            2. Academics contains quiz, assignment, essay, activity, and document tools.
            3. Developer contains the code error fixer.
            4. Every export downloads as a zip file with the correct Hugyoku folder structure.
            """
        )


def render_academics_hub() -> None:
    render_page_intro("academics")
    paths = folder_path_lines()

    render_tool_hub_card(
        "Quiz Solver Workspace",
        "Read quiz files, summarize the task, and generate guided response support in a dedicated workspace.",
        paths["quiz"],
        "Open Quiz Solver",
        "quiz",
        True,
    )
    render_tool_hub_card(
        "Assignment Solver Workspace",
        "Analyze assignments from pasted text or uploaded files, then generate a guided draft or response plan.",
        paths["assignment"],
        "Open Assignment Solver",
        "assignment",
    )
    render_tool_hub_card(
        "Essay Generator Workspace",
        "Build long-form essay drafts with output options, optional name overrides, and cleaner export formatting.",
        paths["essay"],
        "Open Essay Generator",
        "essay",
    )
    render_tool_hub_card(
        "Activity Generator Workspace",
        "Generate worksheets, reflections, drills, or classroom activities from a topic and instruction set.",
        paths["activity"],
        "Open Activity Generator",
        "activity",
    )
    render_tool_hub_card(
        "Document Generator Workspace",
        "Create structured school documents like handouts, reports, reviewers, and formal academic materials.",
        paths["document"],
        "Open Document Generator",
        "document",
    )


def render_developer_hub() -> None:
    render_page_intro("developer")
    paths = folder_path_lines()
    render_tool_hub_card(
        "Code Error Fixer Workspace",
        "Paste code, explain the bug or error, and get a cleaner fix plus a short explanation of what changed.",
        paths["codefix"],
        "Open Code Error Fixer",
        "codefix",
        True,
    )

def render_quiz_page(ai_ready: bool) -> None:
    render_page_intro("quiz")
    render_back_button("academics", "Back To Academics")
    st.caption(folder_path_lines()["quiz"])

    left, right = st.columns(2)
    with left:
        render_card_header("Quiz Intake", "Upload a file or paste quiz content directly. The app reads it first before generating support.")
        uploaded = st.file_uploader("Upload quiz file", type=["docx", "pdf", "txt", "md"], key="quiz_upload_widget")
        load_col, clear_col = st.columns(2)
        if load_col.button("Load Uploaded File", key="quiz_load_file", use_container_width=True):
            if uploaded is None:
                st.warning("Upload a file first.")
            else:
                try:
                    st.session_state.quiz_source_text = read_uploaded_document(uploaded)
                    st.session_state.quiz_upload_name = uploaded.name
                    st.session_state.quiz_summary = ""
                    st.session_state.quiz_response = ""
                    st.session_state.quiz_prompt = ""
                    st.success(f"Loaded {uploaded.name} into the quiz workspace.")
                except Exception as exc:
                    st.error(str(exc))
        if clear_col.button("Clear Quiz Input", key="quiz_clear_input", use_container_width=True):
            clear_quiz_workspace()
            st.rerun()
        st.caption(st.session_state.quiz_upload_name)
        st.caption(f"{count_words(st.session_state.quiz_source_text)} words | {len(st.session_state.quiz_source_text.strip())} characters")
        st.text_area("Quiz content", key="quiz_source_text", height=360)
        if st.button("Analyze Quiz", key="quiz_analyze", use_container_width=True, type="secondary", disabled=not ai_ready):
            source = st.session_state.quiz_source_text.strip()
            if not source:
                st.warning("Upload or paste quiz content before analyzing it.")
            else:
                prompt = (
                    "Analyze the following academic material for educational review.\n"
                    f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                    f"Source content:\n{source}\n\n"
                    "Return these sections clearly:\n"
                    "1. Summary\n"
                    "2. What task or instructions seem to be required\n"
                    "3. Important topics, clues, or constraints\n"
                    "4. Best next step for the student"
                )
                result = run_generation(prompt, "quiz analysis")
                if result:
                    st.session_state.quiz_summary = result
                    st.success("Quiz analysis complete.")

    with right:
        render_card_header("Summary And Task Reading", "The assistant summarizes the uploaded content first so you can understand the task before continuing.")
        st.text_area("Summary preview", value=st.session_state.quiz_summary, height=220, disabled=True, key="quiz_summary_view")
        quiz_mode = st.radio("Response mode", options=["complete", "specific"], index=0 if st.session_state.quiz_mode == "complete" else 1, horizontal=True, format_func=lambda item: "Do it in full" if item == "complete" else "Use specific prompt")
        st.session_state.quiz_mode = quiz_mode
        st.text_area("Specific prompt", key="quiz_prompt", height=120)
        gen_col, clear_col = st.columns(2)
        if gen_col.button("Generate Quiz Guidance", key="quiz_generate", use_container_width=True, type="primary", disabled=not ai_ready):
            source = st.session_state.quiz_source_text.strip()
            summary = st.session_state.quiz_summary.strip()
            custom_prompt = st.session_state.quiz_prompt.strip()
            if not source:
                st.warning("Upload or paste quiz content before generating a response.")
            elif not summary:
                st.warning("Run the summary step first so the app can read the task before generating a response.")
            elif st.session_state.quiz_mode == "specific" and not custom_prompt:
                st.warning("Add a specific prompt or switch to full mode.")
            else:
                instructions = (
                    "Create a complete educational response based on the content. If there are visible questions, answer them in order and include short explanations. If the file describes a task instead of direct questions, complete the task as fully as possible."
                    if st.session_state.quiz_mode == "complete"
                    else "Follow the user's specific prompt while staying grounded in the uploaded content and analysis.\nSpecific prompt: " + custom_prompt
                )
                prompt = (
                    "Use the uploaded quiz/activity content for educational support.\n"
                    f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                    f"Analysis already prepared:\n{summary}\n\n"
                    f"Source content:\n{source}\n\n"
                    f"Instructions:\n{instructions}\n\n"
                    "Structure the answer with:\n"
                    "1. Task understanding\n"
                    "2. Response\n"
                    "3. Short explanation or rationale\n\n"
                    "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
                )
                result = run_generation(prompt, "quiz response")
                if result:
                    st.session_state.quiz_response = result
                    st.success("Quiz guidance ready.")
        if clear_col.button("Clear Result", key="quiz_clear_result", use_container_width=True):
            st.session_state.quiz_response = ""
            st.rerun()
        st.text_area("Generated response", value=st.session_state.quiz_response, height=300, disabled=True, key="quiz_response_view")
        if st.session_state.quiz_response.strip():
            export_title = f"Quiz Support: {Path(st.session_state.quiz_upload_name).stem.replace('_', ' ').title()}" if st.session_state.quiz_upload_name != "No file loaded yet" else "Quiz Support"
            export_name = sanitize_filename(Path(st.session_state.quiz_upload_name).stem if st.session_state.quiz_upload_name != "No file loaded yet" else "quiz_support")
            render_download_button(export_title, st.session_state.quiz_response, export_name, "quiz", clear_quiz_workspace)


def render_assignment_page(ai_ready: bool) -> None:
    render_page_intro("assignment")
    render_back_button("academics", "Back To Academics")
    st.caption(folder_path_lines()["assignment"])

    left, right = st.columns(2)
    with left:
        render_card_header("Assignment Intake", "Upload a document or paste the assignment manually, then let the app analyze the task before generating guidance.")
        uploaded = st.file_uploader("Upload assignment file", type=["docx", "pdf", "txt", "md"], key="assignment_upload_widget")
        load_col, clear_col = st.columns(2)
        if load_col.button("Load Uploaded File", key="assignment_load_file", use_container_width=True):
            if uploaded is None:
                st.warning("Upload an assignment file first.")
            else:
                try:
                    st.session_state.assignment_source_text = read_uploaded_document(uploaded)
                    st.session_state.assignment_upload_name = uploaded.name
                    st.session_state.assignment_summary = ""
                    st.session_state.assignment_response = ""
                    st.success(f"Loaded {uploaded.name} into the assignment workspace.")
                except Exception as exc:
                    st.error(str(exc))
        if clear_col.button("Clear Assignment Input", key="assignment_clear_input", use_container_width=True):
            clear_assignment_workspace()
            st.rerun()
        st.caption(st.session_state.assignment_upload_name)
        st.caption(f"{count_words(st.session_state.assignment_source_text)} words | {len(st.session_state.assignment_source_text.strip())} characters")
        st.text_area("Assignment content", key="assignment_source_text", height=320)
        assignment_mode = st.radio("Assignment mode", options=["guided", "complete"], index=0 if st.session_state.assignment_mode == "guided" else 1, horizontal=True, format_func=lambda item: "Guided response" if item == "guided" else "Full draft")
        st.session_state.assignment_mode = assignment_mode
        st.text_area("Specific prompt (optional)", key="assignment_prompt", height=110)
        action_a, action_b = st.columns(2)
        if action_a.button("Analyze Assignment", key="assignment_analyze", use_container_width=True, type="secondary", disabled=not ai_ready):
            source = st.session_state.assignment_source_text.strip()
            if not source:
                st.warning("Upload or paste assignment content before analyzing it.")
            else:
                prompt = (
                    "Analyze the following assignment for educational review.\n"
                    f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                    f"Assignment content:\n{source}\n\n"
                    "Return these sections clearly:\n"
                    "1. Assignment summary\n"
                    "2. What is being asked\n"
                    "3. Important requirements or constraints\n"
                    "4. Best approach for the student"
                )
                result = run_generation(prompt, "assignment analysis")
                if result:
                    st.session_state.assignment_summary = result
                    st.success("Assignment analysis complete.")
        if action_b.button("Generate Assignment Guidance", key="assignment_generate", use_container_width=True, type="primary", disabled=not ai_ready):
            source = st.session_state.assignment_source_text.strip()
            summary = st.session_state.assignment_summary.strip()
            custom_prompt = st.session_state.assignment_prompt.strip()
            if not source:
                st.warning("Upload or paste assignment content before generating a response.")
            elif not summary:
                st.warning("Run the assignment analysis first.")
            else:
                instructions = (
                    "Create a full draft response based on the assignment while keeping it clear, organized, and educational."
                    if st.session_state.assignment_mode == "complete"
                    else "Create a guided response that explains how to approach the assignment and includes a sample draft the student can review."
                )
                if custom_prompt:
                    instructions = f"{instructions}\nSpecific prompt: {custom_prompt}"
                prompt = (
                    "Use the uploaded assignment for educational support.\n"
                    f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                    f"Assignment analysis:\n{summary}\n\n"
                    f"Assignment content:\n{source}\n\n"
                    f"Instructions:\n{instructions}\n\n"
                    "Structure the answer with:\n"
                    "1. Task understanding\n"
                    "2. Solution plan\n"
                    "3. Sample answer or draft\n"
                    "4. Notes to review\n\n"
                    "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
                )
                result = run_generation(prompt, "assignment response")
                if result:
                    st.session_state.assignment_response = result
                    st.success("Assignment guidance ready.")

    with right:
        render_card_header("Assignment Response", "The assignment summary appears first, followed by the generated guided response or sample draft.")
        st.text_area("Assignment summary", value=st.session_state.assignment_summary, height=220, disabled=True, key="assignment_summary_view")
        st.text_area("Generated response", value=st.session_state.assignment_response, height=280, disabled=True, key="assignment_response_view")
        if st.button("Clear Result", key="assignment_clear_result", use_container_width=True):
            st.session_state.assignment_response = ""
            st.rerun()
        if st.session_state.assignment_response.strip():
            export_title = f"Assignment Support: {Path(st.session_state.assignment_upload_name).stem.replace('_', ' ').title()}" if st.session_state.assignment_upload_name != "No assignment file loaded yet" else "Assignment Support"
            export_name = sanitize_filename(Path(st.session_state.assignment_upload_name).stem if st.session_state.assignment_upload_name != "No assignment file loaded yet" else "assignment_support")
            render_download_button(export_title, st.session_state.assignment_response, export_name, "assignment", clear_assignment_workspace)

def render_essay_page(ai_ready: bool) -> None:
    render_page_intro("essay")
    render_back_button("academics", "Back To Academics")
    st.caption(folder_path_lines()["essay"])

    left, right = st.columns(2)
    with left:
        render_card_header(
            "Essay Builder",
            "Set the title, add an optional custom instruction, choose the target length, and pick English, Tagalog, or both for Taglish.",
        )
        st.text_input("Essay title", key="essay_title")
        st.text_area("Specific prompt (optional)", key="essay_prompt", height=160)
        st.number_input("Target word count", min_value=100, max_value=3000, step=50, key="essay_word_count")
        lang_a, lang_b = st.columns(2)
        lang_a.checkbox("Tagalog", key="essay_tagalog")
        lang_b.checkbox("English", key="essay_english")
        render_card_header(
            "Essay Export Format",
            "Adjust what gets attached to the final essay output. The specific export name is optional and can override the saved dashboard name for this essay only.",
        )
        st.text_input("Specific export name (optional)", key="essay_specific_name")
        st.checkbox("Include saved or specific name in export", key="profile_output_include_name_input")
        st.checkbox("Include date in export", key="profile_output_include_date_input")
        st.checkbox("Include essay heading suggestion", key="profile_essay_include_heading_input")
        st.checkbox("Include self-check tip", key="profile_essay_include_tip_input")
        st.info("Save Profile from the dashboard if you want these export settings to become the active saved defaults.")
        action_a, action_b = st.columns(2)
        if action_a.button("Generate Essay", key="essay_generate", use_container_width=True, type="primary", disabled=not ai_ready):
            title = st.session_state.essay_title.strip()
            prompt_text = st.session_state.essay_prompt.strip()
            language = essay_language()
            if not title:
                st.warning("Enter the essay title before generating.")
            elif language is None:
                st.warning("Select Tagalog, English, or both for Taglish.")
            else:
                prompt = (
                    "Write a polished educational essay draft.\n"
                    f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                    f"Essay title: {title}\n"
                    f"Target length: about {int(st.session_state.essay_word_count)} words\n"
                    f"Language: {language}\n"
                    f"Specific prompt: {prompt_text or 'No extra prompt provided'}\n\n"
                    "Return these plain-text sections in this exact order:\n"
                    "Heading Suggestion:\n"
                    "Essay Body:\n"
                    "Self-Check Tip:\n\n"
                    "Do not use markdown symbols like ### or ####."
                )
                result = run_generation(prompt, "essay draft")
                if result:
                    st.session_state.essay_response = result
                    st.success("Essay draft ready.")
        if action_b.button("Clear Essay Form", key="essay_clear_form", use_container_width=True):
            clear_essay_form()
            clear_essay_result()
            st.rerun()

    with right:
        render_card_header(
            "Essay Preview",
            "The draft appears here once generated. Export it to Word when you are happy with the result.",
        )
        st.text_area("Essay output", value=st.session_state.essay_response, height=520, disabled=True, key="essay_response_view")
        if st.button("Clear Result", key="essay_clear_result", use_container_width=True):
            clear_essay_result()
            st.rerun()
        if st.session_state.essay_response.strip():
            temp_saved = {
                "output_include_name": st.session_state.output_include_name,
                "output_include_date": st.session_state.output_include_date,
                "essay_include_heading": st.session_state.essay_include_heading,
                "essay_include_tip": st.session_state.essay_include_tip,
            }
            st.session_state.output_include_name = st.session_state.profile_output_include_name_input
            st.session_state.output_include_date = st.session_state.profile_output_include_date_input
            st.session_state.essay_include_heading = st.session_state.profile_essay_include_heading_input
            st.session_state.essay_include_tip = st.session_state.profile_essay_include_tip_input
            try:
                render_download_button(
                    st.session_state.essay_title.strip() or "Essay Draft",
                    st.session_state.essay_response,
                    sanitize_filename((st.session_state.essay_title.strip() or "essay_draft").lower().replace(" ", "_")),
                    "essay",
                    lambda: (clear_essay_form(), clear_essay_result()),
                    name_override=st.session_state.essay_specific_name.strip(),
                )
            finally:
                st.session_state.output_include_name = temp_saved["output_include_name"]
                st.session_state.output_include_date = temp_saved["output_include_date"]
                st.session_state.essay_include_heading = temp_saved["essay_include_heading"]
                st.session_state.essay_include_tip = temp_saved["essay_include_tip"]


def render_activity_page(ai_ready: bool) -> None:
    render_page_intro("activity")
    render_back_button("academics", "Back To Academics")
    st.caption(folder_path_lines()["activity"])

    left, right = st.columns(2)
    with left:
        render_card_header(
            "Activity Builder",
            "Set the topic, activity type, and optional level details, then generate a ready-to-use school activity.",
        )
        st.text_input("Activity topic or title", key="activity_title")
        st.text_input("Activity type", key="activity_type")
        st.text_input("Level or class (optional)", key="activity_level")
        st.text_area("Specific instructions", key="activity_prompt", height=180)
        action_a, action_b = st.columns(2)
        if action_a.button("Generate Activity", key="activity_generate", use_container_width=True, type="primary", disabled=not ai_ready):
            title = st.session_state.activity_title.strip()
            activity_type = st.session_state.activity_type.strip() or "Activity"
            level = st.session_state.activity_level.strip()
            request = st.session_state.activity_prompt.strip()
            if not title:
                st.warning("Enter the activity topic or title before generating.")
            else:
                prompt = (
                    "Create a structured educational activity for offline study use.\n"
                    f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                    f"Activity title or topic: {title}\n"
                    f"Activity type: {activity_type}\n"
                    f"Level or class: {level or 'Not specified'}\n"
                    f"Specific request: {request or 'No extra request provided'}\n\n"
                    "Return these plain-text sections in this exact order:\n"
                    "Activity Title:\n"
                    "Objective:\n"
                    "Instructions:\n"
                    "Activity Proper:\n"
                    "Answer Guide:\n\n"
                    "Do not use markdown symbols like ### or ####."
                )
                result = run_generation(prompt, "activity draft")
                if result:
                    st.session_state.activity_response = result
                    st.success("Activity ready.")
        if action_b.button("Clear Activity Form", key="activity_clear_form", use_container_width=True):
            clear_activity_form()
            clear_activity_result()
            st.rerun()

    with right:
        render_card_header(
            "Activity Preview",
            "The generated activity appears here in a clean structure you can export straight to Word.",
        )
        st.text_area("Generated activity", value=st.session_state.activity_response, height=520, disabled=True, key="activity_response_view")
        if st.button("Clear Result", key="activity_clear_result", use_container_width=True):
            clear_activity_result()
            st.rerun()
        if st.session_state.activity_response.strip():
            render_download_button(
                st.session_state.activity_title.strip() or "Activity Draft",
                st.session_state.activity_response,
                sanitize_filename((st.session_state.activity_title.strip() or "activity_draft").lower().replace(" ", "_")),
                "activity",
                lambda: (clear_activity_form(), clear_activity_result()),
            )


def render_document_page(ai_ready: bool) -> None:
    render_page_intro("document")
    render_back_button("academics", "Back To Academics")
    st.caption(folder_path_lines()["document"])

    left, right = st.columns(2)
    with left:
        render_card_header(
            "Document Builder",
            "Generate a structured school document such as a handout, report, reviewer, or formal academic write-up.",
        )
        st.text_input("Document title", key="document_title")
        st.text_input("Document type", key="document_type")
        st.text_input("Audience or purpose (optional)", key="document_audience")
        st.text_area("Specific content request", key="document_prompt", height=180)
        action_a, action_b = st.columns(2)
        if action_a.button("Generate Document", key="document_generate", use_container_width=True, type="primary", disabled=not ai_ready):
            title = st.session_state.document_title.strip()
            doc_type = st.session_state.document_type.strip() or "Study Handout"
            audience = st.session_state.document_audience.strip()
            request = st.session_state.document_prompt.strip()
            if not title:
                st.warning("Enter the document title before generating.")
            elif not request:
                st.warning("Describe what the document should contain before generating.")
            else:
                prompt = (
                    "Create a structured educational document.\n"
                    f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                    f"Document title: {title}\n"
                    f"Document type: {doc_type}\n"
                    f"Audience or purpose: {audience or 'Not specified'}\n"
                    f"Content request: {request}\n\n"
                    "Return these plain-text sections in this exact order:\n"
                    "Document Type:\n"
                    "Purpose:\n"
                    "Main Content:\n\n"
                    "Do not use markdown symbols like ### or ####."
                )
                result = run_generation(prompt, "document draft")
                if result:
                    st.session_state.document_response = result
                    st.success("Document ready.")
        if action_b.button("Clear Document Form", key="document_clear_form", use_container_width=True):
            clear_document_form()
            clear_document_result()
            st.rerun()

    with right:
        render_card_header(
            "Document Preview",
            "The generated document draft appears here and can be exported directly to the document subfolder.",
        )
        st.text_area("Generated document", value=st.session_state.document_response, height=520, disabled=True, key="document_response_view")
        if st.button("Clear Result", key="document_clear_result", use_container_width=True):
            clear_document_result()
            st.rerun()
        if st.session_state.document_response.strip():
            render_download_button(
                st.session_state.document_title.strip() or "Document Draft",
                st.session_state.document_response,
                sanitize_filename((st.session_state.document_title.strip() or "document_draft").lower().replace(" ", "_")),
                "document",
                lambda: (clear_document_form(), clear_document_result()),
            )


def render_codefix_page(ai_ready: bool) -> None:
    render_page_intro("codefix")
    render_back_button("developer", "Back To Developer")
    st.caption(folder_path_lines()["codefix"])

    left, right = st.columns(2)
    with left:
        render_card_header(
            "Code Fix Builder",
            "Paste the broken code, add the error or symptom, and explain the expected behavior so the app can generate a cleaner fix.",
        )
        st.text_input("Issue title (optional)", key="codefix_title")
        st.text_input("Language or stack", key="codefix_language")
        st.text_area("Error message or symptoms", key="codefix_error", height=140)
        st.text_area("Code snippet", key="codefix_source", height=240)
        st.text_area("Expected behavior (optional)", key="codefix_expectation", height=110)
        action_a, action_b = st.columns(2)
        if action_a.button("Generate Code Fix", key="codefix_generate", use_container_width=True, type="primary", disabled=not ai_ready):
            title = st.session_state.codefix_title.strip()
            language = st.session_state.codefix_language.strip() or "Code"
            error_text = st.session_state.codefix_error.strip()
            source = st.session_state.codefix_source.strip()
            expectation = st.session_state.codefix_expectation.strip()
            if not source:
                st.warning("Paste the code snippet before generating a fix.")
            elif not error_text:
                st.warning("Add the error message or symptoms before generating a fix.")
            else:
                prompt = (
                    "Fix the following code issue for study and debugging support.\n\n"
                    f"Language or stack: {language}\n"
                    f"Issue title: {title or 'No title provided'}\n"
                    f"Error message or symptoms:\n{error_text}\n\n"
                    f"Code snippet:\n{source}\n\n"
                    f"Expected behavior:\n{expectation or 'Not provided'}\n\n"
                    "Return these plain-text sections in this exact order:\n"
                    "Issue Summary:\n"
                    "Root Cause:\n"
                    "Fixed Version:\n"
                    "Why It Works:\n"
                    "Next Checks:\n\n"
                    "Do not use markdown symbols like ### or ####."
                )
                result = run_generation(prompt, "code fix")
                if result:
                    st.session_state.codefix_response = result
                    st.success("Code fix ready.")
        if action_b.button("Clear Code Fixer", key="codefix_clear_form", use_container_width=True):
            clear_codefix_form()
            clear_codefix_result()
            st.rerun()

    with right:
        render_card_header(
            "Code Fix Preview",
            "The cleaned-up fix and explanation appear here. Export it if you want to save the debugging notes to Word.",
        )
        st.text_area("Fixed result", value=st.session_state.codefix_response, height=520, disabled=True, key="codefix_response_view")
        if st.button("Clear Result", key="codefix_clear_result", use_container_width=True):
            clear_codefix_result()
            st.rerun()
        if st.session_state.codefix_response.strip():
            export_title = st.session_state.codefix_title.strip() or f"{st.session_state.codefix_language.strip() or 'Code'} Code Fix"
            render_download_button(
                export_title,
                st.session_state.codefix_response,
                sanitize_filename(export_title.lower().replace(" ", "_")),
                "codefix",
                lambda: (clear_codefix_form(), clear_codefix_result()),
            )


def main() -> None:
    st.set_page_config(page_title="Hugyoku | Premium Academics Suite", layout="wide", initial_sidebar_state="expanded")
    st.markdown(THEME_CSS, unsafe_allow_html=True)
    ensure_state()

    _client, model, error = load_client()
    ai_ready = error is None
    model_label = model or "No model configured"
    ai_message = "AI ready for academics" if ai_ready else error

    render_sidebar(ai_ready, model_label, ai_message)
    render_header(ai_ready, model_label, ai_message)

    page = st.session_state.active_page
    if page == "dashboard":
        render_dashboard()
    elif page == "academics":
        render_academics_hub()
    elif page == "developer":
        render_developer_hub()
    elif page == "quiz":
        render_quiz_page(ai_ready)
    elif page == "assignment":
        render_assignment_page(ai_ready)
    elif page == "essay":
        render_essay_page(ai_ready)
    elif page == "activity":
        render_activity_page(ai_ready)
    elif page == "document":
        render_document_page(ai_ready)
    elif page == "codefix":
        render_codefix_page(ai_ready)
    else:
        go("dashboard")


if __name__ == "__main__":
    main()
