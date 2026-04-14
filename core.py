from __future__ import annotations

import io
import json
import os
import re
import sqlite3
import subprocess
import sys
import difflib
import hashlib
import hmac
import textwrap
from datetime import datetime
from html import escape
from pathlib import Path

import streamlit as st
from streamlit.components.v1 import html as components_html
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv(*_args, **_kwargs):
        return False

try:
    from huggingface_hub import InferenceClient
except ImportError:
    InferenceClient = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import numpy as np
except ImportError:
    np = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from rapidocr_onnxruntime import RapidOCR
except ImportError:
    RapidOCR = None

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
except ImportError:
    LETTER = None
    canvas = None

try:
    # Tkinter is only used for the optional local folder picker.
    # Streamlit Cloud / headless environments may not have a usable Tk runtime.
    import tkinter as tk
    from tkinter import filedialog
except Exception:
    tk = None
    filedialog = None

try:
    from streamlit_mic_recorder import speech_to_text
except Exception:
    speech_to_text = None

try:
    import speech_recognition as sr
except Exception:
    sr = None


APP_DIR = Path(__file__).resolve().parent.parent
ENV_PATH = APP_DIR / ".env"
AUTH_DB_PATH = APP_DIR / "hugyoku_auth.db"
_OCR_ENGINE: object | None = None

ROLE_PERMISSIONS = {
    "super_admin": {
        "hugyoku": True,
        "hugyoku_chat": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": True,
        "history": True,
        "settings": True,
        "admin": True,
    },
    "admin": {
        "hugyoku": True,
        "hugyoku_chat": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": True,
        "history": True,
        "settings": True,
        "admin": True,
    },
    "member": {
        "hugyoku": True,
        "hugyoku_chat": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": True,
        "history": True,
        "settings": True,
        "admin": False,
    },
    "viewer": {
        "hugyoku": True,
        "hugyoku_chat": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": False,
        "history": True,
        "settings": True,
        "admin": False,
    },
}

LOCAL_ACCESS_USERNAME = "local"
LOCAL_ACCESS_DISPLAY_NAME = "Owner"
LOCAL_ACCESS_ROLE = "member"


def utc_timestamp() -> str:
    return datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")


def get_auth_connection() -> sqlite3.Connection:
    AUTH_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    connection = sqlite3.connect(AUTH_DB_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def hash_auth_password(password: str, *, salt: str | None = None, iterations: int = 200000) -> str:
    if salt is None:
        salt = os.urandom(16).hex()
    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        bytes.fromhex(salt),
        iterations,
    ).hex()
    return f"pbkdf2_sha256${iterations}${salt}${digest}"


def verify_auth_password(password: str, stored_hash: str) -> bool:
    try:
        algorithm, iteration_text, salt, _expected = stored_hash.split("$", 3)
        if algorithm != "pbkdf2_sha256":
            return False
        iterations = int(iteration_text)
    except Exception:
        return False
    candidate = hash_auth_password(password, salt=salt, iterations=iterations)
    return hmac.compare_digest(candidate, stored_hash)


def normalized_permissions(role: str, permissions: dict[str, bool] | None = None) -> dict[str, bool]:
    base = dict(ROLE_PERMISSIONS.get(role, ROLE_PERMISSIONS["viewer"]))
    if permissions:
        for key, value in permissions.items():
            if key in base:
                base[key] = bool(value)
    return base


def auth_permissions_json(role: str, permissions: dict[str, bool] | None = None) -> str:
    return json.dumps(normalized_permissions(role, permissions), sort_keys=True)


def parse_permissions_json(role: str, raw_value: str | None) -> dict[str, bool]:
    if not raw_value:
        return normalized_permissions(role)
    try:
        loaded = json.loads(raw_value)
        if isinstance(loaded, dict):
            return normalized_permissions(role, {str(k): bool(v) for k, v in loaded.items()})
    except Exception:
        pass
    return normalized_permissions(role)


def auth_log_event(event_type: str, details: str, username: str = "") -> None:
    created_at = utc_timestamp()
    with get_auth_connection() as connection:
        connection.execute(
            """
            INSERT INTO auth_logs (username, event_type, details, created_at)
            VALUES (?, ?, ?, ?)
            """,
            (username.strip(), event_type.strip(), details.strip(), created_at),
        )
        connection.commit()


def initialize_auth_storage() -> None:
    with get_auth_connection() as connection:
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                display_name TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL,
                is_active INTEGER NOT NULL DEFAULT 1,
                permissions_json TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                last_login_at TEXT
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS auth_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT,
                event_type TEXT NOT NULL,
                details TEXT NOT NULL,
                created_at TEXT NOT NULL
            )
            """
        )
        connection.commit()


def auth_user_row_to_dict(row: sqlite3.Row | dict[str, object] | None) -> dict[str, object] | None:
    if row is None:
        return None
    data = dict(row)
    role = str(data.get("role", "viewer") or "viewer")
    return {
        "id": int(data.get("id", 0) or 0),
        "username": str(data.get("username", "") or ""),
        "display_name": str(data.get("display_name", "") or ""),
        "password_hash": str(data.get("password_hash", "") or ""),
        "role": role,
        "is_active": bool(int(data.get("is_active", 0) or 0)),
        "permissions": parse_permissions_json(role, str(data.get("permissions_json", "") or "")),
        "created_at": str(data.get("created_at", "") or ""),
        "updated_at": str(data.get("updated_at", "") or ""),
        "last_login_at": str(data.get("last_login_at", "") or ""),
    }


def auth_user_count() -> int:
    initialize_auth_storage()
    with get_auth_connection() as connection:
        row = connection.execute("SELECT COUNT(*) AS count_value FROM users").fetchone()
    return int(row["count_value"] if row else 0)


def list_auth_users() -> list[dict[str, object]]:
    initialize_auth_storage()
    with get_auth_connection() as connection:
        rows = connection.execute("SELECT * FROM users ORDER BY username COLLATE NOCASE").fetchall()
    users: list[dict[str, object]] = []
    for row in rows:
        user = auth_user_row_to_dict(row)
        if user:
            users.append(user)
    return users


def get_user_by_id(user_id: int) -> dict[str, object] | None:
    initialize_auth_storage()
    with get_auth_connection() as connection:
        row = connection.execute("SELECT * FROM users WHERE id = ?", (int(user_id),)).fetchone()
    return auth_user_row_to_dict(row)


def get_user_by_username(username: str) -> dict[str, object] | None:
    cleaned = username.strip()
    if not cleaned:
        return None
    initialize_auth_storage()
    with get_auth_connection() as connection:
        row = connection.execute("SELECT * FROM users WHERE lower(username) = lower(?)", (cleaned,)).fetchone()
    return auth_user_row_to_dict(row)


def active_admin_count() -> int:
    initialize_auth_storage()
    with get_auth_connection() as connection:
        row = connection.execute(
            """
            SELECT COUNT(*) AS count_value
            FROM users
            WHERE is_active = 1 AND role IN ('admin', 'super_admin')
            """
        ).fetchone()
    return int(row["count_value"] if row else 0)


def create_auth_user(username: str, display_name: str, password: str, role: str = "member") -> tuple[bool, str]:
    initialize_auth_storage()
    cleaned_username = username.strip()
    cleaned_display = display_name.strip()
    cleaned_role = role.strip() or "member"
    if cleaned_role not in ROLE_PERMISSIONS:
        return False, "Choose a valid role."
    if not cleaned_username:
        return False, "Username is required."
    if not cleaned_display:
        return False, "Display name is required."
    if len(password or "") < 4:
        return False, "Password must be at least 4 characters."
    if get_user_by_username(cleaned_username):
        return False, "That username already exists."

    now = utc_timestamp()
    with get_auth_connection() as connection:
        connection.execute(
            """
            INSERT INTO users (
                username, display_name, password_hash, role, is_active,
                permissions_json, created_at, updated_at, last_login_at
            )
            VALUES (?, ?, ?, ?, 1, ?, ?, ?, NULL)
            """,
            (
                cleaned_username,
                cleaned_display,
                hash_auth_password(password),
                cleaned_role,
                auth_permissions_json(cleaned_role),
                now,
                now,
            ),
        )
        connection.commit()
    auth_log_event("create_user", f"Created {cleaned_role} account.", cleaned_username)
    return True, f"User '{cleaned_username}' created."


def build_local_access_user() -> dict[str, object]:
    role = LOCAL_ACCESS_ROLE
    return {
        "id": 0,
        "username": LOCAL_ACCESS_USERNAME,
        "display_name": LOCAL_ACCESS_DISPLAY_NAME,
        "password_hash": "",
        "role": role,
        "is_active": True,
        "permissions": normalized_permissions(role),
        "created_at": "",
        "updated_at": "",
        "last_login_at": "",
        "access_mode": "local",
    }


def enable_local_access_mode() -> None:
    current_mode = str(st.session_state.get("auth_access_mode", "") or "")
    if current_mode == "local" and bool(st.session_state.get("is_authenticated", False)):
        return
    apply_auth_session(build_local_access_user())


def using_local_access_mode() -> bool:
    return str(st.session_state.get("auth_access_mode", "") or "") == "local"


def apply_auth_session(user: dict[str, object] | None) -> None:
    if not user:
        st.session_state.is_authenticated = False
        st.session_state.auth_user_id = 0
        st.session_state.auth_username = ""
        st.session_state.auth_display_name = ""
        st.session_state.auth_role = ""
        st.session_state.auth_permissions = {}
        st.session_state.auth_access_mode = "signed_out"
        return
    st.session_state.is_authenticated = True
    st.session_state.auth_user_id = int(user.get("id", 0) or 0)
    st.session_state.auth_username = str(user.get("username", "") or "")
    st.session_state.auth_display_name = str(user.get("display_name", "") or "")
    st.session_state.auth_role = str(user.get("role", "") or "")
    st.session_state.auth_permissions = dict(user.get("permissions", {}) or {})
    st.session_state.auth_access_mode = str(user.get("access_mode", "account") or "account")


def authenticate_user(username: str, password: str) -> tuple[bool, str]:
    initialize_auth_storage()
    cleaned_username = username.strip()
    if not cleaned_username or not password:
        return False, "Username and password are required."
    user = get_user_by_username(cleaned_username)
    if not user:
        auth_log_event("login_failed", "Unknown username.", cleaned_username)
        return False, "Invalid username or password."
    if not bool(user.get("is_active", False)):
        auth_log_event("login_blocked", "Inactive account.", cleaned_username)
        return False, "This account is inactive."
    if not verify_auth_password(password, str(user.get("password_hash", "") or "")):
        auth_log_event("login_failed", "Incorrect password.", cleaned_username)
        return False, "Invalid username or password."

    now = utc_timestamp()
    with get_auth_connection() as connection:
        connection.execute(
            "UPDATE users SET last_login_at = ?, updated_at = ? WHERE id = ?",
            (now, now, int(user["id"])),
        )
        connection.commit()
    refreshed = get_user_by_id(int(user["id"])) or user
    apply_auth_session(refreshed)
    auth_log_event("login_success", "Signed in successfully.", cleaned_username)
    return True, "Login successful."


def auth_log_rows(limit: int = 20) -> list[dict[str, str]]:
    initialize_auth_storage()
    with get_auth_connection() as connection:
        rows = connection.execute(
            "SELECT username, event_type, details, created_at FROM auth_logs ORDER BY id DESC LIMIT ?",
            (int(limit),),
        ).fetchall()
    return [
        {
            "username": str(row["username"] or "-"),
            "event_type": str(row["event_type"] or ""),
            "details": str(row["details"] or ""),
            "created_at": str(row["created_at"] or ""),
        }
        for row in rows
    ]


def can_access_page(page: str) -> bool:
    if not st.session_state.get("is_authenticated", False):
        return False
    permissions = dict(st.session_state.get("auth_permissions", {}) or {})
    return bool(permissions.get(page, False))


def update_auth_user(
    user_id: int,
    display_name: str,
    role: str,
    is_active: bool,
    permissions: dict[str, bool],
) -> tuple[bool, str]:
    initialize_auth_storage()
    user = get_user_by_id(int(user_id))
    if not user:
        return False, "User not found."
    cleaned_role = role.strip()
    if cleaned_role not in ROLE_PERMISSIONS:
        return False, "Choose a valid role."
    cleaned_display = display_name.strip()
    if not cleaned_display:
        return False, "Display name is required."
    if user["role"] == "super_admin" and st.session_state.get("auth_role") != "super_admin":
        return False, "Only a super admin can edit another super admin."
    if not is_active and user["role"] in {"admin", "super_admin"} and active_admin_count() <= 1:
        return False, "You cannot disable the last active admin."
    if cleaned_role not in {"admin", "super_admin"} and user["role"] in {"admin", "super_admin"} and active_admin_count() <= 1:
        return False, "You cannot demote the last active admin."

    now = utc_timestamp()
    with get_auth_connection() as connection:
        connection.execute(
            """
            UPDATE users
            SET display_name = ?, role = ?, is_active = ?, permissions_json = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                cleaned_display,
                cleaned_role,
                1 if is_active else 0,
                auth_permissions_json(cleaned_role, permissions),
                now,
                int(user_id),
            ),
        )
        connection.commit()
    auth_log_event("update_user", f"Updated account settings for {user['username']}.", str(user["username"]))
    return True, f"Updated '{user['username']}'."


def update_auth_password(user_id: int, new_password: str) -> tuple[bool, str]:
    initialize_auth_storage()
    user = get_user_by_id(int(user_id))
    if not user:
        return False, "User not found."
    if len(new_password or "") < 4:
        return False, "Password must be at least 4 characters."
    now = utc_timestamp()
    with get_auth_connection() as connection:
        connection.execute(
            "UPDATE users SET password_hash = ?, updated_at = ? WHERE id = ?",
            (hash_auth_password(new_password), now, int(user_id)),
        )
        connection.commit()
    auth_log_event("password_reset", "Password updated.", str(user["username"]))
    return True, f"Password updated for '{user['username']}'."


def logout_current_user() -> None:
    if using_local_access_mode():
        st.session_state.flash_message = "Direct local access is enabled. Logout is disabled."
        st.session_state.flash_level = "info"
        st.rerun()

    username = str(st.session_state.get("auth_username", "") or "")
    if username:
        auth_log_event("logout", "Signed out.", username)
    apply_auth_session(None)
    st.session_state.active_page = "hugyoku_chat"
    st.session_state.flash_message = "Signed out successfully."
    st.session_state.flash_level = "success"
    st.rerun()

SYSTEM_PROMPT = (
    "You are Hugyoku, an educational assistant for study use. Help the user analyze "
    "activities, quizzes, assignments, uploaded files, documents, essays, and code "
    "issues. You may produce sample answers, draft essays, structured summaries, "
    "guided explanations, generated study documents, and debugging help for offline "
    "learning and review. Keep outputs clear and well-organized, and avoid assisting "
    "with live cheating or deceptive conduct."
)

TOOL_FOLDERS = {
    "quiz": "quiz_solver",
    "assignment": "assignment_solver",
    "essay": "essay_generator",
    "activity": "activity_generator",
    "document": "document_generator",
    "codefix": "code_error_fixer",
}

KNOWN_EXPORT_LABELS = {
    "heading suggestion",
    "title suggestion",
    "suggested heading",
    "essay body",
    "essay",
    "self-check tip",
    "quick self-check tip",
    "closing self-check tip",
    "task understanding",
    "response",
    "short explanation or rationale",
    "summary",
    "objective",
    "instructions",
    "activity proper",
    "answer guide",
    "document type",
    "purpose",
    "main content",
    "solution plan",
    "sample answer or draft",
    "notes to review",
    "issue summary",
    "root cause",
    "fixed version",
    "why it works",
    "next checks",
}

STATE_DEFAULTS = {
    "active_page": "dashboard",
    "saved_name": "",
    "saved_include_date": False,
    "save_destination_mode": "browser",
    "main_folder_name": "hugyoku_exports",
    "export_root_path": "",
    "output_include_name": True,
    "output_include_date": False,
    "essay_include_heading": True,
    "essay_include_tip": True,
    "profile_name_input": "",
    "profile_include_date_input": False,
    "profile_save_destination_mode_input": "browser",
    "profile_main_folder_input": "hugyoku_exports",
    "profile_export_root_path_input": "",
    "profile_output_include_name_input": True,
    "profile_output_include_date_input": False,
    "profile_essay_include_heading_input": True,
    "profile_essay_include_tip_input": True,
    "quiz_upload_name": "No file loaded yet",
    "quiz_source_text": "",
    "quiz_summary": "",
    "quiz_prompt": "",
    "quiz_response": "",
    "quiz_mode": "complete",
    "assignment_upload_name": "No assignment file loaded yet",
    "assignment_source_text": "",
    "assignment_summary": "",
    "assignment_prompt": "",
    "assignment_response": "",
    "assignment_mode": "guided",
    "essay_title": "",
    "essay_prompt": "",
    "essay_word_count": 500,
    "essay_tagalog": False,
    "essay_english": True,
    "essay_specific_name": "",
    "essay_response": "",
    "activity_title": "",
    "activity_type": "Worksheet",
    "activity_level": "",
    "activity_prompt": "",
    "activity_response": "",
    "document_title": "",
    "document_type": "Study Handout",
    "document_audience": "",
    "document_prompt": "",
    "document_response": "",
    "codefix_title": "",
    "codefix_language": "Python",
    "codefix_error": "",
    "codefix_source": "",
    "codefix_expectation": "",
    "codefix_response": "",
    "pending_reset_action": "",
    "pending_reset_notice": "",
    "pending_reset_level": "success",
    "pending_export_root_selection": None,
    "flash_message": "",
    "flash_level": "success",
}

PAGE_DETAILS = {
    "dashboard": {
        "title": "Dashboard",
        "subtitle": "Save your identity, choose one export folder, and keep all Hugyoku outputs organized for local saves or browser downloads.",
    },
    "academics": {
        "title": "Academics",
        "subtitle": "Choose a focused school workspace. Each tool keeps the same Hugyoku content flow but is adapted for browser-based use.",
    },
    "developer": {
        "title": "Developer",
        "subtitle": "A separate hub for code support so the academic tools stay clean and easier to scan.",
    },
    "quiz": {
        "title": "Quiz Solver",
        "subtitle": "Read quiz files, summarize the task, and generate guided response support in a dedicated web workspace.",
    },
    "assignment": {
        "title": "Assignment Solver",
        "subtitle": "Analyze assignment instructions, understand the task first, then generate a guided response or full draft.",
    },
    "essay": {
        "title": "Essay Generator",
        "subtitle": "Build essay drafts with the same Hugyoku output controls, optional export name, and Word export formatting.",
    },
    "activity": {
        "title": "Activity Generator",
        "subtitle": "Generate worksheets, reflections, drills, or school activities from a topic and instruction set.",
    },
    "document": {
        "title": "Document Generator",
        "subtitle": "Create structured school documents like handouts, reports, reviewers, and formal academic materials.",
    },
    "codefix": {
        "title": "Code Error Fixer",
        "subtitle": "Paste code and the error details, then generate a cleaner fix and explanation in a separate developer workspace.",
    },
}

THEME_CSS = """
<style>
:root {
  --bg-main: #0a1018;
  --bg-top: #18233a;
  --panel: rgba(18, 25, 37, 0.94);
  --panel-soft: rgba(25, 35, 52, 0.9);
  --panel-muted: rgba(31, 41, 59, 0.76);
  --border: rgba(100, 118, 148, 0.34);
  --border-strong: rgba(146, 177, 232, 0.36);
  --text: #f5efe1;
  --muted: #b9c4d8;
  --accent: #ff8b72;
  --accent-soft: #f4dca7;
  --success: #53c795;
  --shadow: 0 20px 50px rgba(2, 8, 18, 0.34);
}

html,
body,
[class*="css"] {
  color: var(--text);
}

.stApp {
  background: radial-gradient(circle at top left, var(--bg-top) 0%, var(--bg-main) 48%, #090d14 100%);
  color: var(--text);
  overflow-x: hidden;
}

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #101722 0%, #0c121b 100%);
  border-right: 1px solid rgba(100, 118, 148, 0.22);
}

[data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
  gap: 0.95rem;
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, rgba(22, 30, 44, 0.95) 0%, rgba(15, 21, 33, 0.95) 100%);
  border: 1px solid rgba(100, 118, 148, 0.25) !important;
  border-radius: 22px !important;
  box-shadow: none;
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: 0.95rem 1rem;
}

.main .block-container {
  max-width: 1480px;
  padding-top: clamp(0.75rem, 1.5vw, 1.1rem);
  padding-bottom: 1.8rem;
  padding-left: clamp(0.8rem, 2vw, 1.4rem);
  padding-right: clamp(0.8rem, 2vw, 1.4rem);
}

[data-testid="stHorizontalBlock"] {
  gap: 0.8rem;
  align-items: stretch;
  flex-wrap: wrap;
}

[data-testid="column"] {
  min-width: min(100%, 280px) !important;
  flex: 1 1 300px !important;
}

[data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, var(--panel) 0%, rgba(14, 20, 31, 0.97) 100%);
  border: 1px solid var(--border) !important;
  border-radius: 24px !important;
  box-shadow: var(--shadow);
  backdrop-filter: blur(16px);
  height: 100%;
  transition: transform 0.18s ease, border-color 0.18s ease, box-shadow 0.18s ease;
}

[data-testid="stVerticalBlockBorderWrapper"]:hover {
  transform: translateY(-2px);
  border-color: var(--border-strong) !important;
  box-shadow: 0 24px 56px rgba(0, 0, 0, 0.28);
}

[data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: clamp(0.8rem, 0.85vw + 0.62rem, 1.08rem);
  height: 100%;
}

.stMarkdown,
.stText,
p,
li,
label,
span,
div {
  overflow-wrap: anywhere;
}

p,
li,
label,
.stCaption,
.stMarkdown {
  color: var(--muted);
}

h1,
h2,
h3,
h4,
h5,
h6 {
  color: var(--text);
  letter-spacing: -0.02em;
}

.app-hero-title {
  font-size: clamp(2rem, 4vw, 3.2rem);
  line-height: 1.1;
  color: var(--text);
  font-weight: 800;
  margin-bottom: 0.45rem;
}

.app-card-header {
  display: flex;
  flex-direction: column;
  gap: 0.28rem;
  margin-bottom: 0.72rem;
}

.app-kicker {
  font-size: 0.75rem;
  text-transform: uppercase;
  letter-spacing: 0.16em;
  font-weight: 800;
  color: var(--accent-soft);
}

.app-card-title {
  font-size: clamp(1.2rem, 2vw, 1.85rem);
  line-height: 1.16;
  font-weight: 800;
  color: var(--text);
}

.app-card-subtitle {
  font-size: 0.98rem;
  line-height: 1.65;
  color: var(--muted);
  margin: 0;
}

.app-chip-row {
  display: flex;
  flex-wrap: wrap;
  gap: 0.45rem;
  margin-top: 0.75rem;
}

.app-chip {
  display: inline-flex;
  align-items: center;
  padding: 0.38rem 0.72rem;
  border-radius: 999px;
  border: 1px solid rgba(243, 215, 154, 0.2);
  background: rgba(243, 215, 154, 0.1);
  color: var(--accent-soft);
  font-size: 0.82rem;
  font-weight: 700;
}

.app-status-pill {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  padding: 0.45rem 0.82rem;
  border-radius: 999px;
  font-size: 0.84rem;
  font-weight: 800;
  margin-bottom: 1rem;
}

.app-status-pill.ready {
  background: rgba(83, 199, 149, 0.18);
  color: #d6ffee;
  border: 1px solid rgba(83, 199, 149, 0.26);
}

.app-status-pill.waiting {
  background: rgba(228, 182, 94, 0.16);
  color: #ffe7bb;
  border: 1px solid rgba(228, 182, 94, 0.22);
}

.app-status-pill.offline {
  background: rgba(227, 123, 116, 0.16);
  color: #ffd4cf;
  border: 1px solid rgba(227, 123, 116, 0.2);
}

.app-meta-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
  gap: 0.75rem;
  margin-bottom: 1rem;
  width: 100%;
}

.app-meta-item {
  padding: 0.72rem 0.82rem;
  border-radius: 16px;
  background: var(--panel-muted);
  border: 1px solid rgba(100, 118, 148, 0.22);
  min-width: 0;
  height: 100%;
}

.app-meta-label {
  font-size: 0.74rem;
  text-transform: uppercase;
  letter-spacing: 0.14em;
  color: var(--accent-soft);
  margin-bottom: 0.35rem;
}

.app-meta-value {
  color: var(--text);
  font-weight: 700;
  line-height: 1.45;
  word-break: break-word;
  overflow-wrap: anywhere;
}

.app-route-block {
  padding: 0.72rem 0.86rem;
  border-radius: 18px;
  background: var(--panel-muted);
  border: 1px solid rgba(100, 118, 148, 0.22);
  margin-bottom: 0.62rem;
}

.app-route-label {
  font-size: 0.75rem;
  text-transform: uppercase;
  letter-spacing: 0.14em;
  color: var(--accent-soft);
  margin-bottom: 0.4rem;
}

.app-route-body {
  color: var(--muted);
  font-family: Consolas, "Courier New", monospace;
  font-size: 0.88rem;
  line-height: 1.6;
  white-space: pre-wrap;
}

.app-guide-list {
  margin: 0;
  padding-left: 1.1rem;
  color: var(--muted);
}

.app-guide-list li {
  margin-bottom: 0.55rem;
  line-height: 1.65;
}

.app-soft-note {
  color: var(--muted);
  font-size: 0.9rem;
  line-height: 1.6;
}

.app-section-gap {
  height: 0.35rem;
}

.app-sidebar-brand {
  font-size: 1.8rem;
  font-weight: 800;
  letter-spacing: -0.03em;
  color: var(--text);
  margin-bottom: 0.35rem;
}

.app-sidebar-copy {
  color: var(--muted);
  line-height: 1.7;
  margin-bottom: 0.4rem;
}

code {
  color: #ffe9bd;
}

.stButton > button,
.stDownloadButton > button {
  width: 100%;
  min-height: 2.85rem;
  border-radius: 16px;
  border: 1px solid var(--border);
  background: linear-gradient(180deg, rgba(36, 47, 67, 0.96) 0%, rgba(20, 28, 42, 0.98) 100%);
  color: var(--text);
  font-weight: 700;
  transition: transform 0.16s ease, border-color 0.16s ease, box-shadow 0.16s ease, background 0.16s ease;
}

.stButton > button:hover,
.stDownloadButton > button:hover {
  transform: translateY(-1px);
  border-color: rgba(244, 220, 167, 0.28);
  box-shadow: 0 14px 26px rgba(0, 0, 0, 0.24);
}

.stButton > button:focus,
.stDownloadButton > button:focus {
  box-shadow: 0 0 0 0.12rem rgba(244, 220, 167, 0.18);
}

div[data-testid="stBaseButton-primary"] > button,
.stDownloadButton > button[kind="primary"] {
  background: linear-gradient(135deg, var(--accent) 0%, var(--accent-soft) 100%);
  color: #111723;
  border: none;
  box-shadow: 0 14px 28px rgba(255, 139, 114, 0.28);
}

div[data-testid="stBaseButton-primary"] > button:hover,
.stDownloadButton > button[kind="primary"]:hover {
  box-shadow: 0 18px 34px rgba(255, 139, 114, 0.32);
}

.stButton > button p,
.stDownloadButton > button p {
  color: inherit !important;
}

[data-baseweb="input"] > div,
[data-baseweb="textarea"] > div,
[data-baseweb="select"] > div,
.stNumberInput > div > div {
  background: rgba(12, 18, 28, 0.88) !important;
  border: 1px solid rgba(90, 108, 137, 0.34) !important;
  border-radius: 16px !important;
}

[data-baseweb="input"] input,
[data-baseweb="textarea"] textarea,
.stNumberInput input {
  color: var(--text) !important;
  font-size: 0.98rem !important;
  line-height: 1.55 !important;
}

[data-baseweb="input"] input::placeholder,
[data-baseweb="textarea"] textarea::placeholder {
  color: #7f8aa0 !important;
}

[data-testid="stFileUploader"] section {
  background: rgba(12, 18, 28, 0.72);
  border: 1px dashed rgba(100, 118, 148, 0.4);
  border-radius: 18px;
}

[data-testid="stTabs"] [data-baseweb="tab-list"] {
  gap: 0.45rem;
  flex-wrap: wrap;
}

[data-testid="stTabs"] [data-baseweb="tab"] {
  min-height: 2.55rem;
  border-radius: 999px;
  padding-inline: 0.95rem;
}

div[role="radiogroup"] {
  gap: 0.55rem;
  flex-wrap: wrap;
}

div[role="radiogroup"] label {
  border: 1px solid rgba(100, 118, 148, 0.3);
  border-radius: 999px;
  padding: 0.3rem 0.8rem;
  background: rgba(18, 25, 38, 0.75);
}

@media (max-width: 1024px) {
  .main .block-container {
    padding-left: 0.85rem;
    padding-right: 0.85rem;
  }
}

@media (max-width: 920px) {
  [data-testid="stHorizontalBlock"] {
    flex-direction: column !important;
    align-items: stretch !important;
    gap: 0.8rem !important;
  }

  [data-testid="column"] {
    width: 100% !important;
    max-width: 100% !important;
    min-width: 100% !important;
    flex-basis: 100% !important;
    flex: 1 1 100% !important;
  }

  .app-hero-title {
    font-size: clamp(1.85rem, 8vw, 2.7rem);
  }

  .app-meta-grid {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }
}

@media (max-width: 768px) {
  [data-testid="stSidebar"] {
    min-width: 15rem;
    max-width: 15rem;
  }

  [data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] > div {
    padding: 0.75rem 0.8rem;
  }

  .app-card-title {
    font-size: 1.08rem;
  }

  .app-card-subtitle,
  .app-soft-note,
  .app-route-body {
    font-size: 0.88rem;
  }

  .stButton > button,
  .stDownloadButton > button {
    min-height: 3rem;
    font-size: 0.94rem;
  }

  .app-sidebar-brand {
    font-size: 1.45rem;
  }

  .app-sidebar-copy {
    font-size: 0.8rem;
    line-height: 1.45;
  }

  [data-testid="stTabs"] [data-baseweb="tab"] {
    width: 100%;
    justify-content: center;
  }

  .app-chip {
    width: 100%;
    justify-content: center;
  }

  .app-meta-grid {
    grid-template-columns: 1fr;
  }

  .app-meta-item,
  .app-route-block {
    padding: 0.7rem 0.78rem;
  }
}

@media (max-width: 520px) {
  .main .block-container {
    padding-left: 0.65rem;
    padding-right: 0.65rem;
    padding-top: 0.65rem;
  }

  [data-testid="stVerticalBlockBorderWrapper"] > div {
    padding: 0.7rem;
  }

  .app-hero-title {
    font-size: 1.85rem;
  }

  .app-card-subtitle,
  .app-route-body,
  .app-soft-note,
  .app-sidebar-copy {
    font-size: 0.84rem;
    line-height: 1.5;
  }

  .app-chip-row {
    gap: 0.35rem;
  }

  .app-chip {
    padding: 0.38rem 0.62rem;
    font-size: 0.76rem;
  }

  [data-baseweb="input"] input,
  [data-baseweb="textarea"] textarea,
  .stNumberInput input {
    font-size: 0.92rem !important;
  }
}

/* Premium 90+ polish overrides */

.stApp {
  background:
    radial-gradient(circle at 8% 8%, rgba(255, 166, 114, 0.12) 0%, rgba(255, 166, 114, 0.03) 20%, transparent 42%),
    radial-gradient(circle at 92% 12%, rgba(116, 176, 255, 0.12) 0%, rgba(116, 176, 255, 0.03) 24%, transparent 44%),
    linear-gradient(180deg, #0c111b 0%, #080d15 55%, #060a12 100%);
}

.main .block-container {
  max-width: 1380px;
  padding-top: clamp(0.78rem, 1vw, 1.02rem);
  padding-bottom: 2rem;
}

[data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, rgba(17, 24, 37, 0.97) 0%, rgba(10, 15, 24, 0.98) 100%);
  border: 1px solid rgba(118, 136, 168, 0.26) !important;
  border-radius: 22px !important;
  box-shadow: 0 16px 34px rgba(2, 8, 18, 0.24);
}

[data-testid="stVerticalBlockBorderWrapper"]:hover {
  transform: translateY(-2px);
  border-color: rgba(182, 202, 236, 0.42) !important;
  box-shadow: 0 20px 42px rgba(2, 8, 18, 0.28);
}

[data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: clamp(0.82rem, 0.82vw + 0.6rem, 1.04rem);
}

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, rgba(10, 15, 24, 0.98) 0%, rgba(7, 11, 18, 0.99) 100%);
  border-right: 1px solid rgba(96, 114, 148, 0.16);
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, rgba(18, 25, 38, 0.94) 0%, rgba(12, 18, 29, 0.96) 100%);
  border-radius: 18px !important;
}

.app-sidebar-brand-row {
  display: flex;
  align-items: flex-start;
  gap: 0.8rem;
  margin-bottom: 0.8rem;
}

.app-sidebar-brand {
  font-size: 1.38rem;
  line-height: 1.05;
  margin-bottom: 0.14rem;
}

.app-sidebar-copy {
  font-size: 0.84rem;
  line-height: 1.45;
  margin-bottom: 0;
}

.app-sidebar-quiet {
  color: rgba(185, 196, 216, 0.66);
  font-size: 0.75rem;
  line-height: 1.4;
}

.app-nav-section {
  margin: 0.18rem 0 0.48rem;
  font-size: 0.72rem;
  text-transform: uppercase;
  letter-spacing: 0.18em;
  color: rgba(244, 220, 167, 0.78);
  font-weight: 800;
}

.app-card-topline {
  display: flex;
  align-items: center;
  gap: 0.55rem;
  flex-wrap: wrap;
  margin-bottom: 0.52rem;
}

.app-anchor-badge {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-width: 2rem;
  height: 2rem;
  padding: 0 0.55rem;
  border-radius: 12px;
  background: linear-gradient(135deg, rgba(255, 139, 114, 0.18) 0%, rgba(244, 220, 167, 0.18) 100%);
  border: 1px solid rgba(244, 220, 167, 0.26);
  color: var(--text);
  font-size: 0.76rem;
  font-weight: 900;
  letter-spacing: 0.08em;
  flex-shrink: 0;
}

.app-tier-pill {
  display: inline-flex;
  align-items: center;
  padding: 0.26rem 0.66rem;
  border-radius: 999px;
  font-size: 0.72rem;
  font-weight: 800;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}

.app-tier-pill.primary {
  color: #111723;
  background: linear-gradient(135deg, var(--accent) 0%, var(--accent-soft) 100%);
}

.app-tier-pill.secondary {
  color: #f7e7c0;
  background: rgba(244, 220, 167, 0.1);
  border: 1px solid rgba(244, 220, 167, 0.2);
}

.app-tier-pill.tertiary {
  color: #d4dded;
  background: rgba(118, 136, 168, 0.14);
  border: 1px solid rgba(118, 136, 168, 0.2);
}

.app-card-header {
  gap: 0.2rem;
  margin-bottom: 0.68rem;
}

.app-card-title {
  font-size: clamp(1.12rem, 1.4vw, 1.58rem);
  line-height: 1.15;
}

.app-card-subtitle {
  max-width: 74ch;
  font-size: 0.93rem;
  line-height: 1.58;
}

.app-card-subtitle.compact {
  font-size: 0.86rem;
  line-height: 1.48;
}

.app-hero-shell {
  display: flex;
  flex-direction: column;
  gap: 0.9rem;
}

.app-hero-title {
  font-size: clamp(2.18rem, 4vw, 3.5rem);
  line-height: 1.02;
  letter-spacing: -0.03em;
  max-width: 12ch;
  margin-bottom: 0;
}

.app-hero-lead {
  color: var(--muted);
  font-size: 0.98rem;
  line-height: 1.62;
  max-width: 72ch;
}

.app-kpi-row {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
  gap: 0.62rem;
  margin-top: 0.12rem;
}

.app-kpi-pill {
  min-width: 0;
  padding: 0.74rem 0.8rem;
  border-radius: 16px;
  background: linear-gradient(180deg, rgba(25, 34, 50, 0.94) 0%, rgba(17, 24, 36, 0.96) 100%);
  border: 1px solid rgba(112, 130, 162, 0.18);
}

.app-kpi-label {
  display: block;
  margin-bottom: 0.3rem;
  color: rgba(244, 220, 167, 0.84);
  font-size: 0.7rem;
  letter-spacing: 0.14em;
  text-transform: uppercase;
  font-weight: 800;
}

.app-kpi-value {
  display: block;
  color: var(--text);
  font-size: 0.94rem;
  font-weight: 800;
  line-height: 1.4;
  overflow-wrap: anywhere;
}

.app-chip-row {
  gap: 0.42rem;
  margin-top: 0.1rem;
}

.app-chip {
  padding: 0.42rem 0.78rem;
  background: linear-gradient(180deg, rgba(33, 43, 62, 0.94) 0%, rgba(22, 30, 45, 0.96) 100%);
  border: 1px solid rgba(243, 215, 154, 0.18);
  transition: transform 0.18s ease, border-color 0.18s ease, background 0.18s ease;
}

.app-chip:hover {
  transform: translateY(-1px);
  border-color: rgba(243, 215, 154, 0.28);
  background: linear-gradient(180deg, rgba(44, 57, 80, 0.96) 0%, rgba(26, 35, 52, 0.98) 100%);
}

.app-status-pill {
  padding: 0.44rem 0.86rem;
  margin-bottom: 0.82rem;
}

.app-meta-grid {
  gap: 0.62rem;
  margin-bottom: 0.82rem;
}

.app-meta-item {
  padding: 0.74rem 0.82rem;
  background: linear-gradient(180deg, rgba(34, 45, 65, 0.9) 0%, rgba(26, 35, 52, 0.96) 100%);
  border-radius: 18px;
}

.app-route-block {
  padding: 0.74rem 0.86rem;
  background: linear-gradient(180deg, rgba(28, 38, 56, 0.92) 0%, rgba(20, 28, 42, 0.96) 100%);
  border-radius: 18px;
  margin-bottom: 0.55rem;
}

.app-route-label {
  margin-bottom: 0.36rem;
}

.app-route-body {
  font-size: 0.84rem;
  line-height: 1.56;
}

[data-testid="stExpander"] {
  border: 1px solid rgba(100, 118, 148, 0.18);
  border-radius: 18px;
  background: rgba(16, 22, 34, 0.58);
  overflow: hidden;
}

[data-testid="stExpander"] details {
  border: none !important;
}

[data-testid="stExpander"] summary {
  font-weight: 800 !important;
  color: var(--text) !important;
}

.stButton > button,
.stDownloadButton > button {
  min-height: 2.92rem;
  border-radius: 14px;
  border: 1px solid rgba(128, 146, 180, 0.22);
  background: linear-gradient(180deg, rgba(33, 44, 63, 0.98) 0%, rgba(19, 27, 40, 0.98) 100%);
  box-shadow: 0 10px 22px rgba(3, 9, 18, 0.18);
}

.stButton > button:hover,
.stDownloadButton > button:hover {
  transform: translateY(-1px) scale(1.002);
  border-color: rgba(194, 213, 245, 0.3);
  box-shadow: 0 14px 28px rgba(3, 9, 18, 0.22);
}

div[data-testid="stBaseButton-primary"] > button,
.stDownloadButton > button[kind="primary"] {
  background: linear-gradient(135deg, #ff8f74 0%, #ffd88f 100%);
  color: #0d1420;
  font-weight: 900;
  letter-spacing: 0.01em;
  box-shadow: 0 14px 28px rgba(255, 139, 114, 0.24);
}

div[data-testid="stBaseButton-primary"] > button:hover,
.stDownloadButton > button[kind="primary"]:hover {
  box-shadow: 0 18px 34px rgba(255, 139, 114, 0.3);
}

[data-baseweb="input"] > div,
[data-baseweb="textarea"] > div,
[data-baseweb="select"] > div,
.stNumberInput > div > div {
  border-radius: 14px !important;
  background: rgba(11, 17, 27, 0.94) !important;
}

[data-testid="stTabs"] [data-baseweb="tab"] {
  min-height: 2.45rem;
  padding-inline: 0.88rem;
  font-weight: 700;
}

@media (max-width: 920px) {
  .app-hero-title {
    max-width: none;
  }

  .app-kpi-row {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }
}

@media (max-width: 768px) {
  .app-sidebar-brand-row {
    gap: 0.65rem;
  }

  .app-anchor-badge {
    min-width: 1.75rem;
    height: 1.75rem;
    border-radius: 10px;
    font-size: 0.7rem;
  }

  .app-kpi-row {
    grid-template-columns: 1fr;
  }

  .app-card-subtitle {
    max-width: none;
  }
}

@media (max-width: 520px) {
  .app-tier-pill {
    font-size: 0.66rem;
    letter-spacing: 0.1em;
  }

  .app-hero-title {
    font-size: 1.98rem;
  }

  .app-kpi-pill,
  .app-meta-item,
  .app-route-block {
    padding: 0.68rem 0.72rem;
  }
}

/* Elite 95+ product-grade overrides */

.main .block-container {
  max-width: 1440px;
  padding-top: clamp(0.68rem, 0.72vw, 0.9rem);
}

[data-testid="stVerticalBlockBorderWrapper"] {
  background:
    linear-gradient(180deg, rgba(15, 22, 34, 0.985) 0%, rgba(9, 14, 24, 0.985) 100%),
    linear-gradient(135deg, rgba(255, 143, 116, 0.05) 0%, transparent 42%);
  border: 1px solid rgba(128, 147, 181, 0.28) !important;
  border-radius: 24px !important;
  box-shadow:
    0 20px 44px rgba(2, 8, 18, 0.28),
    inset 0 1px 0 rgba(255, 255, 255, 0.03);
}

[data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: clamp(0.84rem, 0.78vw + 0.5rem, 1.08rem);
}

[data-testid="stVerticalBlockBorderWrapper"]:hover {
  box-shadow:
    0 24px 52px rgba(2, 8, 18, 0.3),
    inset 0 1px 0 rgba(255, 255, 255, 0.04);
}

[data-testid="stSidebar"] {
  background:
    radial-gradient(circle at 18% 0%, rgba(255, 147, 117, 0.08) 0%, transparent 28%),
    linear-gradient(180deg, rgba(10, 15, 24, 0.99) 0%, rgba(6, 10, 17, 0.995) 100%);
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
  border-radius: 20px !important;
  box-shadow: 0 14px 28px rgba(0, 0, 0, 0.18);
}

[data-testid="stSidebar"] .stButton > button {
  min-height: 2.62rem;
  border-radius: 15px;
  font-size: 0.92rem;
  font-weight: 760;
}

.app-hero-shell {
  display: grid;
  gap: 0.88rem;
}

.app-hero-title {
  max-width: 10ch;
  font-size: clamp(3rem, 4.45vw, 4.65rem);
  line-height: 0.92;
  letter-spacing: -0.055em;
  font-weight: 900;
  color: #fff4df;
  text-wrap: balance;
  margin-bottom: 0;
}

.app-hero-lead {
  max-width: 58ch;
  font-size: 1rem;
  line-height: 1.6;
  color: rgba(223, 231, 246, 0.9);
}

.app-hero-note {
  display: inline-flex;
  align-items: center;
  gap: 0.45rem;
  padding: 0.34rem 0.7rem;
  width: fit-content;
  border-radius: 999px;
  background: rgba(255, 255, 255, 0.03);
  border: 1px solid rgba(141, 160, 193, 0.2);
  color: rgba(209, 219, 238, 0.8);
  font-size: 0.8rem;
  font-weight: 700;
}

.app-luxury-rule {
  width: 100%;
  height: 1px;
  background: linear-gradient(90deg, rgba(255, 143, 116, 0.4) 0%, rgba(244, 220, 167, 0.2) 32%, rgba(244, 220, 167, 0.06) 72%, transparent 100%);
  margin: 0.1rem 0 0.15rem;
}

.app-kpi-row {
  gap: 0.72rem;
}

.app-kpi-pill {
  gap: 0.26rem;
  min-height: 4.4rem;
  border-radius: 18px;
  background:
    linear-gradient(180deg, rgba(28, 38, 56, 0.9) 0%, rgba(20, 28, 42, 0.94) 100%);
  border: 1px solid rgba(122, 143, 177, 0.26);
}

.app-kpi-label {
  color: rgba(244, 220, 167, 0.82);
  font-size: 0.68rem;
}

.app-kpi-value {
  font-size: 1.04rem;
  line-height: 1.28;
  color: #fff6e7;
}

.app-card-title {
  font-size: clamp(1.34rem, 1.25vw, 1.72rem);
  letter-spacing: -0.03em;
}

.app-card-subtitle {
  max-width: 62ch;
  color: rgba(207, 218, 237, 0.84);
}

.app-card-subtitle.compact {
  max-width: 58ch;
  font-size: 0.89rem;
}

.app-card-topline {
  margin-bottom: 0.62rem;
}

.app-anchor-badge {
  min-width: 2.05rem;
  height: 2.05rem;
  border-radius: 13px;
  box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.04);
}

.app-tier-pill {
  border-radius: 999px;
  font-size: 0.7rem;
  font-weight: 820;
}

.app-tier-pill.primary {
  box-shadow: 0 8px 20px rgba(255, 139, 114, 0.18);
}

.app-tier-pill.secondary,
.app-tier-pill.tertiary {
  background: rgba(255, 255, 255, 0.035);
  border: 1px solid rgba(244, 220, 167, 0.16);
}

.app-chip-row {
  gap: 0.42rem;
}

.app-chip {
  padding: 0.42rem 0.72rem;
  border-radius: 999px;
  background: rgba(255, 255, 255, 0.04);
  border: 1px solid rgba(244, 220, 167, 0.15);
  color: #f5e3bd;
  font-size: 0.76rem;
  font-weight: 760;
}

.app-meta-grid {
  gap: 0.7rem;
}

.app-meta-item {
  min-height: 5.2rem;
  border-radius: 18px;
  background: linear-gradient(180deg, rgba(30, 40, 59, 0.88) 0%, rgba(24, 32, 47, 0.94) 100%);
  border: 1px solid rgba(120, 139, 173, 0.24);
}

.app-meta-value {
  font-size: 1rem;
  line-height: 1.34;
  color: #fff4df;
}

.app-route-block {
  border-radius: 18px;
  border: 1px solid rgba(122, 143, 177, 0.24);
  background:
    linear-gradient(180deg, rgba(31, 42, 62, 0.88) 0%, rgba(22, 31, 47, 0.94) 100%);
}

.app-route-label {
  font-size: 0.7rem;
}

.app-route-body {
  color: rgba(232, 238, 248, 0.92);
}

.stButton > button,
.stDownloadButton > button {
  min-height: 2.92rem;
  border-radius: 17px;
  border: 1px solid rgba(124, 143, 176, 0.26);
  background:
    linear-gradient(180deg, rgba(38, 49, 70, 0.98) 0%, rgba(22, 30, 45, 0.98) 100%);
  box-shadow:
    0 12px 24px rgba(2, 8, 18, 0.18),
    inset 0 1px 0 rgba(255, 255, 255, 0.04);
  font-weight: 760;
  letter-spacing: 0.01em;
}

.stButton > button:hover,
.stDownloadButton > button:hover {
  transform: translateY(-2px);
  box-shadow:
    0 18px 30px rgba(2, 8, 18, 0.24),
    inset 0 1px 0 rgba(255, 255, 255, 0.05);
}

div[data-testid="stBaseButton-primary"] > button,
.stDownloadButton > button[kind="primary"] {
  background: linear-gradient(135deg, #ff936f 0%, #ffd08a 100%);
  color: #10161f;
  box-shadow:
    0 16px 32px rgba(255, 139, 114, 0.26),
    inset 0 1px 0 rgba(255, 255, 255, 0.18);
}

div[data-testid="stBaseButton-primary"] > button:hover,
.stDownloadButton > button[kind="primary"]:hover {
  box-shadow:
    0 20px 38px rgba(255, 139, 114, 0.32),
    inset 0 1px 0 rgba(255, 255, 255, 0.2);
}

[data-testid="stExpander"] {
  border-radius: 18px;
  border: 1px solid rgba(124, 143, 176, 0.22);
  background: rgba(255, 255, 255, 0.022);
  overflow: hidden;
}

[data-testid="stExpander"] summary {
  font-weight: 760;
}

@media (max-width: 1100px) {
  .app-hero-title {
    max-width: none;
    font-size: clamp(2.55rem, 6.2vw, 4rem);
  }
}

@media (max-width: 920px) {
  .app-meta-grid {
    grid-template-columns: 1fr 1fr;
  }

  .app-kpi-row {
    grid-template-columns: 1fr 1fr;
  }
}

@media (max-width: 768px) {
  [data-testid="stSidebar"] {
    min-width: 14rem;
    max-width: 14rem;
  }

  .app-hero-title {
    font-size: 2.34rem;
    line-height: 0.96;
  }

  .app-hero-lead {
    font-size: 0.92rem;
  }

  .app-meta-grid,
  .app-kpi-row {
    grid-template-columns: 1fr;
  }
}

@media (max-width: 520px) {
  .app-hero-title {
    font-size: 2.02rem;
  }

  .app-hero-note {
    width: 100%;
    justify-content: center;
    font-size: 0.75rem;
  }

  .app-chip {
    width: 100%;
    justify-content: center;
  }
}

</style>
"""


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", name).strip(" ._")
    return cleaned or "document"


def normalize_package_folder_name(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return "hugyoku_exports"
    if any(sep in raw for sep in ("\\", "/")):
        parts = [part for part in re.split(r"[\\/]+", raw) if part.strip()]
        raw = parts[-1] if parts else raw
    if re.fullmatch(r"[A-Za-z]:", raw):
        raw = "hugyoku_exports"
    return sanitize_filename(raw)


def looks_like_local_path(value: str) -> bool:
    raw = (value or "").strip()
    return "\\" in raw or "/" in raw or bool(re.match(r"^[A-Za-z]:", raw))


def pick_local_folder() -> str | None:
    if tk is None or filedialog is None:
        raise RuntimeError("Local folder picking needs Tkinter support on this machine.")

    root: object | None = None
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        selected = filedialog.askdirectory(title="Choose Hugyoku save folder", mustexist=False)
        return selected or None
    except Exception as exc:
        raise RuntimeError("Could not open the local folder picker in this environment.") from exc
    finally:
        if root is not None:
            try:
                root.destroy()
            except Exception:
                pass


def open_in_file_manager(path: str) -> None:
    target = str(Path(path))
    if os.name == "nt":
        os.startfile(target)
    elif sys.platform == "darwin":
        subprocess.Popen(["open", target])
    else:
        subprocess.Popen(["xdg-open", target])


def current_package_root_name() -> str:
    export_root = (st.session_state.get("export_root_path") or "").strip()
    if export_root:
        return sanitize_filename(Path(export_root).name)
    return sanitize_filename(st.session_state.get("main_folder_name") or "hugyoku_exports")


def local_folder_picker_available() -> bool:
    return tk is not None and filedialog is not None


def save_destination_mode() -> str:
    mode = str(st.session_state.get("save_destination_mode") or "browser").strip().lower()
    return mode if mode in {"browser", "local"} else "browser"


def local_save_active() -> bool:
    return save_destination_mode() == "local" and bool((st.session_state.get("export_root_path") or "").strip())


def browser_export_label() -> str:
    return "Browser-managed download (.docx file)"


def today_string() -> str:
    now = datetime.now()
    return f"{now.strftime('%B')} {now.day}, {now.year}"


def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text))


def collapse_paragraph_lines(lines: list[str]) -> str:
    return " ".join(line.strip() for line in lines if line.strip()).strip()


def split_text_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n\s*\n", text.strip()) if block.strip()]


def strip_heading_markers(text: str) -> str:
    cleaned = re.sub(r"^#+\s*", "", text.strip())
    return cleaned.strip().strip(":").strip()


def heading_level_from_line(text: str) -> int:
    stripped = text.lstrip()
    if stripped.startswith("#"):
        return min(max(len(stripped) - len(stripped.lstrip("#")), 1), 4)
    return 2


def normalize_section_label(text: str) -> str:
    cleaned = strip_heading_markers(text)
    cleaned = re.sub(r"^\d+\.\s*", "", cleaned)
    return re.sub(r"\s+", " ", cleaned).lower().strip()


def parse_structured_blocks(text: str) -> list[dict[str, object]]:
    items: list[dict[str, object]] = []
    for block in split_text_blocks(text):
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        first = lines[0].strip()
        normalized_first = normalize_section_label(first)
        is_heading = first.startswith("#") or first.endswith(":") or normalized_first in KNOWN_EXPORT_LABELS
        if is_heading:
            items.append(
                {
                    "type": "section",
                    "heading": strip_heading_markers(first),
                    "level": heading_level_from_line(first),
                    "content": collapse_paragraph_lines(lines[1:]),
                }
            )
        else:
            items.append({"type": "paragraph", "content": collapse_paragraph_lines(lines)})
    return items


def add_body_paragraph(document: Document, text: str, indent_first_line: bool = False, italic: bool = False) -> None:
    if not text.strip():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(8)
    if indent_first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
    run = paragraph.add_run(text.strip())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = italic


def render_docx_bytes(
    title: str,
    body: str,
    category: str = "generic",
    metadata_lines: list[str] | None = None,
    output_options: dict[str, bool] | None = None,
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    metadata_lines = metadata_lines or []
    output_options = output_options or {}
    export_title = title.strip() or "Untitled Document"

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(export_title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(48, 87, 138)

    if metadata_lines:
        meta_paragraph = document.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_paragraph.paragraph_format.space_after = Pt(10)
        meta_run = meta_paragraph.add_run(" | ".join(metadata_lines))
        meta_run.italic = True
        meta_run.font.name = "Times New Roman"
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(88, 88, 88)

    parsed_blocks = parse_structured_blocks(body)
    essay_subtitle: str | None = None
    essay_tip: str | None = None
    body_items: list[dict[str, object]] = []

    for item in parsed_blocks:
        if item["type"] == "section":
            label = normalize_section_label(str(item["heading"]))
            content = str(item.get("content", "")).strip()
            if category == "essay":
                if label in {"heading suggestion", "title suggestion", "suggested heading"}:
                    if content:
                        essay_subtitle = content
                    continue
                if label in {"essay body", "essay"}:
                    if content:
                        body_items.append({"type": "paragraph", "content": content})
                    continue
                if label in {"self-check tip", "quick self-check tip", "closing self-check tip"}:
                    if content:
                        essay_tip = content
                    continue
            body_items.append(item)
        else:
            body_items.append(item)

    if category == "essay" and output_options.get("essay_include_heading", True) and essay_subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.paragraph_format.space_after = Pt(12)
        subtitle_run = subtitle_paragraph.add_run(essay_subtitle)
        subtitle_run.italic = True
        subtitle_run.font.name = "Times New Roman"
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    for item in body_items:
        if item["type"] == "section":
            heading_text = str(item["heading"]).strip().rstrip(":")
            if heading_text:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(heading_text)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(60, 76, 104)
            content = str(item.get("content", "")).strip()
            if content:
                add_body_paragraph(document, content, indent_first_line=(category == "essay"))
        else:
            add_body_paragraph(document, str(item["content"]), indent_first_line=(category == "essay"))

    if category == "essay" and output_options.get("essay_include_tip", True) and essay_tip:
        tip_heading = document.add_paragraph()
        tip_heading.paragraph_format.space_before = Pt(10)
        tip_heading.paragraph_format.space_after = Pt(2)
        tip_run = tip_heading.add_run("Self-Check Tip")
        tip_run.bold = True
        tip_run.font.name = "Times New Roman"
        tip_run.font.size = Pt(13)
        tip_run.font.color.rgb = RGBColor(60, 76, 104)
        add_body_paragraph(document, essay_tip, italic=True)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def extract_completion_text(message: object) -> str | None:
    content = getattr(message, "content", None)
    if isinstance(content, str):
        return content.strip() or None
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, str) and item.strip():
                parts.append(item.strip())
            elif isinstance(item, dict):
                text = item.get("text") or item.get("content")
                if isinstance(text, str) and text.strip():
                    parts.append(text.strip())
            else:
                text = getattr(item, "text", None) or getattr(item, "content", None)
                if isinstance(text, str) and text.strip():
                    parts.append(text.strip())
        joined = "\n\n".join(parts).strip()
        return joined or None
    return None


def read_secret(name: str) -> str:
    if ENV_PATH.exists():
        load_dotenv(ENV_PATH)
    try:
        secret_value = st.secrets.get(name)
    except Exception:
        secret_value = None
    if isinstance(secret_value, str) and secret_value.strip():
        return secret_value.strip()
    return os.getenv(name, "").strip()


def load_client() -> tuple[object | None, str | None, str | None]:
    token = read_secret("HF_TOKEN")
    model = read_secret("HF_MODEL")

    if InferenceClient is None:
        return None, None, "Install the packages in requirements.txt to enable AI features."
    if not token or not model:
        return None, None, "Add HF_TOKEN and HF_MODEL to Streamlit secrets or local .env to enable AI features."

    return InferenceClient(api_key=token), model, None


def generate_text(prompt: str, label: str, client: object | None = None, model: str | None = None) -> str:
    local_client = client
    local_model = model
    if local_client is None or local_model is None:
        local_client, local_model, error = load_client()
        if error:
            raise RuntimeError(error)

    try:
        completion = local_client.chat.completions.create(
            model=local_model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1400,
            temperature=0.45,
        )
    except Exception as exc:
        raise RuntimeError(f"Could not generate {label}: {exc}") from exc

    message = completion.choices[0].message
    result = extract_completion_text(message)
    if not result:
        raise RuntimeError("The AI model returned an empty response.")
    return result


def read_docx_text(file_obj: io.BytesIO) -> str:
    document = Document(file_obj)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs).strip()


def read_pdf_text(file_obj: io.BytesIO) -> str:
    if PdfReader is None:
        raise RuntimeError("PDF reading needs the pypdf package. Install requirements.txt first.")

    reader = PdfReader(file_obj)
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    joined = "\n\n".join(parts).strip()
    if not joined:
        raise RuntimeError("The PDF did not contain readable text.")
    return joined


def read_text_file(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


def read_uploaded_document(uploaded_file: object) -> str:
    name = getattr(uploaded_file, "name", "uploaded.txt")
    suffix = Path(name).suffix.lower()
    data = uploaded_file.getvalue()
    if suffix == ".docx":
        return read_docx_text(io.BytesIO(data))
    if suffix == ".pdf":
        return read_pdf_text(io.BytesIO(data))
    if suffix in {".txt", ".md"}:
        return read_text_file(data)
    raise RuntimeError("Unsupported file type. Use .docx, .pdf, .txt, or .md.")

def ensure_state() -> None:
    for key, value in STATE_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value


def go(page: str) -> None:
    st.session_state.active_page = page
    st.rerun()


def save_profile() -> None:
    st.session_state.saved_name = st.session_state.profile_name_input.strip()
    st.session_state.saved_include_date = bool(st.session_state.profile_include_date_input)
    mode = str(st.session_state.profile_save_destination_mode_input or "browser").strip().lower()
    if mode not in {"browser", "local"}:
        mode = "browser"
    st.session_state.save_destination_mode = mode
    selected_path = (st.session_state.profile_export_root_path_input or "").strip()
    if mode == "local" and selected_path:
        normalized_folder = sanitize_filename(Path(selected_path).name)
        st.session_state.export_root_path = selected_path
    else:
        normalized_folder = normalize_package_folder_name(st.session_state.profile_main_folder_input)
        st.session_state.export_root_path = ""
    st.session_state.main_folder_name = normalized_folder
    st.session_state.profile_main_folder_input = normalized_folder
    st.session_state.output_include_name = bool(st.session_state.profile_output_include_name_input)
    st.session_state.output_include_date = bool(st.session_state.profile_output_include_date_input)
    st.session_state.essay_include_heading = bool(st.session_state.profile_essay_include_heading_input)
    st.session_state.essay_include_tip = bool(st.session_state.profile_essay_include_tip_input)


def clear_profile() -> None:
    st.session_state.saved_name = ""
    st.session_state.saved_include_date = False
    st.session_state.save_destination_mode = "browser"
    st.session_state.main_folder_name = "hugyoku_exports"
    st.session_state.export_root_path = ""
    st.session_state.output_include_name = False
    st.session_state.output_include_date = False
    st.session_state.essay_include_heading = True
    st.session_state.essay_include_tip = True
    st.session_state.profile_name_input = ""
    st.session_state.profile_include_date_input = False
    st.session_state.profile_save_destination_mode_input = "browser"
    st.session_state.profile_main_folder_input = "hugyoku_exports"
    st.session_state.profile_export_root_path_input = ""
    st.session_state.profile_output_include_name_input = False
    st.session_state.profile_output_include_date_input = False
    st.session_state.profile_essay_include_heading_input = True
    st.session_state.profile_essay_include_tip_input = True


def identity_block() -> str:
    lines: list[str] = []
    if st.session_state.saved_name:
        lines.append(f"Name: {st.session_state.saved_name}")
    if st.session_state.saved_include_date:
        lines.append(f"Date: {today_string()}")
    return "\n".join(lines)


def current_output_settings() -> dict[str, bool]:
    return {
        "include_name": bool(st.session_state.output_include_name),
        "include_date": bool(st.session_state.output_include_date),
        "essay_include_heading": bool(st.session_state.essay_include_heading),
        "essay_include_tip": bool(st.session_state.essay_include_tip),
    }


def export_metadata_lines(category: str = "generic", name_override: str | None = None) -> list[str]:
    lines: list[str] = []
    selected_name = st.session_state.saved_name
    if category == "essay" and name_override and name_override.strip():
        selected_name = name_override.strip()
    if st.session_state.output_include_name and selected_name:
        lines.append(selected_name)
    if st.session_state.output_include_date:
        lines.append(today_string())
    return lines


def folder_path_lines() -> dict[str, str]:
    export_root = (st.session_state.export_root_path or "").strip() if save_destination_mode() == "local" else ""
    paths = {"main": export_root or browser_export_label()}
    for category, folder in TOOL_FOLDERS.items():
        if export_root:
            paths[category] = str(Path(export_root) / folder)
        else:
            paths[category] = browser_export_label()
    return paths


def build_export_document(
    title: str,
    body: str,
    default_name: str,
    category: str,
    name_override: str | None = None,
) -> tuple[bytes, str, str | None]:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = sanitize_filename(default_name)
    docx_filename = f"{stem}_{stamp}.docx"
    internal_folder = TOOL_FOLDERS.get(category, category)
    docx_bytes = render_docx_bytes(
        title,
        body.strip(),
        category=category,
        metadata_lines=export_metadata_lines(category=category, name_override=name_override),
        output_options=current_output_settings(),
    )

    export_root = (st.session_state.export_root_path or "").strip() if save_destination_mode() == "local" else ""
    local_docx_path = str(Path(export_root) / internal_folder / docx_filename) if export_root else None
    return docx_bytes, docx_filename, local_docx_path


def save_local_export_file(docx_bytes: bytes, local_docx_path: str) -> str:
    target_path = Path(local_docx_path)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_bytes(docx_bytes)
    return str(target_path)


def apply_selected_export_folder(selected_folder: str | None) -> None:
    if selected_folder:
        normalized_folder = sanitize_filename(Path(selected_folder).name)
        st.session_state.profile_export_root_path_input = selected_folder
        st.session_state.export_root_path = selected_folder
        st.session_state.save_destination_mode = "local"
        st.session_state.profile_save_destination_mode_input = "local"
        st.session_state.main_folder_name = normalized_folder
        st.session_state.profile_main_folder_input = normalized_folder
    else:
        st.session_state.export_root_path = ""
        st.session_state.save_destination_mode = "browser"
        st.session_state.profile_save_destination_mode_input = "browser"
        st.session_state.main_folder_name = "hugyoku_exports"
        st.session_state.profile_main_folder_input = "hugyoku_exports"


def queue_reset(action: str, notice: str = "", level: str = "success") -> None:
    st.session_state.pending_reset_action = action
    st.session_state.pending_reset_notice = notice
    st.session_state.pending_reset_level = level


def queue_export_root_selection(selected_folder: str | None) -> None:
    st.session_state.pending_export_root_selection = "__CLEAR__" if selected_folder is None else selected_folder


def clear_quiz_workspace() -> None:
    st.session_state.quiz_upload_name = "No file loaded yet"
    st.session_state.quiz_source_text = ""
    st.session_state.quiz_summary = ""
    st.session_state.quiz_prompt = ""
    st.session_state.quiz_response = ""
    st.session_state.quiz_mode = "complete"


def clear_assignment_workspace() -> None:
    st.session_state.assignment_upload_name = "No assignment file loaded yet"
    st.session_state.assignment_source_text = ""
    st.session_state.assignment_summary = ""
    st.session_state.assignment_prompt = ""
    st.session_state.assignment_response = ""
    st.session_state.assignment_mode = "guided"


def clear_essay_form() -> None:
    st.session_state.essay_title = ""
    st.session_state.essay_prompt = ""
    st.session_state.essay_word_count = 500
    st.session_state.essay_tagalog = False
    st.session_state.essay_english = True
    st.session_state.essay_specific_name = ""


def clear_essay_result() -> None:
    st.session_state.essay_response = ""


def clear_quiz_result() -> None:
    st.session_state.quiz_response = ""


def clear_assignment_result() -> None:
    st.session_state.assignment_response = ""


def clear_activity_form() -> None:
    st.session_state.activity_title = ""
    st.session_state.activity_type = "Worksheet"
    st.session_state.activity_level = ""
    st.session_state.activity_prompt = ""


def clear_activity_result() -> None:
    st.session_state.activity_response = ""


def clear_document_form() -> None:
    st.session_state.document_title = ""
    st.session_state.document_type = "Study Handout"
    st.session_state.document_audience = ""
    st.session_state.document_prompt = ""


def clear_document_result() -> None:
    st.session_state.document_response = ""


def clear_codefix_form() -> None:
    st.session_state.codefix_title = ""
    st.session_state.codefix_language = "Python"
    st.session_state.codefix_stack_profile = "Python"
    st.session_state.codefix_custom_stack = ""
    st.session_state.codefix_error = ""
    st.session_state.codefix_source = ""
    st.session_state.codefix_expectation = ""


def clear_codefix_result() -> None:
    st.session_state.codefix_response = ""


def clear_login_form() -> None:
    st.session_state.login_username_input = ""
    st.session_state.login_password_input = ""


def clear_bootstrap_form() -> None:
    st.session_state.bootstrap_username_input = ""
    st.session_state.bootstrap_display_name_input = ""
    st.session_state.bootstrap_password_input = ""
    st.session_state.bootstrap_confirm_password_input = ""


def clear_admin_create_user_form() -> None:
    st.session_state.admin_new_username = ""
    st.session_state.admin_new_display_name = ""
    st.session_state.admin_new_password = ""
    st.session_state.admin_new_role = "member"


def clear_admin_password_reset() -> None:
    st.session_state.admin_reset_password = ""


RESET_ACTIONS: dict[str, object] = {
    "clear_profile": clear_profile,
    "clear_quiz_workspace": clear_quiz_workspace,
    "clear_quiz_result": clear_quiz_result,
    "clear_assignment_workspace": clear_assignment_workspace,
    "clear_assignment_result": clear_assignment_result,
    "clear_essay_workspace": lambda: (clear_essay_form(), clear_essay_result()),
    "clear_essay_result": clear_essay_result,
    "clear_activity_workspace": lambda: (clear_activity_form(), clear_activity_result()),
    "clear_activity_result": clear_activity_result,
    "clear_document_workspace": lambda: (clear_document_form(), clear_document_result()),
    "clear_document_result": clear_document_result,
    "clear_codefix_workspace": lambda: (clear_codefix_form(), clear_codefix_result()),
    "clear_codefix_result": clear_codefix_result,
    "clear_login_form": clear_login_form,
    "clear_bootstrap_form": clear_bootstrap_form,
    "clear_admin_create_user_form": clear_admin_create_user_form,
    "clear_admin_password_reset": clear_admin_password_reset,
}


def apply_pending_state_actions() -> None:
    pending_folder = st.session_state.get("pending_export_root_selection")
    if pending_folder is not None:
        apply_selected_export_folder(None if pending_folder == "__CLEAR__" else str(pending_folder))
        st.session_state.pending_export_root_selection = None

    action_name = (st.session_state.get("pending_reset_action") or "").strip()
    if action_name:
        reset_action = RESET_ACTIONS.get(action_name)
        if callable(reset_action):
            reset_action()
        notice = st.session_state.get("pending_reset_notice", "").strip()
        level = st.session_state.get("pending_reset_level", "success")
        if notice:
            st.session_state.flash_message = notice
            st.session_state.flash_level = level
        st.session_state.pending_reset_action = ""
        st.session_state.pending_reset_notice = ""
        st.session_state.pending_reset_level = "success"


def render_flash_message() -> None:
    message = (st.session_state.get("flash_message") or "").strip()
    if not message:
        return
    level = (st.session_state.get("flash_level") or "success").strip().lower()
    if level == "error":
        st.error(message)
    elif level == "warning":
        st.warning(message)
    elif level == "info":
        st.info(message)
    else:
        st.success(message)
    st.session_state.flash_message = ""
    st.session_state.flash_level = "success"


def run_generation(prompt: str, label: str) -> str | None:
    result, _used_model = run_generation_with_details(prompt, label)
    return result


def essay_language() -> str | None:
    tagalog = bool(st.session_state.essay_tagalog)
    english = bool(st.session_state.essay_english)
    if tagalog and english:
        return "Taglish"
    if tagalog:
        return "Tagalog"
    if english:
        return "English"
    return None


def html_text(value: object) -> str:
    if value is None:
        return ""
    return escape(str(value)).replace("\n", "<br>")


# Shared card helpers keep the layout and typography consistent across pages.
def render_card_header(
    title: str,
    subtitle: str,
    kicker: str | None = None,
    *,
    anchor: str | None = None,
    tier: str = "secondary",
    compact: bool = False,
) -> None:
    kicker_text = (kicker or "").strip()
    anchor_html = f"<span class='app-anchor-badge'>{html_text(anchor)}</span>" if anchor else ""
    tier_class = tier if tier in {"primary", "secondary", "tertiary"} else "secondary"
    tier_html = f"<span class='app-tier-pill {tier_class}'>{html_text(kicker_text)}</span>" if kicker_text else ""
    topline_html = f"<div class='app-card-topline'>{anchor_html}{tier_html}</div>" if anchor_html or tier_html else ""
    subtitle_class = "app-card-subtitle compact" if compact else "app-card-subtitle"
    st.markdown(
        f"""
        <div class="app-card-header">
          {topline_html}
          <div class="app-card-title">{html_text(title)}</div>
          <div class="{subtitle_class}">{html_text(subtitle)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_tag_row(tags: list[str]) -> None:
    if not tags:
        return
    chip_html = "".join(f"<span class='app-chip'>{html_text(tag)}</span>" for tag in tags)
    st.markdown(f"<div class='app-chip-row'>{chip_html}</div>", unsafe_allow_html=True)


def render_route_block(title: str, body: str) -> None:
    st.markdown(
        f"""
        <div class="app-route-block">
          <div class="app-route-label">{html_text(title)}</div>
          <div class="app-route-body">{html_text(body)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_meta_grid(items: list[tuple[str, str]]) -> None:
    if not items:
        return
    cards_html = "".join(
        f"<div class='app-meta-item'><div class='app-meta-label'>{html_text(label)}</div><div class='app-meta-value'>{html_text(value)}</div></div>"
        for label, value in items
    )
    st.markdown(f"<div class='app-meta-grid'>{cards_html}</div>", unsafe_allow_html=True)


def render_kpi_row(items: list[tuple[str, str]]) -> None:
    if not items:
        return
    pills_html = "".join(
        f"<div class='app-kpi-pill'><span class='app-kpi-label'>{html_text(label)}</span><span class='app-kpi-value'>{html_text(value)}</span></div>"
        for label, value in items
    )
    st.markdown(f"<div class='app-kpi-row'>{pills_html}</div>", unsafe_allow_html=True)


def render_preview_panel(
    title: str,
    subtitle: str,
    field_label: str,
    value: str,
    *,
    height: int = 360,
    empty_title: str = "Preview will appear here",
    empty_body: str = "Use the workspace on the left to generate a result. The preview and export actions will appear once there is real content to review.",
    anchor: str = "PV",
    tier: str = "tertiary",
) -> bool:
    render_card_header(title, subtitle, "Preview", anchor=anchor, tier=tier, compact=True)
    content = (value or "").strip()
    if content:
        st.text_area(field_label, value=value, height=height, disabled=True)
        return True
    render_route_block(empty_title, empty_body)
    return False






def render_page_intro(page_key: str) -> None:
    details = PAGE_DETAILS[page_key]
    intro_tags = {
        "dashboard": ["Profile settings", "Responsive cards", "Local or browser export"],
        "academics": ["Quiz", "Assignment", "Essay", "Activity", "Document"],
        "developer": ["Code fix", "Root cause", "Debug notes"],
    }
    with st.container(border=True):
        render_card_header(details["title"], details["subtitle"], "Section")
        render_tag_row(intro_tags.get(page_key, []))


def render_tool_hub_card(title: str, subtitle: str, folder_text: str, button_label: str, target: str, primary: bool = False) -> None:
    with st.container(border=True):
        render_card_header(title, subtitle, "Tool Workspace")
        render_route_block("Save route", folder_text)
        if st.button(
            button_label,
            key=f"open_{target}",
            use_container_width=True,
            type="primary" if primary else "secondary",
        ):
            go(target)


def render_back_button(target: str, label: str) -> None:
    if st.button(label, key=f"back_{target}_{st.session_state.active_page}", use_container_width=True):
        go(target)


def render_workspace_header(page_key: str, back_target: str, back_label: str) -> None:
    intro_col, action_col = st.columns([2.2, 1], gap="large")
    with intro_col:
        with st.container(border=True):
            render_card_header(PAGE_DETAILS[page_key]["title"], PAGE_DETAILS[page_key]["subtitle"], "Workspace")
    with action_col:
        with st.container(border=True):
            render_route_block("Save route", folder_path_lines()[page_key])
            render_back_button(back_target, back_label)


def render_download_button(title: str, body: str, default_name: str, category: str, reset_action: str, name_override: str | None = None) -> None:
    docx_bytes, docx_filename, local_docx_path = build_export_document(
        title,
        body,
        default_name,
        category,
        name_override=name_override,
    )
    if local_docx_path:
        render_route_block("Local save route", local_docx_path)
        local_col, open_col = st.columns(2, gap="small")
        if local_col.button(
            "Save To Selected Folder",
            key=f"save_local_{category}_{sanitize_filename(default_name)}",
            use_container_width=True,
            type="primary",
        ):
            try:
                saved_path = save_local_export_file(docx_bytes, local_docx_path)
                queue_reset(reset_action, f"Saved to {saved_path}")
                st.rerun()
            except Exception as exc:
                st.error(f"Could not save to the selected folder: {exc}")
        if open_col.button(
            "Open Selected Folder",
            key=f"open_local_{category}_{sanitize_filename(default_name)}",
            use_container_width=True,
        ):
            try:
                target_folder = Path(local_docx_path).parent
                target_folder.mkdir(parents=True, exist_ok=True)
                open_in_file_manager(str(target_folder))
            except Exception as exc:
                st.error(f"Could not open the selected folder: {exc}")
    else:
        render_route_block("Browser download", f"{docx_filename}\nSaved directly by your browser.")
    st.download_button(
        "Download Word File",
        data=docx_bytes,
        file_name=docx_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_{category}_{sanitize_filename(default_name)}",
        use_container_width=True,
        on_click=queue_reset,
        args=(reset_action, f"Prepared download: {docx_filename}", "info"),
        type="secondary" if local_docx_path else "primary",
    )







def render_quiz_page(ai_ready: bool) -> None:
    render_workspace_header("quiz", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Quiz Intake",
                "Upload a file or paste quiz content directly. The app reads it first before generating support.",
                "Input",
            )
            uploaded = st.file_uploader("Upload quiz file", type=["docx", "pdf", "txt", "md"], key="quiz_upload_widget")
            load_col, clear_col = st.columns(2, gap="small")
            if load_col.button("Load Uploaded File", key="quiz_load_file", use_container_width=True):
                if uploaded is None:
                    st.warning("Upload a file first.")
                else:
                    try:
                        st.session_state.quiz_source_text = read_uploaded_document(uploaded)
                        st.session_state.quiz_upload_name = uploaded.name
                        st.session_state.quiz_summary = ""
                        st.session_state.quiz_response = ""
                        st.session_state.quiz_prompt = ""
                        st.success(f"Loaded {uploaded.name} into the quiz workspace.")
                    except Exception as exc:
                        st.error(str(exc))
            if clear_col.button("Clear Quiz Input", key="quiz_clear_input", use_container_width=True):
                queue_reset("clear_quiz_workspace", "Quiz workspace cleared.")
                st.rerun()

            render_meta_grid(
                [
                    ("Loaded file", st.session_state.quiz_upload_name),
                    (
                        "Content size",
                        f"{count_words(st.session_state.quiz_source_text)} words\n{len(st.session_state.quiz_source_text.strip())} characters",
                    ),
                ]
            )
            st.text_area("Quiz content", key="quiz_source_text", height=320)
            if st.button("Analyze Quiz", key="quiz_analyze", use_container_width=True, disabled=not ai_ready):
                source = st.session_state.quiz_source_text.strip()
                if not source:
                    st.warning("Upload or paste quiz content before analyzing it.")
                else:
                    prompt = (
                        "Analyze the following academic material for educational review.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Source content:\n{source}\n\n"
                        "Return these sections clearly:\n"
                        "1. Summary\n"
                        "2. What task or instructions seem to be required\n"
                        "3. Important topics, clues, or constraints\n"
                        "4. Best next step for the student"
                    )
                    result = run_generation(prompt, "quiz analysis")
                    if result:
                        st.session_state.quiz_summary = result
                        st.success("Quiz analysis complete.")

    with right:
        with st.container(border=True):
            render_card_header(
                "Summary And Task Reading",
                "The assistant summarizes the uploaded content first so you can understand the task before continuing.",
                "Analysis",
            )
            st.text_area("Summary preview", value=st.session_state.quiz_summary, height=220, disabled=True)

        with st.container(border=True):
            render_card_header(
                "Quiz Guidance",
                "Choose whether to solve the whole task or steer the output with a specific prompt.",
                "Output",
            )
            quiz_mode = st.radio(
                "Response mode",
                options=["complete", "specific"],
                index=0 if st.session_state.quiz_mode == "complete" else 1,
                horizontal=True,
                format_func=lambda item: "Do it in full" if item == "complete" else "Use specific prompt",
            )
            st.session_state.quiz_mode = quiz_mode
            st.text_area("Specific prompt", key="quiz_prompt", height=120)
            gen_col, clear_col = st.columns(2, gap="small")
            if gen_col.button("Generate Quiz Guidance", key="quiz_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                source = st.session_state.quiz_source_text.strip()
                summary = st.session_state.quiz_summary.strip()
                custom_prompt = st.session_state.quiz_prompt.strip()
                if not source:
                    st.warning("Upload or paste quiz content before generating a response.")
                elif not summary:
                    st.warning("Run the summary step first so the app can read the task before generating a response.")
                elif st.session_state.quiz_mode == "specific" and not custom_prompt:
                    st.warning("Add a specific prompt or switch to full mode.")
                else:
                    instructions = (
                        "Create a complete educational response based on the content. If there are visible questions, answer them in order and include short explanations. If the file describes a task instead of direct questions, complete the task as fully as possible."
                        if st.session_state.quiz_mode == "complete"
                        else "Follow the user's specific prompt while staying grounded in the uploaded content and analysis.\nSpecific prompt: " + custom_prompt
                    )
                    prompt = (
                        "Use the uploaded quiz/activity content for educational support.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Analysis already prepared:\n{summary}\n\n"
                        f"Source content:\n{source}\n\n"
                        f"Instructions:\n{instructions}\n\n"
                        "Structure the answer with:\n"
                        "1. Task understanding\n"
                        "2. Response\n"
                        "3. Short explanation or rationale\n\n"
                        "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
                    )
                    result = run_generation(prompt, "quiz response")
                    if result:
                        st.session_state.quiz_response = result
                        st.success("Quiz guidance ready.")
            if clear_col.button("Clear Result", key="quiz_clear_result", use_container_width=True):
                queue_reset("clear_quiz_result", "Quiz result cleared.")
                st.rerun()
            st.text_area("Generated response", value=st.session_state.quiz_response, height=300, disabled=True)
            if st.session_state.quiz_response.strip():
                export_title = f"Quiz Support: {Path(st.session_state.quiz_upload_name).stem.replace('_', ' ').title()}" if st.session_state.quiz_upload_name != "No file loaded yet" else "Quiz Support"
                export_name = sanitize_filename(Path(st.session_state.quiz_upload_name).stem if st.session_state.quiz_upload_name != "No file loaded yet" else "quiz_support")
                render_download_button(export_title, st.session_state.quiz_response, export_name, "quiz", "clear_quiz_workspace")


def render_assignment_page(ai_ready: bool) -> None:
    render_workspace_header("assignment", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Assignment Intake",
                "Upload a document or paste the assignment manually, then analyze the task before generating guidance.",
                "Input",
            )
            uploaded = st.file_uploader("Upload assignment file", type=["docx", "pdf", "txt", "md"], key="assignment_upload_widget")
            load_col, clear_col = st.columns(2, gap="small")
            if load_col.button("Load Uploaded File", key="assignment_load_file", use_container_width=True):
                if uploaded is None:
                    st.warning("Upload an assignment file first.")
                else:
                    try:
                        st.session_state.assignment_source_text = read_uploaded_document(uploaded)
                        st.session_state.assignment_upload_name = uploaded.name
                        st.session_state.assignment_summary = ""
                        st.session_state.assignment_response = ""
                        st.success(f"Loaded {uploaded.name} into the assignment workspace.")
                    except Exception as exc:
                        st.error(str(exc))
            if clear_col.button("Clear Assignment Input", key="assignment_clear_input", use_container_width=True):
                queue_reset("clear_assignment_workspace", "Assignment workspace cleared.")
                st.rerun()

            render_meta_grid(
                [
                    ("Loaded file", st.session_state.assignment_upload_name),
                    (
                        "Content size",
                        f"{count_words(st.session_state.assignment_source_text)} words\n{len(st.session_state.assignment_source_text.strip())} characters",
                    ),
                ]
            )
            st.text_area("Assignment content", key="assignment_source_text", height=300)
            assignment_mode = st.radio(
                "Assignment mode",
                options=["guided", "complete"],
                index=0 if st.session_state.assignment_mode == "guided" else 1,
                horizontal=True,
                format_func=lambda item: "Guided response" if item == "guided" else "Full draft",
            )
            st.session_state.assignment_mode = assignment_mode
            st.text_area("Specific prompt (optional)", key="assignment_prompt", height=110)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Analyze Assignment", key="assignment_analyze", use_container_width=True, disabled=not ai_ready):
                source = st.session_state.assignment_source_text.strip()
                if not source:
                    st.warning("Upload or paste assignment content before analyzing it.")
                else:
                    prompt = (
                        "Analyze the following assignment for educational review.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Assignment content:\n{source}\n\n"
                        "Return these sections clearly:\n"
                        "1. Assignment summary\n"
                        "2. What is being asked\n"
                        "3. Important requirements or constraints\n"
                        "4. Best approach for the student"
                    )
                    result = run_generation(prompt, "assignment analysis")
                    if result:
                        st.session_state.assignment_summary = result
                        st.success("Assignment analysis complete.")
            if action_b.button("Generate Assignment Guidance", key="assignment_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                source = st.session_state.assignment_source_text.strip()
                summary = st.session_state.assignment_summary.strip()
                custom_prompt = st.session_state.assignment_prompt.strip()
                if not source:
                    st.warning("Upload or paste assignment content before generating a response.")
                elif not summary:
                    st.warning("Run the assignment analysis first.")
                else:
                    instructions = (
                        "Create a full draft response based on the assignment while keeping it clear, organized, and educational."
                        if st.session_state.assignment_mode == "complete"
                        else "Create a guided response that explains how to approach the assignment and includes a sample draft the student can review."
                    )
                    if custom_prompt:
                        instructions = f"{instructions}\nSpecific prompt: {custom_prompt}"
                    prompt = (
                        "Use the uploaded assignment for educational support.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Assignment analysis:\n{summary}\n\n"
                        f"Assignment content:\n{source}\n\n"
                        f"Instructions:\n{instructions}\n\n"
                        "Structure the answer with:\n"
                        "1. Task understanding\n"
                        "2. Solution plan\n"
                        "3. Sample answer or draft\n"
                        "4. Notes to review\n\n"
                        "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
                    )
                    result = run_generation(prompt, "assignment response")
                    if result:
                        st.session_state.assignment_response = result
                        st.success("Assignment guidance ready.")

    with right:
        with st.container(border=True):
            render_card_header(
                "Assignment Summary",
                "The uploaded task is summarized first so the response stays grounded in the actual requirements.",
                "Analysis",
            )
            st.text_area("Assignment summary", value=st.session_state.assignment_summary, height=220, disabled=True)

        with st.container(border=True):
            render_card_header(
                "Assignment Response",
                "The guided response or full draft appears here and can be exported directly to Word.",
                "Output",
            )
            st.text_area("Generated response", value=st.session_state.assignment_response, height=280, disabled=True)
            if st.button("Clear Result", key="assignment_clear_result", use_container_width=True):
                queue_reset("clear_assignment_result", "Assignment result cleared.")
                st.rerun()
            if st.session_state.assignment_response.strip():
                export_title = f"Assignment Support: {Path(st.session_state.assignment_upload_name).stem.replace('_', ' ').title()}" if st.session_state.assignment_upload_name != "No assignment file loaded yet" else "Assignment Support"
                export_name = sanitize_filename(Path(st.session_state.assignment_upload_name).stem if st.session_state.assignment_upload_name != "No assignment file loaded yet" else "assignment_support")
                render_download_button(export_title, st.session_state.assignment_response, export_name, "assignment", "clear_assignment_workspace")

def render_essay_page(ai_ready: bool) -> None:
    render_workspace_header("essay", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Essay Builder",
                "Set the title, add an optional custom instruction, choose the target length, and pick English, Tagalog, or both for Taglish.",
                "Compose",
            )
            st.text_input("Essay title", key="essay_title")
            st.text_area("Specific prompt (optional)", key="essay_prompt", height=160)
            st.number_input("Target word count", min_value=100, max_value=3000, step=50, key="essay_word_count")
            lang_a, lang_b = st.columns(2, gap="small")
            lang_a.checkbox("Tagalog", key="essay_tagalog")
            lang_b.checkbox("English", key="essay_english")
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Essay", key="essay_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.essay_title.strip()
                prompt_text = st.session_state.essay_prompt.strip()
                language = essay_language()
                if not title:
                    st.warning("Enter the essay title before generating.")
                elif language is None:
                    st.warning("Select Tagalog, English, or both for Taglish.")
                else:
                    prompt = (
                        "Write a polished educational essay draft.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Essay title: {title}\n"
                        f"Target length: about {int(st.session_state.essay_word_count)} words\n"
                        f"Language: {language}\n"
                        f"Specific prompt: {prompt_text or 'No extra prompt provided'}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Heading Suggestion:\n"
                        "Essay Body:\n"
                        "Self-Check Tip:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "essay draft")
                    if result:
                        st.session_state.essay_response = result
                        st.success("Essay draft ready.")
            if action_b.button("Clear Essay Form", key="essay_clear_form", use_container_width=True):
                queue_reset("clear_essay_workspace", "Essay workspace cleared.")
                st.rerun()

        with st.container(border=True):
            render_card_header(
                "Essay Export Format",
                "Adjust what gets attached to the final essay output. The specific export name can override the saved dashboard name for this essay only.",
                "Output Settings",
            )
            st.text_input("Specific export name (optional)", key="essay_specific_name")
            opt_left, opt_right = st.columns(2, gap="small")
            with opt_left:
                st.checkbox("Include saved or specific name in export", key="profile_output_include_name_input")
                st.checkbox("Include date in export", key="profile_output_include_date_input")
            with opt_right:
                st.checkbox("Include essay heading suggestion", key="profile_essay_include_heading_input")
                st.checkbox("Include self-check tip", key="profile_essay_include_tip_input")
            st.caption("Save Profile from the dashboard if you want these export settings to become the active session defaults.")

    with right:
        with st.container(border=True):
            render_card_header(
                "Essay Preview",
                "The draft appears here once generated. Export it when the structure and tone look right.",
                "Preview",
            )
            st.text_area("Essay output", value=st.session_state.essay_response, height=560, disabled=True)
            if st.button("Clear Result", key="essay_clear_result", use_container_width=True):
                queue_reset("clear_essay_result", "Essay result cleared.")
                st.rerun()
            if st.session_state.essay_response.strip():
                temp_saved = {
                    "output_include_name": st.session_state.output_include_name,
                    "output_include_date": st.session_state.output_include_date,
                    "essay_include_heading": st.session_state.essay_include_heading,
                    "essay_include_tip": st.session_state.essay_include_tip,
                }
                st.session_state.output_include_name = st.session_state.profile_output_include_name_input
                st.session_state.output_include_date = st.session_state.profile_output_include_date_input
                st.session_state.essay_include_heading = st.session_state.profile_essay_include_heading_input
                st.session_state.essay_include_tip = st.session_state.profile_essay_include_tip_input
                try:
                    render_download_button(
                        st.session_state.essay_title.strip() or "Essay Draft",
                        st.session_state.essay_response,
                        sanitize_filename((st.session_state.essay_title.strip() or "essay_draft").lower().replace(" ", "_")),
                        "essay",
                        "clear_essay_workspace",
                        name_override=st.session_state.essay_specific_name.strip(),
                    )
                finally:
                    st.session_state.output_include_name = temp_saved["output_include_name"]
                    st.session_state.output_include_date = temp_saved["output_include_date"]
                    st.session_state.essay_include_heading = temp_saved["essay_include_heading"]
                    st.session_state.essay_include_tip = temp_saved["essay_include_tip"]


def render_activity_page(ai_ready: bool) -> None:
    render_workspace_header("activity", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Activity Builder",
                "Set the topic, activity type, and optional level details, then generate a ready-to-use school activity.",
                "Compose",
            )
            st.text_input("Activity topic or title", key="activity_title")
            st.text_input("Activity type", key="activity_type")
            st.text_input("Level or class (optional)", key="activity_level")
            st.text_area("Specific instructions", key="activity_prompt", height=180)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Activity", key="activity_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.activity_title.strip()
                activity_type = st.session_state.activity_type.strip() or "Activity"
                level = st.session_state.activity_level.strip()
                request = st.session_state.activity_prompt.strip()
                if not title:
                    st.warning("Enter the activity topic or title before generating.")
                else:
                    prompt = (
                        "Create a structured educational activity for offline study use.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Activity title or topic: {title}\n"
                        f"Activity type: {activity_type}\n"
                        f"Level or class: {level or 'Not specified'}\n"
                        f"Specific request: {request or 'No extra request provided'}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Activity Title:\n"
                        "Objective:\n"
                        "Instructions:\n"
                        "Activity Proper:\n"
                        "Answer Guide:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "activity draft")
                    if result:
                        st.session_state.activity_response = result
                        st.success("Activity ready.")
            if action_b.button("Clear Activity Form", key="activity_clear_form", use_container_width=True):
                queue_reset("clear_activity_workspace", "Activity workspace cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Activity Preview",
                "The generated activity appears here in a clean structure you can export directly to Word.",
                "Preview",
            )
            st.text_area("Generated activity", value=st.session_state.activity_response, height=560, disabled=True)
            if st.button("Clear Result", key="activity_clear_result", use_container_width=True):
                queue_reset("clear_activity_result", "Activity result cleared.")
                st.rerun()
            if st.session_state.activity_response.strip():
                render_download_button(
                    st.session_state.activity_title.strip() or "Activity Draft",
                    st.session_state.activity_response,
                    sanitize_filename((st.session_state.activity_title.strip() or "activity_draft").lower().replace(" ", "_")),
                    "activity",
                    "clear_activity_workspace",
                )


def render_document_page(ai_ready: bool) -> None:
    render_workspace_header("document", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Document Builder",
                "Generate a structured school document such as a handout, report, reviewer, or formal academic write-up.",
                "Compose",
            )
            st.text_input("Document title", key="document_title")
            st.text_input("Document type", key="document_type")
            st.text_input("Audience or purpose (optional)", key="document_audience")
            st.text_area("Specific content request", key="document_prompt", height=180)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Document", key="document_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.document_title.strip()
                doc_type = st.session_state.document_type.strip() or "Study Handout"
                audience = st.session_state.document_audience.strip()
                request = st.session_state.document_prompt.strip()
                if not title:
                    st.warning("Enter the document title before generating.")
                elif not request:
                    st.warning("Describe what the document should contain before generating.")
                else:
                    prompt = (
                        "Create a structured educational document.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Document title: {title}\n"
                        f"Document type: {doc_type}\n"
                        f"Audience or purpose: {audience or 'Not specified'}\n"
                        f"Content request: {request}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Document Type:\n"
                        "Purpose:\n"
                        "Main Content:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "document draft")
                    if result:
                        st.session_state.document_response = result
                        st.success("Document ready.")
            if action_b.button("Clear Document Form", key="document_clear_form", use_container_width=True):
                queue_reset("clear_document_workspace", "Document workspace cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Document Preview",
                "The generated document draft appears here and can be exported directly to Word.",
                "Preview",
            )
            st.text_area("Generated document", value=st.session_state.document_response, height=560, disabled=True)
            if st.button("Clear Result", key="document_clear_result", use_container_width=True):
                queue_reset("clear_document_result", "Document result cleared.")
                st.rerun()
            if st.session_state.document_response.strip():
                render_download_button(
                    st.session_state.document_title.strip() or "Document Draft",
                    st.session_state.document_response,
                    sanitize_filename((st.session_state.document_title.strip() or "document_draft").lower().replace(" ", "_")),
                    "document",
                    "clear_document_workspace",
                )


def render_codefix_page(ai_ready: bool) -> None:
    render_workspace_header("codefix", "developer", "Back To Developer")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Code Fix Builder",
                "Paste the broken code, add the error or symptom, and explain the expected behavior so the app can generate a cleaner fix.",
                "Compose",
            )
            st.text_input("Issue title (optional)", key="codefix_title")
            st.text_input("Language or stack", key="codefix_language")
            st.text_area("Error message or symptoms", key="codefix_error", height=140)
            st.text_area("Code snippet", key="codefix_source", height=240)
            st.text_area("Expected behavior (optional)", key="codefix_expectation", height=110)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Code Fix", key="codefix_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.codefix_title.strip()
                language = st.session_state.codefix_language.strip() or "Code"
                error_text = st.session_state.codefix_error.strip()
                source = st.session_state.codefix_source.strip()
                expectation = st.session_state.codefix_expectation.strip()
                if not source:
                    st.warning("Paste the code snippet before generating a fix.")
                elif not error_text:
                    st.warning("Add the error message or symptoms before generating a fix.")
                else:
                    prompt = (
                        "Fix the following code issue for study and debugging support.\n\n"
                        f"Language or stack: {language}\n"
                        f"Issue title: {title or 'No title provided'}\n"
                        f"Error message or symptoms:\n{error_text}\n\n"
                        f"Code snippet:\n{source}\n\n"
                        f"Expected behavior:\n{expectation or 'Not provided'}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Issue Summary:\n"
                        "Root Cause:\n"
                        "Fixed Version:\n"
                        "Why It Works:\n"
                        "Next Checks:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "code fix")
                    if result:
                        st.session_state.codefix_response = result
                        st.success("Code fix ready.")
            if action_b.button("Clear Code Fixer", key="codefix_clear_form", use_container_width=True):
                queue_reset("clear_codefix_workspace", "Code fixer workspace cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Code Fix Preview",
                "The cleaned-up fix and explanation appear here. Export it if you want to keep debugging notes in Word format.",
                "Preview",
            )
            st.text_area("Fixed result", value=st.session_state.codefix_response, height=560, disabled=True)
            if st.button("Clear Result", key="codefix_clear_result", use_container_width=True):
                queue_reset("clear_codefix_result", "Code fix result cleared.")
                st.rerun()
            if st.session_state.codefix_response.strip():
                export_title = st.session_state.codefix_title.strip() or f"{st.session_state.codefix_language.strip() or 'Code'} Code Fix"
                render_download_button(
                    export_title,
                    st.session_state.codefix_response,
                    sanitize_filename(export_title.lower().replace(" ", "_")),
                    "codefix",
                    "clear_codefix_workspace",
                )


STATE_DEFAULTS.update(
    {
        "workspaces": {},
        "active_workspace_id": "",
        "workspace_new_name_input": "",
        "workspace_new_description_input": "",
        "workspace_analysis_result": "",
        "hugyoku_task_input": "",
        "hugyoku_attachment_note": "",
        "hugyoku_stage": 1,
        "hugyoku_understanding": "",
        "hugyoku_refinement_prompt": "",
        "hugyoku_refinement_round": 0,
        "hugyoku_output_sections": [],
        "hugyoku_output_title": "",
        "hugyoku_output_format": "docx",
        "hugyoku_output_raw": "",
        "hugyoku_result_prompt": "",
        "hugyoku_generation_note": "",
        "hugyoku_last_bundle": "",
        "hugyoku_last_ocr_status": "",
        "hugyoku_chat_messages": [],
        "hugyoku_chat_draft": "",
        "hugyoku_chat_attachment_note": "",
        "hugyoku_chat_controls_open": False,
        "hugyoku_chat_reasoning_mode": "Balanced",
        "hugyoku_chat_model_choice": "Active Session Model",
        "hugyoku_chat_model_custom": "",
        "hugyoku_chat_voice_language": "en-US",
        "hugyoku_chat_upload_nonce": 0,
        "hugyoku_chat_pending_transcript": "",
        "hugyoku_chat_pending_job": None,
        "hugyoku_chat_show_full_history": False,
        "hugyoku_chat_edit_mode": False,
        "hugyoku_chat_edit_text": "",
        "hugyoku_chat_last_bundle": "",
        "hugyoku_chat_last_ocr_status": "",
        "hugyoku_chat_last_used_model": "",
        "hugyoku_chat_last_transcript_text": "",
        "source_lab_question_input": "",
        "source_lab_answer": "",
        "source_lab_analysis": "",
        "reviewer_request_input": "",
        "reviewer_response": "",
        "flashcard_request_input": "",
        "flashcard_response": "",
        "practice_test_request_input": "",
        "practice_test_response": "",
        "answer_checker_question_input": "",
        "answer_checker_reference_input": "",
        "answer_checker_user_answer_input": "",
        "answer_checker_response": "",
        "rubric_title_input": "",
        "rubric_prompt_input": "",
        "rubric_focus_input": ["Clarity", "Structure", "Depth"],
        "rubric_response": "",
        "batch_topic_input": "",
        "batch_request_input": "",
        "batch_response": "",
        "history_entries": [],
        "model_override": "",
        "model_choice_input": "HF Secret Default",
        "model_custom_input": "",
        "is_authenticated": False,
        "auth_user_id": 0,
        "auth_username": "",
        "auth_display_name": "",
        "auth_role": "",
        "auth_permissions": {},
        "auth_access_mode": "signed_out",
        "login_username_input": "",
        "login_password_input": "",
        "bootstrap_username_input": "",
        "bootstrap_display_name_input": "",
        "bootstrap_password_input": "",
        "bootstrap_confirm_password_input": "",
        "admin_new_username": "",
        "admin_new_display_name": "",
        "admin_new_password": "",
        "admin_new_role": "member",
        "admin_selected_user_id": 0,
        "admin_reset_password": "",
        "admin_filter_role": "All",
        "last_generation_model": "",
        "last_generation_label": "",
        "last_generation_status": "",
        "last_generation_time": "",
        "last_generation_note": "",
        "verification_last_model": "",
        "verification_last_status": "",
        "verification_last_time": "",
        "verification_last_note": "",
        "codegen_title": "",
        "codegen_stack_profile": "Python",
        "codegen_custom_stack": "",
        "codegen_description": "",
        "codegen_expectation": "",
        "codegen_attachment_note": "",
        "codegen_response": "",
        "codefix_stack_profile": "Python",
        "codefix_custom_stack": "",
        "selftest_stack_profile": "Python",
        "selftest_custom_stack": "",
        "selftest_response": "",
        "selftest_last_model": "",
        "compare_model_a": "Active Session Model",
        "compare_model_b": "Qwen/Qwen2.5-Coder-7B-Instruct",
        "compare_custom_model_a": "",
        "compare_custom_model_b": "",
        "compare_prompt": "",
        "compare_output_a": "",
        "compare_output_b": "",
        "compare_used_model_a": "",
        "compare_used_model_b": "",
        "export_template": "Academic Classic",
        "settings_history_limit": 120,
    }
)

TOOL_FOLDERS.update(
    {
        "workspace_analysis": "workspace_analysis",
        "grounded_answer": "source_grounded_answers",
        "documentqa": "document_qna",
        "reviewer": "reviewer_generator",
        "flashcards": "flashcard_generator",
        "practice_test": "practice_test_generator",
        "answer_checker": "answer_checker",
        "rubric": "rubric_writer",
        "batch": "batch_output_generator",
        "codegen": "code_generator",
        "hugyoku": "hugyoku_universal",
    }
)

PAGE_DETAILS.update(
    {
        "hugyoku": {
            "title": "Hugyoku",
            "subtitle": "Universal academic command workflow that reads the task, confirms the AI understanding, loops on corrections, then generates an editable final result and export.",
        },
        "hugyoku_chat": {
            "title": "Hugyoku Chat",
            "subtitle": "A responsive chat workspace for free-form academic and developer conversations with files, images, model selection, reasoning controls, and voice-to-text support.",
        },
        "workspaces": {
            "title": "Workspaces",
            "subtitle": "Create project workspaces, collect source files and screenshots, and keep all generated outputs grouped by task.",
        },
        "history": {
            "title": "History",
            "subtitle": "Review recent uploads, generations, and exports across the active session.",
        },
        "settings": {
            "title": "Settings",
            "subtitle": "Control the active AI model, export template, and session-level app behavior.",
        },
        "admin": {
            "title": "Admin",
            "subtitle": "Control user access, roles, and account status for the app before anyone reaches the main workspace.",
        },
    }
)

THEME_CSS += """
<style>
div[data-testid="stChatMessage"] {
  border: 0;
  border-radius: 1.2rem;
  background: linear-gradient(180deg, rgba(18, 22, 31, 0.52), rgba(14, 19, 29, 0.78));
  box-shadow: none;
  padding: 0.12rem 0.18rem;
  max-width: 860px;
  margin: 0 auto 0.95rem;
}

div[data-testid="stChatMessage"]:has(.hugyoku-chat-role-label.user) {
  background: linear-gradient(180deg, rgba(26, 22, 30, 0.7), rgba(17, 20, 31, 0.9));
}

div[data-testid="stChatMessage"]:has(.hugyoku-chat-role-label.assistant) {
  background: linear-gradient(180deg, rgba(18, 24, 36, 0.74), rgba(14, 19, 29, 0.9));
}

div[data-testid="stChatMessage"] [data-testid="stMarkdownContainer"] p,
div[data-testid="stChatMessage"] [data-testid="stMarkdownContainer"] li {
  line-height: 1.72;
}

div[data-testid="stChatMessage"] [data-testid="stMarkdownContainer"] {
  max-width: 760px;
}

.hugyoku-chat-topbar {
  display: flex;
  justify-content: space-between;
  align-items: center;
  gap: 0.8rem;
  padding: 0.15rem 0 0.5rem;
}

.hugyoku-chat-topbar-left,
.hugyoku-chat-topbar-right,
.hugyoku-chat-thread-meta,
.hugyoku-chat-toolbar,
.hugyoku-chat-attachment-list {
  display: flex;
  align-items: center;
  gap: 0.5rem;
  flex-wrap: wrap;
}

.hugyoku-chat-pill,
.hugyoku-chat-chip {
  display: inline-flex;
  align-items: center;
  gap: 0.35rem;
  padding: 0.38rem 0.74rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 214, 160, 0.14);
  background: rgba(255, 196, 126, 0.08);
  color: #f7d2a2;
  font-size: 0.72rem;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}

.hugyoku-chat-thread-tag {
  color: rgba(231, 238, 250, 0.9);
  font-size: 0.78rem;
  font-weight: 600;
}

.hugyoku-chat-topbar-note {
  color: rgba(172, 186, 212, 0.74);
  font-size: 0.84rem;
}

.hugyoku-chat-minibar {
  display: flex;
  justify-content: space-between;
  align-items: center;
  gap: 0.72rem;
  flex-wrap: wrap;
  flex: 0 0 auto;
  margin-bottom: 0.4rem;
}

.hugyoku-chat-kicker {
  color: rgba(160, 174, 196, 0.78);
  font-size: 0.84rem;
}

.hugyoku-chat-empty {
  min-height: 24vh;
  display: flex;
  flex-direction: column;
  align-items: center;
  justify-content: center;
  gap: 0.55rem;
  text-align: center;
  padding: 1.25rem 0.9rem 1.5rem;
}

.hugyoku-chat-empty-icon {
  width: 3rem;
  height: 3rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.16);
  display: inline-flex;
  align-items: center;
  justify-content: center;
  color: #f5f7fb;
  font-size: 1.45rem;
  background: radial-gradient(circle at 30% 30%, rgba(255, 214, 160, 0.12), rgba(255, 255, 255, 0.02));
  box-shadow: 0 20px 44px rgba(0, 0, 0, 0.26);
}

.hugyoku-chat-empty-title {
  color: #f8f3eb;
  font-size: clamp(1.65rem, 3.7vw, 2.8rem);
  font-weight: 700;
  line-height: 1.04;
  letter-spacing: -0.04em;
}

.hugyoku-chat-empty-subtitle {
  color: rgba(194, 204, 223, 0.88);
  font-size: clamp(0.92rem, 1.7vw, 1.1rem);
  line-height: 1.35;
}

.hugyoku-chat-thread-meta {
  margin-bottom: 0.42rem;
}

.hugyoku-chat-thread-meta span {
  color: rgba(155, 170, 197, 0.82);
  font-size: 0.76rem;
  letter-spacing: 0.03em;
}

.hugyoku-chat-role-label {
  display: inline-flex;
  align-items: center;
  min-height: 1.6rem;
  padding: 0.18rem 0.58rem;
  border-radius: 999px;
  font-size: 0.74rem;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
  margin-bottom: 0.32rem;
}

.hugyoku-chat-role-label.user {
  background: rgba(255, 103, 103, 0.12);
  border: 1px solid rgba(255, 103, 103, 0.18);
  color: #ffb0b0;
}

.hugyoku-chat-role-label.assistant {
  background: rgba(255, 198, 110, 0.11);
  border: 1px solid rgba(255, 198, 110, 0.18);
  color: #ffd7a1;
}

.hugyoku-chat-shell {
  display: grid;
  gap: 1rem;
}

.hugyoku-chat-composer-note {
  color: rgba(189, 202, 225, 0.7);
  font-size: 0.84rem;
  line-height: 1.5;
}

.hugyoku-chat-toolbar-grid {
  margin-top: 0.08rem;
  overflow-x: hidden;
  overflow-y: hidden;
  padding-bottom: 0.08rem;
}

.hugyoku-chat-toolbar-grid [data-testid="stHorizontalBlock"] {
  align-items: stretch;
  flex-wrap: nowrap !important;
  gap: 0.32rem;
}

.hugyoku-chat-toolbar-grid [data-testid="column"] {
  min-width: 0 !important;
}

.hugyoku-chat-toolbar-grid [data-testid="column"]:first-child {
  flex: 0 0 auto !important;
  width: auto !important;
}

.hugyoku-chat-toolbar-grid [data-testid="column"]:last-child {
  flex: 1 1 auto !important;
  width: 0 !important;
}

.hugyoku-chat-toolbar-grid [data-testid="column"]:first-child {
  flex: 0 0 auto !important;
  width: auto !important;
}

.hugyoku-chat-toolbar-grid [data-testid="column"]:last-child {
  flex: 1 1 auto !important;
  width: 0 !important;
}

.hugyoku-chat-attachment {
  display: inline-flex;
  align-items: center;
  gap: 0.35rem;
  padding: 0.44rem 0.68rem;
  border-radius: 0.78rem;
  border: 1px solid rgba(116, 134, 168, 0.14);
  background: rgba(17, 23, 37, 0.88);
  color: #dce7fa;
  font-size: 0.8rem;
}

.hugyoku-chat-composer-shell {
  max-width: 820px;
  margin: 0 auto;
  display: grid;
  gap: 0.46rem;
}

.hugyoku-chat-composer-sticky {
  position: relative;
  z-index: 20;
  padding: 0.54rem;
  border-radius: 1.1rem;
  background: linear-gradient(180deg, rgba(10, 14, 22, 0.94), rgba(10, 14, 22, 0.985));
  border: 1px solid rgba(86, 101, 132, 0.18);
  box-shadow: 0 -14px 32px rgba(2, 6, 13, 0.32);
  backdrop-filter: blur(18px);
}

.hugyoku-chat-thread-shell {
  max-width: 840px;
  margin: 0 auto;
  display: flex;
  flex-direction: column;
  min-height: 0;
  flex: 1 1 auto;
  width: 100%;
}

.hugyoku-chat-scroll-region {
  flex: 1 1 auto;
  min-height: 0;
  height: 100%;
  overflow-y: auto;
  overflow-x: hidden;
  padding-right: 0.2rem;
  scrollbar-width: thin;
}

.hugyoku-chat-surface {
  display: grid;
  grid-template-rows: minmax(0, 1fr) auto;
  gap: 0.7rem;
  min-height: 0;
  height: auto;
  flex: 1 1 auto;
  overflow: hidden;
}

body:has(.hugyoku-chat-surface) {
  overflow: hidden;
  min-height: 100dvh;
}

[data-testid="stAppViewContainer"]:has(.hugyoku-chat-surface),
[data-testid="stMain"]:has(.hugyoku-chat-surface),
.main:has(.hugyoku-chat-surface) {
  overflow: hidden;
  min-height: 100dvh;
}

[data-testid="block-container"]:has(.hugyoku-chat-surface),
.main .block-container:has(.hugyoku-chat-surface) {
  display: flex;
  flex-direction: column;
  min-height: calc(100dvh - 0.9rem);
  overflow: hidden;
  padding-bottom: 0.35rem;
  gap: 0.55rem;
}

  .hugyoku-chat-message-counter {
    display: none;
  }

  .hugyoku-chat-thread-meta {
    margin-bottom: 0.28rem;
  }

  .hugyoku-chat-message-counter {
    color: rgba(163, 177, 203, 0.72);
    font-size: 0.62rem;
    margin: 0 0 0.3rem;
  }

  .hugyoku-chat-surface [data-testid="stTextArea"] textarea {
    min-height: 62px;
    border-radius: 1rem;
    padding-top: 0.55rem;
    font-size: 0.9rem;
  }

.hugyoku-chat-surface [data-testid="stFileUploader"] section {
  border-radius: 1rem;
}

.hugyoku-chat-surface [data-testid="stTextArea"] > div,
.hugyoku-chat-surface [data-testid="stTextInput"] > div {
  border-radius: 1.25rem;
}

.hugyoku-chat-surface [data-testid="stPopover"] button {
  min-height: 2.45rem;
  border-radius: 0.82rem;
}

.hugyoku-chat-surface [data-testid="stButton"] button[kind="primary"] {
  min-height: 2.45rem;
}

.hugyoku-chat-surface [data-testid="stButton"] button,
.hugyoku-chat-surface [data-testid="stPopover"] button {
  background: linear-gradient(180deg, rgba(24, 32, 48, 0.96), rgba(18, 25, 38, 0.96));
  border-color: rgba(120, 141, 184, 0.18);
  box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.04);
}

.hugyoku-chat-surface [data-testid="stButton"] button:hover,
.hugyoku-chat-surface [data-testid="stPopover"] button:hover {
  border-color: rgba(255, 214, 160, 0.24);
  transform: translateY(-1px);
}

.hugyoku-chat-surface [data-testid="stSelectbox"] > div,
.hugyoku-chat-surface [data-testid="stTextArea"] > div,
.hugyoku-chat-surface [data-testid="stFileUploader"] section {
  background: rgba(14, 20, 31, 0.94);
}

@media (max-width: 900px) {
  .hugyoku-chat-topbar {
    flex-direction: column;
    align-items: flex-start;
  }

  .hugyoku-chat-minibar {
    align-items: flex-start;
  }

  .hugyoku-chat-toolbar-grid [data-testid="stHorizontalBlock"] {
    gap: 0.42rem;
  }
}

@media (max-width: 640px) {
  .hugyoku-chat-empty {
    min-height: 8vh;
    gap: 0.26rem;
    padding: 0.12rem 0.3rem 0.18rem;
  }

  .hugyoku-chat-composer-shell,
  .hugyoku-chat-thread-shell {
    max-width: 100%;
  }

  .hugyoku-chat-minibar {
    gap: 0.2rem;
    margin-bottom: 0.08rem;
  }

  .hugyoku-chat-topbar-note {
    display: none;
  }

  .hugyoku-chat-topbar-right {
    display: none;
  }

  .hugyoku-chat-topbar-left {
    gap: 0.16rem;
  }

  .hugyoku-chat-pill,
  .hugyoku-chat-chip {
    padding: 0.18rem 0.42rem;
    font-size: 0.58rem;
    letter-spacing: 0.05em;
  }

  .hugyoku-chat-thread-tag {
    font-size: 0.66rem;
  }

  .hugyoku-chat-empty-icon {
    width: 2rem;
    height: 2rem;
    font-size: 1rem;
  }

  .hugyoku-chat-empty-title {
    font-size: 0.98rem;
  }

  .hugyoku-chat-empty-subtitle {
    font-size: 0.68rem;
  }

  .hugyoku-chat-surface {
    height: auto;
    gap: 0.35rem;
  }

  .hugyoku-chat-scroll-region {
    padding-right: 0;
  }

  div[data-testid="stChatMessage"] {
    border-radius: 1rem;
    margin-bottom: 0.7rem;
  }

  div[data-testid="stChatMessage"] [data-testid="stMarkdownContainer"] {
    max-width: 100%;
  }

  .hugyoku-chat-role-label {
    font-size: 0.64rem;
    padding: 0.12rem 0.4rem;
  }

  .hugyoku-chat-thread-meta span {
    font-size: 0.64rem;
  }

  .hugyoku-chat-composer-note {
    display: none;
  }

  .hugyoku-chat-toolbar-grid [data-testid="stHorizontalBlock"] {
    gap: 0.18rem !important;
  }

  .hugyoku-chat-toolbar-grid [data-testid="column"] {
    min-width: 0 !important;
  }

  .hugyoku-chat-toolbar-grid [data-testid="column"]:first-child {
    flex: 0 0 2.9rem !important;
    width: 2.9rem !important;
  }

  .hugyoku-chat-toolbar-grid [data-testid="column"]:last-child {
    flex: 1 1 auto !important;
    width: 0 !important;
  }

  .hugyoku-chat-surface [data-testid="stButton"] button,
  .hugyoku-chat-surface [data-testid="stPopover"] button {
    min-height: 1.5rem;
    padding-left: 0.24rem;
    padding-right: 0.24rem;
    font-size: 0.58rem;
    border-radius: 0.6rem;
  }

  .hugyoku-chat-surface [data-testid="stTextArea"] textarea {
    min-height: 40px;
    padding-top: 0.42rem;
    font-size: 0.76rem;
  }

  .hugyoku-chat-composer-sticky {
    border-radius: 0.85rem 0.85rem 0 0;
    padding: 0.2rem 0.2rem 0.24rem;
  }

  [data-testid="block-container"]:has(.hugyoku-chat-surface),
  .main .block-container:has(.hugyoku-chat-surface) {
    min-height: calc(100dvh - 0.45rem);
    padding-bottom: 0.2rem;
    gap: 0.2rem;
  }
}

/* Hugyoku Chat native mobile override */
body:has(.hugyoku-chat-page.hugyoku-chat-custom) {
  overflow: hidden;
  min-height: 100dvh;
}

[data-testid="stAppViewContainer"]:has(.hugyoku-chat-page.hugyoku-chat-custom),
[data-testid="stMain"]:has(.hugyoku-chat-page.hugyoku-chat-custom),
.main:has(.hugyoku-chat-page.hugyoku-chat-custom) {
  overflow: hidden;
  min-height: 100dvh;
}

[data-testid="block-container"]:has(.hugyoku-chat-page.hugyoku-chat-custom),
.main .block-container:has(.hugyoku-chat-page.hugyoku-chat-custom) {
  display: flex;
  flex-direction: column;
  min-height: calc(100dvh - 0.5rem);
  overflow: hidden;
  padding-bottom: 0.18rem;
  gap: 0.18rem;
}

.hugyoku-chat-page.hugyoku-chat-custom .hugyoku-chat-surface {
  display: grid;
  grid-template-rows: minmax(0, 1fr) auto;
  flex: 1 1 auto;
  min-height: 0;
  height: 100%;
  overflow: hidden;
}

.hugyoku-chat-page.hugyoku-chat-custom .hugyoku-chat-thread-shell {
  max-width: 780px;
  min-height: 0;
  height: 100%;
}

.hugyoku-chat-page.hugyoku-chat-custom .hugyoku-chat-scroll-region {
  display: block;
  min-height: 0;
  height: 100%;
  overflow-y: auto;
  overflow-x: hidden;
  padding-right: 0;
}

.hugyoku-chat-page.hugyoku-chat-custom .hugyoku-chat-composer-shell {
  display: grid !important;
}

.hugyoku-chat-page.hugyoku-chat-custom .hugyoku-chat-toolbar-grid {
  display: block !important;
}

.hugyoku-chat-page.hugyoku-chat-custom .hugyoku-chat-composer-note {
  display: none !important;
}

.hugyoku-chat-page {
  max-width: 860px;
  margin: 0 auto;
  display: flex;
  flex-direction: column;
  gap: 0.62rem;
}

.hugyoku-chat-header {
  display: flex;
  justify-content: flex-start;
  align-items: center;
  gap: 0.42rem;
  padding: 0.04rem 0 0.08rem;
}

.hugyoku-chat-header-left {
  display: flex;
  align-items: center;
  gap: 0.42rem;
  flex-wrap: wrap;
}

.hugyoku-chat-header-right [data-testid="stHorizontalBlock"] {
  gap: 0.35rem;
  align-items: center;
}

.hugyoku-chat-control-row [data-testid="stHorizontalBlock"] {
  gap: 0.4rem;
  align-items: center;
}

.hugyoku-chat-workspace-label {
  color: rgba(229, 236, 248, 0.94);
  font-size: 0.82rem;
  font-weight: 600;
}

.hugyoku-chat-thread-shell {
  max-width: 780px;
  width: 100%;
  margin: 0 auto;
}

.hugyoku-chat-empty {
  min-height: 28vh;
  gap: 0.38rem;
  padding: 0.2rem 0 0.1rem;
}

.hugyoku-chat-empty-icon {
  width: 2.5rem;
  height: 2.5rem;
  font-size: 1.2rem;
  box-shadow: 0 12px 26px rgba(0, 0, 0, 0.22);
}

.hugyoku-chat-empty-title {
  font-size: clamp(1.25rem, 2.6vw, 2rem);
}

.hugyoku-chat-empty-subtitle {
  font-size: 0.85rem;
}

.hugyoku-chat-message-counter {
  color: rgba(163, 177, 203, 0.72);
  font-size: 0.72rem;
  margin: 0 0 0.4rem;
}

div[data-testid="stChatMessage"] {
  background: transparent;
  border: none;
  box-shadow: none;
  padding: 0;
  margin: 0 0 0.72rem;
}

div[data-testid="stChatMessage"] [data-testid="stMarkdownContainer"] {
  max-width: 100%;
}

.hugyoku-chat-bubble {
  max-width: min(100%, 42rem);
  border-radius: 1.1rem;
  border: 1px solid rgba(92, 109, 145, 0.14);
  box-shadow: 0 12px 32px rgba(4, 9, 18, 0.18);
  padding: 0.76rem 0.92rem 0.82rem;
}

.hugyoku-chat-bubble.user {
  margin-left: auto;
  background: linear-gradient(180deg, rgba(25, 18, 26, 0.96), rgba(19, 14, 24, 0.98));
  border-color: rgba(255, 108, 108, 0.1);
}

.hugyoku-chat-bubble.assistant {
  margin-right: auto;
  background: linear-gradient(180deg, rgba(15, 22, 35, 0.98), rgba(10, 16, 28, 0.985));
  border-color: rgba(255, 198, 110, 0.1);
}

.hugyoku-chat-bubble .hugyoku-chat-thread-meta {
  display: flex;
  flex-wrap: wrap;
  gap: 0.48rem;
  margin-bottom: 0.36rem;
}

.hugyoku-chat-bubble .hugyoku-chat-thread-meta span {
  font-size: 0.72rem;
}

.hugyoku-chat-bubble .hugyoku-chat-role-label {
  margin-bottom: 0.28rem;
}

.hugyoku-chat-attachment-list {
  margin-top: 0.45rem;
}

.hugyoku-chat-attachment {
  padding: 0.28rem 0.58rem;
  border-radius: 0.72rem;
  font-size: 0.74rem;
}

.hugyoku-chat-control-row [data-testid="stPopover"] button,
.hugyoku-chat-control-row [data-testid="stButton"] button {
  min-height: 1.92rem;
  padding: 0.18rem 0.62rem;
  font-size: 0.72rem;
  border-radius: 0.72rem;
  background: linear-gradient(180deg, rgba(22, 30, 44, 0.96), rgba(16, 23, 35, 0.98));
  border-color: rgba(120, 141, 184, 0.18);
}

@media (max-width: 640px) {
  [data-testid="block-container"]:has(.hugyoku-chat-page.hugyoku-chat-custom),
  .main .block-container:has(.hugyoku-chat-page.hugyoku-chat-custom) {
    padding: 0.04rem 0.06rem 0.06rem !important;
  }

  .hugyoku-chat-page {
    gap: 0.08rem;
  }

  .hugyoku-chat-header {
    gap: 0.08rem;
    padding: 0;
  }

  .hugyoku-chat-header-left {
    gap: 0.08rem;
  }

  .hugyoku-chat-pill {
    padding: 0.08rem 0.26rem;
    font-size: 0.48rem;
    letter-spacing: 0.05em;
  }

  .hugyoku-chat-workspace-label {
    font-size: 0.58rem;
  }

  .hugyoku-chat-message-counter {
    display: none;
  }

  .hugyoku-chat-empty {
    min-height: 0;
    gap: 0.06rem;
    padding: 0;
  }

  .hugyoku-chat-empty-icon {
    width: 1.1rem;
    height: 1.1rem;
    font-size: 0.68rem;
  }

  .hugyoku-chat-empty-title {
    font-size: 0.68rem;
  }

  .hugyoku-chat-empty-subtitle {
    font-size: 0.5rem;
  }

  .hugyoku-chat-bubble {
    max-width: 100%;
    border-radius: 0.82rem;
    padding: 0.48rem 0.58rem 0.54rem;
  }

  .hugyoku-chat-bubble .hugyoku-chat-thread-meta {
    gap: 0.32rem;
    margin-bottom: 0.28rem;
  }

  .hugyoku-chat-bubble .hugyoku-chat-thread-meta span {
    font-size: 0.62rem;
  }

  .hugyoku-chat-role-label {
    min-height: 1.12rem;
    padding: 0.06rem 0.26rem;
    font-size: 0.54rem;
  }

  .hugyoku-chat-composer-shell {
    max-width: 100%;
    gap: 0.22rem;
  }

  .hugyoku-chat-composer-sticky {
    padding: 0.24rem;
    border-radius: 0.82rem;
  }

  .hugyoku-chat-composer-sticky [data-testid="stTextArea"] textarea {
    min-height: 56px !important;
    font-size: 0.76rem;
    padding-top: 0.2rem;
    padding-bottom: 0.2rem;
  }

  .hugyoku-chat-toolbar-grid [data-testid="stHorizontalBlock"] {
    gap: 0.16rem !important;
  }

  .hugyoku-chat-toolbar-grid [data-testid="stButton"] button,
  .hugyoku-chat-toolbar-grid [data-testid="stPopover"] button {
    min-height: 1.48rem;
    padding: 0.04rem 0.3rem;
    font-size: 0.58rem;
    border-radius: 0.6rem;
  }
}
</style>
"""

HUGYOKU_CHAT_V2_CSS = """
<style>
body:has(.hugyoku-chat-v2-root),
[data-testid="stAppViewContainer"]:has(.hugyoku-chat-v2-root),
[data-testid="stMain"]:has(.hugyoku-chat-v2-root),
.main:has(.hugyoku-chat-v2-root) {
  overflow: hidden;
  min-height: 100dvh;
  background: #0b1018;
}

[data-testid="block-container"]:has(.hugyoku-chat-v2-root),
.main .block-container:has(.hugyoku-chat-v2-root) {
  max-width: 100% !important;
  width: 100% !important;
  min-height: 100dvh !important;
  height: 100dvh !important;
  padding: 0 !important;
  overflow: hidden;
}

.hugyoku-chat-v2-root {
  min-height: 100dvh;
  height: 100dvh;
  display: flex;
  flex-direction: column;
  background:
    radial-gradient(circle at 12% 0%, rgba(255, 146, 104, 0.09), transparent 24%),
    radial-gradient(circle at 88% 0%, rgba(72, 116, 214, 0.1), transparent 30%),
    linear-gradient(180deg, #0d121b 0%, #0a0f17 58%, #090d14 100%);
  color: #eef3fb;
}

.hugyoku-chat-v2-topbar {
  position: sticky;
  top: 0;
  z-index: 40;
  padding: 0.58rem 0.9rem;
  border-bottom: 1px solid rgba(86, 101, 128, 0.14);
  background: rgba(9, 13, 20, 0.9);
  backdrop-filter: blur(16px);
}

.hugyoku-chat-v2-topbar-row {
  max-width: 680px;
  margin: 0 auto;
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 0.65rem;
}

.hugyoku-chat-v2-titlebar {
  min-width: 0;
  display: flex;
  align-items: center;
  gap: 0.48rem;
}

.hugyoku-chat-v2-topbar-icon {
  color: #d9bf91;
  font-size: 0.94rem;
  line-height: 1;
}

.hugyoku-chat-v2-title {
  color: #f3efe6;
  font-size: 0.84rem;
  font-weight: 700;
  line-height: 1.1;
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}

.hugyoku-chat-v2-menu [data-testid="stPopover"] > button,
.hugyoku-chat-v2-older [data-testid="stButton"] > button {
  min-height: 2rem;
  border-radius: 999px;
  padding: 0.12rem 0.68rem;
  background: rgba(17, 23, 35, 0.94);
  border: 1px solid rgba(94, 109, 138, 0.16);
  color: #edf2fa;
  box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.03);
}

.hugyoku-chat-v2-menu [data-testid="stPopover"] > button {
  min-width: 2.05rem;
  width: 2.05rem;
  padding-inline: 0;
}

.hugyoku-chat-v2-edit {
  max-width: 680px;
  margin: 0 auto;
  width: 100%;
  padding: 0.45rem 0.9rem 0;
}

.hugyoku-chat-v2-edit [data-testid="stVerticalBlockBorderWrapper"] {
  border-radius: 16px !important;
  box-shadow: none;
  background: rgba(13, 19, 29, 0.94);
}

.chat-container {
  flex: 1 1 auto;
  min-height: 0;
  width: 100%;
  max-width: 680px;
  margin: 0 auto;
  padding: 0.15rem 0.9rem 6.1rem;
  overflow: hidden;
}

.chat-scroll {
  height: 100%;
  overflow-y: auto;
  overflow-x: hidden;
  padding: 0.2rem 0 0.45rem;
  scroll-behavior: smooth;
  scrollbar-width: thin;
  overscroll-behavior: contain;
}

.hugyoku-chat-v2-older {
  display: flex;
  justify-content: center;
  margin: 0.1rem 0 0.55rem;
}

.hugyoku-chat-v2-empty-hint {
  color: rgba(154, 168, 193, 0.72);
  font-size: 0.82rem;
  line-height: 1.4;
  padding: 0.45rem 0.2rem 0;
}

div[data-testid="stChatMessage"] {
  background: transparent;
  border: none;
  box-shadow: none;
  padding: 0;
  margin: 0 0 0.68rem;
}

div[data-testid="stChatMessage"] [data-testid="stMarkdownContainer"] {
  max-width: 100%;
}

@keyframes hugyokuChatFade {
  from {
    opacity: 0;
    transform: translateY(8px);
  }
  to {
    opacity: 1;
    transform: translateY(0);
  }
}

.chat-message {
  max-width: min(100%, 42rem);
  border-radius: 1.15rem;
  padding: 0.8rem 0.9rem 0.84rem;
  border: 1px solid rgba(98, 113, 142, 0.12);
  box-shadow: 0 10px 24px rgba(4, 9, 18, 0.16);
  animation: hugyokuChatFade 0.18s ease both;
}

.chat-message.user {
  margin-left: auto;
  background: linear-gradient(180deg, rgba(30, 22, 33, 0.98), rgba(18, 14, 24, 0.98));
  border-color: rgba(255, 112, 112, 0.1);
}

.chat-message.assistant {
  margin-right: auto;
  max-width: min(100%, 38rem);
  background: linear-gradient(180deg, rgba(14, 21, 33, 0.98), rgba(10, 16, 26, 0.99));
  border-color: rgba(255, 203, 122, 0.08);
}

.chat-message-role {
  display: inline-flex;
  align-items: center;
  min-height: 1.34rem;
  padding: 0.08rem 0.5rem;
  border-radius: 999px;
  font-size: 0.63rem;
  font-weight: 800;
  letter-spacing: 0.09em;
  text-transform: uppercase;
  margin-bottom: 0.28rem;
}

.chat-message-role.user {
  color: #ffb7b7;
  background: rgba(255, 103, 103, 0.1);
  border: 1px solid rgba(255, 103, 103, 0.16);
}

.chat-message-role.assistant {
  color: #ffdba7;
  background: rgba(255, 198, 110, 0.08);
  border: 1px solid rgba(255, 198, 110, 0.14);
}

.chat-message-meta {
  display: flex;
  flex-wrap: wrap;
  gap: 0.38rem;
  margin-bottom: 0.28rem;
}

.chat-message-meta span {
  color: rgba(160, 174, 198, 0.82);
  font-size: 0.69rem;
}

.chat-message pre {
  background: rgba(7, 11, 19, 0.92);
  border: 1px solid rgba(94, 108, 134, 0.16);
  border-radius: 14px;
  padding: 0.76rem 0.84rem;
  overflow-x: auto;
}

.chat-message code {
  color: #ffe6b8;
}

.hugyoku-chat-v2-chip {
  display: inline-flex;
  align-items: center;
  padding: 0.15rem 0.5rem;
  border-radius: 999px;
  border: 1px solid rgba(100, 116, 146, 0.14);
  background: rgba(18, 24, 36, 0.84);
  color: #dfe7f6;
  font-size: 0.68rem;
}

.chat-input {
  position: fixed;
  left: 0;
  right: 0;
  bottom: 0;
  z-index: 999;
  padding: 0.38rem 0.9rem calc(env(safe-area-inset-bottom, 0px) + 0.48rem);
  border-top: 1px solid rgba(255, 255, 255, 0.06);
  background: linear-gradient(180deg, rgba(8, 12, 18, 0), rgba(8, 12, 18, 0.84) 24%, rgba(8, 12, 18, 0.97) 50%, rgba(8, 12, 18, 0.995) 100%);
}

.chat-input-inner {
  max-width: 680px;
  margin: 0 auto;
}

.chat-input-status {
  color: rgba(154, 168, 193, 0.74);
  font-size: 0.68rem;
  line-height: 1.3;
  padding: 0 0.12rem 0.28rem;
}

.chat-input [data-testid="stChatInput"] {
  margin: 0;
}

.chat-input [data-testid="stChatInput"] > div {
  background: rgba(13, 19, 30, 0.97);
  border: 1px solid rgba(95, 110, 139, 0.18);
  border-radius: 18px;
  box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.03);
}

.chat-input [data-testid="stChatInput"] textarea {
  color: #eef3fb !important;
  font-size: 0.94rem !important;
  line-height: 1.42 !important;
  min-height: 46px !important;
  max-height: 180px !important;
  padding-top: 0.48rem !important;
  padding-bottom: 0.48rem !important;
}

.chat-input [data-testid="stChatInput"] textarea::placeholder {
  color: rgba(146, 160, 184, 0.76) !important;
}

.chat-input [data-testid="stChatInput"] button {
  min-width: 2.28rem !important;
  width: 2.28rem !important;
  height: 2.28rem !important;
  border-radius: 999px !important;
  background: linear-gradient(180deg, rgba(61, 74, 103, 0.96), rgba(39, 49, 72, 0.98)) !important;
  border: 1px solid rgba(132, 153, 190, 0.16) !important;
}

.chat-input [data-testid="stChatInput"] button:hover {
  border-color: rgba(255, 214, 160, 0.22) !important;
}

@media (max-width: 768px) {
  .hugyoku-chat-v2-topbar {
    padding: 0.52rem 0.72rem;
  }

  .hugyoku-chat-v2-title {
    font-size: 0.8rem;
  }

  .chat-container {
    padding: 0.05rem 0.7rem 5.7rem;
  }

  .chat-input {
    padding: 0.34rem 0.7rem calc(env(safe-area-inset-bottom, 0px) + 0.42rem);
  }

  .chat-message {
    max-width: 100%;
    padding: 0.76rem 0.82rem 0.8rem;
  }
}

@media (max-width: 480px) {
  .hugyoku-chat-v2-topbar {
    padding: 0.48rem 0.58rem;
  }

  .hugyoku-chat-v2-topbar-icon {
    font-size: 0.88rem;
  }

  .hugyoku-chat-v2-title {
    font-size: 0.76rem;
  }

  .hugyoku-chat-v2-edit {
    padding: 0.32rem 0.58rem 0;
  }

  .chat-container {
    padding: 0 0.55rem 5.3rem;
  }

  .chat-scroll {
    padding-top: 0.1rem;
  }

  .chat-message {
    border-radius: 1rem;
    padding: 0.72rem 0.76rem 0.78rem;
    margin-bottom: 0.56rem;
  }

  .chat-message-meta span {
    font-size: 0.64rem;
  }

  .chat-input {
    padding: 0.3rem 0.5rem calc(env(safe-area-inset-bottom, 0px) + 0.34rem);
  }

  .chat-input-status {
    font-size: 0.64rem;
    padding-bottom: 0.2rem;
  }

  .chat-input [data-testid="stChatInput"] > div {
    border-radius: 16px;
  }

  .chat-input [data-testid="stChatInput"] textarea {
    min-height: 42px !important;
    font-size: 0.91rem !important;
    padding-top: 0.42rem !important;
    padding-bottom: 0.42rem !important;
  }

  .chat-input [data-testid="stChatInput"] button {
    min-width: 2.12rem !important;
    width: 2.12rem !important;
    height: 2.12rem !important;
  }
}
</style>
"""


def read_hugyoku_reference_bundle(doc_files: list[object], image_files: list[object], attachment_note: str) -> tuple[str, list[str], str]:
    parts: list[str] = []
    issues: list[str] = []
    ocr_status = "No images attached."
    for uploaded_file in doc_files:
        try:
            text = read_uploaded_document(uploaded_file)
            parts.append(f"[Reference File: {getattr(uploaded_file, 'name', 'reference.txt')}]\n{text}")
        except Exception as exc:
            issues.append(f"{getattr(uploaded_file, 'name', 'file')}: {exc}")
    if image_files:
        if ocr_supported():
            extracted_count = 0
            for uploaded_image in image_files:
                try:
                    text = extract_image_text_from_bytes(uploaded_image.getvalue())
                    label = getattr(uploaded_image, "name", "image")
                    if text:
                        parts.append(f"[Image OCR: {label}]\n{text}")
                        extracted_count += 1
                    else:
                        parts.append(f"[Image OCR: {label}]\nNo readable text was detected from this image.")
                except Exception as exc:
                    issues.append(f"{getattr(uploaded_image, 'name', 'image')}: OCR failed ({exc})")
            ocr_status = f"OCR processed {extracted_count} of {len(image_files)} image(s)."
        else:
            for uploaded_image in image_files:
                parts.append(
                    f"[Image Attachment: {getattr(uploaded_image, 'name', 'image')}]\n"
                    "Image attached, but OCR is not available in this environment."
                )
            ocr_status = "Images attached, but OCR support is not available here."
    if attachment_note.strip():
        parts.append(f"[Attachment Note]\n{attachment_note.strip()}")
    return "\n\n".join(parts).strip(), issues, ocr_status


def detect_requested_export_format(task_text: str, understanding_text: str = "") -> str:
    explicit = extract_section_value(
        understanding_text,
        ["Requested File Format", "Preferred File Format", "Submission Format", "Output Format"],
    ).strip().lower()
    combined = f"{explicit}\n{task_text}".lower()
    if any(token in combined for token in ["pdf", ".pdf", "portable document"]):
        return "pdf"
    if any(token in combined for token in ["txt", ".txt", "text file", "plain text"]):
        return "txt"
    if any(token in combined for token in ["docx", ".docx", "word file", "microsoft word", "ms word"]):
        return "docx"
    return "docx"


def detect_hugyoku_output_type(task_text: str, understanding_text: str = "") -> str:
    explicit = extract_section_value(
        understanding_text,
        ["Requested Output", "Task Summary", "What You Understand"],
    ).strip().lower()
    combined = f"{task_text}\n{explicit}".lower()
    if any(token in combined for token in ["essay", "reaction paper", "reflection paper"]):
        return "essay"
    if any(token in combined for token in ["summary", "summarize", "reviewer", "review notes"]):
        return "summary"
    if any(token in combined for token in ["analysis", "analyze", "compare", "comparison", "critique"]):
        return "analysis"
    if any(token in combined for token in ["explain", "meaning", "what is", "why", "how"]):
        return "explanation"
    if any(token in combined for token in ["report", "document", "write-up", "paper"]):
        return "report"
    return "academic response"


def compile_hugyoku_sections(sections: list[dict[str, object]]) -> str:
    parts: list[str] = []
    for section in sections:
        heading = str(section.get("heading", "")).strip()
        content = str(section.get("content", "")).strip()
        if heading and content:
            parts.append(f"{heading}:\n{content}")
        elif heading:
            parts.append(heading)
        elif content:
            parts.append(content)
    return "\n\n".join(part for part in parts if part.strip()).strip()


def build_hugyoku_sections(text: str) -> list[dict[str, object]]:
    sections: list[dict[str, object]] = []
    for index, block in enumerate(parse_structured_blocks(text), start=1):
        if block["type"] == "section":
            sections.append(
                {
                    "heading": str(block.get("heading", "")).strip() or f"Section {index}",
                    "content": str(block.get("content", "")).strip(),
                }
            )
        else:
            sections.append(
                {
                    "heading": f"Section {index}",
                    "content": str(block.get("content", "")).strip(),
                }
            )
    if not sections and text.strip():
        sections.append({"heading": "Main Output", "content": text.strip()})
    return sections


def sync_hugyoku_sections_from_widgets() -> None:
    updated: list[dict[str, object]] = []
    for index, section in enumerate(list(st.session_state.get("hugyoku_output_sections", []))):
        heading = str(st.session_state.get(f"hugyoku_section_heading_{index}", section.get("heading", ""))).strip()
        content = str(st.session_state.get(f"hugyoku_section_content_{index}", section.get("content", ""))).strip()
        updated.append({"heading": heading, "content": content})
    st.session_state.hugyoku_output_sections = updated
    st.session_state.hugyoku_output_raw = compile_hugyoku_sections(updated)


def prime_hugyoku_section_widgets(sections: list[dict[str, object]]) -> None:
    for index, section in enumerate(sections):
        st.session_state[f"hugyoku_section_heading_{index}"] = str(section.get("heading", "")).strip()
        st.session_state[f"hugyoku_section_content_{index}"] = str(section.get("content", "")).strip()


def guess_hugyoku_title(task_text: str, understanding_text: str, sections: list[dict[str, object]]) -> str:
    explicit = extract_section_value(understanding_text, ["Requested Output", "Task Summary"]).strip()
    for section in sections:
        heading = str(section.get("heading", "")).strip().lower()
        if heading in {"title", "essay title", "document title"}:
            content = str(section.get("content", "")).strip()
            if content:
                return content
    if explicit:
        return explicit[:90]
    cleaned = " ".join(task_text.strip().split())
    return cleaned[:90] if cleaned else "Hugyoku Output"


def render_txt_bytes(title: str, body: str, metadata_lines: list[str] | None = None) -> bytes:
    lines: list[str] = [title.strip() or "Untitled Document"]
    if metadata_lines:
        lines.append(" | ".join(metadata_lines))
    if body.strip():
        lines.extend(["", body.strip()])
    return "\n".join(lines).encode("utf-8")


def render_pdf_bytes(title: str, body: str, metadata_lines: list[str] | None = None) -> bytes:
    if canvas is None or LETTER is None:
        raise RuntimeError("PDF export requires reportlab in this environment.")
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    left_margin = 72
    top = height - 72
    bottom = 72
    line_height = 14
    y = top

    def write_line(text: str, *, font: str = "Helvetica", size: int = 11) -> None:
        nonlocal y
        if y <= bottom:
            pdf.showPage()
            y = top
        pdf.setFont(font, size)
        pdf.drawString(left_margin, y, text)
        y -= line_height if size <= 11 else line_height + 2

    write_line(title.strip() or "Untitled Document", font="Helvetica-Bold", size=18)
    if metadata_lines:
        for meta_line in metadata_lines:
            for wrapped in textwrap.wrap(meta_line, width=86) or [""]:
                write_line(wrapped, font="Helvetica-Oblique", size=10)
        y -= 4

    for paragraph in compile_hugyoku_sections(build_hugyoku_sections(body)).split("\n"):
        wrapped_lines = textwrap.wrap(paragraph, width=90) or [""]
        for wrapped in wrapped_lines:
            write_line(wrapped)

    pdf.save()
    return buffer.getvalue()


def build_hugyoku_export_payload(
    title: str,
    body: str,
    export_format: str,
    *,
    name_override: str | None = None,
) -> tuple[bytes, str, str, str | None]:
    requested = export_format.lower().strip()
    if requested not in {"docx", "pdf", "txt"}:
        requested = "docx"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = sanitize_filename(name_override or title or "hugyoku_output")
    metadata_lines = export_metadata_lines(category="hugyoku", name_override=name_override)
    export_root = (st.session_state.export_root_path or "").strip() if save_destination_mode() == "local" else ""
    internal_folder = TOOL_FOLDERS.get("hugyoku", "hugyoku_universal")
    if requested == "pdf":
        payload = render_pdf_bytes(title, body, metadata_lines)
        filename = f"{stem}_{stamp}.pdf"
        mime = "application/pdf"
    elif requested == "txt":
        payload = render_txt_bytes(title, body, metadata_lines)
        filename = f"{stem}_{stamp}.txt"
        mime = "text/plain"
    else:
        payload = render_docx_bytes(
            title,
            body,
            category="hugyoku",
            metadata_lines=metadata_lines,
            output_options=current_output_settings(),
        )
        filename = f"{stem}_{stamp}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    local_path = str(Path(export_root) / internal_folder / filename) if export_root else None
    return payload, filename, mime, local_path


def prepare_hugyoku_download(file_name: str) -> None:
    append_history_entry(
        "export",
        f"Prepared Hugyoku download: {file_name}",
        "Browser download prepared for the universal Hugyoku workflow.",
        "hugyoku",
        active_workspace_name(),
        "hugyoku",
    )
    st.session_state.flash_message = f"Prepared download: {file_name}"
    st.session_state.flash_level = "info"


def render_hugyoku_export_button(title: str, body: str, export_format: str) -> None:
    file_bytes, file_name, mime_type, local_path = build_hugyoku_export_payload(title, body, export_format, name_override=title)
    if local_path:
        render_route_block("Local save route", local_path)
        local_col, open_col = st.columns(2, gap="small")
        if local_col.button("Save To Selected Folder", key="save_local_hugyoku_export", use_container_width=True, type="primary"):
            try:
                saved_path = save_local_export_file(file_bytes, local_path)
                append_history_entry("export", f"Saved Hugyoku export: {file_name}", saved_path, "hugyoku", active_workspace_name(), "hugyoku")
                st.success(f"Saved to {saved_path}")
            except Exception as exc:
                st.error(f"Could not save to the selected folder: {exc}")
        if open_col.button("Open Selected Folder", key="open_local_hugyoku_export", use_container_width=True):
            try:
                target_folder = Path(local_path).parent
                target_folder.mkdir(parents=True, exist_ok=True)
                open_in_file_manager(str(target_folder))
            except Exception as exc:
                st.error(f"Could not open the selected folder: {exc}")
    else:
        render_route_block("Browser download", f"{file_name}\nSaved directly by your browser.")
    st.download_button(
        f"Download {export_format.upper()} File",
        data=file_bytes,
        file_name=file_name,
        mime=mime_type,
        key=f"download_hugyoku_{sanitize_filename(title)}_{export_format}",
        use_container_width=True,
        on_click=prepare_hugyoku_download,
        args=(file_name,),
        type="secondary" if local_path else "primary",
    )


def build_hugyoku_understanding_prompt(task_text: str, reference_bundle: str) -> str:
    return (
        "You are Hugyoku AI, a high-precision academic task interpreter.\n"
        "Analyze the user's task before generating anything. Be explicit, structured, and avoid assumptions.\n\n"
        f"User task:\n{task_text or 'No direct task text provided. Infer from references.'}\n\n"
        f"Reference bundle:\n{trim_prompt_source(reference_bundle) if reference_bundle else 'No reference files provided.'}\n\n"
        "Return these plain-text sections in this exact order:\n"
        "Output Type:\n"
        "Task Summary:\n"
        "What You Understand:\n"
        "Requested Output:\n"
        "Requested File Format:\n"
        "Important Requirements:\n"
        "Missing or Uncertain Parts:\n"
        "Next Confirmation Step:\n"
    )


def build_hugyoku_refinement_prompt(task_text: str, reference_bundle: str, current_understanding: str, refinement_prompt: str) -> str:
    return (
        "You are Hugyoku AI, a high-precision academic task interpreter.\n"
        "Revise your understanding of the user's task. Keep it accurate, specific, and ready for confirmation before generation.\n\n"
        f"Original task:\n{task_text or 'No direct task text provided.'}\n\n"
        f"Reference bundle:\n{trim_prompt_source(reference_bundle) if reference_bundle else 'No reference files provided.'}\n\n"
        f"Current understanding:\n{current_understanding}\n\n"
        f"User correction or refinement:\n{refinement_prompt}\n\n"
        "Return these plain-text sections in this exact order:\n"
        "Output Type:\n"
        "Task Summary:\n"
        "What You Understand:\n"
        "Requested Output:\n"
        "Requested File Format:\n"
        "Important Requirements:\n"
        "Missing or Uncertain Parts:\n"
        "Next Confirmation Step:\n"
    )


def build_hugyoku_generation_prompt(task_text: str, reference_bundle: str, understanding_text: str, export_format: str, revision_prompt: str = "") -> str:
    output_type = detect_hugyoku_output_type(task_text, understanding_text)
    return (
        "You are Hugyoku AI, a high-precision academic and analytical writing system designed for factual accuracy, structured reasoning, and zero hallucination.\n\n"
        "PRIMARY OBJECTIVE:\n"
        "Generate a clear, structured, and academically sound response using ONLY the provided context. Every claim must be explicitly supported by the provided context.\n"
        "Always prioritize accuracy over completeness.\n\n"
        f"Required output type: {output_type}\n"
        f"Original task:\n{task_text or 'No direct task text provided.'}\n\n"
        f"Confirmed understanding:\n{understanding_text}\n\n"
        f"CONTEXT:\n{trim_prompt_source(reference_bundle) if reference_bundle else 'No reliable reference context provided.'}\n\n"
        f"Requested file format: {export_format.upper()}\n"
        f"Additional revision prompt: {revision_prompt or 'None'}\n\n"
        "STRICT RULES:\n"
        "- Determine whether the context is sufficient, partially sufficient, or insufficient.\n"
        "- If the context is insufficient, output ONLY this exact sentence and nothing else:\n"
        "Insufficient data to provide a reliable answer.\n"
        "- If the context is partially sufficient, answer only the supported parts.\n"
        "- For unsupported parts, use this exact sentence:\n"
        "No reliable information found in the provided context.\n"
        "- Use ONLY the provided context.\n"
        "- Do NOT invent facts.\n"
        "- Do NOT assume missing details.\n"
        "- Do NOT add outside knowledge.\n"
        "- Do NOT create fake quotes.\n"
        "- If something is not explicitly supported by the context, exclude it.\n\n"
        "STYLE RULES:\n"
        "- Formal academic tone.\n"
        "- Clear logical flow.\n"
        "- No filler.\n"
        "- No meta-commentary.\n"
        "- Do not explain your reasoning process.\n\n"
        "OUTPUT FORMAT:\n"
        "If the context is sufficient or partially sufficient, return these plain-text sections in this exact order:\n"
        "Title:\n"
        "Answer:\n"
        "Key Points:\n\n"
        "Key Points rules:\n"
        "- Include only directly supported points.\n"
        "- If fewer than 3 reliable points exist, include fewer.\n"
        "- For missing aspects in a partially supported task, explicitly write: No reliable information found in the provided context.\n"
        "- Do not use markdown fence blocks."
    )


def clear_workspace_analysis_workspace() -> None:
    st.session_state.workspace_analysis_result = ""


def clear_source_lab_workspace() -> None:
    st.session_state.source_lab_question_input = ""
    st.session_state.source_lab_answer = ""
    st.session_state.source_lab_analysis = ""


def clear_reviewer_workspace() -> None:
    st.session_state.reviewer_request_input = ""
    st.session_state.reviewer_response = ""


def clear_flashcard_workspace() -> None:
    st.session_state.flashcard_request_input = ""
    st.session_state.flashcard_response = ""


def clear_practice_test_workspace() -> None:
    st.session_state.practice_test_request_input = ""
    st.session_state.practice_test_response = ""


def clear_answer_checker_workspace() -> None:
    st.session_state.answer_checker_question_input = ""
    st.session_state.answer_checker_reference_input = ""
    st.session_state.answer_checker_user_answer_input = ""
    st.session_state.answer_checker_response = ""


def clear_rubric_workspace() -> None:
    st.session_state.rubric_title_input = ""
    st.session_state.rubric_prompt_input = ""
    st.session_state.rubric_focus_input = ["Clarity", "Structure", "Depth"]
    st.session_state.rubric_response = ""


def clear_batch_workspace() -> None:
    st.session_state.batch_topic_input = ""
    st.session_state.batch_request_input = ""
    st.session_state.batch_response = ""


def clear_history_entries() -> None:
    st.session_state.history_entries = []


def clear_model_selection() -> None:
    st.session_state.model_choice_input = "HF Secret Default"
    st.session_state.model_override = ""
    st.session_state.model_custom_input = ""


def clear_codegen_workspace() -> None:
    st.session_state.codegen_title = ""
    st.session_state.codegen_stack_profile = "Python"
    st.session_state.codegen_custom_stack = ""
    st.session_state.codegen_description = ""
    st.session_state.codegen_expectation = ""
    st.session_state.codegen_attachment_note = ""
    st.session_state.codegen_response = ""


def clear_selftest_workspace() -> None:
    st.session_state.selftest_stack_profile = "Python"
    st.session_state.selftest_custom_stack = ""
    st.session_state.selftest_response = ""
    st.session_state.selftest_last_model = ""


def clear_compare_workspace() -> None:
    st.session_state.compare_model_a = "Active Session Model"
    st.session_state.compare_model_b = "Qwen/Qwen2.5-Coder-7B-Instruct"
    st.session_state.compare_custom_model_a = ""
    st.session_state.compare_custom_model_b = ""
    st.session_state.compare_prompt = ""
    st.session_state.compare_output_a = ""
    st.session_state.compare_output_b = ""
    st.session_state.compare_used_model_a = ""
    st.session_state.compare_used_model_b = ""


def clear_hugyoku_workspace() -> None:
    st.session_state.hugyoku_task_input = ""
    st.session_state.hugyoku_attachment_note = ""
    st.session_state.hugyoku_stage = 1
    st.session_state.hugyoku_understanding = ""
    st.session_state.hugyoku_refinement_prompt = ""
    st.session_state.hugyoku_refinement_round = 0
    st.session_state.hugyoku_output_sections = []
    st.session_state.hugyoku_output_title = ""
    st.session_state.hugyoku_output_format = "docx"
    st.session_state.hugyoku_output_raw = ""
    st.session_state.hugyoku_result_prompt = ""
    st.session_state.hugyoku_generation_note = ""
    st.session_state.hugyoku_last_bundle = ""
    st.session_state.hugyoku_last_ocr_status = ""


def clear_hugyoku_chat_workspace() -> None:
    st.session_state.hugyoku_chat_messages = []
    st.session_state.hugyoku_chat_draft = ""
    st.session_state.hugyoku_chat_attachment_note = ""
    st.session_state.hugyoku_chat_controls_open = False
    st.session_state.hugyoku_chat_pending_job = None
    st.session_state.hugyoku_chat_show_full_history = False
    st.session_state.hugyoku_chat_edit_mode = False
    st.session_state.hugyoku_chat_edit_text = ""
    st.session_state.hugyoku_chat_reasoning_mode = "Balanced"
    st.session_state.hugyoku_chat_model_choice = "Active Session Model"
    st.session_state.hugyoku_chat_model_custom = ""
    st.session_state.hugyoku_chat_voice_language = "en-US"
    st.session_state.hugyoku_chat_pending_transcript = ""
    st.session_state.hugyoku_chat_last_bundle = ""
    st.session_state.hugyoku_chat_last_ocr_status = ""
    st.session_state.hugyoku_chat_last_used_model = ""
    st.session_state.hugyoku_chat_last_transcript_text = ""
    st.session_state.hugyoku_chat_upload_nonce = int(st.session_state.get("hugyoku_chat_upload_nonce", 0) or 0) + 1


def clear_hugyoku_chat_composer() -> None:
    st.session_state.hugyoku_chat_draft = ""
    st.session_state.hugyoku_chat_attachment_note = ""
    st.session_state.hugyoku_chat_pending_transcript = ""
    st.session_state.hugyoku_chat_last_transcript_text = ""
    st.session_state.hugyoku_chat_upload_nonce = int(st.session_state.get("hugyoku_chat_upload_nonce", 0) or 0) + 1


def render_hugyoku_chat_autoscroll(message_count: int) -> None:
    components_html(
        f"""
        <script>
        const run = () => {{
          const root = window.parent.document;
          const pane = root.querySelector('.hugyoku-chat-v2-root .chat-scroll');
          if (pane) {{
            pane.scrollTo({{ top: pane.scrollHeight, behavior: 'smooth' }});
          }}
        }};
        setTimeout(run, 40);
        </script>
        """,
        height=0,
        width=0,
        key=f"hugyoku-chat-autoscroll-{message_count}",
    )


RESET_ACTIONS.update(
    {
        "clear_workspace_analysis_workspace": clear_workspace_analysis_workspace,
        "clear_source_lab_workspace": clear_source_lab_workspace,
        "clear_reviewer_workspace": clear_reviewer_workspace,
        "clear_flashcard_workspace": clear_flashcard_workspace,
        "clear_practice_test_workspace": clear_practice_test_workspace,
        "clear_answer_checker_workspace": clear_answer_checker_workspace,
        "clear_rubric_workspace": clear_rubric_workspace,
        "clear_batch_workspace": clear_batch_workspace,
        "clear_history_entries": clear_history_entries,
        "clear_model_selection": clear_model_selection,
        "clear_codegen_workspace": clear_codegen_workspace,
        "clear_selftest_workspace": clear_selftest_workspace,
        "clear_compare_workspace": clear_compare_workspace,
        "clear_hugyoku_workspace": clear_hugyoku_workspace,
        "clear_hugyoku_chat_workspace": clear_hugyoku_chat_workspace,
        "clear_hugyoku_chat_composer": clear_hugyoku_chat_composer,
    }
)


def render_history_snippets(limit: int = 5, page_filter: str | None = None) -> None:
    entries = list(st.session_state.get("history_entries", []))
    if page_filter:
        entries = [entry for entry in entries if entry.get("page") == page_filter]
    if not entries:
        st.caption("No history entries yet.")
        return
    for entry in entries[:limit]:
        with st.container(border=True):
            render_card_header(
                str(entry.get("title", "History Entry")),
                str(entry.get("details", "")),
                f"{entry.get('timestamp', '')} • {entry.get('kind', 'event')}",
                anchor="EV",
                tier="tertiary",
                compact=True,
            )
            render_meta_grid(
                [
                    ("Page", str(entry.get("page", "-"))),
                    ("Workspace", str(entry.get("workspace", "-"))),
                    ("Category", str(entry.get("category", "-")) or "-"),
                ]
            )


def render_workspace_outputs(limit: int = 6) -> None:
    workspace = active_workspace()
    outputs = list(workspace.get("outputs", []))
    if not outputs:
        st.caption("No generated outputs saved in this workspace yet.")
        return
    for output in outputs[:limit]:
        with st.container(border=True):
            render_card_header(
                str(output.get("title", "Saved Output")),
                str(output.get("preview", "")),
                f"{output.get('timestamp', '')} • {output.get('kind', 'output')}",
                anchor="OP",
                tier="tertiary",
                compact=True,
            )
            st.caption(f"Category: {output.get('category', '-')}")


def current_identity_name() -> str:
    return st.session_state.saved_name or st.session_state.auth_display_name or st.session_state.auth_username or "Guest"
