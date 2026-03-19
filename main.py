from __future__ import annotations

import json
import os
import re
import threading
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from huggingface_hub import InferenceClient
except ImportError:
    InferenceClient = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None


APP_DIR = Path(__file__).resolve().parent
EXPORT_DIR = APP_DIR / "exports"
ENV_PATH = APP_DIR / ".env"
PROFILE_PATH = APP_DIR / "profile.json"

SYSTEM_PROMPT = (
    "You are Hugyoku, an educational assistant for study use. Help the user analyze "
    "activities, quizzes, assignments, uploaded files, documents, essays, and code "
    "issues. You may produce sample answers, draft essays, structured summaries, "
    "guided explanations, generated study documents, and debugging help for offline "
    "learning and review. Keep outputs clear and well-organized, and avoid assisting "
    "with live cheating or deceptive conduct."
)

PALETTE = {
    "bg": "#0C1017",
    "sidebar": "#121924",
    "panel": "#1B2331",
    "panel_alt": "#253043",
    "input": "#121A26",
    "border": "#3A465B",
    "text": "#F6F0E5",
    "muted": "#A9B2C3",
    "accent": "#C7A86A",
    "accent_soft": "#F3D79A",
    "accent_dark": "#7B6234",
    "success": "#50B58E",
    "warning": "#E4B65E",
    "danger": "#E37B74",
    "info": "#7CB8FF",
    "glass": "#1A2433",
    "glass_inner": "#202C3E",
    "glass_border": "#50607A",
    "glass_shadow": "#080B11",
    "glass_shadow_lift": "#121823",
    "glass_hover": "#2A3850",
    "glass_chip": "#293548",
}

KNOWN_EXPORT_LABELS = {
    "heading suggestion",
    "title suggestion",
    "suggested heading",
    "essay body",
    "essay",
    "self-check tip",
    "quick self-check tip",
    "closing self-check tip",
    "task understanding",
    "response",
    "short explanation or rationale",
    "summary",
    "objective",
    "instructions",
    "activity proper",
    "answer guide",
    "document type",
    "purpose",
    "main content",
    "solution plan",
    "sample answer or draft",
    "notes to review",
    "issue summary",
    "root cause",
    "fixed version",
    "why it works",
    "next checks",
}


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", name).strip(" ._")
    return cleaned or "document"


def today_string() -> str:
    now = datetime.now()
    return f"{now.strftime('%B')} {now.day}, {now.year}"


def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text))


def collapse_paragraph_lines(lines: list[str]) -> str:
    return " ".join(line.strip() for line in lines if line.strip()).strip()


def split_text_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n\s*\n", text.strip()) if block.strip()]


def strip_heading_markers(text: str) -> str:
    cleaned = re.sub(r"^#+\s*", "", text.strip())
    return cleaned.strip().strip(":").strip()


def heading_level_from_line(text: str) -> int:
    stripped = text.lstrip()
    if stripped.startswith("#"):
        return min(max(len(stripped) - len(stripped.lstrip("#")), 1), 4)
    return 2


def normalize_section_label(text: str) -> str:
    cleaned = strip_heading_markers(text)
    cleaned = re.sub(r"^\d+\.\s*", "", cleaned)
    return re.sub(r"\s+", " ", cleaned).lower().strip()


def parse_structured_blocks(text: str) -> list[dict[str, object]]:
    items: list[dict[str, object]] = []
    for block in split_text_blocks(text):
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        first = lines[0].strip()
        normalized_first = normalize_section_label(first)
        is_heading = first.startswith("#") or first.endswith(":") or normalized_first in KNOWN_EXPORT_LABELS
        if is_heading:
            items.append(
                {
                    "type": "section",
                    "heading": strip_heading_markers(first),
                    "level": heading_level_from_line(first),
                    "content": collapse_paragraph_lines(lines[1:]),
                }
            )
        else:
            items.append({"type": "paragraph", "content": collapse_paragraph_lines(lines)})
    return items


def add_body_paragraph(document: Document, text: str, indent_first_line: bool = False, italic: bool = False) -> None:
    if not text.strip():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(8)
    if indent_first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
    run = paragraph.add_run(text.strip())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = italic


def save_docx(
    title: str,
    body: str,
    output_path: Path,
    category: str = "generic",
    metadata_lines: list[str] | None = None,
    output_options: dict[str, bool] | None = None,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    metadata_lines = metadata_lines or []
    output_options = output_options or {}
    export_title = title.strip() or "Untitled Document"

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(export_title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(48, 87, 138)

    if metadata_lines:
        meta_paragraph = document.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_paragraph.paragraph_format.space_after = Pt(10)
        meta_run = meta_paragraph.add_run(" | ".join(metadata_lines))
        meta_run.italic = True
        meta_run.font.name = "Times New Roman"
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(88, 88, 88)

    parsed_blocks = parse_structured_blocks(body)
    essay_subtitle: str | None = None
    essay_tip: str | None = None
    body_items: list[dict[str, object]] = []

    for item in parsed_blocks:
        if item["type"] == "section":
            label = normalize_section_label(str(item["heading"]))
            content = str(item.get("content", "")).strip()
            if category == "essay":
                if label in {"heading suggestion", "title suggestion", "suggested heading"}:
                    if content:
                        essay_subtitle = content
                    continue
                if label in {"essay body", "essay"}:
                    if content:
                        body_items.append({"type": "paragraph", "content": content})
                    continue
                if label in {"self-check tip", "quick self-check tip", "closing self-check tip"}:
                    if content:
                        essay_tip = content
                    continue
            body_items.append(item)
        else:
            body_items.append(item)

    if category == "essay" and output_options.get("essay_include_heading", True) and essay_subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.paragraph_format.space_after = Pt(12)
        subtitle_run = subtitle_paragraph.add_run(essay_subtitle)
        subtitle_run.italic = True
        subtitle_run.font.name = "Times New Roman"
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    for item in body_items:
        if item["type"] == "section":
            heading_text = str(item["heading"]).strip().rstrip(":")
            if heading_text:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(heading_text)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(60, 76, 104)
            content = str(item.get("content", "")).strip()
            if content:
                add_body_paragraph(document, content, indent_first_line=(category == "essay"))
        else:
            add_body_paragraph(document, str(item["content"]), indent_first_line=(category == "essay"))

    if category == "essay" and output_options.get("essay_include_tip", True) and essay_tip:
        tip_heading = document.add_paragraph()
        tip_heading.paragraph_format.space_before = Pt(10)
        tip_heading.paragraph_format.space_after = Pt(2)
        tip_run = tip_heading.add_run("Self-Check Tip")
        tip_run.bold = True
        tip_run.font.name = "Times New Roman"
        tip_run.font.size = Pt(13)
        tip_run.font.color.rgb = RGBColor(60, 76, 104)
        add_body_paragraph(document, essay_tip, italic=True)

    document.save(output_path)


def extract_completion_text(message: object) -> str | None:
    content = getattr(message, "content", None)
    if isinstance(content, str):
        return content.strip() or None
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, str) and item.strip():
                parts.append(item.strip())
            elif isinstance(item, dict):
                text = item.get("text") or item.get("content")
                if isinstance(text, str) and text.strip():
                    parts.append(text.strip())
        joined = "\n\n".join(parts).strip()
        return joined or None
    return None


def load_client() -> tuple[object | None, str | None, str | None]:
    load_dotenv(ENV_PATH)
    token = os.getenv("HF_TOKEN", "").strip()
    model = os.getenv("HF_MODEL", "").strip()

    if InferenceClient is None:
        return None, None, "Install the packages in requirements.txt to enable AI features."
    if not token or not model:
        return None, None, "Add HF_TOKEN and HF_MODEL to .env to enable AI features."

    return InferenceClient(api_key=token), model, None


def generate_text(prompt: str, label: str) -> str:
    client, model, error = load_client()
    if error:
        raise RuntimeError(error)

    try:
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1400,
            temperature=0.45,
        )
    except Exception as exc:
        raise RuntimeError(f"Could not generate {label}: {exc}") from exc

    message = completion.choices[0].message
    result = extract_completion_text(message)
    if not result:
        raise RuntimeError("The AI model returned an empty response.")
    return result


def load_profile_settings() -> dict[str, object]:
    if not PROFILE_PATH.exists():
        return {
            "name": "",
            "include_date": False,
            "base_folder": "",
            "output_include_name": True,
            "output_include_date": False,
            "essay_include_heading_suggestion": True,
            "essay_include_self_check_tip": True,
        }
    try:
        data = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return {
            "name": "",
            "include_date": False,
            "base_folder": "",
            "output_include_name": True,
            "output_include_date": False,
            "essay_include_heading_suggestion": True,
            "essay_include_self_check_tip": True,
        }
    return {
        "name": str(data.get("name", "")).strip(),
        "include_date": bool(data.get("include_date", False)),
        "base_folder": str(data.get("base_folder", "")).strip(),
        "output_include_name": bool(data.get("output_include_name", True)),
        "output_include_date": bool(data.get("output_include_date", bool(data.get("include_date", False)))),
        "essay_include_heading_suggestion": bool(data.get("essay_include_heading_suggestion", True)),
        "essay_include_self_check_tip": bool(data.get("essay_include_self_check_tip", True)),
    }


def save_profile_settings(
    name: str,
    include_date: bool,
    base_folder: str,
    output_include_name: bool,
    output_include_date: bool,
    essay_include_heading_suggestion: bool,
    essay_include_self_check_tip: bool,
) -> None:
    payload = {
        "name": name.strip(),
        "include_date": include_date,
        "base_folder": base_folder.strip(),
        "output_include_name": output_include_name,
        "output_include_date": output_include_date,
        "essay_include_heading_suggestion": essay_include_heading_suggestion,
        "essay_include_self_check_tip": essay_include_self_check_tip,
    }
    PROFILE_PATH.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def read_docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs).strip()


def read_pdf_text(path: Path) -> str:
    if PdfReader is None:
        raise RuntimeError("PDF reading needs the pypdf package. Install requirements.txt first.")

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    joined = "\n\n".join(parts).strip()
    if not joined:
        raise RuntimeError("The PDF did not contain readable text.")
    return joined


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def read_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(path)
    if suffix == ".pdf":
        return read_pdf_text(path)
    if suffix in {".txt", ".md"}:
        return read_text_file(path)
    raise RuntimeError("Unsupported file type. Use .docx, .pdf, .txt, or .md.")


class PremiumStudyAssistant(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Hugyoku | Premium Academics Suite")
        self.geometry("1360x860")
        self.minsize(980, 640)
        self.configure(bg=PALETTE["bg"])

        EXPORT_DIR.mkdir(parents=True, exist_ok=True)

        profile = load_profile_settings()
        self.saved_name = str(profile.get("name", "")).strip()
        self.saved_include_date = bool(profile.get("include_date", False))
        saved_folder = str(profile.get("base_folder", "")).strip()
        self.base_export_dir = Path(saved_folder).expanduser() if saved_folder else EXPORT_DIR
        self.quiz_loaded_path: Path | None = None
        self.busy = False
        self.ai_ready = False
        self.pages: dict[str, tk.Frame] = {}
        self.nav_buttons: dict[str, tk.Button] = {}
        self.academic_buttons: dict[str, tk.Button] = {}
        self.academic_tools: dict[str, tk.Frame] = {}
        self.ai_action_buttons: list[tk.Button] = []
        self._responsive_after_id: str | None = None
        self.dashboard_layout: dict[str, object] = {}
        self.quiz_layout: dict[str, object] = {}
        self.essay_layout: dict[str, object] = {}

        self.status_var = tk.StringVar(value="Workspace ready.")
        self.ai_chip_var = tk.StringVar(value="Checking AI status...")
        self.model_var = tk.StringVar(value="No model configured")
        self.export_var = tk.StringVar(value=str(EXPORT_DIR))
        self.header_title_var = tk.StringVar()
        self.header_subtitle_var = tk.StringVar()
        self.profile_summary_var = tk.StringVar()
        self.profile_date_var = tk.StringVar()
        self.profile_note_var = tk.StringVar()
        self.main_folder_var = tk.StringVar()
        self.quiz_folder_var = tk.StringVar()
        self.essay_folder_var = tk.StringVar()
        self.activity_folder_var = tk.StringVar()
        self.document_folder_var = tk.StringVar()
        self.assignment_folder_var = tk.StringVar()
        self.codefix_folder_var = tk.StringVar()
        self.academics_folders_var = tk.StringVar()
        self.developer_folder_var = tk.StringVar()
        self.folder_note_var = tk.StringVar()

        self.name_input_var = tk.StringVar(value=self.saved_name)
        self.include_date_var = tk.BooleanVar(value=self.saved_include_date)
        self.output_include_name_var = tk.BooleanVar(value=bool(profile.get("output_include_name", True)))
        self.output_include_date_var = tk.BooleanVar(value=bool(profile.get("output_include_date", self.saved_include_date)))
        self.essay_include_heading_var = tk.BooleanVar(
            value=bool(profile.get("essay_include_heading_suggestion", True))
        )
        self.essay_include_tip_var = tk.BooleanVar(value=bool(profile.get("essay_include_self_check_tip", True)))
        self.academic_tool_var = tk.StringVar(value="quiz")

        self.quiz_file_var = tk.StringVar(value="No file loaded yet")
        self.quiz_stats_var = tk.StringVar(value="0 words   |   0 characters")
        self.quiz_progress_var = tk.StringVar(value="Upload a file or paste text, then analyze it.")
        self.quiz_mode_var = tk.StringVar(value="complete")
        self.quiz_summary_cache = ""
        self.quiz_response_cache = ""
        self.quiz_export_title = "Quiz Support"
        self.quiz_export_name = "quiz_support"

        self.essay_title_var = tk.StringVar()
        self.essay_specific_name_var = tk.StringVar()
        self.essay_word_count_var = tk.StringVar(value="500")
        self.essay_tagalog_var = tk.BooleanVar(value=False)
        self.essay_english_var = tk.BooleanVar(value=True)
        self.essay_progress_var = tk.StringVar(value="Set the essay details, then generate.")
        self.essay_export_title = "Essay Draft"
        self.essay_export_name = "essay_draft"
        self.essay_response_cache = ""

        self.activity_title_var = tk.StringVar()
        self.activity_type_var = tk.StringVar(value="Worksheet")
        self.activity_level_var = tk.StringVar()
        self.activity_progress_var = tk.StringVar(value="Set the activity details, then generate.")
        self.activity_response_cache = ""
        self.activity_export_title = "Activity Draft"
        self.activity_export_name = "activity_draft"

        self.document_title_var = tk.StringVar()
        self.document_type_var = tk.StringVar(value="Study Handout")
        self.document_audience_var = tk.StringVar()
        self.document_progress_var = tk.StringVar(value="Set the document request, then generate.")
        self.document_response_cache = ""
        self.document_export_title = "Document Draft"
        self.document_export_name = "document_draft"

        self.assignment_loaded_path: Path | None = None
        self.assignment_file_var = tk.StringVar(value="No assignment file loaded yet")
        self.assignment_stats_var = tk.StringVar(value="0 words   |   0 characters")
        self.assignment_progress_var = tk.StringVar(value="Upload or paste the assignment, then analyze it.")
        self.assignment_mode_var = tk.StringVar(value="guided")
        self.assignment_summary_cache = ""
        self.assignment_response_cache = ""
        self.assignment_export_title = "Assignment Support"
        self.assignment_export_name = "assignment_support"

        self.codefix_title_var = tk.StringVar()
        self.codefix_language_var = tk.StringVar(value="Python")
        self.codefix_progress_var = tk.StringVar(value="Paste the code and error details, then generate a fix.")
        self.codefix_response_cache = ""
        self.codefix_export_title = "Code Fix"
        self.codefix_export_name = "code_fix"

        self.heading_font = ("Palatino Linotype", 28, "bold")
        self.page_title_font = ("Palatino Linotype", 21, "bold")
        self.card_title_font = ("Palatino Linotype", 15, "bold")
        self.body_font = ("Segoe UI", 11)
        self.label_font = ("Segoe UI Semibold", 10)
        self.small_font = ("Segoe UI", 9)
        self.button_font = ("Segoe UI Semibold", 10)
        self.metric_font = ("Segoe UI Semibold", 13)

        self.page_icon_map = {
            "Dashboard": "◈",
            "Academics": "◎",
            "Developer": "⌘",
            "Quiz Solver": "◆",
            "Essay Generator": "✦",
            "Activity Generator": "▣",
            "Document Generator": "▤",
            "Assignment Solver": "◌",
            "Code Error Fixer": "⌁",
        }
        self.card_icon_map = {
            "Profile Section": "◈",
            "Saved Identity Preview": "◇",
            "Export Folder Routing": "▣",
            "Quick Launch": "➜",
            "How This Flow Works": "◎",
            "Academic Navigation": "◎",
            "Developer Navigation": "⌘",
            "Quiz Solver Workspace": "◆",
            "Essay Generator Workspace": "✦",
            "Activity Generator Workspace": "▣",
            "Document Generator Workspace": "▤",
            "Assignment Solver Workspace": "◌",
            "Code Error Fixer Workspace": "⌁",
            "Quiz Intake": "◆",
            "Summary And Task Reading": "◌",
            "Response Builder": "➤",
            "Essay Builder": "✦",
            "Essay Export Format": "◇",
            "Essay Preview": "✎",
            "Activity Builder": "▣",
            "Activity Preview": "✎",
            "Document Builder": "▤",
            "Document Preview": "✎",
            "Assignment Intake": "◌",
            "Assignment Summary": "◇",
            "Assignment Response": "➤",
            "Code Fix Builder": "⌁",
            "Code Fix Preview": "✎",
        }
        self.button_icon_map = {
            "01  Dashboard": "◈",
            "02  Academics": "◎",
            "03  Developer": "⌘",
            "Refresh AI Status": "↻",
            "Save Profile": "✦",
            "Clear Saved Profile": "⊖",
            "Choose Main Folder": "▣",
            "Open Main Folder": "↗",
            "Open Academics Hub": "◎",
            "Open Developer Hub": "⌘",
            "Open Quiz Solver": "◆",
            "Open Essay Generator": "✦",
            "Open Activity Generator": "▣",
            "Open Document Generator": "▤",
            "Open Assignment Solver": "◌",
            "Open Code Error Fixer": "⌁",
            "Back To Academics": "↩",
            "Back To Developer": "↩",
            "Upload File": "↑",
            "Upload Assignment File": "↑",
            "Clear Quiz Input": "⊖",
            "Clear Assignment Input": "⊖",
            "Analyze Quiz": "◌",
            "Analyze Assignment": "◌",
            "Generate Quiz Guidance": "➤",
            "Generate Assignment Guidance": "➤",
            "Export Result": "↓",
            "Clear Result": "⊖",
            "Generate Essay": "✦",
            "Clear Essay Form": "⊖",
            "Generate Activity": "▣",
            "Clear Activity Form": "⊖",
            "Generate Document": "▤",
            "Clear Document Form": "⊖",
            "Generate Code Fix": "⌁",
            "Clear Code Fixer": "⊖",
        }

        self.option_add("*Font", self.body_font)
        self._ensure_export_directories()
        self._refresh_identity_labels()
        self._build_shell()
        self.refresh_ai_status()
        self.show_page("dashboard")
        self.bind("<Configure>", self._schedule_responsive_refresh, add="+")
        self.after(120, self._apply_responsive_layout)

    def _build_shell(self) -> None:
        shell = tk.Frame(self, bg=PALETTE["bg"])
        shell.pack(fill="both", expand=True, padx=18, pady=18)

        sidebar = tk.Frame(
            shell,
            bg=PALETTE["sidebar"],
            width=250,
            highlightthickness=1,
            highlightbackground=PALETTE["border"],
        )
        sidebar.pack(side="left", fill="y")
        sidebar.pack_propagate(False)

        content = tk.Frame(shell, bg=PALETTE["bg"])
        content.pack(side="left", fill="both", expand=True, padx=(18, 0))

        self._build_sidebar(sidebar)
        self._build_header(content)

        self.page_container = tk.Frame(content, bg=PALETTE["bg"])
        self.page_container.pack(fill="both", expand=True)

        footer = tk.Frame(
            content,
            bg=PALETTE["panel"],
            highlightthickness=1,
            highlightbackground=PALETTE["border"],
        )
        footer.pack(fill="x", pady=(16, 0))
        tk.Label(
            footer,
            textvariable=self.status_var,
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.small_font,
            padx=18,
            pady=14,
            anchor="w",
        ).pack(fill="x")

        self._build_dashboard_page()
        self._build_academics_page()
        self._build_developer_page()
        self._build_quiz_page()
        self._build_essay_page()
        self._build_activity_page()
        self._build_document_page()
        self._build_assignment_page()
        self._build_codefix_page()

    def _build_sidebar(self, parent: tk.Frame) -> None:
        logo_wrap = tk.Frame(parent, bg=PALETTE["sidebar"])
        logo_wrap.pack(fill="x", padx=22, pady=(22, 14))

        emblem = tk.Frame(logo_wrap, bg=PALETTE["accent"], width=56, height=56)
        emblem.pack(anchor="w")
        emblem.pack_propagate(False)
        tk.Label(
            emblem,
            text="H",
            bg=PALETTE["accent"],
            fg=PALETTE["bg"],
            font=("Palatino Linotype", 22, "bold"),
        ).pack(expand=True)

        tk.Label(
            logo_wrap,
            text="Hugyoku",
            bg=PALETTE["sidebar"],
            fg=PALETTE["text"],
            font=("Palatino Linotype", 24, "bold"),
            pady=10,
        ).pack(anchor="w")
        tk.Label(
            logo_wrap,
            text="Premium study and utility desk for academic generators, solvers, and code-fixing support.",
            bg=PALETTE["sidebar"],
            fg=PALETTE["muted"],
            font=self.body_font,
            justify="left",
            wraplength=230,
        ).pack(anchor="w")

        nav = tk.Frame(parent, bg=PALETTE["sidebar"])
        nav.pack(fill="x", padx=18, pady=(18, 12))

        items = [
            ("dashboard", "01  Dashboard"),
            ("academics", "02  Academics"),
            ("developer", "03  Developer"),
        ]
        for key, label in items:
            button = self._make_button(nav, label, lambda page=key: self.show_page(page), variant="nav", anchor="w")
            button.pack(fill="x", pady=6)
            self.nav_buttons[key] = button

        card = tk.Frame(
            parent,
            bg=PALETTE["panel"],
            highlightthickness=1,
            highlightbackground=PALETTE["border"],
        )
        card.pack(fill="x", padx=18, pady=(18, 10))
        tk.Frame(card, bg=PALETTE["accent"], height=3).pack(fill="x")
        inner = tk.Frame(card, bg=PALETTE["panel"])
        inner.pack(fill="both", expand=True, padx=16, pady=16)
        tk.Label(
            inner,
            text="Educational use only",
            bg=PALETTE["panel"],
            fg=PALETTE["text"],
            font=self.card_title_font,
        ).pack(anchor="w")
        tk.Label(
            inner,
            text="Use these flows for studying, drafting, and understanding tasks from your files. Review outputs before submitting anything.",
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            wraplength=228,
            justify="left",
        ).pack(anchor="w", pady=(8, 0))

        tk.Label(
            parent,
            text="Deluxe premium academic workflow",
            bg=PALETTE["sidebar"],
            fg=PALETTE["accent_soft"],
            font=self.small_font,
            pady=18,
        ).pack(side="bottom")

    def _build_header(self, parent: tk.Frame) -> None:
        header = tk.Frame(
            parent,
            bg=PALETTE["panel"],
            highlightthickness=1,
            highlightbackground=PALETTE["border"],
        )
        header.pack(fill="x", pady=(0, 16))

        left = tk.Frame(header, bg=PALETTE["panel"])
        left.pack(side="left", fill="x", expand=True, padx=22, pady=18)
        tk.Label(
            left,
            textvariable=self.header_title_var,
            bg=PALETTE["panel"],
            fg=PALETTE["text"],
            font=self.heading_font,
            anchor="w",
        ).pack(anchor="w")
        tk.Label(
            left,
            textvariable=self.header_subtitle_var,
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.body_font,
            anchor="w",
            wraplength=760,
            justify="left",
        ).pack(anchor="w", pady=(6, 0))

        right = tk.Frame(header, bg=PALETTE["panel"])
        right.pack(side="right", padx=22, pady=18)

        self.ai_chip = tk.Label(
            right,
            textvariable=self.ai_chip_var,
            bg=PALETTE["panel_alt"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
            padx=14,
            pady=8,
        )
        self.ai_chip.pack(anchor="e")

        tk.Label(
            right,
            textvariable=self.profile_summary_var,
            bg=PALETTE["panel"],
            fg=PALETTE["accent_soft"],
            font=self.small_font,
            anchor="e",
            justify="right",
        ).pack(anchor="e", pady=(8, 0))
        tk.Label(
            right,
            textvariable=self.profile_date_var,
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.small_font,
            anchor="e",
            justify="right",
        ).pack(anchor="e", pady=(4, 0))
        tk.Label(
            right,
            textvariable=self.model_var,
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.small_font,
            anchor="e",
            justify="right",
            wraplength=340,
        ).pack(anchor="e", pady=(4, 0))
        self._make_button(right, "Refresh AI Status", self.refresh_ai_status, variant="secondary").pack(
            anchor="e", pady=(12, 0)
        )

    def _build_page_header(self, parent: tk.Frame, title: str, subtitle: str) -> None:
        tk.Label(
            parent,
            text=self._page_title_text(title),
            bg=PALETTE["bg"],
            fg=PALETTE["text"],
            font=self.page_title_font,
            anchor="w",
        ).pack(anchor="w")
        tk.Label(
            parent,
            text=subtitle,
            bg=PALETTE["bg"],
            fg=PALETTE["muted"],
            font=self.body_font,
            wraplength=980,
            justify="left",
            anchor="w",
        ).pack(anchor="w", pady=(6, 18))

    def _create_page(self, key: str) -> tk.Frame:
        page = tk.Frame(self.page_container, bg=PALETTE["bg"])
        page.place(relx=0, rely=0, relwidth=1, relheight=1)
        self.pages[key] = page
        return page

    def _bind_mousewheel(self, canvas: tk.Canvas) -> None:
        def on_mousewheel(event: tk.Event) -> None:
            delta = getattr(event, "delta", 0)
            if delta:
                canvas.yview_scroll(int(-delta / 120), "units")

        def bind_canvas(_event: tk.Event | None = None) -> None:
            canvas.bind_all("<MouseWheel>", on_mousewheel)

        def unbind_canvas(_event: tk.Event | None = None) -> None:
            canvas.unbind_all("<MouseWheel>")

        canvas.bind("<Enter>", bind_canvas, add="+")
        canvas.bind("<Leave>", unbind_canvas, add="+")

    def _create_scrollable_section(self, parent: tk.Widget, bg: str | None = None) -> tuple[tk.Frame, tk.Frame, tk.Canvas]:
        section_bg = bg or PALETTE["bg"]
        shell = tk.Frame(parent, bg=section_bg)
        canvas = tk.Canvas(
            shell,
            bg=section_bg,
            highlightthickness=0,
            bd=0,
            relief="flat",
            yscrollincrement=24,
        )
        scrollbar = tk.Scrollbar(
            shell,
            orient="vertical",
            command=canvas.yview,
            troughcolor=section_bg,
            activebackground=PALETTE["accent"],
            bg=PALETTE["panel_alt"],
            width=12,
        )
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        body = tk.Frame(canvas, bg=section_bg)
        window_id = canvas.create_window((0, 0), window=body, anchor="nw")

        def on_body_configure(_event: tk.Event | None = None) -> None:
            canvas.configure(scrollregion=canvas.bbox("all"))

        def on_canvas_configure(event: tk.Event) -> None:
            canvas.itemconfigure(window_id, width=event.width)
            canvas.configure(scrollregion=canvas.bbox("all"))

        body.bind("<Configure>", on_body_configure, add="+")
        canvas.bind("<Configure>", on_canvas_configure, add="+")
        self._bind_mousewheel(canvas)
        return shell, body, canvas

    def _page_title_text(self, title: str) -> str:
        icon = self.page_icon_map.get(title)
        return f"{icon}  {title}" if icon else title

    def _card_title_text(self, title: str) -> tuple[str, str | None]:
        return title, self.card_icon_map.get(title)

    def _button_text(self, text: str) -> str:
        icon = self.button_icon_map.get(text)
        return f"{icon}  {text}" if icon else text

    def _pointer_inside(self, widget: tk.Widget) -> bool:
        try:
            x = self.winfo_pointerx()
            y = self.winfo_pointery()
            left = widget.winfo_rootx()
            top = widget.winfo_rooty()
            right = left + widget.winfo_width()
            bottom = top + widget.winfo_height()
        except tk.TclError:
            return False
        return left <= x <= right and top <= y <= bottom

    def _bind_hover_group(self, root: tk.Widget, widgets: list[tk.Widget], on_enter: object, on_leave: object) -> None:
        def handle_enter(_event: object | None = None) -> None:
            on_enter()

        def handle_leave(_event: object | None = None) -> None:
            self.after(25, lambda: on_enter() if self._pointer_inside(root) else on_leave())

        for widget in widgets:
            widget.bind("<Enter>", handle_enter, add="+")
            widget.bind("<Leave>", handle_leave, add="+")

    def _apply_button_palette(
        self,
        button: tk.Button,
        bg: str,
        fg: str,
        hover_bg: str,
        hover_fg: str,
        border: str,
        hover_border: str,
    ) -> None:
        button._base_bg = bg
        button._base_fg = fg
        button._hover_bg = hover_bg
        button._hover_fg = hover_fg
        button._base_border = border
        button._hover_border = hover_border
        button.configure(
            bg=bg,
            fg=fg,
            activebackground=hover_bg,
            activeforeground=hover_fg,
            highlightbackground=border,
            highlightcolor=hover_border,
        )

    def _create_card(self, parent: tk.Widget, title: str, subtitle: str, accent: str) -> tuple[tk.Frame, tk.Frame]:
        display_title, icon = self._card_title_text(title)
        outer = tk.Frame(parent, bg=PALETTE["glass_shadow"], highlightthickness=0)
        card = tk.Frame(
            outer,
            bg=PALETTE["glass"],
            highlightthickness=1,
            highlightbackground=PALETTE["glass_border"],
        )
        card.pack(fill="both", expand=True, padx=(0, 3), pady=(0, 3))
        accent_bar = tk.Frame(card, bg=accent, height=3)
        accent_bar.pack(fill="x")

        shell = tk.Frame(card, bg=PALETTE["glass_inner"])
        shell.pack(fill="both", expand=True, padx=20, pady=18)
        header_row = tk.Frame(shell, bg=PALETTE["glass_inner"])
        header_row.pack(fill="x")
        badge = tk.Label(
            header_row,
            text=icon or "◇",
            bg=PALETTE["glass_chip"],
            fg=accent,
            font=("Segoe UI Symbol", 11, "bold"),
            padx=10,
            pady=6,
        )
        badge.pack(side="left")
        title_label = tk.Label(
            header_row,
            text=display_title,
            bg=PALETTE["glass_inner"],
            fg=PALETTE["text"],
            font=self.card_title_font,
            anchor="w",
        )
        title_label.pack(side="left", padx=(12, 0))
        subtitle_label = tk.Label(
            shell,
            text=subtitle,
            bg=PALETTE["glass_inner"],
            fg=PALETTE["muted"],
            justify="left",
            wraplength=620,
            anchor="w",
        )
        subtitle_label.pack(anchor="w", pady=(10, 14))
        body = tk.Frame(shell, bg=PALETTE["glass_inner"])
        body.pack(fill="both", expand=True)

        def on_enter() -> None:
            outer.configure(bg=PALETTE["glass_shadow_lift"])
            card.configure(bg=PALETTE["glass_hover"], highlightbackground=accent)
            shell.configure(bg=PALETTE["glass_hover"])
            header_row.configure(bg=PALETTE["glass_hover"])
            title_label.configure(bg=PALETTE["glass_hover"], fg=PALETTE["accent_soft"])
            subtitle_label.configure(bg=PALETTE["glass_hover"])
            body.configure(bg=PALETTE["glass_hover"])
            badge.configure(bg=accent, fg=PALETTE["bg"])

        def on_leave() -> None:
            outer.configure(bg=PALETTE["glass_shadow"])
            card.configure(bg=PALETTE["glass"], highlightbackground=PALETTE["glass_border"])
            shell.configure(bg=PALETTE["glass_inner"])
            header_row.configure(bg=PALETTE["glass_inner"])
            title_label.configure(bg=PALETTE["glass_inner"], fg=PALETTE["text"])
            subtitle_label.configure(bg=PALETTE["glass_inner"])
            body.configure(bg=PALETTE["glass_inner"])
            badge.configure(bg=PALETTE["glass_chip"], fg=accent)

        self._bind_hover_group(
            outer,
            [outer, card, accent_bar, shell, header_row, badge, title_label, subtitle_label, body],
            on_enter,
            on_leave,
        )
        return outer, body

    def _make_button(
        self,
        parent: tk.Widget,
        text: str,
        command: object,
        variant: str = "primary",
        anchor: str = "center",
    ) -> tk.Button:
        styles = {
            "primary": {
                "bg": PALETTE["accent"],
                "fg": PALETTE["bg"],
                "hover_bg": PALETTE["accent_soft"],
                "hover_fg": PALETTE["bg"],
                "border": PALETTE["accent_dark"],
                "hover_border": PALETTE["accent_soft"],
            },
            "secondary": {
                "bg": PALETTE["panel_alt"],
                "fg": PALETTE["text"],
                "hover_bg": PALETTE["glass_hover"],
                "hover_fg": PALETTE["accent_soft"],
                "border": PALETTE["glass_border"],
                "hover_border": PALETTE["accent"],
            },
            "ghost": {
                "bg": PALETTE["bg"],
                "fg": PALETTE["accent_soft"],
                "hover_bg": PALETTE["panel_alt"],
                "hover_fg": PALETTE["text"],
                "border": PALETTE["border"],
                "hover_border": PALETTE["accent_soft"],
            },
            "nav": {
                "bg": PALETTE["sidebar"],
                "fg": PALETTE["text"],
                "hover_bg": PALETTE["panel_alt"],
                "hover_fg": PALETTE["accent_soft"],
                "border": PALETTE["sidebar"],
                "hover_border": PALETTE["accent_dark"],
            },
        }
        config = styles[variant]
        button = tk.Button(
            parent,
            text=self._button_text(text),
            command=command,
            disabledforeground="#727C8F",
            relief="flat",
            bd=0,
            highlightthickness=1,
            cursor="hand2",
            font=self.button_font,
            padx=18,
            pady=12,
            anchor=anchor,
        )
        self._apply_button_palette(
            button,
            config["bg"],
            config["fg"],
            config["hover_bg"],
            config["hover_fg"],
            config["border"],
            config["hover_border"],
        )

        def on_enter(_event: object | None = None) -> None:
            if str(button.cget("state")) == "disabled":
                return
            button.configure(
                bg=button._hover_bg,
                fg=button._hover_fg,
                highlightbackground=button._hover_border,
            )

        def on_leave(_event: object | None = None) -> None:
            if str(button.cget("state")) == "disabled":
                return
            button.configure(
                bg=button._base_bg,
                fg=button._base_fg,
                highlightbackground=button._base_border,
            )

        button.bind("<Enter>", on_enter, add="+")
        button.bind(
            "<Leave>",
            lambda _event: self.after(20, lambda: on_enter() if self._pointer_inside(button) else on_leave()),
            add="+",
        )
        return button

    def _register_ai_button(self, button: tk.Button) -> tk.Button:
        self.ai_action_buttons.append(button)
        return button

    def _styled_checkbutton(self, parent: tk.Widget, text: str, variable: tk.BooleanVar) -> tk.Checkbutton:
        return tk.Checkbutton(
            parent,
            text=text,
            variable=variable,
            bg=PALETTE["panel"],
            fg=PALETTE["text"],
            activebackground=PALETTE["panel"],
            activeforeground=PALETTE["text"],
            selectcolor=PALETTE["panel_alt"],
            highlightthickness=0,
            bd=0,
            font=self.body_font,
        )

    def _styled_radiobutton(self, parent: tk.Widget, text: str, variable: tk.StringVar, value: str) -> tk.Radiobutton:
        return tk.Radiobutton(
            parent,
            text=text,
            variable=variable,
            value=value,
            bg=PALETTE["panel"],
            fg=PALETTE["text"],
            activebackground=PALETTE["panel"],
            activeforeground=PALETTE["text"],
            selectcolor=PALETTE["panel_alt"],
            highlightthickness=0,
            bd=0,
            font=self.body_font,
        )

    def _labeled_entry(self, parent: tk.Widget, label: str, variable: tk.StringVar, hint: str) -> tk.Entry:
        block = tk.Frame(parent, bg=PALETTE["panel"])
        block.pack(fill="x", pady=(0, 16))
        tk.Label(block, text=label, bg=PALETTE["panel"], fg=PALETTE["text"], font=self.label_font).pack(anchor="w")
        shell = tk.Frame(block, bg=PALETTE["border"])
        shell.pack(fill="x", pady=(8, 0))
        entry = tk.Entry(
            shell,
            textvariable=variable,
            bg=PALETTE["input"],
            fg=PALETTE["text"],
            insertbackground=PALETTE["text"],
            relief="flat",
            bd=0,
            highlightthickness=0,
            font=self.body_font,
            selectbackground=PALETTE["accent_dark"],
            selectforeground=PALETTE["text"],
        )
        entry.pack(fill="x", padx=1, pady=1, ipady=10)
        tk.Label(block, text=hint, bg=PALETTE["panel"], fg=PALETTE["muted"], font=self.small_font).pack(
            anchor="w", pady=(6, 0)
        )
        return entry

    def _labeled_text(
        self,
        parent: tk.Widget,
        label: str,
        height: int,
        hint: str,
        read_only: bool = False,
    ) -> ScrolledText:
        block = tk.Frame(parent, bg=PALETTE["panel"])
        block.pack(fill="both", expand=True, pady=(0, 16))
        tk.Label(block, text=label, bg=PALETTE["panel"], fg=PALETTE["text"], font=self.label_font).pack(anchor="w")
        shell = tk.Frame(block, bg=PALETTE["border"])
        shell.pack(fill="both", expand=True, pady=(8, 0))
        text = ScrolledText(
            shell,
            height=height,
            wrap="word",
            bg=PALETTE["input"],
            fg=PALETTE["text"],
            insertbackground=PALETTE["text"],
            relief="flat",
            bd=0,
            highlightthickness=0,
            font=self.body_font,
            padx=14,
            pady=14,
            selectbackground=PALETTE["accent_dark"],
            selectforeground=PALETTE["text"],
        )
        text.pack(fill="both", expand=True, padx=1, pady=1)
        if read_only:
            text.config(state="disabled")
        tk.Label(block, text=hint, bg=PALETTE["panel"], fg=PALETTE["muted"], font=self.small_font).pack(
            anchor="w", pady=(6, 0)
        )
        return text

    def _metric_tile(self, parent: tk.Widget, label: str, variable: tk.StringVar, value_color: str) -> tk.Frame:
        tile = tk.Frame(
            parent,
            bg=PALETTE["glass_inner"],
            highlightthickness=1,
            highlightbackground=PALETTE["glass_border"],
        )
        label_widget = tk.Label(
            tile,
            text=label,
            bg=PALETTE["glass_inner"],
            fg=PALETTE["muted"],
            font=self.small_font,
        )
        label_widget.pack(anchor="w", padx=14, pady=(12, 4))
        value_widget = tk.Label(
            tile,
            textvariable=variable,
            bg=PALETTE["glass_inner"],
            fg=value_color,
            font=self.metric_font,
            justify="left",
            wraplength=360,
        )
        value_widget.pack(anchor="w", padx=14, pady=(0, 12))

        def on_enter() -> None:
            tile.configure(bg=PALETTE["glass_hover"], highlightbackground=PALETTE["accent"])
            label_widget.configure(bg=PALETTE["glass_hover"])
            value_widget.configure(bg=PALETTE["glass_hover"])

        def on_leave() -> None:
            tile.configure(bg=PALETTE["glass_inner"], highlightbackground=PALETTE["glass_border"])
            label_widget.configure(bg=PALETTE["glass_inner"])
            value_widget.configure(bg=PALETTE["glass_inner"])

        self._bind_hover_group(tile, [tile, label_widget, value_widget], on_enter, on_leave)
        return tile

    def _bullet(self, parent: tk.Widget, text: str) -> tk.Label:
        return tk.Label(
            parent,
            text=f"- {text}",
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            justify="left",
            wraplength=430,
            anchor="w",
        )

    def _schedule_responsive_refresh(self, _event: tk.Event | None = None) -> None:
        if self._responsive_after_id is not None:
            self.after_cancel(self._responsive_after_id)
        self._responsive_after_id = self.after(80, self._apply_responsive_layout)

    def _apply_responsive_layout(self) -> None:
        self._responsive_after_id = None
        self._update_dashboard_layout()
        self._update_quiz_layout()
        self._update_essay_layout()

    def _update_dashboard_layout(self) -> None:
        if not self.dashboard_layout:
            return

        grid = self.dashboard_layout["grid"]
        width = grid.winfo_width() or self.page_container.winfo_width() or self.winfo_width()
        compact = width < 1160

        profile = self.dashboard_layout["profile"]
        preview = self.dashboard_layout["preview"]
        launch = self.dashboard_layout["launch"]
        notes = self.dashboard_layout["notes"]

        for widget in (profile, preview, launch, notes):
            widget.grid_forget()

        if compact:
            grid.grid_columnconfigure(0, weight=1)
            grid.grid_columnconfigure(1, weight=0)
            profile.grid(row=0, column=0, sticky="nsew", padx=0, pady=(0, 10))
            preview.grid(row=1, column=0, sticky="nsew", padx=0, pady=(0, 10))
            launch.grid(row=2, column=0, sticky="nsew", padx=0, pady=(0, 10))
            notes.grid(row=3, column=0, sticky="nsew", padx=0, pady=(0, 10))
        else:
            grid.grid_columnconfigure(0, weight=5)
            grid.grid_columnconfigure(1, weight=4)
            profile.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=(0, 10))
            preview.grid(row=0, column=1, sticky="nsew", padx=(10, 0), pady=(0, 10))
            launch.grid(row=1, column=0, sticky="nsew", padx=(0, 10), pady=(10, 0))
            notes.grid(row=1, column=1, sticky="nsew", padx=(10, 0), pady=(10, 0))

    def _update_quiz_layout(self) -> None:
        if not self.quiz_layout:
            return

        grid = self.quiz_layout["grid"]
        width = grid.winfo_width() or self.page_container.winfo_width() or self.winfo_width()
        compact = width < 1160

        intake = self.quiz_layout["intake"]
        right = self.quiz_layout["right"]

        intake.grid_forget()
        right.grid_forget()

        if compact:
            grid.grid_columnconfigure(0, weight=1)
            grid.grid_columnconfigure(1, weight=0)
            intake.grid(row=0, column=0, sticky="nsew", padx=0, pady=(0, 10))
            right.grid(row=1, column=0, sticky="nsew", padx=0, pady=(0, 10))
        else:
            grid.grid_columnconfigure(0, weight=5)
            grid.grid_columnconfigure(1, weight=5)
            intake.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=0)
            right.grid(row=0, column=1, sticky="nsew", padx=(10, 0), pady=0)

    def _update_essay_layout(self) -> None:
        if not self.essay_layout:
            return

        grid = self.essay_layout["grid"]
        width = grid.winfo_width() or self.page_container.winfo_width() or self.winfo_width()
        compact = width < 1160

        build = self.essay_layout["build"]
        result = self.essay_layout["result"]

        build.grid_forget()
        result.grid_forget()

        if compact:
            grid.grid_columnconfigure(0, weight=1)
            grid.grid_columnconfigure(1, weight=0)
            build.grid(row=0, column=0, sticky="nsew", padx=0, pady=(0, 10))
            result.grid(row=1, column=0, sticky="nsew", padx=0, pady=(0, 10))
        else:
            grid.grid_columnconfigure(0, weight=4)
            grid.grid_columnconfigure(1, weight=5)
            build.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=0)
            result.grid(row=0, column=1, sticky="nsew", padx=(10, 0), pady=0)

    def _build_dashboard_page(self) -> None:
        page = self._create_page("dashboard")
        self._build_page_header(
            page,
            "Dashboard",
            "Save your identity, choose one main save folder, and let the app auto-create dedicated subfolders for every academic and developer tool.",
        )

        scroll_shell, scroll_body, _dashboard_canvas = self._create_scrollable_section(page)
        scroll_shell.pack(fill="both", expand=True)

        grid = tk.Frame(scroll_body, bg=PALETTE["bg"])
        grid.pack(fill="both", expand=True, padx=(0, 4))
        grid.grid_columnconfigure(0, weight=5)
        grid.grid_columnconfigure(1, weight=4)
        grid.grid_rowconfigure(0, weight=1)
        grid.grid_rowconfigure(1, weight=1)

        profile_card, profile_body = self._create_card(
            grid,
            "Profile Section",
            "Enter your name, decide whether to stamp today's date, then choose one main folder where all generated files will be organized automatically.",
            PALETTE["accent"],
        )
        profile_card.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=(0, 10))
        self._labeled_entry(
            profile_body,
            "Enter your name",
            self.name_input_var,
            "This gets remembered locally in the app once you hit save.",
        )
        check_row = tk.Frame(profile_body, bg=PALETTE["panel"])
        check_row.pack(fill="x", pady=(0, 12))
        self._styled_checkbutton(check_row, "Add date today", self.include_date_var).pack(side="left")
        action_row = tk.Frame(profile_body, bg=PALETTE["panel"])
        action_row.pack(fill="x", pady=(8, 12))
        self._make_button(action_row, "Save Profile", self.save_profile, variant="primary").pack(side="left", padx=(0, 10))
        self._make_button(action_row, "Clear Saved Profile", self.clear_profile, variant="ghost").pack(side="left")
        folder_row = tk.Frame(profile_body, bg=PALETTE["panel"])
        folder_row.pack(fill="x", pady=(0, 10))
        self._make_button(folder_row, "Choose Main Folder", self.choose_main_folder, variant="secondary").pack(side="left", padx=(0, 10))
        self._make_button(folder_row, "Open Main Folder", self.open_main_folder, variant="ghost").pack(side="left")
        tk.Label(
            profile_body,
            textvariable=self.main_folder_var,
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.small_font,
            justify="left",
            wraplength=560,
        ).pack(anchor="w")
        tk.Label(
            profile_body,
            text="Output Options",
            bg=PALETTE["panel"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
        ).pack(anchor="w", pady=(16, 8))
        output_options = tk.Frame(profile_body, bg=PALETTE["panel"])
        output_options.pack(fill="x", pady=(0, 6))
        self._styled_checkbutton(
            output_options,
            "Include saved name in export",
            self.output_include_name_var,
        ).grid(row=0, column=0, sticky="w", padx=(0, 18), pady=(0, 8))
        self._styled_checkbutton(
            output_options,
            "Include date in export",
            self.output_include_date_var,
        ).grid(row=0, column=1, sticky="w", pady=(0, 8))
        self._styled_checkbutton(
            output_options,
            "Include essay heading suggestion",
            self.essay_include_heading_var,
        ).grid(row=1, column=0, sticky="w", padx=(0, 18))
        self._styled_checkbutton(
            output_options,
            "Include self-check tip",
            self.essay_include_tip_var,
        ).grid(row=1, column=1, sticky="w")
        output_options.grid_columnconfigure(0, weight=1)
        output_options.grid_columnconfigure(1, weight=1)
        tk.Label(
            profile_body,
            text="These checkboxes affect exported Word output so you can keep the document clean or more detailed.",
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.small_font,
            justify="left",
            wraplength=560,
        ).pack(anchor="w", pady=(6, 0))

        preview_card, preview_body = self._create_card(
            grid,
            "Export Folder Routing",
            "Choose one main folder and the app will automatically create separate export directories for all active tools inside it.",
            PALETTE["info"],
        )
        preview_card.grid(row=0, column=1, sticky="nsew", padx=(10, 0), pady=(0, 10))
        self._metric_tile(preview_body, "Main folder", self.main_folder_var, PALETTE["accent_soft"]).pack(fill="x", pady=(0, 12))
        self._metric_tile(preview_body, "Academics suite folders", self.academics_folders_var, PALETTE["text"]).pack(fill="x", pady=(0, 12))
        self._metric_tile(preview_body, "Developer suite folder", self.developer_folder_var, PALETTE["text"]).pack(fill="x", pady=(0, 12))
        self._metric_tile(preview_body, "Routing note", self.folder_note_var, PALETTE["text"]).pack(fill="x")

        launch_card, launch_body = self._create_card(
            grid,
            "Quick Launch",
            "The suite is now grouped into clean hubs so the menu stays minimal even with more tools available.",
            PALETTE["success"],
        )
        launch_card.grid(row=1, column=0, sticky="nsew", padx=(0, 10), pady=(10, 0))
        tk.Label(
            launch_body,
            text="Go straight where you need",
            bg=PALETTE["panel"],
            fg=PALETTE["accent_soft"],
            font=("Segoe UI Semibold", 12),
        ).pack(anchor="w")
        tk.Label(
            launch_body,
            text="Use the Academics hub for school tools and the Developer hub for code fixing. Each workspace has its own page and autosave folder.",
            bg=PALETTE["panel"],
            fg=PALETTE["text"],
            wraplength=520,
            justify="left",
            font=("Segoe UI", 13),
        ).pack(anchor="w", pady=(8, 18))
        launch_buttons = tk.Frame(launch_body, bg=PALETTE["panel"])
        launch_buttons.pack(anchor="w")
        self._make_button(launch_buttons, "Open Academics Hub", lambda: self.show_page("academics"), variant="primary").pack(
            side="left", padx=(0, 10)
        )
        self._make_button(launch_buttons, "Open Developer Hub", lambda: self.show_page("developer"), variant="secondary").pack(
            side="left"
        )

        notes_card, notes_body = self._create_card(
            grid,
            "How This Flow Works",
            "The dashboard now manages both your identity settings and your export folder routing.",
            PALETTE["warning"],
        )
        notes_card.grid(row=1, column=1, sticky="nsew", padx=(10, 0), pady=(10, 0))
        for line in [
            "1. Save your name and optional date from this dashboard once.",
            "2. Choose a single main folder and the app will create subfolders for quiz, assignment, essay, activity, document, and code-fix exports.",
            "3. Open a hub, choose one focused tool page, and the export will route itself to the correct subfolder automatically.",
        ]:
            self._bullet(notes_body, line).pack(anchor="w", fill="x", pady=(0, 10))

        self.dashboard_layout = {
            "grid": grid,
            "profile": profile_card,
            "preview": preview_card,
            "launch": launch_card,
            "notes": notes_card,
        }

    def _hub_tool_card(
        self,
        parent: tk.Frame,
        card_title: str,
        number: str,
        description: str,
        folder_var: tk.StringVar,
        button_text: str,
        command: object,
        accent: str,
        variant: str = "secondary",
    ) -> tk.Frame:
        card, body = self._create_card(parent, card_title, description, accent)
        card.pack(fill="x", pady=(0, 14))
        tk.Label(
            body,
            text=number,
            bg=PALETTE["glass_inner"],
            fg=PALETTE["accent_soft"],
            font=("Segoe UI Semibold", 13),
        ).pack(anchor="w")
        tk.Label(
            body,
            textvariable=folder_var,
            bg=PALETTE["glass_inner"],
            fg=PALETTE["muted"],
            wraplength=900,
            justify="left",
            font=self.small_font,
        ).pack(anchor="w", pady=(8, 12))
        self._make_button(body, button_text, command, variant=variant).pack(anchor="w")
        return card

    def _build_tool_page_shell(
        self,
        key: str,
        title: str,
        subtitle: str,
        folder_var: tk.StringVar,
        back_label: str,
        back_target: str,
    ) -> tk.Frame:
        page = self._create_page(key)
        self._build_page_header(page, title, subtitle)
        self._make_button(page, back_label, lambda target=back_target: self.show_page(target), variant="ghost").pack(
            anchor="w", pady=(0, 10)
        )
        tk.Label(
            page,
            textvariable=folder_var,
            bg=PALETTE["bg"],
            fg=PALETTE["muted"],
            font=self.small_font,
            wraplength=980,
            justify="left",
            anchor="w",
        ).pack(anchor="w", pady=(0, 12))
        tool_host = tk.Frame(page, bg=PALETTE["bg"])
        tool_host.pack(fill="both", expand=True)
        return tool_host

    def _build_academics_page(self) -> None:
        page = self._create_page("academics")
        self._build_page_header(
            page,
            "Academics",
            "Choose a focused school workspace. Each tool now has its own page and autosave subfolder so the layout stays minimal, clearer, and easier to use.",
        )

        scroll_shell, scroll_body, _academics_canvas = self._create_scrollable_section(page)
        scroll_shell.pack(fill="both", expand=True)

        hub = tk.Frame(scroll_body, bg=PALETTE["bg"])
        hub.pack(fill="both", expand=True, padx=(0, 4))

        intro_card, intro_body = self._create_card(
            hub,
            "Academic Navigation",
            "Open the exact tool you need. Quiz, assignment, essay, activity, and document flows are now separated into their own workspaces.",
            PALETTE["accent"],
        )
        intro_card.pack(fill="x", pady=(0, 14))
        tk.Label(
            intro_body,
            textvariable=self.academics_folders_var,
            bg=PALETTE["glass_inner"],
            fg=PALETTE["muted"],
            wraplength=900,
            justify="left",
            font=self.small_font,
        ).pack(anchor="w")

        self._hub_tool_card(
            hub,
            "Quiz Solver Workspace",
            "01  Quiz Solver",
            "Read quiz files, summarize the task, and generate guided response support in a dedicated workspace.",
            self.quiz_folder_var,
            "Open Quiz Solver",
            lambda: self.show_academic_tool("quiz"),
            PALETTE["accent"],
            "primary",
        )
        self._hub_tool_card(
            hub,
            "Assignment Solver Workspace",
            "02  Assignment Solver",
            "Analyze assignments from pasted text or uploaded files, then generate a guided draft or response plan.",
            self.assignment_folder_var,
            "Open Assignment Solver",
            lambda: self.show_academic_tool("assignment"),
            PALETTE["info"],
        )
        self._hub_tool_card(
            hub,
            "Essay Generator Workspace",
            "03  Essay Generator",
            "Build long-form essay drafts with output options, optional name overrides, and cleaner export formatting.",
            self.essay_folder_var,
            "Open Essay Generator",
            lambda: self.show_academic_tool("essay"),
            PALETTE["warning"],
        )
        self._hub_tool_card(
            hub,
            "Activity Generator Workspace",
            "04  Activity Generator",
            "Generate worksheets, reflections, or classroom activities from a topic, level, and instruction set.",
            self.activity_folder_var,
            "Open Activity Generator",
            lambda: self.show_academic_tool("activity"),
            PALETTE["success"],
        )
        self._hub_tool_card(
            hub,
            "Document Generator Workspace",
            "05  Document Generator",
            "Create structured school documents like handouts, reports, and study materials from a simple request.",
            self.document_folder_var,
            "Open Document Generator",
            lambda: self.show_academic_tool("document"),
            PALETTE["accent_soft"],
        )

        guide_card, guide_body = self._create_card(
            hub,
            "How This Flow Works",
            "The structure is now cleaner: enter a hub, choose a tool, work on a single task, export, then start again with a fresh workspace.",
            PALETTE["info"],
        )
        guide_card.pack(fill="x")
        for line in [
            "1. Dashboard stores your profile and main save folder.",
            "2. Academics contains school-focused generators and solvers.",
            "3. Each academic tool autosaves into its own dedicated subfolder.",
        ]:
            tk.Label(
                guide_body,
                text=line,
                bg=PALETTE["glass_inner"],
                fg=PALETTE["muted"],
                justify="left",
                wraplength=900,
                anchor="w",
            ).pack(anchor="w", fill="x", pady=(0, 10))

    def _build_developer_page(self) -> None:
        page = self._create_page("developer")
        self._build_page_header(
            page,
            "Developer",
            "A separate hub for coding support. Right now this contains a focused code error fixer so it does not clutter the academics tools.",
        )

        scroll_shell, scroll_body, _developer_canvas = self._create_scrollable_section(page)
        scroll_shell.pack(fill="both", expand=True)

        hub = tk.Frame(scroll_body, bg=PALETTE["bg"])
        hub.pack(fill="both", expand=True, padx=(0, 4))

        intro_card, intro_body = self._create_card(
            hub,
            "Developer Navigation",
            "Developer tools live separately from school tools so the menu stays cleaner and easier to scan on smaller windows.",
            PALETTE["warning"],
        )
        intro_card.pack(fill="x", pady=(0, 14))
        tk.Label(
            intro_body,
            textvariable=self.developer_folder_var,
            bg=PALETTE["glass_inner"],
            fg=PALETTE["muted"],
            wraplength=900,
            justify="left",
            font=self.small_font,
        ).pack(anchor="w")

        self._hub_tool_card(
            hub,
            "Code Error Fixer Workspace",
            "01  Code Error Fixer",
            "Paste code, explain the bug or error, and get a cleaner fix plus a short explanation of what changed.",
            self.codefix_folder_var,
            "Open Code Error Fixer",
            lambda: self.show_page("codefix"),
            PALETTE["warning"],
            "primary",
        )

    def _build_quiz_page(self) -> None:
        tool_host = self._build_tool_page_shell(
            "quiz",
            "Quiz Solver",
            "A dedicated quiz solver workspace for reading files, understanding the task, and generating a guided response without mixing in other tools.",
            self.quiz_folder_var,
            "Back To Academics",
            "academics",
        )
        self._build_quiz_tool(tool_host)

    def _build_essay_page(self) -> None:
        tool_host = self._build_tool_page_shell(
            "essay",
            "Essay Generator",
            "A dedicated essay generator workspace for building drafts, customizing export formatting, and saving to the essay subfolder.",
            self.essay_folder_var,
            "Back To Academics",
            "academics",
        )
        self._build_essay_tool(tool_host)

    def _build_activity_page(self) -> None:
        tool_host = self._build_tool_page_shell(
            "activity",
            "Activity Generator",
            "Generate a classroom activity, worksheet, reflection prompt, or guided exercise from a topic and instruction set.",
            self.activity_folder_var,
            "Back To Academics",
            "academics",
        )
        self._build_activity_tool(tool_host)

    def _build_document_page(self) -> None:
        tool_host = self._build_tool_page_shell(
            "document",
            "Document Generator",
            "Generate structured academic documents like study handouts, reports, or simple formal school papers from a clean prompt.",
            self.document_folder_var,
            "Back To Academics",
            "academics",
        )
        self._build_document_tool(tool_host)

    def _build_assignment_page(self) -> None:
        tool_host = self._build_tool_page_shell(
            "assignment",
            "Assignment Solver",
            "Use a dedicated assignment page for uploaded tasks, guided analysis, and a structured sample response or completion plan.",
            self.assignment_folder_var,
            "Back To Academics",
            "academics",
        )
        self._build_assignment_tool(tool_host)

    def _build_codefix_page(self) -> None:
        tool_host = self._build_tool_page_shell(
            "codefix",
            "Code Error Fixer",
            "Use a separate code workspace to paste code, error messages, and expected behavior, then generate a fix with a short explanation.",
            self.codefix_folder_var,
            "Back To Developer",
            "developer",
        )
        self._build_codefix_tool(tool_host)

    def _build_quiz_tool(self, parent: tk.Frame) -> None:
        scroll_shell, scroll_body, _quiz_canvas = self._create_scrollable_section(parent)
        scroll_shell.pack(fill="both", expand=True)

        grid = tk.Frame(scroll_body, bg=PALETTE["bg"])
        grid.pack(fill="both", expand=True, padx=(0, 4))
        grid.grid_columnconfigure(0, weight=5)
        grid.grid_columnconfigure(1, weight=5)
        grid.grid_rowconfigure(0, weight=1)

        intake_card, intake_body = self._create_card(
            grid,
            "Quiz Intake",
            "Upload a Word, PDF, text, or markdown file, or paste the content directly. The app reads it first, then you can choose how to continue.",
            PALETTE["accent"],
        )
        intake_card.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        file_row = tk.Frame(intake_body, bg=PALETTE["panel"])
        file_row.pack(fill="x", pady=(0, 14))
        self._make_button(file_row, "Upload File", self.load_quiz_file, variant="primary").pack(side="left", padx=(0, 10))
        self._make_button(file_row, "Clear Quiz Input", self.clear_quiz_workspace, variant="ghost").pack(side="left")
        tk.Label(
            intake_body,
            textvariable=self.quiz_file_var,
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.small_font,
            justify="left",
            wraplength=580,
        ).pack(anchor="w", pady=(0, 10))
        self._metric_tile(intake_body, "Content stats", self.quiz_stats_var, PALETTE["accent_soft"]).pack(fill="x", pady=(0, 14))
        self.quiz_source_text = self._labeled_text(
            intake_body,
            "Quiz content",
            20,
            "If you do not upload a file, you can paste the full content here manually.",
        )
        self.quiz_source_text.bind("<KeyRelease>", lambda _event: self.refresh_quiz_stats())
        analyze_button = self._make_button(intake_body, "Analyze Quiz", self.start_quiz_analysis, variant="secondary")
        self._register_ai_button(analyze_button).pack(anchor="w")

        right = tk.Frame(grid, bg=PALETTE["bg"])
        right.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        right.grid_rowconfigure(0, weight=4)
        right.grid_rowconfigure(1, weight=5)
        right.grid_columnconfigure(0, weight=1)

        summary_card, summary_body = self._create_card(
            right,
            "Summary And Task Reading",
            "The assistant first summarizes the uploaded content and explains what task appears to be requested.",
            PALETTE["info"],
        )
        summary_card.grid(row=0, column=0, sticky="nsew", pady=(0, 10))
        tk.Label(
            summary_body,
            textvariable=self.quiz_progress_var,
            bg=PALETTE["panel_alt"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
            padx=14,
            pady=8,
        ).pack(anchor="w", pady=(0, 12))
        self.quiz_summary_text = self._labeled_text(
            summary_body,
            "Summary preview",
            11,
            "This first pass should help you understand the file before asking for a full response.",
            read_only=True,
        )

        response_card, response_body = self._create_card(
            right,
            "Response Builder",
            "After the summary, choose either a full response or a custom prompt tied to the uploaded content.",
            PALETTE["success"],
        )
        response_card.grid(row=1, column=0, sticky="nsew", pady=(10, 0))
        mode_row = tk.Frame(response_body, bg=PALETTE["panel"])
        mode_row.pack(fill="x", pady=(0, 10))
        self._styled_radiobutton(mode_row, "Do it in full", self.quiz_mode_var, "complete").pack(side="left", padx=(0, 16))
        self._styled_radiobutton(mode_row, "Use specific prompt", self.quiz_mode_var, "specific").pack(side="left")
        self.quiz_prompt_text = self._labeled_text(
            response_body,
            "Specific prompt",
            5,
            "Use this when you want a custom instruction after the summary, for example: answer only part 2 or explain each item one by one.",
        )
        actions = tk.Frame(response_body, bg=PALETTE["panel"])
        actions.pack(fill="x", pady=(0, 12))
        self.quiz_generate_button = self._register_ai_button(
            self._make_button(actions, "Generate Quiz Guidance", self.start_quiz_response, variant="primary")
        )
        self.quiz_generate_button.pack(side="left", padx=(0, 10))
        self.quiz_export_button = self._make_button(actions, "Export Result", self.export_quiz_result, variant="secondary")
        self.quiz_export_button.pack(side="left", padx=(0, 10))
        self.quiz_export_button.config(state="disabled")
        self._make_button(actions, "Clear Result", self.clear_quiz_result, variant="ghost").pack(side="left")
        self.quiz_result_text = self._labeled_text(
            response_body,
            "Generated response",
            11,
            "Full response or prompt-based response appears here.",
            read_only=True,
        )

        self.quiz_layout = {
            "grid": grid,
            "intake": intake_card,
            "right": right,
        }

    def _build_essay_tool(self, parent: tk.Frame) -> None:
        scroll_shell, scroll_body, _essay_canvas = self._create_scrollable_section(parent)
        scroll_shell.pack(fill="both", expand=True)

        grid = tk.Frame(scroll_body, bg=PALETTE["bg"])
        grid.pack(fill="both", expand=True, padx=(0, 4))
        grid.grid_columnconfigure(0, weight=4)
        grid.grid_columnconfigure(1, weight=5)
        grid.grid_rowconfigure(0, weight=1)

        build_card, build_body = self._create_card(
            grid,
            "Essay Builder",
            "Set the title, add an optional custom instruction, choose the target length, and pick English, Tagalog, or both for Taglish.",
            PALETTE["warning"],
        )
        build_card.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        self._labeled_entry(
            build_body,
            "Essay title",
            self.essay_title_var,
            "This becomes the essay topic and export title by default.",
        )
        self.essay_prompt_text = self._labeled_text(
            build_body,
            "Specific prompt (optional)",
            6,
            "Examples: persuasive angle, reflective tone, or points that must be included.",
        )
        word_row = tk.Frame(build_body, bg=PALETTE["panel"])
        word_row.pack(fill="x", pady=(0, 16))
        tk.Label(word_row, text="Target word count", bg=PALETTE["panel"], fg=PALETTE["text"], font=self.label_font).pack(anchor="w")
        spin_shell = tk.Frame(word_row, bg=PALETTE["border"])
        spin_shell.pack(anchor="w", pady=(8, 0))
        self.essay_word_spinbox = tk.Spinbox(
            spin_shell,
            from_=100,
            to=3000,
            increment=50,
            textvariable=self.essay_word_count_var,
            bg=PALETTE["input"],
            fg=PALETTE["text"],
            insertbackground=PALETTE["text"],
            buttonbackground=PALETTE["panel_alt"],
            relief="flat",
            bd=0,
            highlightthickness=0,
            width=12,
            font=self.body_font,
        )
        self.essay_word_spinbox.pack(padx=1, pady=1, ipady=8)
        tk.Label(word_row, text="Choose a target and the model will aim around that length.", bg=PALETTE["panel"], fg=PALETTE["muted"], font=self.small_font).pack(anchor="w", pady=(6, 0))

        language_box = tk.Frame(build_body, bg=PALETTE["panel_alt"], highlightthickness=1, highlightbackground=PALETTE["border"])
        language_box.pack(fill="x", pady=(0, 16))
        tk.Label(language_box, text="Language selection", bg=PALETTE["panel_alt"], fg=PALETTE["text"], font=self.label_font).pack(anchor="w", padx=14, pady=(12, 8))
        checks = tk.Frame(language_box, bg=PALETTE["panel_alt"])
        checks.pack(anchor="w", padx=12, pady=(0, 12))
        self._styled_checkbutton(checks, "Tagalog", self.essay_tagalog_var).pack(side="left", padx=(0, 16))
        self._styled_checkbutton(checks, "English", self.essay_english_var).pack(side="left")

        format_box = tk.Frame(
            build_body,
            bg=PALETTE["panel_alt"],
            highlightthickness=1,
            highlightbackground=PALETTE["border"],
        )
        format_box.pack(fill="x", pady=(0, 16))
        tk.Label(
            format_box,
            text="Essay export format",
            bg=PALETTE["panel_alt"],
            fg=PALETTE["text"],
            font=self.label_font,
        ).pack(anchor="w", padx=14, pady=(12, 8))
        tk.Label(
            format_box,
            text="Adjust what gets attached to the final essay output. The specific export name is optional and can override the saved dashboard name for this essay only.",
            bg=PALETTE["panel_alt"],
            fg=PALETTE["muted"],
            wraplength=500,
            justify="left",
            font=self.small_font,
        ).pack(anchor="w", padx=14)

        specific_name_wrap = tk.Frame(format_box, bg=PALETTE["panel_alt"])
        specific_name_wrap.pack(fill="x", padx=14, pady=(12, 10))
        tk.Label(
            specific_name_wrap,
            text="Specific export name (optional)",
            bg=PALETTE["panel_alt"],
            fg=PALETTE["text"],
            font=self.label_font,
        ).pack(anchor="w")
        specific_name_shell = tk.Frame(specific_name_wrap, bg=PALETTE["border"])
        specific_name_shell.pack(fill="x", pady=(8, 0))
        tk.Entry(
            specific_name_shell,
            textvariable=self.essay_specific_name_var,
            bg=PALETTE["input"],
            fg=PALETTE["text"],
            insertbackground=PALETTE["text"],
            relief="flat",
            bd=0,
            highlightthickness=0,
            font=self.body_font,
            selectbackground=PALETTE["accent_dark"],
            selectforeground=PALETTE["text"],
        ).pack(fill="x", padx=1, pady=1, ipady=10)
        tk.Label(
            specific_name_wrap,
            text="Leave this blank if you want the essay export to use the saved dashboard name instead.",
            bg=PALETTE["panel_alt"],
            fg=PALETTE["muted"],
            font=self.small_font,
        ).pack(anchor="w", pady=(6, 0))

        options_wrap = tk.Frame(format_box, bg=PALETTE["panel_alt"])
        options_wrap.pack(fill="x", padx=14, pady=(0, 12))
        for text, variable in [
            ("Include saved or specific name in export", self.output_include_name_var),
            ("Include date in export", self.output_include_date_var),
            ("Include essay heading suggestion", self.essay_include_heading_var),
            ("Include self-check tip", self.essay_include_tip_var),
        ]:
            tk.Checkbutton(
                options_wrap,
                text=text,
                variable=variable,
                bg=PALETTE["panel_alt"],
                fg=PALETTE["text"],
                activebackground=PALETTE["panel_alt"],
                activeforeground=PALETTE["text"],
                selectcolor=PALETTE["glass_inner"],
                highlightthickness=0,
                bd=0,
                font=self.body_font,
            ).pack(anchor="w", pady=(0, 8))

        action_row = tk.Frame(build_body, bg=PALETTE["panel"])
        action_row.pack(fill="x")
        self.essay_generate_button = self._register_ai_button(
            self._make_button(action_row, "Generate Essay", self.start_essay_generation, variant="primary")
        )
        self.essay_generate_button.pack(side="left", padx=(0, 10))
        self._make_button(action_row, "Clear Essay Form", self.clear_essay_form, variant="ghost").pack(side="left")

        result_card, result_body = self._create_card(
            grid,
            "Essay Preview",
            "The draft appears here once generated. Export it to Word when you are happy with the result.",
            PALETTE["info"],
        )
        result_card.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        tk.Label(
            result_body,
            textvariable=self.essay_progress_var,
            bg=PALETTE["panel_alt"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
            padx=14,
            pady=8,
        ).pack(anchor="w", pady=(0, 12))
        self.essay_result_text = self._labeled_text(
            result_body,
            "Essay output",
            24,
            "Generated essay draft appears here.",
            read_only=True,
        )
        essay_actions = tk.Frame(result_body, bg=PALETTE["panel"])
        essay_actions.pack(fill="x")
        self.essay_export_button = self._make_button(essay_actions, "Export Result", self.export_essay_result, variant="secondary")
        self.essay_export_button.pack(side="left", padx=(0, 10))
        self.essay_export_button.config(state="disabled")
        self._make_button(essay_actions, "Clear Result", self.clear_essay_result, variant="ghost").pack(side="left")

        self.essay_layout = {
            "grid": grid,
            "build": build_card,
            "result": result_card,
        }

    def _build_activity_tool(self, parent: tk.Frame) -> None:
        scroll_shell, scroll_body, _activity_canvas = self._create_scrollable_section(parent)
        scroll_shell.pack(fill="both", expand=True)

        build_card, build_body = self._create_card(
            scroll_body,
            "Activity Builder",
            "Set the topic, activity type, and optional level details, then generate a ready-to-use school activity.",
            PALETTE["success"],
        )
        build_card.pack(fill="x", pady=(0, 14))
        self._labeled_entry(
            build_body,
            "Activity topic or title",
            self.activity_title_var,
            "Example: Climate change awareness, reading comprehension, or crypto basics.",
        )
        self._labeled_entry(
            build_body,
            "Activity type",
            self.activity_type_var,
            "Examples: Worksheet, reflection, pair work, quiz drill, seatwork.",
        )
        self._labeled_entry(
            build_body,
            "Level or class (optional)",
            self.activity_level_var,
            "Example: Grade 10, senior high, college, or section name.",
        )
        self.activity_prompt_text = self._labeled_text(
            build_body,
            "Specific instructions",
            8,
            "Add goals, constraints, style, or topics that must appear in the generated activity.",
        )
        activity_actions = tk.Frame(build_body, bg=PALETTE["panel"])
        activity_actions.pack(fill="x")
        self.activity_generate_button = self._register_ai_button(
            self._make_button(activity_actions, "Generate Activity", self.start_activity_generation, variant="primary")
        )
        self.activity_generate_button.pack(side="left", padx=(0, 10))
        self._make_button(activity_actions, "Clear Activity Form", self.clear_activity_form, variant="ghost").pack(side="left")

        result_card, result_body = self._create_card(
            scroll_body,
            "Activity Preview",
            "The generated activity appears here in a clean structure you can export straight to Word.",
            PALETTE["info"],
        )
        result_card.pack(fill="both", expand=True)
        tk.Label(
            result_body,
            textvariable=self.activity_progress_var,
            bg=PALETTE["panel_alt"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
            padx=14,
            pady=8,
        ).pack(anchor="w", pady=(0, 12))
        self.activity_result_text = self._labeled_text(
            result_body,
            "Generated activity",
            18,
            "Your activity draft appears here.",
            read_only=True,
        )
        result_actions = tk.Frame(result_body, bg=PALETTE["panel"])
        result_actions.pack(fill="x")
        self.activity_export_button = self._make_button(result_actions, "Export Result", self.export_activity_result, variant="secondary")
        self.activity_export_button.pack(side="left", padx=(0, 10))
        self.activity_export_button.config(state="disabled")
        self._make_button(result_actions, "Clear Result", self.clear_activity_result, variant="ghost").pack(side="left")

    def _build_document_tool(self, parent: tk.Frame) -> None:
        scroll_shell, scroll_body, _document_canvas = self._create_scrollable_section(parent)
        scroll_shell.pack(fill="both", expand=True)

        build_card, build_body = self._create_card(
            scroll_body,
            "Document Builder",
            "Generate a structured school document such as a handout, report, reviewer, or formal academic write-up.",
            PALETTE["accent"],
        )
        build_card.pack(fill="x", pady=(0, 14))
        self._labeled_entry(
            build_body,
            "Document title",
            self.document_title_var,
            "This becomes the export title by default.",
        )
        self._labeled_entry(
            build_body,
            "Document type",
            self.document_type_var,
            "Examples: Study handout, report, reviewer, request letter, lesson guide.",
        )
        self._labeled_entry(
            build_body,
            "Audience or purpose (optional)",
            self.document_audience_var,
            "Example: teacher, classmates, groupmates, beginners, or school office.",
        )
        self.document_prompt_text = self._labeled_text(
            build_body,
            "Specific content request",
            8,
            "Add the main ideas, sections, tone, or requirements that the document should include.",
        )
        document_actions = tk.Frame(build_body, bg=PALETTE["panel"])
        document_actions.pack(fill="x")
        self.document_generate_button = self._register_ai_button(
            self._make_button(document_actions, "Generate Document", self.start_document_generation, variant="primary")
        )
        self.document_generate_button.pack(side="left", padx=(0, 10))
        self._make_button(document_actions, "Clear Document Form", self.clear_document_form, variant="ghost").pack(side="left")

        result_card, result_body = self._create_card(
            scroll_body,
            "Document Preview",
            "The generated document draft appears here and can be exported directly to the document subfolder.",
            PALETTE["info"],
        )
        result_card.pack(fill="both", expand=True)
        tk.Label(
            result_body,
            textvariable=self.document_progress_var,
            bg=PALETTE["panel_alt"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
            padx=14,
            pady=8,
        ).pack(anchor="w", pady=(0, 12))
        self.document_result_text = self._labeled_text(
            result_body,
            "Generated document",
            18,
            "Your structured document appears here.",
            read_only=True,
        )
        result_actions = tk.Frame(result_body, bg=PALETTE["panel"])
        result_actions.pack(fill="x")
        self.document_export_button = self._make_button(result_actions, "Export Result", self.export_document_result, variant="secondary")
        self.document_export_button.pack(side="left", padx=(0, 10))
        self.document_export_button.config(state="disabled")
        self._make_button(result_actions, "Clear Result", self.clear_document_result, variant="ghost").pack(side="left")

    def _build_assignment_tool(self, parent: tk.Frame) -> None:
        scroll_shell, scroll_body, _assignment_canvas = self._create_scrollable_section(parent)
        scroll_shell.pack(fill="both", expand=True)

        intake_card, intake_body = self._create_card(
            scroll_body,
            "Assignment Intake",
            "Upload a document or paste the assignment manually, then let the app analyze the task before generating a guided response.",
            PALETTE["warning"],
        )
        intake_card.pack(fill="x", pady=(0, 14))
        file_row = tk.Frame(intake_body, bg=PALETTE["panel"])
        file_row.pack(fill="x", pady=(0, 14))
        self._make_button(file_row, "Upload Assignment File", self.load_assignment_file, variant="primary").pack(
            side="left", padx=(0, 10)
        )
        self._make_button(file_row, "Clear Assignment Input", self.clear_assignment_workspace, variant="ghost").pack(
            side="left"
        )
        tk.Label(
            intake_body,
            textvariable=self.assignment_file_var,
            bg=PALETTE["panel"],
            fg=PALETTE["muted"],
            font=self.small_font,
            justify="left",
            wraplength=760,
        ).pack(anchor="w", pady=(0, 10))
        self._metric_tile(intake_body, "Assignment stats", self.assignment_stats_var, PALETTE["accent_soft"]).pack(
            fill="x", pady=(0, 14)
        )
        self.assignment_source_text = self._labeled_text(
            intake_body,
            "Assignment content",
            15,
            "Paste the assignment instructions here if you are not uploading a file.",
        )
        self.assignment_source_text.bind("<KeyRelease>", lambda _event: self.refresh_assignment_stats())
        mode_row = tk.Frame(intake_body, bg=PALETTE["panel"])
        mode_row.pack(fill="x", pady=(0, 12))
        self._styled_radiobutton(mode_row, "Guided response", self.assignment_mode_var, "guided").pack(side="left", padx=(0, 16))
        self._styled_radiobutton(mode_row, "Full draft", self.assignment_mode_var, "complete").pack(side="left")
        self.assignment_prompt_text = self._labeled_text(
            intake_body,
            "Specific prompt (optional)",
            5,
            "Examples: answer only parts 1 to 3, make it formal, or explain step by step.",
        )
        action_row = tk.Frame(intake_body, bg=PALETTE["panel"])
        action_row.pack(fill="x")
        self.assignment_analyze_button = self._register_ai_button(
            self._make_button(action_row, "Analyze Assignment", self.start_assignment_analysis, variant="secondary")
        )
        self.assignment_analyze_button.pack(side="left", padx=(0, 10))
        self.assignment_generate_button = self._register_ai_button(
            self._make_button(action_row, "Generate Assignment Guidance", self.start_assignment_response, variant="primary")
        )
        self.assignment_generate_button.pack(side="left")

        result_card, result_body = self._create_card(
            scroll_body,
            "Assignment Response",
            "The assignment summary appears first, followed by the generated guided response or sample draft.",
            PALETTE["info"],
        )
        result_card.pack(fill="both", expand=True)
        tk.Label(
            result_body,
            textvariable=self.assignment_progress_var,
            bg=PALETTE["panel_alt"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
            padx=14,
            pady=8,
        ).pack(anchor="w", pady=(0, 12))
        self.assignment_summary_text = self._labeled_text(
            result_body,
            "Assignment summary",
            8,
            "The app reads the assignment first before generating the response.",
            read_only=True,
        )
        self.assignment_result_text = self._labeled_text(
            result_body,
            "Generated response",
            12,
            "The generated assignment support appears here.",
            read_only=True,
        )
        result_actions = tk.Frame(result_body, bg=PALETTE["panel"])
        result_actions.pack(fill="x")
        self.assignment_export_button = self._make_button(
            result_actions, "Export Result", self.export_assignment_result, variant="secondary"
        )
        self.assignment_export_button.pack(side="left", padx=(0, 10))
        self.assignment_export_button.config(state="disabled")
        self._make_button(result_actions, "Clear Result", self.clear_assignment_result, variant="ghost").pack(side="left")

    def _build_codefix_tool(self, parent: tk.Frame) -> None:
        scroll_shell, scroll_body, _codefix_canvas = self._create_scrollable_section(parent)
        scroll_shell.pack(fill="both", expand=True)

        build_card, build_body = self._create_card(
            scroll_body,
            "Code Fix Builder",
            "Paste the broken code, add the error or symptom, and explain the expected behavior so the app can generate a cleaner fix.",
            PALETTE["warning"],
        )
        build_card.pack(fill="x", pady=(0, 14))
        self._labeled_entry(
            build_body,
            "Issue title (optional)",
            self.codefix_title_var,
            "Example: Login bug, syntax error, API request failure, or Tkinter freeze.",
        )
        self._labeled_entry(
            build_body,
            "Language or stack",
            self.codefix_language_var,
            "Examples: Python, JavaScript, Tkinter, React, Node, Flask.",
        )
        self.codefix_error_text = self._labeled_text(
            build_body,
            "Error message or symptoms",
            6,
            "Paste the exact error message or explain what is going wrong.",
        )
        self.codefix_source_text = self._labeled_text(
            build_body,
            "Code snippet",
            12,
            "Paste only the relevant code block if possible so the response stays focused.",
        )
        self.codefix_expectation_text = self._labeled_text(
            build_body,
            "Expected behavior (optional)",
            4,
            "Describe what the code should do once fixed.",
        )
        action_row = tk.Frame(build_body, bg=PALETTE["panel"])
        action_row.pack(fill="x")
        self.codefix_generate_button = self._register_ai_button(
            self._make_button(action_row, "Generate Code Fix", self.start_codefix_generation, variant="primary")
        )
        self.codefix_generate_button.pack(side="left", padx=(0, 10))
        self._make_button(action_row, "Clear Code Fixer", self.clear_codefix_form, variant="ghost").pack(side="left")

        result_card, result_body = self._create_card(
            scroll_body,
            "Code Fix Preview",
            "The cleaned-up fix and explanation appear here. Export it if you want to save the debugging notes to Word.",
            PALETTE["info"],
        )
        result_card.pack(fill="both", expand=True)
        tk.Label(
            result_body,
            textvariable=self.codefix_progress_var,
            bg=PALETTE["panel_alt"],
            fg=PALETTE["accent_soft"],
            font=self.label_font,
            padx=14,
            pady=8,
        ).pack(anchor="w", pady=(0, 12))
        self.codefix_result_text = self._labeled_text(
            result_body,
            "Fixed result",
            18,
            "The code fix output appears here.",
            read_only=True,
        )
        result_actions = tk.Frame(result_body, bg=PALETTE["panel"])
        result_actions.pack(fill="x")
        self.codefix_export_button = self._make_button(result_actions, "Export Result", self.export_codefix_result, variant="secondary")
        self.codefix_export_button.pack(side="left", padx=(0, 10))
        self.codefix_export_button.config(state="disabled")
        self._make_button(result_actions, "Clear Result", self.clear_codefix_result, variant="ghost").pack(side="left")

    def _refresh_identity_labels(self) -> None:
        display_name = self.saved_name or "Guest"
        self.header_title_var.set(f"{display_name}'s Premium Study Desk")
        self.header_subtitle_var.set(
            "A cleaner dashboard for identity settings, grouped academic workspaces, and a separate developer hub for code-fixing support."
        )
        if self.saved_name:
            self.profile_summary_var.set(self.saved_name)
        else:
            self.profile_summary_var.set("No saved name yet")

        if self.saved_include_date:
            self.profile_date_var.set(f"Date today enabled: {today_string()}")
        else:
            self.profile_date_var.set("Date today disabled")

        if self.saved_name and self.saved_include_date:
            self.profile_note_var.set("Name and today's date are ready to be attached to exported work.")
        elif self.saved_name:
            self.profile_note_var.set("Only the saved name will be remembered for this workspace.")
        else:
            self.profile_note_var.set("Save a profile if you want a personalized dashboard header.")

    def _persist_profile_settings(self) -> None:
        save_profile_settings(
            self.saved_name,
            self.saved_include_date,
            str(self.base_export_dir),
            self.output_include_name_var.get(),
            self.output_include_date_var.get(),
            self.essay_include_heading_var.get(),
            self.essay_include_tip_var.get(),
        )

    def _current_output_settings(self) -> dict[str, bool]:
        return {
            "include_name": self.output_include_name_var.get(),
            "include_date": self.output_include_date_var.get(),
            "essay_include_heading": self.essay_include_heading_var.get(),
            "essay_include_tip": self.essay_include_tip_var.get(),
        }

    def _refresh_export_folder_labels(self) -> None:
        self.export_var.set(str(self.base_export_dir))
        self.main_folder_var.set(str(self.base_export_dir))
        self.quiz_folder_var.set(str(self.quiz_export_dir))
        self.essay_folder_var.set(str(self.essay_export_dir))
        self.activity_folder_var.set(str(self.activity_export_dir))
        self.document_folder_var.set(str(self.document_export_dir))
        self.assignment_folder_var.set(str(self.assignment_export_dir))
        self.codefix_folder_var.set(str(self.codefix_export_dir))
        self.academics_folders_var.set(
            "Quiz Solver\n"
            "Assignment Solver\n"
            "Essay Generator\n"
            "Activity Generator\n"
            "Document Generator"
        )
        self.developer_folder_var.set(f"Code Error Fixer\n{self.codefix_export_dir}")
        self.folder_note_var.set("Each tool now autosaves into its own dedicated subfolder inside the chosen main folder.")

    def _ensure_export_directories(self) -> None:
        self.base_export_dir.mkdir(parents=True, exist_ok=True)
        self.export_dirs = {
            "quiz": self.base_export_dir / "quiz_solver",
            "essay": self.base_export_dir / "essay_generator",
            "activity": self.base_export_dir / "activity_generator",
            "document": self.base_export_dir / "document_generator",
            "assignment": self.base_export_dir / "assignment_solver",
            "codefix": self.base_export_dir / "code_error_fixer",
        }
        self.quiz_export_dir = self.export_dirs["quiz"]
        self.essay_export_dir = self.export_dirs["essay"]
        self.activity_export_dir = self.export_dirs["activity"]
        self.document_export_dir = self.export_dirs["document"]
        self.assignment_export_dir = self.export_dirs["assignment"]
        self.codefix_export_dir = self.export_dirs["codefix"]
        for folder in self.export_dirs.values():
            folder.mkdir(parents=True, exist_ok=True)
        self._refresh_export_folder_labels()

    def choose_main_folder(self) -> None:
        selected = filedialog.askdirectory(title="Choose Main Save Folder", initialdir=str(self.base_export_dir))
        if not selected:
            return
        self.base_export_dir = Path(selected)
        self._ensure_export_directories()
        self._persist_profile_settings()
        self.status_var.set("Main save folder updated.")
        messagebox.showinfo(
            "Folder Ready",
            "Main folder updated and the dedicated tool subfolders are ready.\n\n"
            f"Main: {self.base_export_dir}\n"
            f"Quiz Solver: {self.quiz_export_dir}\n"
            f"Assignment Solver: {self.assignment_export_dir}\n"
            f"Essay Generator: {self.essay_export_dir}\n"
            f"Activity Generator: {self.activity_export_dir}\n"
            f"Document Generator: {self.document_export_dir}\n"
            f"Code Error Fixer: {self.codefix_export_dir}",
        )

    def open_main_folder(self) -> None:
        self._ensure_export_directories()
        try:
            os.startfile(self.base_export_dir)
        except AttributeError:
            self._notify("Open Folder", f"Main folder path:\n{self.base_export_dir}")
        except OSError as exc:
            self._notify("Open Folder Error", str(exc), level="error")

    def save_profile(self) -> None:
        name = self.name_input_var.get().strip()
        include_date = self.include_date_var.get()
        self.saved_name = name
        self.saved_include_date = include_date
        self._ensure_export_directories()
        self._persist_profile_settings()
        self._refresh_identity_labels()
        self.status_var.set("Dashboard profile saved.")
        messagebox.showinfo("Profile Saved", "Your dashboard profile, output options, and save-folder setup have been saved.")

    def clear_profile(self) -> None:
        self.name_input_var.set("")
        self.include_date_var.set(False)
        self.output_include_name_var.set(False)
        self.output_include_date_var.set(False)
        self.essay_include_heading_var.set(True)
        self.essay_include_tip_var.set(True)
        self.base_export_dir = EXPORT_DIR
        self.saved_name = ""
        self.saved_include_date = False
        self._ensure_export_directories()
        self._persist_profile_settings()
        self._refresh_identity_labels()
        self.status_var.set("Saved profile cleared and folder reset to default exports.")

    def show_page(self, key: str) -> None:
        self.pages[key].lift()
        academic_pages = {"academics", "quiz", "essay", "activity", "document", "assignment"}
        developer_pages = {"developer", "codefix"}
        if key in academic_pages:
            active_nav = "academics"
        elif key in developer_pages:
            active_nav = "developer"
        else:
            active_nav = key
        for name, button in self.nav_buttons.items():
            if name == active_nav:
                self._apply_button_palette(
                    button,
                    PALETTE["panel_alt"],
                    PALETTE["accent_soft"],
                    PALETTE["glass_hover"],
                    PALETTE["accent_soft"],
                    PALETTE["accent_dark"],
                    PALETTE["accent"],
                )
            else:
                self._apply_button_palette(
                    button,
                    PALETTE["sidebar"],
                    PALETTE["text"],
                    PALETTE["panel_alt"],
                    PALETTE["accent_soft"],
                    PALETTE["sidebar"],
                    PALETTE["accent_dark"],
                )
        self.status_var.set(f"Viewing {key.title()}.")
        self._schedule_responsive_refresh()

    def open_academics(self, tool: str) -> None:
        self.show_academic_tool(tool)

    def show_academic_tool(self, tool: str) -> None:
        self.academic_tool_var.set(tool)
        if tool in {"quiz", "essay", "activity", "document", "assignment"}:
            self.show_page(tool)
        else:
            self.show_page("academics")
        self.status_var.set(f"Academics tool ready: {tool.replace('_', ' ').title()}.")
        self._schedule_responsive_refresh()

    def refresh_ai_status(self) -> None:
        load_dotenv(ENV_PATH)
        model = os.getenv("HF_MODEL", "").strip()
        token = os.getenv("HF_TOKEN", "").strip()

        if InferenceClient is None:
            self.ai_ready = False
            self.ai_chip_var.set("AI offline: missing dependency")
            self.model_var.set("Install requirements.txt to enable AI")
            self.ai_chip.configure(bg=PALETTE["danger"], fg=PALETTE["bg"])
        elif token and model:
            self.ai_ready = True
            self.ai_chip_var.set("AI ready for academics")
            self.model_var.set(model)
            self.ai_chip.configure(bg=PALETTE["success"], fg=PALETTE["bg"])
        else:
            self.ai_ready = False
            self.ai_chip_var.set("AI waiting for .env setup")
            self.model_var.set("Add HF_TOKEN and HF_MODEL to .env")
            self.ai_chip.configure(bg=PALETTE["warning"], fg=PALETTE["bg"])

        state = "normal" if self.ai_ready and not self.busy else "disabled"
        for button in self.ai_action_buttons:
            button.configure(state=state)
        self.status_var.set("AI settings refreshed.")

    def _set_ai_buttons_state(self, enabled: bool) -> None:
        state = "normal" if enabled and self.ai_ready else "disabled"
        for button in self.ai_action_buttons:
            button.configure(state=state)

    def _get_text(self, widget: ScrolledText) -> str:
        return widget.get("1.0", "end-1c")

    def _set_text(self, widget: ScrolledText, value: str, read_only: bool = False) -> None:
        widget.config(state="normal")
        widget.delete("1.0", "end")
        if value:
            widget.insert("1.0", value)
        if read_only:
            widget.config(state="disabled")

    def _notify(self, title: str, message: str, level: str = "info") -> None:
        self.status_var.set(message)
        if level == "error":
            messagebox.showerror(title, message)
        elif level == "warning":
            messagebox.showwarning(title, message)
        else:
            messagebox.showinfo(title, message)

    def _identity_block(self) -> str:
        lines: list[str] = []
        if self.saved_name:
            lines.append(f"Name: {self.saved_name}")
        if self.saved_include_date:
            lines.append(f"Date: {today_string()}")
        return "\n".join(lines)

    def _export_metadata_lines(self, category: str = "generic", name_override: str | None = None) -> list[str]:
        lines: list[str] = []
        selected_name = self.saved_name
        if category == "essay" and name_override and name_override.strip():
            selected_name = name_override.strip()
        if self.output_include_name_var.get() and selected_name:
            lines.append(selected_name)
        if self.output_include_date_var.get():
            lines.append(today_string())
        return lines

    def _auto_output_path(self, folder: Path, default_name: str) -> Path:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = sanitize_filename(default_name)
        candidate = folder / f"{stem}_{stamp}.docx"
        counter = 1
        while candidate.exists():
            candidate = folder / f"{stem}_{stamp}_{counter}.docx"
            counter += 1
        return candidate

    def _export_text_output(
        self,
        title: str,
        body: str,
        default_name: str,
        category: str,
        name_override: str | None = None,
    ) -> Path | None:
        final_body = body.strip()

        self._ensure_export_directories()
        target_dir = self.export_dirs.get(category, self.base_export_dir)
        output_path = self._auto_output_path(target_dir, default_name)
        save_docx(
            title,
            final_body,
            output_path,
            category=category,
            metadata_lines=self._export_metadata_lines(category=category, name_override=name_override),
            output_options=self._current_output_settings(),
        )
        self.status_var.set(f"Saved {output_path.name} to {target_dir.name}.")
        messagebox.showinfo("Saved", f"Document saved automatically to:\n{output_path}")
        return output_path

    def load_quiz_file(self) -> None:
        selected = filedialog.askopenfilename(
            title="Choose Quiz File",
            filetypes=[
                ("Supported Files", "*.docx *.pdf *.txt *.md"),
                ("Word Document", "*.docx"),
                ("PDF File", "*.pdf"),
                ("Text File", "*.txt"),
                ("Markdown", "*.md"),
            ],
        )
        if not selected:
            return

        path = Path(selected)
        try:
            text = read_document_text(path)
        except Exception as exc:
            self._notify("File Read Error", str(exc), level="error")
            return

        self.quiz_loaded_path = path
        self.quiz_file_var.set(str(path))
        self._set_text(self.quiz_source_text, text)
        self.refresh_quiz_stats()
        self.clear_quiz_result()
        self._set_text(self.quiz_summary_text, "", read_only=True)
        self.quiz_summary_cache = ""
        self.quiz_progress_var.set("File loaded. Analyze it when ready.")
        self.status_var.set(f"Loaded {path.name} into the quiz workspace.")

    def refresh_quiz_stats(self) -> None:
        text = self._get_text(self.quiz_source_text)
        self.quiz_stats_var.set(f"{count_words(text)} words   |   {len(text.strip())} characters")

    def load_assignment_file(self) -> None:
        selected = filedialog.askopenfilename(
            title="Choose Assignment File",
            filetypes=[
                ("Supported Files", "*.docx *.pdf *.txt *.md"),
                ("Word Document", "*.docx"),
                ("PDF File", "*.pdf"),
                ("Text File", "*.txt"),
                ("Markdown", "*.md"),
            ],
        )
        if not selected:
            return

        path = Path(selected)
        try:
            text = read_document_text(path)
        except Exception as exc:
            self._notify("File Read Error", str(exc), level="error")
            return

        self.assignment_loaded_path = path
        self.assignment_file_var.set(str(path))
        self._set_text(self.assignment_source_text, text)
        self.refresh_assignment_stats()
        self.clear_assignment_result()
        self._set_text(self.assignment_summary_text, "", read_only=True)
        self.assignment_summary_cache = ""
        self.assignment_progress_var.set("Assignment file loaded. Analyze it when ready.")
        self.status_var.set(f"Loaded {path.name} into the assignment workspace.")

    def refresh_assignment_stats(self) -> None:
        text = self._get_text(self.assignment_source_text)
        self.assignment_stats_var.set(f"{count_words(text)} words   |   {len(text.strip())} characters")

    def clear_quiz_workspace(self) -> None:
        self.quiz_loaded_path = None
        self.quiz_file_var.set("No file loaded yet")
        self._set_text(self.quiz_source_text, "")
        self._set_text(self.quiz_summary_text, "", read_only=True)
        self._set_text(self.quiz_prompt_text, "")
        self.clear_quiz_result()
        self.quiz_summary_cache = ""
        self.quiz_progress_var.set("Upload a file or paste text, then analyze it.")
        self.refresh_quiz_stats()
        self.status_var.set("Quiz workspace cleared.")

    def clear_quiz_result(self) -> None:
        self.quiz_response_cache = ""
        self.quiz_export_title = "Quiz Support"
        self.quiz_export_name = "quiz_support"
        self._set_text(self.quiz_result_text, "", read_only=True)
        self.quiz_export_button.configure(state="disabled")
        if not self.quiz_summary_cache:
            self.quiz_progress_var.set("Upload a file or paste text, then analyze it.")

    def clear_essay_form(self) -> None:
        self.essay_title_var.set("")
        self.essay_specific_name_var.set("")
        self.essay_word_count_var.set("500")
        self.essay_tagalog_var.set(False)
        self.essay_english_var.set(True)
        self._set_text(self.essay_prompt_text, "")
        self.status_var.set("Essay form cleared.")

    def clear_essay_result(self) -> None:
        self.essay_response_cache = ""
        self.essay_export_title = "Essay Draft"
        self.essay_export_name = "essay_draft"
        self._set_text(self.essay_result_text, "", read_only=True)
        self.essay_export_button.configure(state="disabled")
        self.essay_progress_var.set("Set the essay details, then generate.")

    def clear_activity_form(self) -> None:
        self.activity_title_var.set("")
        self.activity_type_var.set("Worksheet")
        self.activity_level_var.set("")
        self._set_text(self.activity_prompt_text, "")
        self.status_var.set("Activity form cleared.")

    def clear_activity_result(self) -> None:
        self.activity_response_cache = ""
        self.activity_export_title = "Activity Draft"
        self.activity_export_name = "activity_draft"
        self._set_text(self.activity_result_text, "", read_only=True)
        self.activity_export_button.configure(state="disabled")
        self.activity_progress_var.set("Set the activity details, then generate.")

    def clear_document_form(self) -> None:
        self.document_title_var.set("")
        self.document_type_var.set("Study Handout")
        self.document_audience_var.set("")
        self._set_text(self.document_prompt_text, "")
        self.status_var.set("Document form cleared.")

    def clear_document_result(self) -> None:
        self.document_response_cache = ""
        self.document_export_title = "Document Draft"
        self.document_export_name = "document_draft"
        self._set_text(self.document_result_text, "", read_only=True)
        self.document_export_button.configure(state="disabled")
        self.document_progress_var.set("Set the document request, then generate.")

    def clear_assignment_workspace(self) -> None:
        self.assignment_loaded_path = None
        self.assignment_file_var.set("No assignment file loaded yet")
        self._set_text(self.assignment_source_text, "")
        self._set_text(self.assignment_summary_text, "", read_only=True)
        self._set_text(self.assignment_prompt_text, "")
        self.clear_assignment_result()
        self.assignment_summary_cache = ""
        self.assignment_progress_var.set("Upload or paste the assignment, then analyze it.")
        self.refresh_assignment_stats()
        self.status_var.set("Assignment workspace cleared.")

    def clear_assignment_result(self) -> None:
        self.assignment_response_cache = ""
        self.assignment_export_title = "Assignment Support"
        self.assignment_export_name = "assignment_support"
        self._set_text(self.assignment_result_text, "", read_only=True)
        self.assignment_export_button.configure(state="disabled")
        if not self.assignment_summary_cache:
            self.assignment_progress_var.set("Upload or paste the assignment, then analyze it.")

    def clear_codefix_form(self) -> None:
        self.codefix_title_var.set("")
        self.codefix_language_var.set("Python")
        self._set_text(self.codefix_error_text, "")
        self._set_text(self.codefix_source_text, "")
        self._set_text(self.codefix_expectation_text, "")
        self.status_var.set("Code fixer form cleared.")

    def clear_codefix_result(self) -> None:
        self.codefix_response_cache = ""
        self.codefix_export_title = "Code Fix"
        self.codefix_export_name = "code_fix"
        self._set_text(self.codefix_result_text, "", read_only=True)
        self.codefix_export_button.configure(state="disabled")
        self.codefix_progress_var.set("Paste the code and error details, then generate a fix.")

    def _start_ai_job(self, label: str, prompt: str, callback: object) -> None:
        if self.busy:
            self._notify("Please Wait", "Another AI request is still running.", level="warning")
            return
        if not self.ai_ready:
            self._notify("AI Not Ready", "Configure .env and install dependencies before using AI features.", level="warning")
            return

        self.busy = True
        self._set_ai_buttons_state(False)
        self.status_var.set(f"Generating {label}...")
        worker = threading.Thread(target=self._run_ai_job, args=(label, prompt, callback), daemon=True)
        worker.start()

    def _run_ai_job(self, label: str, prompt: str, callback: object) -> None:
        try:
            result = generate_text(prompt, label)
        except Exception as exc:
            self.after(0, lambda: self._finish_ai_job_error(str(exc)))
            return
        self.after(0, lambda: self._finish_ai_job_success(result, callback))

    def _finish_ai_job_success(self, result: str, callback: object) -> None:
        self.busy = False
        self._set_ai_buttons_state(True)
        callback(result)

    def _finish_ai_job_error(self, error: str) -> None:
        self.busy = False
        self._set_ai_buttons_state(True)
        self._notify("Generation Error", error, level="error")
    def start_quiz_analysis(self) -> None:
        source = self._get_text(self.quiz_source_text).strip()
        if not source:
            self._notify("Missing Content", "Upload or paste quiz content before analyzing it.", level="warning")
            return

        self.quiz_progress_var.set("Reading the file and summarizing the task...")
        prompt = (
            "Analyze the following academic material for educational review.\n"
            f"Saved dashboard profile:\n{self._identity_block() or 'No saved profile'}\n\n"
            f"Source content:\n{source}\n\n"
            "Return these sections clearly:\n"
            "1. Summary\n"
            "2. What task or instructions seem to be required\n"
            "3. Important topics, clues, or constraints\n"
            "4. Best next step for the student"
        )
        self._start_ai_job("quiz analysis", prompt, self.finish_quiz_analysis)

    def finish_quiz_analysis(self, result: str) -> None:
        self.quiz_summary_cache = result
        self._set_text(self.quiz_summary_text, result, read_only=True)
        self.quiz_progress_var.set("Summary ready. Choose full response or specific prompt next.")
        self.status_var.set("Quiz analysis complete.")

    def start_quiz_response(self) -> None:
        source = self._get_text(self.quiz_source_text).strip()
        if not source:
            self._notify("Missing Content", "Upload or paste quiz content before generating a response.", level="warning")
            return
        if not self.quiz_summary_cache.strip():
            self._notify("Analyze First", "Run the summary step first so the app can read the task before generating a response.", level="warning")
            return

        mode = self.quiz_mode_var.get()
        custom_prompt = self._get_text(self.quiz_prompt_text).strip()
        if mode == "specific" and not custom_prompt:
            self._notify("Missing Prompt", "Add a specific prompt or switch to full mode.", level="warning")
            return

        source_name = self.quiz_loaded_path.stem if self.quiz_loaded_path else "quiz_activity"
        self.quiz_export_title = f"Quiz Support: {source_name.replace('_', ' ').title()}"
        self.quiz_export_name = sanitize_filename(source_name)
        self.quiz_progress_var.set("Generating quiz response...")

        if mode == "complete":
            instructions = (
                "Create a complete educational response based on the content. If there are visible questions, answer them in order and include short explanations. If the file describes a task instead of direct questions, complete the task as fully as possible."
            )
        else:
            instructions = (
                "Follow the user's specific prompt while staying grounded in the uploaded content and analysis.\n"
                f"Specific prompt: {custom_prompt}"
            )

        prompt = (
            "Use the uploaded quiz/activity content for educational support.\n"
            f"Saved dashboard profile:\n{self._identity_block() or 'No saved profile'}\n\n"
            f"Analysis already prepared:\n{self.quiz_summary_cache}\n\n"
            f"Source content:\n{source}\n\n"
            f"Instructions:\n{instructions}\n\n"
            "Structure the answer with:\n"
            "1. Task understanding\n"
            "2. Response\n"
            "3. Short explanation or rationale\n\n"
            "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
        )
        self._start_ai_job("quiz response", prompt, self.finish_quiz_response)

    def finish_quiz_response(self, result: str) -> None:
        self.quiz_response_cache = result
        self._set_text(self.quiz_result_text, result, read_only=True)
        self.quiz_export_button.configure(state="normal")
        self.quiz_progress_var.set("Quiz response ready. Export if you want a Word copy.")
        self.status_var.set("Quiz response complete.")

    def export_quiz_result(self) -> None:
        if not self.quiz_response_cache.strip():
            self._notify("Nothing To Export", "Generate a quiz response first.", level="warning")
            return
        output_path = self._export_text_output(self.quiz_export_title, self.quiz_response_cache, self.quiz_export_name, "quiz")
        if output_path is not None:
            self.clear_quiz_workspace()
            self.status_var.set(f"Saved {output_path.name} and cleared the quiz workspace.")

    def _essay_language(self) -> str | None:
        tagalog = self.essay_tagalog_var.get()
        english = self.essay_english_var.get()
        if tagalog and english:
            return "Taglish"
        if tagalog:
            return "Tagalog"
        if english:
            return "English"
        return None

    def start_essay_generation(self) -> None:
        title = self.essay_title_var.get().strip()
        prompt_text = self._get_text(self.essay_prompt_text).strip()
        language = self._essay_language()
        if not title:
            self._notify("Missing Title", "Enter the essay title before generating.", level="warning")
            return
        if language is None:
            self._notify("Missing Language", "Select Tagalog, English, or both for Taglish.", level="warning")
            return

        try:
            target_words = int(self.essay_word_count_var.get().strip())
        except ValueError:
            self._notify("Invalid Word Count", "Enter a valid number for the target word count.", level="warning")
            return
        if target_words < 100:
            self._notify("Word Count Too Low", "Use at least 100 words for the essay target.", level="warning")
            return

        self.essay_export_title = title
        self.essay_export_name = sanitize_filename(title.lower().replace(" ", "_"))
        self.essay_progress_var.set("Generating essay draft...")
        prompt = (
            "Write a polished educational essay draft.\n"
            f"Saved dashboard profile:\n{self._identity_block() or 'No saved profile'}\n\n"
            f"Essay title: {title}\n"
            f"Target length: about {target_words} words\n"
            f"Language: {language}\n"
            f"Specific prompt: {prompt_text or 'No extra prompt provided'}\n\n"
            "Return these plain-text sections in this exact order:\n"
            "Heading Suggestion:\n"
            "Essay Body:\n"
            "Self-Check Tip:\n\n"
            "Do not use markdown symbols like ### or ####."
        )
        self._start_ai_job("essay draft", prompt, self.finish_essay_generation)

    def finish_essay_generation(self, result: str) -> None:
        self.essay_response_cache = result
        self._set_text(self.essay_result_text, result, read_only=True)
        self.essay_export_button.configure(state="normal")
        self.essay_progress_var.set("Essay ready. Export it if you want a Word copy.")
        self.status_var.set("Essay generation complete.")

    def export_essay_result(self) -> None:
        if not self.essay_response_cache.strip():
            self._notify("Nothing To Export", "Generate an essay first.", level="warning")
            return
        output_path = self._export_text_output(
            self.essay_export_title,
            self.essay_response_cache,
            self.essay_export_name,
            "essay",
            name_override=self.essay_specific_name_var.get().strip(),
        )
        if output_path is not None:
            self.clear_essay_form()
            self.clear_essay_result()
            self.status_var.set(f"Saved {output_path.name} and cleared the essay workspace.")

    def start_activity_generation(self) -> None:
        title = self.activity_title_var.get().strip()
        activity_type = self.activity_type_var.get().strip() or "Activity"
        level = self.activity_level_var.get().strip()
        request = self._get_text(self.activity_prompt_text).strip()
        if not title:
            self._notify("Missing Topic", "Enter the activity topic or title before generating.", level="warning")
            return

        self.activity_export_title = title
        self.activity_export_name = sanitize_filename(title.lower().replace(" ", "_"))
        self.activity_progress_var.set("Generating activity draft...")
        prompt = (
            "Create a structured educational activity for offline study use.\n"
            f"Saved dashboard profile:\n{self._identity_block() or 'No saved profile'}\n\n"
            f"Activity title or topic: {title}\n"
            f"Activity type: {activity_type}\n"
            f"Level or class: {level or 'Not specified'}\n"
            f"Specific request: {request or 'No extra request provided'}\n\n"
            "Return these plain-text sections in this exact order:\n"
            "Activity Title:\n"
            "Objective:\n"
            "Instructions:\n"
            "Activity Proper:\n"
            "Answer Guide:\n\n"
            "Do not use markdown symbols like ### or ####."
        )
        self._start_ai_job("activity draft", prompt, self.finish_activity_generation)

    def finish_activity_generation(self, result: str) -> None:
        self.activity_response_cache = result
        self._set_text(self.activity_result_text, result, read_only=True)
        self.activity_export_button.configure(state="normal")
        self.activity_progress_var.set("Activity ready. Export it if you want a Word copy.")
        self.status_var.set("Activity generation complete.")

    def export_activity_result(self) -> None:
        if not self.activity_response_cache.strip():
            self._notify("Nothing To Export", "Generate an activity first.", level="warning")
            return
        output_path = self._export_text_output(
            self.activity_export_title,
            self.activity_response_cache,
            self.activity_export_name,
            "activity",
        )
        if output_path is not None:
            self.clear_activity_form()
            self.clear_activity_result()
            self.status_var.set(f"Saved {output_path.name} and cleared the activity workspace.")

    def start_document_generation(self) -> None:
        title = self.document_title_var.get().strip()
        doc_type = self.document_type_var.get().strip() or "Study Handout"
        audience = self.document_audience_var.get().strip()
        request = self._get_text(self.document_prompt_text).strip()
        if not title:
            self._notify("Missing Title", "Enter the document title before generating.", level="warning")
            return
        if not request:
            self._notify("Missing Request", "Describe what the document should contain before generating.", level="warning")
            return

        self.document_export_title = title
        self.document_export_name = sanitize_filename(title.lower().replace(" ", "_"))
        self.document_progress_var.set("Generating document draft...")
        prompt = (
            "Create a structured educational document.\n"
            f"Saved dashboard profile:\n{self._identity_block() or 'No saved profile'}\n\n"
            f"Document title: {title}\n"
            f"Document type: {doc_type}\n"
            f"Audience or purpose: {audience or 'Not specified'}\n"
            f"Content request: {request}\n\n"
            "Return these plain-text sections in this exact order:\n"
            "Document Type:\n"
            "Purpose:\n"
            "Main Content:\n\n"
            "Do not use markdown symbols like ### or ####."
        )
        self._start_ai_job("document draft", prompt, self.finish_document_generation)

    def finish_document_generation(self, result: str) -> None:
        self.document_response_cache = result
        self._set_text(self.document_result_text, result, read_only=True)
        self.document_export_button.configure(state="normal")
        self.document_progress_var.set("Document ready. Export it if you want a Word copy.")
        self.status_var.set("Document generation complete.")

    def export_document_result(self) -> None:
        if not self.document_response_cache.strip():
            self._notify("Nothing To Export", "Generate a document first.", level="warning")
            return
        output_path = self._export_text_output(
            self.document_export_title,
            self.document_response_cache,
            self.document_export_name,
            "document",
        )
        if output_path is not None:
            self.clear_document_form()
            self.clear_document_result()
            self.status_var.set(f"Saved {output_path.name} and cleared the document workspace.")

    def start_assignment_analysis(self) -> None:
        source = self._get_text(self.assignment_source_text).strip()
        if not source:
            self._notify("Missing Content", "Upload or paste assignment content before analyzing it.", level="warning")
            return

        self.assignment_progress_var.set("Reading the assignment and summarizing the task...")
        prompt = (
            "Analyze the following assignment for educational review.\n"
            f"Saved dashboard profile:\n{self._identity_block() or 'No saved profile'}\n\n"
            f"Assignment content:\n{source}\n\n"
            "Return these sections clearly:\n"
            "1. Assignment summary\n"
            "2. What is being asked\n"
            "3. Important requirements or constraints\n"
            "4. Best approach for the student"
        )
        self._start_ai_job("assignment analysis", prompt, self.finish_assignment_analysis)

    def finish_assignment_analysis(self, result: str) -> None:
        self.assignment_summary_cache = result
        self._set_text(self.assignment_summary_text, result, read_only=True)
        self.assignment_progress_var.set("Assignment summary ready. Generate the guided response next.")
        self.status_var.set("Assignment analysis complete.")

    def start_assignment_response(self) -> None:
        source = self._get_text(self.assignment_source_text).strip()
        if not source:
            self._notify("Missing Content", "Upload or paste assignment content before generating a response.", level="warning")
            return
        if not self.assignment_summary_cache.strip():
            self._notify("Analyze First", "Run the assignment analysis first.", level="warning")
            return

        mode = self.assignment_mode_var.get()
        custom_prompt = self._get_text(self.assignment_prompt_text).strip()
        source_name = self.assignment_loaded_path.stem if self.assignment_loaded_path else "assignment_task"
        self.assignment_export_title = f"Assignment Support: {source_name.replace('_', ' ').title()}"
        self.assignment_export_name = sanitize_filename(source_name)
        self.assignment_progress_var.set("Generating assignment guidance...")

        if mode == "complete":
            instructions = (
                "Create a full draft response based on the assignment while keeping it clear, organized, and educational."
            )
        else:
            instructions = (
                "Create a guided response that explains how to approach the assignment and includes a sample draft the student can review."
            )
        if custom_prompt:
            instructions = f"{instructions}\nSpecific prompt: {custom_prompt}"

        prompt = (
            "Use the uploaded assignment for educational support.\n"
            f"Saved dashboard profile:\n{self._identity_block() or 'No saved profile'}\n\n"
            f"Assignment analysis:\n{self.assignment_summary_cache}\n\n"
            f"Assignment content:\n{source}\n\n"
            f"Instructions:\n{instructions}\n\n"
            "Structure the answer with:\n"
            "1. Task understanding\n"
            "2. Solution plan\n"
            "3. Sample answer or draft\n"
            "4. Notes to review\n\n"
            "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
        )
        self._start_ai_job("assignment response", prompt, self.finish_assignment_response)

    def finish_assignment_response(self, result: str) -> None:
        self.assignment_response_cache = result
        self._set_text(self.assignment_result_text, result, read_only=True)
        self.assignment_export_button.configure(state="normal")
        self.assignment_progress_var.set("Assignment guidance ready. Export if you want a Word copy.")
        self.status_var.set("Assignment response complete.")

    def export_assignment_result(self) -> None:
        if not self.assignment_response_cache.strip():
            self._notify("Nothing To Export", "Generate assignment guidance first.", level="warning")
            return
        output_path = self._export_text_output(
            self.assignment_export_title,
            self.assignment_response_cache,
            self.assignment_export_name,
            "assignment",
        )
        if output_path is not None:
            self.clear_assignment_workspace()
            self.status_var.set(f"Saved {output_path.name} and cleared the assignment workspace.")

    def start_codefix_generation(self) -> None:
        title = self.codefix_title_var.get().strip()
        language = self.codefix_language_var.get().strip() or "Code"
        error_text = self._get_text(self.codefix_error_text).strip()
        source = self._get_text(self.codefix_source_text).strip()
        expectation = self._get_text(self.codefix_expectation_text).strip()
        if not source:
            self._notify("Missing Code", "Paste the code snippet before generating a fix.", level="warning")
            return
        if not error_text:
            self._notify("Missing Error Details", "Add the error message or symptoms before generating a fix.", level="warning")
            return

        export_title = title or f"{language} Code Fix"
        self.codefix_export_title = export_title
        self.codefix_export_name = sanitize_filename(export_title.lower().replace(" ", "_"))
        self.codefix_progress_var.set("Generating code fix...")
        prompt = (
            "Fix the following code issue for study and debugging support.\n\n"
            f"Language or stack: {language}\n"
            f"Issue title: {title or 'No title provided'}\n"
            f"Error message or symptoms:\n{error_text}\n\n"
            f"Code snippet:\n{source}\n\n"
            f"Expected behavior:\n{expectation or 'Not provided'}\n\n"
            "Return these plain-text sections in this exact order:\n"
            "Issue Summary:\n"
            "Root Cause:\n"
            "Fixed Version:\n"
            "Why It Works:\n"
            "Next Checks:\n\n"
            "Do not use markdown symbols like ### or ####."
        )
        self._start_ai_job("code fix", prompt, self.finish_codefix_generation)

    def finish_codefix_generation(self, result: str) -> None:
        self.codefix_response_cache = result
        self._set_text(self.codefix_result_text, result, read_only=True)
        self.codefix_export_button.configure(state="normal")
        self.codefix_progress_var.set("Code fix ready. Export it if you want a Word copy.")
        self.status_var.set("Code fix generation complete.")

    def export_codefix_result(self) -> None:
        if not self.codefix_response_cache.strip():
            self._notify("Nothing To Export", "Generate a code fix first.", level="warning")
            return
        output_path = self._export_text_output(
            self.codefix_export_title,
            self.codefix_response_cache,
            self.codefix_export_name,
            "codefix",
        )
        if output_path is not None:
            self.clear_codefix_form()
            self.clear_codefix_result()
            self.status_var.set(f"Saved {output_path.name} and cleared the code fixer workspace.")


def main() -> None:
    app = PremiumStudyAssistant()
    app.mainloop()


if __name__ == "__main__":
    main()
