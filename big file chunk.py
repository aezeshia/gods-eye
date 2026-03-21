from __future__ import annotations

import io
import json
import os
import re
import sqlite3
import subprocess
import sys
import difflib
import hashlib
import hmac
import textwrap
from datetime import datetime
from html import escape
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv(*_args, **_kwargs):
        return False

try:
    from huggingface_hub import InferenceClient
except ImportError:
    InferenceClient = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import numpy as np
except ImportError:
    np = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from rapidocr_onnxruntime import RapidOCR
except ImportError:
    RapidOCR = None

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
except ImportError:
    LETTER = None
    canvas = None

try:
    # Tkinter is only used for the optional local folder picker.
    # Streamlit Cloud / headless environments may not have a usable Tk runtime.
    import tkinter as tk
    from tkinter import filedialog
except Exception:
    tk = None
    filedialog = None


APP_DIR = Path(__file__).resolve().parent
ENV_PATH = APP_DIR / ".env"
AUTH_DB_PATH = APP_DIR / "hugyoku_auth.db"
_OCR_ENGINE: object | None = None

ROLE_PERMISSIONS = {
    "super_admin": {
        "hugyoku": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": True,
        "history": True,
        "settings": True,
        "admin": True,
    },
    "admin": {
        "hugyoku": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": True,
        "history": True,
        "settings": True,
        "admin": True,
    },
    "member": {
        "hugyoku": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": True,
        "history": True,
        "settings": True,
        "admin": False,
    },
    "viewer": {
        "hugyoku": True,
        "dashboard": True,
        "workspaces": True,
        "academics": True,
        "developer": False,
        "history": True,
        "settings": True,
        "admin": False,
    },
}

SYSTEM_PROMPT = (
    "You are Hugyoku, an educational assistant for study use. Help the user analyze "
    "activities, quizzes, assignments, uploaded files, documents, essays, and code "
    "issues. You may produce sample answers, draft essays, structured summaries, "
    "guided explanations, generated study documents, and debugging help for offline "
    "learning and review. Keep outputs clear and well-organized, and avoid assisting "
    "with live cheating or deceptive conduct."
)

TOOL_FOLDERS = {
    "quiz": "quiz_solver",
    "assignment": "assignment_solver",
    "essay": "essay_generator",
    "activity": "activity_generator",
    "document": "document_generator",
    "codefix": "code_error_fixer",
}

KNOWN_EXPORT_LABELS = {
    "heading suggestion",
    "title suggestion",
    "suggested heading",
    "essay body",
    "essay",
    "self-check tip",
    "quick self-check tip",
    "closing self-check tip",
    "task understanding",
    "response",
    "short explanation or rationale",
    "summary",
    "objective",
    "instructions",
    "activity proper",
    "answer guide",
    "document type",
    "purpose",
    "main content",
    "solution plan",
    "sample answer or draft",
    "notes to review",
    "issue summary",
    "root cause",
    "fixed version",
    "why it works",
    "next checks",
}

STATE_DEFAULTS = {
    "active_page": "dashboard",
    "saved_name": "",
    "saved_include_date": False,
    "save_destination_mode": "browser",
    "main_folder_name": "hugyoku_exports",
    "export_root_path": "",
    "output_include_name": True,
    "output_include_date": False,
    "essay_include_heading": True,
    "essay_include_tip": True,
    "profile_name_input": "",
    "profile_include_date_input": False,
    "profile_save_destination_mode_input": "browser",
    "profile_main_folder_input": "hugyoku_exports",
    "profile_export_root_path_input": "",
    "profile_output_include_name_input": True,
    "profile_output_include_date_input": False,
    "profile_essay_include_heading_input": True,
    "profile_essay_include_tip_input": True,
    "quiz_upload_name": "No file loaded yet",
    "quiz_source_text": "",
    "quiz_summary": "",
    "quiz_prompt": "",
    "quiz_response": "",
    "quiz_mode": "complete",
    "assignment_upload_name": "No assignment file loaded yet",
    "assignment_source_text": "",
    "assignment_summary": "",
    "assignment_prompt": "",
    "assignment_response": "",
    "assignment_mode": "guided",
    "essay_title": "",
    "essay_prompt": "",
    "essay_word_count": 500,
    "essay_tagalog": False,
    "essay_english": True,
    "essay_specific_name": "",
    "essay_response": "",
    "activity_title": "",
    "activity_type": "Worksheet",
    "activity_level": "",
    "activity_prompt": "",
    "activity_response": "",
    "document_title": "",
    "document_type": "Study Handout",
    "document_audience": "",
    "document_prompt": "",
    "document_response": "",
    "codefix_title": "",
    "codefix_language": "Python",
    "codefix_error": "",
    "codefix_source": "",
    "codefix_expectation": "",
    "codefix_response": "",
    "pending_reset_action": "",
    "pending_reset_notice": "",
    "pending_reset_level": "success",
    "pending_export_root_selection": None,
    "flash_message": "",
    "flash_level": "success",
}

PAGE_DETAILS = {
    "dashboard": {
        "title": "Dashboard",
        "subtitle": "Save your identity, choose one export folder, and keep all Hugyoku outputs organized for local saves or browser downloads.",
    },
    "academics": {
        "title": "Academics",
        "subtitle": "Choose a focused school workspace. Each tool keeps the same Hugyoku content flow but is adapted for browser-based use.",
    },
    "developer": {
        "title": "Developer",
        "subtitle": "A separate hub for code support so the academic tools stay clean and easier to scan.",
    },
    "quiz": {
        "title": "Quiz Solver",
        "subtitle": "Read quiz files, summarize the task, and generate guided response support in a dedicated web workspace.",
    },
    "assignment": {
        "title": "Assignment Solver",
        "subtitle": "Analyze assignment instructions, understand the task first, then generate a guided response or full draft.",
    },
    "essay": {
        "title": "Essay Generator",
        "subtitle": "Build essay drafts with the same Hugyoku output controls, optional export name, and Word export formatting.",
    },
    "activity": {
        "title": "Activity Generator",
        "subtitle": "Generate worksheets, reflections, drills, or school activities from a topic and instruction set.",
    },
    "document": {
        "title": "Document Generator",
        "subtitle": "Create structured school documents like handouts, reports, reviewers, and formal academic materials.",
    },
    "codefix": {
        "title": "Code Error Fixer",
        "subtitle": "Paste code and the error details, then generate a cleaner fix and explanation in a separate developer workspace.",
    },
}

THEME_CSS = """
<style>
:root {
  --bg-main: #0a1018;
  --bg-top: #18233a;
  --panel: rgba(18, 25, 37, 0.94);
  --panel-soft: rgba(25, 35, 52, 0.9);
  --panel-muted: rgba(31, 41, 59, 0.76);
  --border: rgba(100, 118, 148, 0.34);
  --border-strong: rgba(146, 177, 232, 0.36);
  --text: #f5efe1;
  --muted: #b9c4d8;
  --accent: #ff8b72;
  --accent-soft: #f4dca7;
  --success: #53c795;
  --shadow: 0 20px 50px rgba(2, 8, 18, 0.34);
}

html,
body,
[class*="css"] {
  color: var(--text);
}

.stApp {
  background: radial-gradient(circle at top left, var(--bg-top) 0%, var(--bg-main) 48%, #090d14 100%);
  color: var(--text);
  overflow-x: hidden;
}

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #101722 0%, #0c121b 100%);
  border-right: 1px solid rgba(100, 118, 148, 0.22);
}

[data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
  gap: 0.95rem;
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, rgba(22, 30, 44, 0.95) 0%, rgba(15, 21, 33, 0.95) 100%);
  border: 1px solid rgba(100, 118, 148, 0.25) !important;
  border-radius: 22px !important;
  box-shadow: none;
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: 0.95rem 1rem;
}

.main .block-container {
  max-width: 1480px;
  padding-top: clamp(0.75rem, 1.5vw, 1.1rem);
  padding-bottom: 1.8rem;
  padding-left: clamp(0.8rem, 2vw, 1.4rem);
  padding-right: clamp(0.8rem, 2vw, 1.4rem);
}

[data-testid="stHorizontalBlock"] {
  gap: 0.8rem;
  align-items: stretch;
  flex-wrap: wrap;
}

[data-testid="column"] {
  min-width: min(100%, 280px) !important;
  flex: 1 1 300px !important;
}

[data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, var(--panel) 0%, rgba(14, 20, 31, 0.97) 100%);
  border: 1px solid var(--border) !important;
  border-radius: 24px !important;
  box-shadow: var(--shadow);
  backdrop-filter: blur(16px);
  height: 100%;
  transition: transform 0.18s ease, border-color 0.18s ease, box-shadow 0.18s ease;
}

[data-testid="stVerticalBlockBorderWrapper"]:hover {
  transform: translateY(-2px);
  border-color: var(--border-strong) !important;
  box-shadow: 0 24px 56px rgba(0, 0, 0, 0.28);
}

[data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: clamp(0.8rem, 0.85vw + 0.62rem, 1.08rem);
  height: 100%;
}

.stMarkdown,
.stText,
p,
li,
label,
span,
div {
  overflow-wrap: anywhere;
}

p,
li,
label,
.stCaption,
.stMarkdown {
  color: var(--muted);
}

h1,
h2,
h3,
h4,
h5,
h6 {
  color: var(--text);
  letter-spacing: -0.02em;
}

.app-hero-title {
  font-size: clamp(2rem, 4vw, 3.2rem);
  line-height: 1.1;
  color: var(--text);
  font-weight: 800;
  margin-bottom: 0.45rem;
}

.app-card-header {
  display: flex;
  flex-direction: column;
  gap: 0.28rem;
  margin-bottom: 0.72rem;
}

.app-kicker {
  font-size: 0.75rem;
  text-transform: uppercase;
  letter-spacing: 0.16em;
  font-weight: 800;
  color: var(--accent-soft);
}

.app-card-title {
  font-size: clamp(1.2rem, 2vw, 1.85rem);
  line-height: 1.16;
  font-weight: 800;
  color: var(--text);
}

.app-card-subtitle {
  font-size: 0.98rem;
  line-height: 1.65;
  color: var(--muted);
  margin: 0;
}

.app-chip-row {
  display: flex;
  flex-wrap: wrap;
  gap: 0.45rem;
  margin-top: 0.75rem;
}

.app-chip {
  display: inline-flex;
  align-items: center;
  padding: 0.38rem 0.72rem;
  border-radius: 999px;
  border: 1px solid rgba(243, 215, 154, 0.2);
  background: rgba(243, 215, 154, 0.1);
  color: var(--accent-soft);
  font-size: 0.82rem;
  font-weight: 700;
}

.app-status-pill {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  padding: 0.45rem 0.82rem;
  border-radius: 999px;
  font-size: 0.84rem;
  font-weight: 800;
  margin-bottom: 1rem;
}

.app-status-pill.ready {
  background: rgba(83, 199, 149, 0.18);
  color: #d6ffee;
  border: 1px solid rgba(83, 199, 149, 0.26);
}

.app-status-pill.waiting {
  background: rgba(228, 182, 94, 0.16);
  color: #ffe7bb;
  border: 1px solid rgba(228, 182, 94, 0.22);
}

.app-status-pill.offline {
  background: rgba(227, 123, 116, 0.16);
  color: #ffd4cf;
  border: 1px solid rgba(227, 123, 116, 0.2);
}

.app-meta-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
  gap: 0.75rem;
  margin-bottom: 1rem;
  width: 100%;
}

.app-meta-item {
  padding: 0.72rem 0.82rem;
  border-radius: 16px;
  background: var(--panel-muted);
  border: 1px solid rgba(100, 118, 148, 0.22);
  min-width: 0;
  height: 100%;
}

.app-meta-label {
  font-size: 0.74rem;
  text-transform: uppercase;
  letter-spacing: 0.14em;
  color: var(--accent-soft);
  margin-bottom: 0.35rem;
}

.app-meta-value {
  color: var(--text);
  font-weight: 700;
  line-height: 1.45;
  word-break: break-word;
  overflow-wrap: anywhere;
}

.app-route-block {
  padding: 0.72rem 0.86rem;
  border-radius: 18px;
  background: var(--panel-muted);
  border: 1px solid rgba(100, 118, 148, 0.22);
  margin-bottom: 0.62rem;
}

.app-route-label {
  font-size: 0.75rem;
  text-transform: uppercase;
  letter-spacing: 0.14em;
  color: var(--accent-soft);
  margin-bottom: 0.4rem;
}

.app-route-body {
  color: var(--muted);
  font-family: Consolas, "Courier New", monospace;
  font-size: 0.88rem;
  line-height: 1.6;
  white-space: pre-wrap;
}

.app-guide-list {
  margin: 0;
  padding-left: 1.1rem;
  color: var(--muted);
}

.app-guide-list li {
  margin-bottom: 0.55rem;
  line-height: 1.65;
}

.app-soft-note {
  color: var(--muted);
  font-size: 0.9rem;
  line-height: 1.6;
}

.app-section-gap {
  height: 0.35rem;
}

.app-sidebar-brand {
  font-size: 1.8rem;
  font-weight: 800;
  letter-spacing: -0.03em;
  color: var(--text);
  margin-bottom: 0.35rem;
}

.app-sidebar-copy {
  color: var(--muted);
  line-height: 1.7;
  margin-bottom: 0.4rem;
}

code {
  color: #ffe9bd;
}

.stButton > button,
.stDownloadButton > button {
  width: 100%;
  min-height: 2.85rem;
  border-radius: 16px;
  border: 1px solid var(--border);
  background: linear-gradient(180deg, rgba(36, 47, 67, 0.96) 0%, rgba(20, 28, 42, 0.98) 100%);
  color: var(--text);
  font-weight: 700;
  transition: transform 0.16s ease, border-color 0.16s ease, box-shadow 0.16s ease, background 0.16s ease;
}

.stButton > button:hover,
.stDownloadButton > button:hover {
  transform: translateY(-1px);
  border-color: rgba(244, 220, 167, 0.28);
  box-shadow: 0 14px 26px rgba(0, 0, 0, 0.24);
}

.stButton > button:focus,
.stDownloadButton > button:focus {
  box-shadow: 0 0 0 0.12rem rgba(244, 220, 167, 0.18);
}

div[data-testid="stBaseButton-primary"] > button,
.stDownloadButton > button[kind="primary"] {
  background: linear-gradient(135deg, var(--accent) 0%, var(--accent-soft) 100%);
  color: #111723;
  border: none;
  box-shadow: 0 14px 28px rgba(255, 139, 114, 0.28);
}

div[data-testid="stBaseButton-primary"] > button:hover,
.stDownloadButton > button[kind="primary"]:hover {
  box-shadow: 0 18px 34px rgba(255, 139, 114, 0.32);
}

.stButton > button p,
.stDownloadButton > button p {
  color: inherit !important;
}

[data-baseweb="input"] > div,
[data-baseweb="textarea"] > div,
[data-baseweb="select"] > div,
.stNumberInput > div > div {
  background: rgba(12, 18, 28, 0.88) !important;
  border: 1px solid rgba(90, 108, 137, 0.34) !important;
  border-radius: 16px !important;
}

[data-baseweb="input"] input,
[data-baseweb="textarea"] textarea,
.stNumberInput input {
  color: var(--text) !important;
  font-size: 0.98rem !important;
  line-height: 1.55 !important;
}

[data-baseweb="input"] input::placeholder,
[data-baseweb="textarea"] textarea::placeholder {
  color: #7f8aa0 !important;
}

[data-testid="stFileUploader"] section {
  background: rgba(12, 18, 28, 0.72);
  border: 1px dashed rgba(100, 118, 148, 0.4);
  border-radius: 18px;
}

[data-testid="stTabs"] [data-baseweb="tab-list"] {
  gap: 0.45rem;
  flex-wrap: wrap;
}

[data-testid="stTabs"] [data-baseweb="tab"] {
  min-height: 2.55rem;
  border-radius: 999px;
  padding-inline: 0.95rem;
}

div[role="radiogroup"] {
  gap: 0.55rem;
  flex-wrap: wrap;
}

div[role="radiogroup"] label {
  border: 1px solid rgba(100, 118, 148, 0.3);
  border-radius: 999px;
  padding: 0.3rem 0.8rem;
  background: rgba(18, 25, 38, 0.75);
}

@media (max-width: 1024px) {
  .main .block-container {
    padding-left: 0.85rem;
    padding-right: 0.85rem;
  }
}

@media (max-width: 920px) {
  [data-testid="stHorizontalBlock"] {
    flex-direction: column !important;
    align-items: stretch !important;
    gap: 0.8rem !important;
  }

  [data-testid="column"] {
    width: 100% !important;
    max-width: 100% !important;
    min-width: 100% !important;
    flex-basis: 100% !important;
    flex: 1 1 100% !important;
  }

  .app-hero-title {
    font-size: clamp(1.85rem, 8vw, 2.7rem);
  }

  .app-meta-grid {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }
}

@media (max-width: 768px) {
  [data-testid="stSidebar"] {
    min-width: 15rem;
    max-width: 15rem;
  }

  [data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] > div {
    padding: 0.75rem 0.8rem;
  }

  .app-card-title {
    font-size: 1.08rem;
  }

  .app-card-subtitle,
  .app-soft-note,
  .app-route-body {
    font-size: 0.88rem;
  }

  .stButton > button,
  .stDownloadButton > button {
    min-height: 3rem;
    font-size: 0.94rem;
  }

  .app-sidebar-brand {
    font-size: 1.45rem;
  }

  .app-sidebar-copy {
    font-size: 0.8rem;
    line-height: 1.45;
  }

  [data-testid="stTabs"] [data-baseweb="tab"] {
    width: 100%;
    justify-content: center;
  }

  .app-chip {
    width: 100%;
    justify-content: center;
  }

  .app-meta-grid {
    grid-template-columns: 1fr;
  }

  .app-meta-item,
  .app-route-block {
    padding: 0.7rem 0.78rem;
  }
}

@media (max-width: 520px) {
  .main .block-container {
    padding-left: 0.65rem;
    padding-right: 0.65rem;
    padding-top: 0.65rem;
  }

  [data-testid="stVerticalBlockBorderWrapper"] > div {
    padding: 0.7rem;
  }

  .app-hero-title {
    font-size: 1.85rem;
  }

  .app-card-subtitle,
  .app-route-body,
  .app-soft-note,
  .app-sidebar-copy {
    font-size: 0.84rem;
    line-height: 1.5;
  }

  .app-chip-row {
    gap: 0.35rem;
  }

  .app-chip {
    padding: 0.38rem 0.62rem;
    font-size: 0.76rem;
  }

  [data-baseweb="input"] input,
  [data-baseweb="textarea"] textarea,
  .stNumberInput input {
    font-size: 0.92rem !important;
  }
}

/* Premium 90+ polish overrides */

.stApp {
  background:
    radial-gradient(circle at 8% 8%, rgba(255, 166, 114, 0.12) 0%, rgba(255, 166, 114, 0.03) 20%, transparent 42%),
    radial-gradient(circle at 92% 12%, rgba(116, 176, 255, 0.12) 0%, rgba(116, 176, 255, 0.03) 24%, transparent 44%),
    linear-gradient(180deg, #0c111b 0%, #080d15 55%, #060a12 100%);
}

.main .block-container {
  max-width: 1380px;
  padding-top: clamp(0.78rem, 1vw, 1.02rem);
  padding-bottom: 2rem;
}

[data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, rgba(17, 24, 37, 0.97) 0%, rgba(10, 15, 24, 0.98) 100%);
  border: 1px solid rgba(118, 136, 168, 0.26) !important;
  border-radius: 22px !important;
  box-shadow: 0 16px 34px rgba(2, 8, 18, 0.24);
}

[data-testid="stVerticalBlockBorderWrapper"]:hover {
  transform: translateY(-2px);
  border-color: rgba(182, 202, 236, 0.42) !important;
  box-shadow: 0 20px 42px rgba(2, 8, 18, 0.28);
}

[data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: clamp(0.82rem, 0.82vw + 0.6rem, 1.04rem);
}

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, rgba(10, 15, 24, 0.98) 0%, rgba(7, 11, 18, 0.99) 100%);
  border-right: 1px solid rgba(96, 114, 148, 0.16);
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
  background: linear-gradient(180deg, rgba(18, 25, 38, 0.94) 0%, rgba(12, 18, 29, 0.96) 100%);
  border-radius: 18px !important;
}

.app-sidebar-brand-row {
  display: flex;
  align-items: flex-start;
  gap: 0.8rem;
  margin-bottom: 0.8rem;
}

.app-sidebar-brand {
  font-size: 1.38rem;
  line-height: 1.05;
  margin-bottom: 0.14rem;
}

.app-sidebar-copy {
  font-size: 0.84rem;
  line-height: 1.45;
  margin-bottom: 0;
}

.app-sidebar-quiet {
  color: rgba(185, 196, 216, 0.66);
  font-size: 0.75rem;
  line-height: 1.4;
}

.app-nav-section {
  margin: 0.18rem 0 0.48rem;
  font-size: 0.72rem;
  text-transform: uppercase;
  letter-spacing: 0.18em;
  color: rgba(244, 220, 167, 0.78);
  font-weight: 800;
}

.app-card-topline {
  display: flex;
  align-items: center;
  gap: 0.55rem;
  flex-wrap: wrap;
  margin-bottom: 0.52rem;
}

.app-anchor-badge {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-width: 2rem;
  height: 2rem;
  padding: 0 0.55rem;
  border-radius: 12px;
  background: linear-gradient(135deg, rgba(255, 139, 114, 0.18) 0%, rgba(244, 220, 167, 0.18) 100%);
  border: 1px solid rgba(244, 220, 167, 0.26);
  color: var(--text);
  font-size: 0.76rem;
  font-weight: 900;
  letter-spacing: 0.08em;
  flex-shrink: 0;
}

.app-tier-pill {
  display: inline-flex;
  align-items: center;
  padding: 0.26rem 0.66rem;
  border-radius: 999px;
  font-size: 0.72rem;
  font-weight: 800;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}

.app-tier-pill.primary {
  color: #111723;
  background: linear-gradient(135deg, var(--accent) 0%, var(--accent-soft) 100%);
}

.app-tier-pill.secondary {
  color: #f7e7c0;
  background: rgba(244, 220, 167, 0.1);
  border: 1px solid rgba(244, 220, 167, 0.2);
}

.app-tier-pill.tertiary {
  color: #d4dded;
  background: rgba(118, 136, 168, 0.14);
  border: 1px solid rgba(118, 136, 168, 0.2);
}

.app-card-header {
  gap: 0.2rem;
  margin-bottom: 0.68rem;
}

.app-card-title {
  font-size: clamp(1.12rem, 1.4vw, 1.58rem);
  line-height: 1.15;
}

.app-card-subtitle {
  max-width: 74ch;
  font-size: 0.93rem;
  line-height: 1.58;
}

.app-card-subtitle.compact {
  font-size: 0.86rem;
  line-height: 1.48;
}

.app-hero-shell {
  display: flex;
  flex-direction: column;
  gap: 0.9rem;
}

.app-hero-title {
  font-size: clamp(2.18rem, 4vw, 3.5rem);
  line-height: 1.02;
  letter-spacing: -0.03em;
  max-width: 12ch;
  margin-bottom: 0;
}

.app-hero-lead {
  color: var(--muted);
  font-size: 0.98rem;
  line-height: 1.62;
  max-width: 72ch;
}

.app-kpi-row {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
  gap: 0.62rem;
  margin-top: 0.12rem;
}

.app-kpi-pill {
  min-width: 0;
  padding: 0.74rem 0.8rem;
  border-radius: 16px;
  background: linear-gradient(180deg, rgba(25, 34, 50, 0.94) 0%, rgba(17, 24, 36, 0.96) 100%);
  border: 1px solid rgba(112, 130, 162, 0.18);
}

.app-kpi-label {
  display: block;
  margin-bottom: 0.3rem;
  color: rgba(244, 220, 167, 0.84);
  font-size: 0.7rem;
  letter-spacing: 0.14em;
  text-transform: uppercase;
  font-weight: 800;
}

.app-kpi-value {
  display: block;
  color: var(--text);
  font-size: 0.94rem;
  font-weight: 800;
  line-height: 1.4;
  overflow-wrap: anywhere;
}

.app-chip-row {
  gap: 0.42rem;
  margin-top: 0.1rem;
}

.app-chip {
  padding: 0.42rem 0.78rem;
  background: linear-gradient(180deg, rgba(33, 43, 62, 0.94) 0%, rgba(22, 30, 45, 0.96) 100%);
  border: 1px solid rgba(243, 215, 154, 0.18);
  transition: transform 0.18s ease, border-color 0.18s ease, background 0.18s ease;
}

.app-chip:hover {
  transform: translateY(-1px);
  border-color: rgba(243, 215, 154, 0.28);
  background: linear-gradient(180deg, rgba(44, 57, 80, 0.96) 0%, rgba(26, 35, 52, 0.98) 100%);
}

.app-status-pill {
  padding: 0.44rem 0.86rem;
  margin-bottom: 0.82rem;
}

.app-meta-grid {
  gap: 0.62rem;
  margin-bottom: 0.82rem;
}

.app-meta-item {
  padding: 0.74rem 0.82rem;
  background: linear-gradient(180deg, rgba(34, 45, 65, 0.9) 0%, rgba(26, 35, 52, 0.96) 100%);
  border-radius: 18px;
}

.app-route-block {
  padding: 0.74rem 0.86rem;
  background: linear-gradient(180deg, rgba(28, 38, 56, 0.92) 0%, rgba(20, 28, 42, 0.96) 100%);
  border-radius: 18px;
  margin-bottom: 0.55rem;
}

.app-route-label {
  margin-bottom: 0.36rem;
}

.app-route-body {
  font-size: 0.84rem;
  line-height: 1.56;
}

[data-testid="stExpander"] {
  border: 1px solid rgba(100, 118, 148, 0.18);
  border-radius: 18px;
  background: rgba(16, 22, 34, 0.58);
  overflow: hidden;
}

[data-testid="stExpander"] details {
  border: none !important;
}

[data-testid="stExpander"] summary {
  font-weight: 800 !important;
  color: var(--text) !important;
}

.stButton > button,
.stDownloadButton > button {
  min-height: 2.92rem;
  border-radius: 14px;
  border: 1px solid rgba(128, 146, 180, 0.22);
  background: linear-gradient(180deg, rgba(33, 44, 63, 0.98) 0%, rgba(19, 27, 40, 0.98) 100%);
  box-shadow: 0 10px 22px rgba(3, 9, 18, 0.18);
}

.stButton > button:hover,
.stDownloadButton > button:hover {
  transform: translateY(-1px) scale(1.002);
  border-color: rgba(194, 213, 245, 0.3);
  box-shadow: 0 14px 28px rgba(3, 9, 18, 0.22);
}

div[data-testid="stBaseButton-primary"] > button,
.stDownloadButton > button[kind="primary"] {
  background: linear-gradient(135deg, #ff8f74 0%, #ffd88f 100%);
  color: #0d1420;
  font-weight: 900;
  letter-spacing: 0.01em;
  box-shadow: 0 14px 28px rgba(255, 139, 114, 0.24);
}

div[data-testid="stBaseButton-primary"] > button:hover,
.stDownloadButton > button[kind="primary"]:hover {
  box-shadow: 0 18px 34px rgba(255, 139, 114, 0.3);
}

[data-baseweb="input"] > div,
[data-baseweb="textarea"] > div,
[data-baseweb="select"] > div,
.stNumberInput > div > div {
  border-radius: 14px !important;
  background: rgba(11, 17, 27, 0.94) !important;
}

[data-testid="stTabs"] [data-baseweb="tab"] {
  min-height: 2.45rem;
  padding-inline: 0.88rem;
  font-weight: 700;
}

@media (max-width: 920px) {
  .app-hero-title {
    max-width: none;
  }

  .app-kpi-row {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }
}

@media (max-width: 768px) {
  .app-sidebar-brand-row {
    gap: 0.65rem;
  }

  .app-anchor-badge {
    min-width: 1.75rem;
    height: 1.75rem;
    border-radius: 10px;
    font-size: 0.7rem;
  }

  .app-kpi-row {
    grid-template-columns: 1fr;
  }

  .app-card-subtitle {
    max-width: none;
  }
}

@media (max-width: 520px) {
  .app-tier-pill {
    font-size: 0.66rem;
    letter-spacing: 0.1em;
  }

  .app-hero-title {
    font-size: 1.98rem;
  }

  .app-kpi-pill,
  .app-meta-item,
  .app-route-block {
    padding: 0.68rem 0.72rem;
  }
}

/* Elite 95+ product-grade overrides */

.main .block-container {
  max-width: 1440px;
  padding-top: clamp(0.68rem, 0.72vw, 0.9rem);
}

[data-testid="stVerticalBlockBorderWrapper"] {
  background:
    linear-gradient(180deg, rgba(15, 22, 34, 0.985) 0%, rgba(9, 14, 24, 0.985) 100%),
    linear-gradient(135deg, rgba(255, 143, 116, 0.05) 0%, transparent 42%);
  border: 1px solid rgba(128, 147, 181, 0.28) !important;
  border-radius: 24px !important;
  box-shadow:
    0 20px 44px rgba(2, 8, 18, 0.28),
    inset 0 1px 0 rgba(255, 255, 255, 0.03);
}

[data-testid="stVerticalBlockBorderWrapper"] > div {
  padding: clamp(0.84rem, 0.78vw + 0.5rem, 1.08rem);
}

[data-testid="stVerticalBlockBorderWrapper"]:hover {
  box-shadow:
    0 24px 52px rgba(2, 8, 18, 0.3),
    inset 0 1px 0 rgba(255, 255, 255, 0.04);
}

[data-testid="stSidebar"] {
  background:
    radial-gradient(circle at 18% 0%, rgba(255, 147, 117, 0.08) 0%, transparent 28%),
    linear-gradient(180deg, rgba(10, 15, 24, 0.99) 0%, rgba(6, 10, 17, 0.995) 100%);
}

[data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
  border-radius: 20px !important;
  box-shadow: 0 14px 28px rgba(0, 0, 0, 0.18);
}

[data-testid="stSidebar"] .stButton > button {
  min-height: 2.62rem;
  border-radius: 15px;
  font-size: 0.92rem;
  font-weight: 760;
}

.app-hero-shell {
  display: grid;
  gap: 0.88rem;
}

.app-hero-title {
  max-width: 10ch;
  font-size: clamp(3rem, 4.45vw, 4.65rem);
  line-height: 0.92;
  letter-spacing: -0.055em;
  font-weight: 900;
  color: #fff4df;
  text-wrap: balance;
  margin-bottom: 0;
}

.app-hero-lead {
  max-width: 58ch;
  font-size: 1rem;
  line-height: 1.6;
  color: rgba(223, 231, 246, 0.9);
}

.app-hero-note {
  display: inline-flex;
  align-items: center;
  gap: 0.45rem;
  padding: 0.34rem 0.7rem;
  width: fit-content;
  border-radius: 999px;
  background: rgba(255, 255, 255, 0.03);
  border: 1px solid rgba(141, 160, 193, 0.2);
  color: rgba(209, 219, 238, 0.8);
  font-size: 0.8rem;
  font-weight: 700;
}

.app-luxury-rule {
  width: 100%;
  height: 1px;
  background: linear-gradient(90deg, rgba(255, 143, 116, 0.4) 0%, rgba(244, 220, 167, 0.2) 32%, rgba(244, 220, 167, 0.06) 72%, transparent 100%);
  margin: 0.1rem 0 0.15rem;
}

.app-kpi-row {
  gap: 0.72rem;
}

.app-kpi-pill {
  gap: 0.26rem;
  min-height: 4.4rem;
  border-radius: 18px;
  background:
    linear-gradient(180deg, rgba(28, 38, 56, 0.9) 0%, rgba(20, 28, 42, 0.94) 100%);
  border: 1px solid rgba(122, 143, 177, 0.26);
}

.app-kpi-label {
  color: rgba(244, 220, 167, 0.82);
  font-size: 0.68rem;
}

.app-kpi-value {
  font-size: 1.04rem;
  line-height: 1.28;
  color: #fff6e7;
}

.app-card-title {
  font-size: clamp(1.34rem, 1.25vw, 1.72rem);
  letter-spacing: -0.03em;
}

.app-card-subtitle {
  max-width: 62ch;
  color: rgba(207, 218, 237, 0.84);
}

.app-card-subtitle.compact {
  max-width: 58ch;
  font-size: 0.89rem;
}

.app-card-topline {
  margin-bottom: 0.62rem;
}

.app-anchor-badge {
  min-width: 2.05rem;
  height: 2.05rem;
  border-radius: 13px;
  box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.04);
}

.app-tier-pill {
  border-radius: 999px;
  font-size: 0.7rem;
  font-weight: 820;
}

.app-tier-pill.primary {
  box-shadow: 0 8px 20px rgba(255, 139, 114, 0.18);
}

.app-tier-pill.secondary,
.app-tier-pill.tertiary {
  background: rgba(255, 255, 255, 0.035);
  border: 1px solid rgba(244, 220, 167, 0.16);
}

.app-chip-row {
  gap: 0.42rem;
}

.app-chip {
  padding: 0.42rem 0.72rem;
  border-radius: 999px;
  background: rgba(255, 255, 255, 0.04);
  border: 1px solid rgba(244, 220, 167, 0.15);
  color: #f5e3bd;
  font-size: 0.76rem;
  font-weight: 760;
}

.app-meta-grid {
  gap: 0.7rem;
}

.app-meta-item {
  min-height: 5.2rem;
  border-radius: 18px;
  background: linear-gradient(180deg, rgba(30, 40, 59, 0.88) 0%, rgba(24, 32, 47, 0.94) 100%);
  border: 1px solid rgba(120, 139, 173, 0.24);
}

.app-meta-value {
  font-size: 1rem;
  line-height: 1.34;
  color: #fff4df;
}

.app-route-block {
  border-radius: 18px;
  border: 1px solid rgba(122, 143, 177, 0.24);
  background:
    linear-gradient(180deg, rgba(31, 42, 62, 0.88) 0%, rgba(22, 31, 47, 0.94) 100%);
}

.app-route-label {
  font-size: 0.7rem;
}

.app-route-body {
  color: rgba(232, 238, 248, 0.92);
}

.stButton > button,
.stDownloadButton > button {
  min-height: 2.92rem;
  border-radius: 17px;
  border: 1px solid rgba(124, 143, 176, 0.26);
  background:
    linear-gradient(180deg, rgba(38, 49, 70, 0.98) 0%, rgba(22, 30, 45, 0.98) 100%);
  box-shadow:
    0 12px 24px rgba(2, 8, 18, 0.18),
    inset 0 1px 0 rgba(255, 255, 255, 0.04);
  font-weight: 760;
  letter-spacing: 0.01em;
}

.stButton > button:hover,
.stDownloadButton > button:hover {
  transform: translateY(-2px);
  box-shadow:
    0 18px 30px rgba(2, 8, 18, 0.24),
    inset 0 1px 0 rgba(255, 255, 255, 0.05);
}

div[data-testid="stBaseButton-primary"] > button,
.stDownloadButton > button[kind="primary"] {
  background: linear-gradient(135deg, #ff936f 0%, #ffd08a 100%);
  color: #10161f;
  box-shadow:
    0 16px 32px rgba(255, 139, 114, 0.26),
    inset 0 1px 0 rgba(255, 255, 255, 0.18);
}

div[data-testid="stBaseButton-primary"] > button:hover,
.stDownloadButton > button[kind="primary"]:hover {
  box-shadow:
    0 20px 38px rgba(255, 139, 114, 0.32),
    inset 0 1px 0 rgba(255, 255, 255, 0.2);
}

[data-testid="stExpander"] {
  border-radius: 18px;
  border: 1px solid rgba(124, 143, 176, 0.22);
  background: rgba(255, 255, 255, 0.022);
  overflow: hidden;
}

[data-testid="stExpander"] summary {
  font-weight: 760;
}

@media (max-width: 1100px) {
  .app-hero-title {
    max-width: none;
    font-size: clamp(2.55rem, 6.2vw, 4rem);
  }
}

@media (max-width: 920px) {
  .app-meta-grid {
    grid-template-columns: 1fr 1fr;
  }

  .app-kpi-row {
    grid-template-columns: 1fr 1fr;
  }
}

@media (max-width: 768px) {
  [data-testid="stSidebar"] {
    min-width: 14rem;
    max-width: 14rem;
  }

  .app-hero-title {
    font-size: 2.34rem;
    line-height: 0.96;
  }

  .app-hero-lead {
    font-size: 0.92rem;
  }

  .app-meta-grid,
  .app-kpi-row {
    grid-template-columns: 1fr;
  }
}

@media (max-width: 520px) {
  .app-hero-title {
    font-size: 2.02rem;
  }

  .app-hero-note {
    width: 100%;
    justify-content: center;
    font-size: 0.75rem;
  }

  .app-chip {
    width: 100%;
    justify-content: center;
  }
}

</style>
"""


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", name).strip(" ._")
    return cleaned or "document"


def normalize_package_folder_name(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return "hugyoku_exports"
    if any(sep in raw for sep in ("\\", "/")):
        parts = [part for part in re.split(r"[\\/]+", raw) if part.strip()]
        raw = parts[-1] if parts else raw
    if re.fullmatch(r"[A-Za-z]:", raw):
        raw = "hugyoku_exports"
    return sanitize_filename(raw)


def looks_like_local_path(value: str) -> bool:
    raw = (value or "").strip()
    return "\\" in raw or "/" in raw or bool(re.match(r"^[A-Za-z]:", raw))


def pick_local_folder() -> str | None:
    if tk is None or filedialog is None:
        raise RuntimeError("Local folder picking needs Tkinter support on this machine.")

    root: object | None = None
    try:
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        selected = filedialog.askdirectory(title="Choose Hugyoku save folder", mustexist=False)
        return selected or None
    except Exception as exc:
        raise RuntimeError("Could not open the local folder picker in this environment.") from exc
    finally:
        if root is not None:
            try:
                root.destroy()
            except Exception:
                pass


def open_in_file_manager(path: str) -> None:
    target = str(Path(path))
    if os.name == "nt":
        os.startfile(target)
    elif sys.platform == "darwin":
        subprocess.Popen(["open", target])
    else:
        subprocess.Popen(["xdg-open", target])


def current_package_root_name() -> str:
    export_root = (st.session_state.get("export_root_path") or "").strip()
    if export_root:
        return sanitize_filename(Path(export_root).name)
    return sanitize_filename(st.session_state.get("main_folder_name") or "hugyoku_exports")


def local_folder_picker_available() -> bool:
    return tk is not None and filedialog is not None


def save_destination_mode() -> str:
    mode = str(st.session_state.get("save_destination_mode") or "browser").strip().lower()
    return mode if mode in {"browser", "local"} else "browser"


def local_save_active() -> bool:
    return save_destination_mode() == "local" and bool((st.session_state.get("export_root_path") or "").strip())


def browser_export_label() -> str:
    return "Browser-managed download (.docx file)"


def today_string() -> str:
    now = datetime.now()
    return f"{now.strftime('%B')} {now.day}, {now.year}"


def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text))


def collapse_paragraph_lines(lines: list[str]) -> str:
    return " ".join(line.strip() for line in lines if line.strip()).strip()


def split_text_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n\s*\n", text.strip()) if block.strip()]


def strip_heading_markers(text: str) -> str:
    cleaned = re.sub(r"^#+\s*", "", text.strip())
    return cleaned.strip().strip(":").strip()


def heading_level_from_line(text: str) -> int:
    stripped = text.lstrip()
    if stripped.startswith("#"):
        return min(max(len(stripped) - len(stripped.lstrip("#")), 1), 4)
    return 2


def normalize_section_label(text: str) -> str:
    cleaned = strip_heading_markers(text)
    cleaned = re.sub(r"^\d+\.\s*", "", cleaned)
    return re.sub(r"\s+", " ", cleaned).lower().strip()


def parse_structured_blocks(text: str) -> list[dict[str, object]]:
    items: list[dict[str, object]] = []
    for block in split_text_blocks(text):
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        first = lines[0].strip()
        normalized_first = normalize_section_label(first)
        is_heading = first.startswith("#") or first.endswith(":") or normalized_first in KNOWN_EXPORT_LABELS
        if is_heading:
            items.append(
                {
                    "type": "section",
                    "heading": strip_heading_markers(first),
                    "level": heading_level_from_line(first),
                    "content": collapse_paragraph_lines(lines[1:]),
                }
            )
        else:
            items.append({"type": "paragraph", "content": collapse_paragraph_lines(lines)})
    return items


def add_body_paragraph(document: Document, text: str, indent_first_line: bool = False, italic: bool = False) -> None:
    if not text.strip():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(8)
    if indent_first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
    run = paragraph.add_run(text.strip())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = italic


def render_docx_bytes(
    title: str,
    body: str,
    category: str = "generic",
    metadata_lines: list[str] | None = None,
    output_options: dict[str, bool] | None = None,
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    metadata_lines = metadata_lines or []
    output_options = output_options or {}
    export_title = title.strip() or "Untitled Document"

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(export_title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(48, 87, 138)

    if metadata_lines:
        meta_paragraph = document.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_paragraph.paragraph_format.space_after = Pt(10)
        meta_run = meta_paragraph.add_run(" | ".join(metadata_lines))
        meta_run.italic = True
        meta_run.font.name = "Times New Roman"
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(88, 88, 88)

    parsed_blocks = parse_structured_blocks(body)
    essay_subtitle: str | None = None
    essay_tip: str | None = None
    body_items: list[dict[str, object]] = []

    for item in parsed_blocks:
        if item["type"] == "section":
            label = normalize_section_label(str(item["heading"]))
            content = str(item.get("content", "")).strip()
            if category == "essay":
                if label in {"heading suggestion", "title suggestion", "suggested heading"}:
                    if content:
                        essay_subtitle = content
                    continue
                if label in {"essay body", "essay"}:
                    if content:
                        body_items.append({"type": "paragraph", "content": content})
                    continue
                if label in {"self-check tip", "quick self-check tip", "closing self-check tip"}:
                    if content:
                        essay_tip = content
                    continue
            body_items.append(item)
        else:
            body_items.append(item)

    if category == "essay" and output_options.get("essay_include_heading", True) and essay_subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.paragraph_format.space_after = Pt(12)
        subtitle_run = subtitle_paragraph.add_run(essay_subtitle)
        subtitle_run.italic = True
        subtitle_run.font.name = "Times New Roman"
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    for item in body_items:
        if item["type"] == "section":
            heading_text = str(item["heading"]).strip().rstrip(":")
            if heading_text:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(heading_text)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(60, 76, 104)
            content = str(item.get("content", "")).strip()
            if content:
                add_body_paragraph(document, content, indent_first_line=(category == "essay"))
        else:
            add_body_paragraph(document, str(item["content"]), indent_first_line=(category == "essay"))

    if category == "essay" and output_options.get("essay_include_tip", True) and essay_tip:
        tip_heading = document.add_paragraph()
        tip_heading.paragraph_format.space_before = Pt(10)
        tip_heading.paragraph_format.space_after = Pt(2)
        tip_run = tip_heading.add_run("Self-Check Tip")
        tip_run.bold = True
        tip_run.font.name = "Times New Roman"
        tip_run.font.size = Pt(13)
        tip_run.font.color.rgb = RGBColor(60, 76, 104)
        add_body_paragraph(document, essay_tip, italic=True)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def extract_completion_text(message: object) -> str | None:
    content = getattr(message, "content", None)
    if isinstance(content, str):
        return content.strip() or None
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, str) and item.strip():
                parts.append(item.strip())
            elif isinstance(item, dict):
                text = item.get("text") or item.get("content")
                if isinstance(text, str) and text.strip():
                    parts.append(text.strip())
            else:
                text = getattr(item, "text", None) or getattr(item, "content", None)
                if isinstance(text, str) and text.strip():
                    parts.append(text.strip())
        joined = "\n\n".join(parts).strip()
        return joined or None
    return None


def read_secret(name: str) -> str:
    if ENV_PATH.exists():
        load_dotenv(ENV_PATH)
    try:
        secret_value = st.secrets.get(name)
    except Exception:
        secret_value = None
    if isinstance(secret_value, str) and secret_value.strip():
        return secret_value.strip()
    return os.getenv(name, "").strip()


def load_client() -> tuple[object | None, str | None, str | None]:
    token = read_secret("HF_TOKEN")
    model = read_secret("HF_MODEL")

    if InferenceClient is None:
        return None, None, "Install the packages in requirements.txt to enable AI features."
    if not token or not model:
        return None, None, "Add HF_TOKEN and HF_MODEL to Streamlit secrets or local .env to enable AI features."

    return InferenceClient(api_key=token), model, None


def generate_text(prompt: str, label: str, client: object | None = None, model: str | None = None) -> str:
    local_client = client
    local_model = model
    if local_client is None or local_model is None:
        local_client, local_model, error = load_client()
        if error:
            raise RuntimeError(error)

    try:
        completion = local_client.chat.completions.create(
            model=local_model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1400,
            temperature=0.45,
        )
    except Exception as exc:
        raise RuntimeError(f"Could not generate {label}: {exc}") from exc

    message = completion.choices[0].message
    result = extract_completion_text(message)
    if not result:
        raise RuntimeError("The AI model returned an empty response.")
    return result


def read_docx_text(file_obj: io.BytesIO) -> str:
    document = Document(file_obj)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs).strip()


def read_pdf_text(file_obj: io.BytesIO) -> str:
    if PdfReader is None:
        raise RuntimeError("PDF reading needs the pypdf package. Install requirements.txt first.")

    reader = PdfReader(file_obj)
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    joined = "\n\n".join(parts).strip()
    if not joined:
        raise RuntimeError("The PDF did not contain readable text.")
    return joined


def read_text_file(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


def read_uploaded_document(uploaded_file: object) -> str:
    name = getattr(uploaded_file, "name", "uploaded.txt")
    suffix = Path(name).suffix.lower()
    data = uploaded_file.getvalue()
    if suffix == ".docx":
        return read_docx_text(io.BytesIO(data))
    if suffix == ".pdf":
        return read_pdf_text(io.BytesIO(data))
    if suffix in {".txt", ".md"}:
        return read_text_file(data)
    raise RuntimeError("Unsupported file type. Use .docx, .pdf, .txt, or .md.")

def ensure_state() -> None:
    for key, value in STATE_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value


def go(page: str) -> None:
    st.session_state.active_page = page
    st.rerun()


def save_profile() -> None:
    st.session_state.saved_name = st.session_state.profile_name_input.strip()
    st.session_state.saved_include_date = bool(st.session_state.profile_include_date_input)
    mode = str(st.session_state.profile_save_destination_mode_input or "browser").strip().lower()
    if mode not in {"browser", "local"}:
        mode = "browser"
    st.session_state.save_destination_mode = mode
    selected_path = (st.session_state.profile_export_root_path_input or "").strip()
    if mode == "local" and selected_path:
        normalized_folder = sanitize_filename(Path(selected_path).name)
        st.session_state.export_root_path = selected_path
    else:
        normalized_folder = normalize_package_folder_name(st.session_state.profile_main_folder_input)
        st.session_state.export_root_path = ""
    st.session_state.main_folder_name = normalized_folder
    st.session_state.profile_main_folder_input = normalized_folder
    st.session_state.output_include_name = bool(st.session_state.profile_output_include_name_input)
    st.session_state.output_include_date = bool(st.session_state.profile_output_include_date_input)
    st.session_state.essay_include_heading = bool(st.session_state.profile_essay_include_heading_input)
    st.session_state.essay_include_tip = bool(st.session_state.profile_essay_include_tip_input)


def clear_profile() -> None:
    st.session_state.saved_name = ""
    st.session_state.saved_include_date = False
    st.session_state.save_destination_mode = "browser"
    st.session_state.main_folder_name = "hugyoku_exports"
    st.session_state.export_root_path = ""
    st.session_state.output_include_name = False
    st.session_state.output_include_date = False
    st.session_state.essay_include_heading = True
    st.session_state.essay_include_tip = True
    st.session_state.profile_name_input = ""
    st.session_state.profile_include_date_input = False
    st.session_state.profile_save_destination_mode_input = "browser"
    st.session_state.profile_main_folder_input = "hugyoku_exports"
    st.session_state.profile_export_root_path_input = ""
    st.session_state.profile_output_include_name_input = False
    st.session_state.profile_output_include_date_input = False
    st.session_state.profile_essay_include_heading_input = True
    st.session_state.profile_essay_include_tip_input = True


def identity_block() -> str:
    lines: list[str] = []
    if st.session_state.saved_name:
        lines.append(f"Name: {st.session_state.saved_name}")
    if st.session_state.saved_include_date:
        lines.append(f"Date: {today_string()}")
    return "\n".join(lines)


def current_output_settings() -> dict[str, bool]:
    return {
        "include_name": bool(st.session_state.output_include_name),
        "include_date": bool(st.session_state.output_include_date),
        "essay_include_heading": bool(st.session_state.essay_include_heading),
        "essay_include_tip": bool(st.session_state.essay_include_tip),
    }


def export_metadata_lines(category: str = "generic", name_override: str | None = None) -> list[str]:
    lines: list[str] = []
    selected_name = st.session_state.saved_name
    if category == "essay" and name_override and name_override.strip():
        selected_name = name_override.strip()
    if st.session_state.output_include_name and selected_name:
        lines.append(selected_name)
    if st.session_state.output_include_date:
        lines.append(today_string())
    return lines


def folder_path_lines() -> dict[str, str]:
    export_root = (st.session_state.export_root_path or "").strip() if save_destination_mode() == "local" else ""
    paths = {"main": export_root or browser_export_label()}
    for category, folder in TOOL_FOLDERS.items():
        if export_root:
            paths[category] = str(Path(export_root) / folder)
        else:
            paths[category] = browser_export_label()
    return paths


def build_export_document(
    title: str,
    body: str,
    default_name: str,
    category: str,
    name_override: str | None = None,
) -> tuple[bytes, str, str | None]:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = sanitize_filename(default_name)
    docx_filename = f"{stem}_{stamp}.docx"
    internal_folder = TOOL_FOLDERS.get(category, category)
    docx_bytes = render_docx_bytes(
        title,
        body.strip(),
        category=category,
        metadata_lines=export_metadata_lines(category=category, name_override=name_override),
        output_options=current_output_settings(),
    )

    export_root = (st.session_state.export_root_path or "").strip() if save_destination_mode() == "local" else ""
    local_docx_path = str(Path(export_root) / internal_folder / docx_filename) if export_root else None
    return docx_bytes, docx_filename, local_docx_path


def save_local_export_file(docx_bytes: bytes, local_docx_path: str) -> str:
    target_path = Path(local_docx_path)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_bytes(docx_bytes)
    return str(target_path)


def apply_selected_export_folder(selected_folder: str | None) -> None:
    if selected_folder:
        normalized_folder = sanitize_filename(Path(selected_folder).name)
        st.session_state.profile_export_root_path_input = selected_folder
        st.session_state.export_root_path = selected_folder
        st.session_state.save_destination_mode = "local"
        st.session_state.profile_save_destination_mode_input = "local"
        st.session_state.main_folder_name = normalized_folder
        st.session_state.profile_main_folder_input = normalized_folder
    else:
        st.session_state.export_root_path = ""
        st.session_state.save_destination_mode = "browser"
        st.session_state.profile_save_destination_mode_input = "browser"
        st.session_state.main_folder_name = "hugyoku_exports"
        st.session_state.profile_main_folder_input = "hugyoku_exports"


def queue_reset(action: str, notice: str = "", level: str = "success") -> None:
    st.session_state.pending_reset_action = action
    st.session_state.pending_reset_notice = notice
    st.session_state.pending_reset_level = level


def queue_export_root_selection(selected_folder: str | None) -> None:
    st.session_state.pending_export_root_selection = "__CLEAR__" if selected_folder is None else selected_folder


def clear_quiz_workspace() -> None:
    st.session_state.quiz_upload_name = "No file loaded yet"
    st.session_state.quiz_source_text = ""
    st.session_state.quiz_summary = ""
    st.session_state.quiz_prompt = ""
    st.session_state.quiz_response = ""
    st.session_state.quiz_mode = "complete"


def clear_assignment_workspace() -> None:
    st.session_state.assignment_upload_name = "No assignment file loaded yet"
    st.session_state.assignment_source_text = ""
    st.session_state.assignment_summary = ""
    st.session_state.assignment_prompt = ""
    st.session_state.assignment_response = ""
    st.session_state.assignment_mode = "guided"


def clear_essay_form() -> None:
    st.session_state.essay_title = ""
    st.session_state.essay_prompt = ""
    st.session_state.essay_word_count = 500
    st.session_state.essay_tagalog = False
    st.session_state.essay_english = True
    st.session_state.essay_specific_name = ""


def clear_essay_result() -> None:
    st.session_state.essay_response = ""


def clear_quiz_result() -> None:
    st.session_state.quiz_response = ""


def clear_assignment_result() -> None:
    st.session_state.assignment_response = ""


def clear_activity_form() -> None:
    st.session_state.activity_title = ""
    st.session_state.activity_type = "Worksheet"
    st.session_state.activity_level = ""
    st.session_state.activity_prompt = ""


def clear_activity_result() -> None:
    st.session_state.activity_response = ""


def clear_document_form() -> None:
    st.session_state.document_title = ""
    st.session_state.document_type = "Study Handout"
    st.session_state.document_audience = ""
    st.session_state.document_prompt = ""


def clear_document_result() -> None:
    st.session_state.document_response = ""


def clear_codefix_form() -> None:
    st.session_state.codefix_title = ""
    st.session_state.codefix_language = "Python"
    st.session_state.codefix_stack_profile = "Python"
    st.session_state.codefix_custom_stack = ""
    st.session_state.codefix_error = ""
    st.session_state.codefix_source = ""
    st.session_state.codefix_expectation = ""


def clear_codefix_result() -> None:
    st.session_state.codefix_response = ""


def clear_login_form() -> None:
    st.session_state.login_username_input = ""
    st.session_state.login_password_input = ""


def clear_bootstrap_form() -> None:
    st.session_state.bootstrap_username_input = ""
    st.session_state.bootstrap_display_name_input = ""
    st.session_state.bootstrap_password_input = ""
    st.session_state.bootstrap_confirm_password_input = ""


def clear_admin_create_user_form() -> None:
    st.session_state.admin_new_username = ""
    st.session_state.admin_new_display_name = ""
    st.session_state.admin_new_password = ""
    st.session_state.admin_new_role = "member"


def clear_admin_password_reset() -> None:
    st.session_state.admin_reset_password = ""


RESET_ACTIONS: dict[str, object] = {
    "clear_profile": clear_profile,
    "clear_quiz_workspace": clear_quiz_workspace,
    "clear_quiz_result": clear_quiz_result,
    "clear_assignment_workspace": clear_assignment_workspace,
    "clear_assignment_result": clear_assignment_result,
    "clear_essay_workspace": lambda: (clear_essay_form(), clear_essay_result()),
    "clear_essay_result": clear_essay_result,
    "clear_activity_workspace": lambda: (clear_activity_form(), clear_activity_result()),
    "clear_activity_result": clear_activity_result,
    "clear_document_workspace": lambda: (clear_document_form(), clear_document_result()),
    "clear_document_result": clear_document_result,
    "clear_codefix_workspace": lambda: (clear_codefix_form(), clear_codefix_result()),
    "clear_codefix_result": clear_codefix_result,
    "clear_login_form": clear_login_form,
    "clear_bootstrap_form": clear_bootstrap_form,
    "clear_admin_create_user_form": clear_admin_create_user_form,
    "clear_admin_password_reset": clear_admin_password_reset,
}


def apply_pending_state_actions() -> None:
    pending_folder = st.session_state.get("pending_export_root_selection")
    if pending_folder is not None:
        apply_selected_export_folder(None if pending_folder == "__CLEAR__" else str(pending_folder))
        st.session_state.pending_export_root_selection = None

    action_name = (st.session_state.get("pending_reset_action") or "").strip()
    if action_name:
        reset_action = RESET_ACTIONS.get(action_name)
        if callable(reset_action):
            reset_action()
        notice = st.session_state.get("pending_reset_notice", "").strip()
        level = st.session_state.get("pending_reset_level", "success")
        if notice:
            st.session_state.flash_message = notice
            st.session_state.flash_level = level
        st.session_state.pending_reset_action = ""
        st.session_state.pending_reset_notice = ""
        st.session_state.pending_reset_level = "success"


def render_flash_message() -> None:
    message = (st.session_state.get("flash_message") or "").strip()
    if not message:
        return
    level = (st.session_state.get("flash_level") or "success").strip().lower()
    if level == "error":
        st.error(message)
    elif level == "warning":
        st.warning(message)
    elif level == "info":
        st.info(message)
    else:
        st.success(message)
    st.session_state.flash_message = ""
    st.session_state.flash_level = "success"


def run_generation(prompt: str, label: str) -> str | None:
    result, _used_model = run_generation_with_details(prompt, label)
    return result


def essay_language() -> str | None:
    tagalog = bool(st.session_state.essay_tagalog)
    english = bool(st.session_state.essay_english)
    if tagalog and english:
        return "Taglish"
    if tagalog:
        return "Tagalog"
    if english:
        return "English"
    return None


def html_text(value: object) -> str:
    if value is None:
        return ""
    return escape(str(value)).replace("\n", "<br>")


# Shared card helpers keep the layout and typography consistent across pages.
def render_card_header(
    title: str,
    subtitle: str,
    kicker: str | None = None,
    *,
    anchor: str | None = None,
    tier: str = "secondary",
    compact: bool = False,
) -> None:
    kicker_text = (kicker or "").strip()
    anchor_html = f"<span class='app-anchor-badge'>{html_text(anchor)}</span>" if anchor else ""
    tier_class = tier if tier in {"primary", "secondary", "tertiary"} else "secondary"
    tier_html = f"<span class='app-tier-pill {tier_class}'>{html_text(kicker_text)}</span>" if kicker_text else ""
    topline_html = f"<div class='app-card-topline'>{anchor_html}{tier_html}</div>" if anchor_html or tier_html else ""
    subtitle_class = "app-card-subtitle compact" if compact else "app-card-subtitle"
    st.markdown(
        f"""
        <div class="app-card-header">
          {topline_html}
          <div class="app-card-title">{html_text(title)}</div>
          <div class="{subtitle_class}">{html_text(subtitle)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_tag_row(tags: list[str]) -> None:
    if not tags:
        return
    chip_html = "".join(f"<span class='app-chip'>{html_text(tag)}</span>" for tag in tags)
    st.markdown(f"<div class='app-chip-row'>{chip_html}</div>", unsafe_allow_html=True)


def render_route_block(title: str, body: str) -> None:
    st.markdown(
        f"""
        <div class="app-route-block">
          <div class="app-route-label">{html_text(title)}</div>
          <div class="app-route-body">{html_text(body)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_meta_grid(items: list[tuple[str, str]]) -> None:
    if not items:
        return
    cards_html = "".join(
        f"<div class='app-meta-item'><div class='app-meta-label'>{html_text(label)}</div><div class='app-meta-value'>{html_text(value)}</div></div>"
        for label, value in items
    )
    st.markdown(f"<div class='app-meta-grid'>{cards_html}</div>", unsafe_allow_html=True)


def render_kpi_row(items: list[tuple[str, str]]) -> None:
    if not items:
        return
    pills_html = "".join(
        f"<div class='app-kpi-pill'><span class='app-kpi-label'>{html_text(label)}</span><span class='app-kpi-value'>{html_text(value)}</span></div>"
        for label, value in items
    )
    st.markdown(f"<div class='app-kpi-row'>{pills_html}</div>", unsafe_allow_html=True)


def render_preview_panel(
    title: str,
    subtitle: str,
    field_label: str,
    value: str,
    *,
    height: int = 360,
    empty_title: str = "Preview will appear here",
    empty_body: str = "Use the workspace on the left to generate a result. The preview and export actions will appear once there is real content to review.",
    anchor: str = "PV",
    tier: str = "tertiary",
) -> bool:
    render_card_header(title, subtitle, "Preview", anchor=anchor, tier=tier, compact=True)
    content = (value or "").strip()
    if content:
        st.text_area(field_label, value=value, height=height, disabled=True)
        return True
    render_route_block(empty_title, empty_body)
    return False


def render_header(ai_ready: bool, model_label: str, ai_message: str) -> None:
    title = f"{st.session_state.saved_name or 'Guest'}'s Premium Study Desk"
    subtitle = (
        "A cleaner dashboard for identity settings, grouped academic workspaces, "
        "and a separate developer hub for code-fixing support."
    )
    status_text = "AI ready for academics" if ai_ready else ai_message
    chip_class = "ready" if ai_ready else ("waiting" if "Add HF_TOKEN" in ai_message or "Streamlit secrets" in ai_message else "offline")

    left, right = st.columns([2.35, 1], gap="large")
    with left:
        with st.container(border=True):
            st.markdown(
                f"""
                <div class="app-kicker">Hugyoku Workspace</div>
                <div class="app-hero-title">{html_text(title)}</div>
                <div class="app-card-subtitle">{html_text(subtitle)}</div>
                """,
                unsafe_allow_html=True,
            )
            render_tag_row(
                [
                    "Dashboard Identity",
                    "Academic Solvers",
                    "Document Drafting",
                    "Code Fix Support",
                ]
            )
    with right:
        with st.container(border=True):
            st.markdown(
                f"<div class='app-status-pill {chip_class}'>{html_text(status_text)}</div>",
                unsafe_allow_html=True,
            )
            render_meta_grid(
                [
                    ("Saved name", st.session_state.saved_name or "No saved name yet"),
                    ("Date stamp", today_string() if st.session_state.saved_include_date else "Date today disabled"),
                    ("Active model", model_label),
                ]
            )
            if st.button("Refresh AI Status", key="refresh_ai_status", use_container_width=True):
                st.rerun()


def render_sidebar(ai_ready: bool, model_label: str, ai_message: str) -> None:
    nav_items = [
        ("dashboard", "01 Dashboard"),
        ("academics", "02 Academics"),
        ("developer", "03 Developer"),
    ]
    academic_pages = {"academics", "quiz", "assignment", "essay", "activity", "document"}
    developer_pages = {"developer", "codefix"}

    with st.sidebar:
        st.markdown("<div class='app-sidebar-brand'>Hugyoku</div>", unsafe_allow_html=True)
        st.markdown(
            "<div class='app-sidebar-copy'>Premium study and utility desk for academic generators, solvers, and code-fixing support.</div>",
            unsafe_allow_html=True,
        )
        st.caption("The sidebar collapses automatically on smaller screens.")

        for page, label in nav_items:
            active = (
                page == "dashboard" and st.session_state.active_page == "dashboard"
            ) or (
                page == "academics" and st.session_state.active_page in academic_pages
            ) or (
                page == "developer" and st.session_state.active_page in developer_pages
            )
            if st.button(
                label,
                key=f"nav_{page}",
                type="primary" if active else "secondary",
                use_container_width=True,
            ):
                go(page)

        with st.container(border=True):
            st.markdown("**Educational use only**")
            st.caption(
                "Use these flows for studying, drafting, and understanding tasks from your files. "
                "Review outputs before submitting anything."
            )

        with st.container(border=True):
            st.markdown("**AI Status**")
            st.write("Ready" if ai_ready else ai_message)
            st.caption(model_label)


def render_page_intro(page_key: str) -> None:
    details = PAGE_DETAILS[page_key]
    intro_tags = {
        "dashboard": ["Profile settings", "Responsive cards", "Local or browser export"],
        "academics": ["Quiz", "Assignment", "Essay", "Activity", "Document"],
        "developer": ["Code fix", "Root cause", "Debug notes"],
    }
    with st.container(border=True):
        render_card_header(details["title"], details["subtitle"], "Section")
        render_tag_row(intro_tags.get(page_key, []))


def render_tool_hub_card(title: str, subtitle: str, folder_text: str, button_label: str, target: str, primary: bool = False) -> None:
    with st.container(border=True):
        render_card_header(title, subtitle, "Tool Workspace")
        render_route_block("Save route", folder_text)
        if st.button(
            button_label,
            key=f"open_{target}",
            use_container_width=True,
            type="primary" if primary else "secondary",
        ):
            go(target)


def render_back_button(target: str, label: str) -> None:
    if st.button(label, key=f"back_{target}_{st.session_state.active_page}", use_container_width=True):
        go(target)


def render_workspace_header(page_key: str, back_target: str, back_label: str) -> None:
    intro_col, action_col = st.columns([2.2, 1], gap="large")
    with intro_col:
        with st.container(border=True):
            render_card_header(PAGE_DETAILS[page_key]["title"], PAGE_DETAILS[page_key]["subtitle"], "Workspace")
    with action_col:
        with st.container(border=True):
            render_route_block("Save route", folder_path_lines()[page_key])
            render_back_button(back_target, back_label)


def render_download_button(title: str, body: str, default_name: str, category: str, reset_action: str, name_override: str | None = None) -> None:
    docx_bytes, docx_filename, local_docx_path = build_export_document(
        title,
        body,
        default_name,
        category,
        name_override=name_override,
    )
    if local_docx_path:
        render_route_block("Local save route", local_docx_path)
        local_col, open_col = st.columns(2, gap="small")
        if local_col.button(
            "Save To Selected Folder",
            key=f"save_local_{category}_{sanitize_filename(default_name)}",
            use_container_width=True,
            type="primary",
        ):
            try:
                saved_path = save_local_export_file(docx_bytes, local_docx_path)
                queue_reset(reset_action, f"Saved to {saved_path}")
                st.rerun()
            except Exception as exc:
                st.error(f"Could not save to the selected folder: {exc}")
        if open_col.button(
            "Open Selected Folder",
            key=f"open_local_{category}_{sanitize_filename(default_name)}",
            use_container_width=True,
        ):
            try:
                target_folder = Path(local_docx_path).parent
                target_folder.mkdir(parents=True, exist_ok=True)
                open_in_file_manager(str(target_folder))
            except Exception as exc:
                st.error(f"Could not open the selected folder: {exc}")
    else:
        render_route_block("Browser download", f"{docx_filename}\nSaved directly by your browser.")
    st.download_button(
        "Download Word File",
        data=docx_bytes,
        file_name=docx_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_{category}_{sanitize_filename(default_name)}",
        use_container_width=True,
        on_click=queue_reset,
        args=(reset_action, f"Prepared download: {docx_filename}", "info"),
        type="secondary" if local_docx_path else "primary",
    )


def render_dashboard() -> None:
    render_page_intro("dashboard")

    col_a, col_b = st.columns([1.18, 0.82], gap="large")
    with col_a:
        with st.container(border=True):
            render_card_header(
                "Profile Section",
                "Enter your name, decide whether to stamp today's date, then choose one save folder for Hugyoku exports.",
                "Dashboard Identity",
            )
            st.text_input("Enter your name", key="profile_name_input", placeholder="Name to attach to exports")
            st.checkbox("Add date today", key="profile_include_date_input")

            folder_col, pick_col, clear_col = st.columns([1.7, 0.9, 0.8], gap="small")
            folder_col.text_input(
                "Selected save folder",
                value=st.session_state.profile_export_root_path_input or "",
                disabled=True,
                placeholder="No local folder selected yet",
            )
            if pick_col.button("Choose Folder", key="choose_export_folder", use_container_width=True):
                try:
                    selected_folder = pick_local_folder()
                    if selected_folder:
                        queue_export_root_selection(selected_folder)
                        st.rerun()
                    else:
                        st.info("No folder was selected.")
                except Exception as exc:
                    st.warning(str(exc))
            if clear_col.button("Clear Folder", key="clear_export_folder", use_container_width=True):
                queue_export_root_selection(None)
                st.rerun()

            if st.session_state.profile_export_root_path_input.strip():
                st.caption("Local exports save into this folder. Tool subfolders are created automatically when needed.")
            else:
                st.caption("No local save folder selected. Browser downloads will save a single Word file directly.")

            st.markdown("#### Output Options")
            opt_left, opt_right = st.columns(2, gap="small")
            with opt_left:
                st.checkbox("Include saved name in export", key="profile_output_include_name_input")
                st.checkbox("Include date in export", key="profile_output_include_date_input")
            with opt_right:
                st.checkbox("Include essay heading suggestion", key="profile_essay_include_heading_input")
                st.checkbox("Include self-check tip", key="profile_essay_include_tip_input")

            action_left, action_right = st.columns(2, gap="small")
            if action_left.button("Save Profile", use_container_width=True, type="primary"):
                save_profile()
                st.success("Profile saved for this session.")
            if action_right.button("Clear Saved Profile", use_container_width=True):
                queue_reset("clear_profile", "Profile cleared.")
                st.rerun()

    with col_b:
        with st.container(border=True):
            render_card_header(
                "Export Folder Routing",
                "Local exports save into the selected folder with automatic tool subfolders. Browser downloads remain available as single Word files.",
                "Download Structure",
            )
            paths = folder_path_lines()
            render_route_block("Save root folder", paths["main"])
            render_route_block(
                "Academics suite",
                "\n".join(
                    [
                        paths["quiz"],
                        paths["assignment"],
                        paths["essay"],
                        paths["activity"],
                        paths["document"],
                    ]
                ),
            )
            render_route_block("Developer suite", paths["codefix"])

    launch_a, launch_b = st.columns(2, gap="large")
    with launch_a:
        with st.container(border=True):
            render_card_header(
                "Quick Launch",
                "Open the academics hub for school tools or the developer hub for code fixing.",
                "Next Step",
            )
            if st.button("Open Academics Hub", key="open_academics_hub", use_container_width=True, type="primary"):
                go("academics")
            if st.button("Open Developer Hub", key="open_developer_hub", use_container_width=True):
                go("developer")
    with launch_b:
        with st.container(border=True):
            render_card_header(
                "How This Flow Works",
                "The structure stays simple: save your profile, open a hub, choose a tool, then save or download the finished Word file.",
                "Guide",
            )
            st.markdown(
                """
                <ol class="app-guide-list">
                  <li>Dashboard stores your active profile for this session.</li>
                  <li>Academics contains quiz, assignment, essay, activity, and document tools.</li>
                  <li>Developer contains the code error fixer.</li>
                  <li>Every export saves as a Word file, either into your selected folder or through the browser download flow.</li>
                </ol>
                """,
                unsafe_allow_html=True,
            )


def render_academics_hub() -> None:
    render_page_intro("academics")
    paths = folder_path_lines()
    tool_cards = [
        (
            "Quiz Solver Workspace",
            "Read quiz files, summarize the task, and generate guided response support in a dedicated workspace.",
            paths["quiz"],
            "Open Quiz Solver",
            "quiz",
            True,
        ),
        (
            "Assignment Solver Workspace",
            "Analyze assignments from pasted text or uploaded files, then generate a guided draft or response plan.",
            paths["assignment"],
            "Open Assignment Solver",
            "assignment",
            False,
        ),
        (
            "Essay Generator Workspace",
            "Build long-form essay drafts with output options, optional name overrides, and cleaner export formatting.",
            paths["essay"],
            "Open Essay Generator",
            "essay",
            False,
        ),
        (
            "Activity Generator Workspace",
            "Generate worksheets, reflections, drills, or classroom activities from a topic and instruction set.",
            paths["activity"],
            "Open Activity Generator",
            "activity",
            False,
        ),
        (
            "Document Generator Workspace",
            "Create structured school documents like handouts, reports, reviewers, and formal academic materials.",
            paths["document"],
            "Open Document Generator",
            "document",
            False,
        ),
    ]
    for start in range(0, len(tool_cards), 2):
        row = st.columns(2, gap="large")
        for column, card in zip(row, tool_cards[start:start + 2]):
            with column:
                render_tool_hub_card(*card)


def render_developer_hub() -> None:
    render_page_intro("developer")
    paths = folder_path_lines()
    left, right = st.columns(2, gap="large")
    with left:
        render_tool_hub_card(
            "Code Error Fixer Workspace",
            "Paste code, explain the bug or error, and get a cleaner fix plus a short explanation of what changed.",
            paths["codefix"],
            "Open Code Error Fixer",
            "codefix",
            True,
        )
    with right:
        with st.container(border=True):
            render_card_header(
                "What To Paste",
                "The best results come from giving the model the actual error, the relevant code, and what the code should do after the fix.",
                "Developer Guide",
            )
            st.markdown(
                """
                <ol class="app-guide-list">
                  <li>Paste the exact error or symptom first.</li>
                  <li>Include the smallest code snippet that still reproduces the issue.</li>
                  <li>Describe the expected behavior so the fix has a clear target.</li>
                </ol>
                """,
                unsafe_allow_html=True,
            )

def render_quiz_page(ai_ready: bool) -> None:
    render_workspace_header("quiz", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Quiz Intake",
                "Upload a file or paste quiz content directly. The app reads it first before generating support.",
                "Input",
            )
            uploaded = st.file_uploader("Upload quiz file", type=["docx", "pdf", "txt", "md"], key="quiz_upload_widget")
            load_col, clear_col = st.columns(2, gap="small")
            if load_col.button("Load Uploaded File", key="quiz_load_file", use_container_width=True):
                if uploaded is None:
                    st.warning("Upload a file first.")
                else:
                    try:
                        st.session_state.quiz_source_text = read_uploaded_document(uploaded)
                        st.session_state.quiz_upload_name = uploaded.name
                        st.session_state.quiz_summary = ""
                        st.session_state.quiz_response = ""
                        st.session_state.quiz_prompt = ""
                        st.success(f"Loaded {uploaded.name} into the quiz workspace.")
                    except Exception as exc:
                        st.error(str(exc))
            if clear_col.button("Clear Quiz Input", key="quiz_clear_input", use_container_width=True):
                queue_reset("clear_quiz_workspace", "Quiz workspace cleared.")
                st.rerun()

            render_meta_grid(
                [
                    ("Loaded file", st.session_state.quiz_upload_name),
                    (
                        "Content size",
                        f"{count_words(st.session_state.quiz_source_text)} words\n{len(st.session_state.quiz_source_text.strip())} characters",
                    ),
                ]
            )
            st.text_area("Quiz content", key="quiz_source_text", height=320)
            if st.button("Analyze Quiz", key="quiz_analyze", use_container_width=True, disabled=not ai_ready):
                source = st.session_state.quiz_source_text.strip()
                if not source:
                    st.warning("Upload or paste quiz content before analyzing it.")
                else:
                    prompt = (
                        "Analyze the following academic material for educational review.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Source content:\n{source}\n\n"
                        "Return these sections clearly:\n"
                        "1. Summary\n"
                        "2. What task or instructions seem to be required\n"
                        "3. Important topics, clues, or constraints\n"
                        "4. Best next step for the student"
                    )
                    result = run_generation(prompt, "quiz analysis")
                    if result:
                        st.session_state.quiz_summary = result
                        st.success("Quiz analysis complete.")

    with right:
        with st.container(border=True):
            render_card_header(
                "Summary And Task Reading",
                "The assistant summarizes the uploaded content first so you can understand the task before continuing.",
                "Analysis",
            )
            st.text_area("Summary preview", value=st.session_state.quiz_summary, height=220, disabled=True)

        with st.container(border=True):
            render_card_header(
                "Quiz Guidance",
                "Choose whether to solve the whole task or steer the output with a specific prompt.",
                "Output",
            )
            quiz_mode = st.radio(
                "Response mode",
                options=["complete", "specific"],
                index=0 if st.session_state.quiz_mode == "complete" else 1,
                horizontal=True,
                format_func=lambda item: "Do it in full" if item == "complete" else "Use specific prompt",
            )
            st.session_state.quiz_mode = quiz_mode
            st.text_area("Specific prompt", key="quiz_prompt", height=120)
            gen_col, clear_col = st.columns(2, gap="small")
            if gen_col.button("Generate Quiz Guidance", key="quiz_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                source = st.session_state.quiz_source_text.strip()
                summary = st.session_state.quiz_summary.strip()
                custom_prompt = st.session_state.quiz_prompt.strip()
                if not source:
                    st.warning("Upload or paste quiz content before generating a response.")
                elif not summary:
                    st.warning("Run the summary step first so the app can read the task before generating a response.")
                elif st.session_state.quiz_mode == "specific" and not custom_prompt:
                    st.warning("Add a specific prompt or switch to full mode.")
                else:
                    instructions = (
                        "Create a complete educational response based on the content. If there are visible questions, answer them in order and include short explanations. If the file describes a task instead of direct questions, complete the task as fully as possible."
                        if st.session_state.quiz_mode == "complete"
                        else "Follow the user's specific prompt while staying grounded in the uploaded content and analysis.\nSpecific prompt: " + custom_prompt
                    )
                    prompt = (
                        "Use the uploaded quiz/activity content for educational support.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Analysis already prepared:\n{summary}\n\n"
                        f"Source content:\n{source}\n\n"
                        f"Instructions:\n{instructions}\n\n"
                        "Structure the answer with:\n"
                        "1. Task understanding\n"
                        "2. Response\n"
                        "3. Short explanation or rationale\n\n"
                        "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
                    )
                    result = run_generation(prompt, "quiz response")
                    if result:
                        st.session_state.quiz_response = result
                        st.success("Quiz guidance ready.")
            if clear_col.button("Clear Result", key="quiz_clear_result", use_container_width=True):
                queue_reset("clear_quiz_result", "Quiz result cleared.")
                st.rerun()
            st.text_area("Generated response", value=st.session_state.quiz_response, height=300, disabled=True)
            if st.session_state.quiz_response.strip():
                export_title = f"Quiz Support: {Path(st.session_state.quiz_upload_name).stem.replace('_', ' ').title()}" if st.session_state.quiz_upload_name != "No file loaded yet" else "Quiz Support"
                export_name = sanitize_filename(Path(st.session_state.quiz_upload_name).stem if st.session_state.quiz_upload_name != "No file loaded yet" else "quiz_support")
                render_download_button(export_title, st.session_state.quiz_response, export_name, "quiz", "clear_quiz_workspace")


def render_assignment_page(ai_ready: bool) -> None:
    render_workspace_header("assignment", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Assignment Intake",
                "Upload a document or paste the assignment manually, then analyze the task before generating guidance.",
                "Input",
            )
            uploaded = st.file_uploader("Upload assignment file", type=["docx", "pdf", "txt", "md"], key="assignment_upload_widget")
            load_col, clear_col = st.columns(2, gap="small")
            if load_col.button("Load Uploaded File", key="assignment_load_file", use_container_width=True):
                if uploaded is None:
                    st.warning("Upload an assignment file first.")
                else:
                    try:
                        st.session_state.assignment_source_text = read_uploaded_document(uploaded)
                        st.session_state.assignment_upload_name = uploaded.name
                        st.session_state.assignment_summary = ""
                        st.session_state.assignment_response = ""
                        st.success(f"Loaded {uploaded.name} into the assignment workspace.")
                    except Exception as exc:
                        st.error(str(exc))
            if clear_col.button("Clear Assignment Input", key="assignment_clear_input", use_container_width=True):
                queue_reset("clear_assignment_workspace", "Assignment workspace cleared.")
                st.rerun()

            render_meta_grid(
                [
                    ("Loaded file", st.session_state.assignment_upload_name),
                    (
                        "Content size",
                        f"{count_words(st.session_state.assignment_source_text)} words\n{len(st.session_state.assignment_source_text.strip())} characters",
                    ),
                ]
            )
            st.text_area("Assignment content", key="assignment_source_text", height=300)
            assignment_mode = st.radio(
                "Assignment mode",
                options=["guided", "complete"],
                index=0 if st.session_state.assignment_mode == "guided" else 1,
                horizontal=True,
                format_func=lambda item: "Guided response" if item == "guided" else "Full draft",
            )
            st.session_state.assignment_mode = assignment_mode
            st.text_area("Specific prompt (optional)", key="assignment_prompt", height=110)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Analyze Assignment", key="assignment_analyze", use_container_width=True, disabled=not ai_ready):
                source = st.session_state.assignment_source_text.strip()
                if not source:
                    st.warning("Upload or paste assignment content before analyzing it.")
                else:
                    prompt = (
                        "Analyze the following assignment for educational review.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Assignment content:\n{source}\n\n"
                        "Return these sections clearly:\n"
                        "1. Assignment summary\n"
                        "2. What is being asked\n"
                        "3. Important requirements or constraints\n"
                        "4. Best approach for the student"
                    )
                    result = run_generation(prompt, "assignment analysis")
                    if result:
                        st.session_state.assignment_summary = result
                        st.success("Assignment analysis complete.")
            if action_b.button("Generate Assignment Guidance", key="assignment_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                source = st.session_state.assignment_source_text.strip()
                summary = st.session_state.assignment_summary.strip()
                custom_prompt = st.session_state.assignment_prompt.strip()
                if not source:
                    st.warning("Upload or paste assignment content before generating a response.")
                elif not summary:
                    st.warning("Run the assignment analysis first.")
                else:
                    instructions = (
                        "Create a full draft response based on the assignment while keeping it clear, organized, and educational."
                        if st.session_state.assignment_mode == "complete"
                        else "Create a guided response that explains how to approach the assignment and includes a sample draft the student can review."
                    )
                    if custom_prompt:
                        instructions = f"{instructions}\nSpecific prompt: {custom_prompt}"
                    prompt = (
                        "Use the uploaded assignment for educational support.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Assignment analysis:\n{summary}\n\n"
                        f"Assignment content:\n{source}\n\n"
                        f"Instructions:\n{instructions}\n\n"
                        "Structure the answer with:\n"
                        "1. Task understanding\n"
                        "2. Solution plan\n"
                        "3. Sample answer or draft\n"
                        "4. Notes to review\n\n"
                        "Keep the output clean and readable. Avoid markdown symbols like ### unless a heading is truly needed."
                    )
                    result = run_generation(prompt, "assignment response")
                    if result:
                        st.session_state.assignment_response = result
                        st.success("Assignment guidance ready.")

    with right:
        with st.container(border=True):
            render_card_header(
                "Assignment Summary",
                "The uploaded task is summarized first so the response stays grounded in the actual requirements.",
                "Analysis",
            )
            st.text_area("Assignment summary", value=st.session_state.assignment_summary, height=220, disabled=True)

        with st.container(border=True):
            render_card_header(
                "Assignment Response",
                "The guided response or full draft appears here and can be exported directly to Word.",
                "Output",
            )
            st.text_area("Generated response", value=st.session_state.assignment_response, height=280, disabled=True)
            if st.button("Clear Result", key="assignment_clear_result", use_container_width=True):
                queue_reset("clear_assignment_result", "Assignment result cleared.")
                st.rerun()
            if st.session_state.assignment_response.strip():
                export_title = f"Assignment Support: {Path(st.session_state.assignment_upload_name).stem.replace('_', ' ').title()}" if st.session_state.assignment_upload_name != "No assignment file loaded yet" else "Assignment Support"
                export_name = sanitize_filename(Path(st.session_state.assignment_upload_name).stem if st.session_state.assignment_upload_name != "No assignment file loaded yet" else "assignment_support")
                render_download_button(export_title, st.session_state.assignment_response, export_name, "assignment", "clear_assignment_workspace")

def render_essay_page(ai_ready: bool) -> None:
    render_workspace_header("essay", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Essay Builder",
                "Set the title, add an optional custom instruction, choose the target length, and pick English, Tagalog, or both for Taglish.",
                "Compose",
            )
            st.text_input("Essay title", key="essay_title")
            st.text_area("Specific prompt (optional)", key="essay_prompt", height=160)
            st.number_input("Target word count", min_value=100, max_value=3000, step=50, key="essay_word_count")
            lang_a, lang_b = st.columns(2, gap="small")
            lang_a.checkbox("Tagalog", key="essay_tagalog")
            lang_b.checkbox("English", key="essay_english")
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Essay", key="essay_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.essay_title.strip()
                prompt_text = st.session_state.essay_prompt.strip()
                language = essay_language()
                if not title:
                    st.warning("Enter the essay title before generating.")
                elif language is None:
                    st.warning("Select Tagalog, English, or both for Taglish.")
                else:
                    prompt = (
                        "Write a polished educational essay draft.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Essay title: {title}\n"
                        f"Target length: about {int(st.session_state.essay_word_count)} words\n"
                        f"Language: {language}\n"
                        f"Specific prompt: {prompt_text or 'No extra prompt provided'}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Heading Suggestion:\n"
                        "Essay Body:\n"
                        "Self-Check Tip:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "essay draft")
                    if result:
                        st.session_state.essay_response = result
                        st.success("Essay draft ready.")
            if action_b.button("Clear Essay Form", key="essay_clear_form", use_container_width=True):
                queue_reset("clear_essay_workspace", "Essay workspace cleared.")
                st.rerun()

        with st.container(border=True):
            render_card_header(
                "Essay Export Format",
                "Adjust what gets attached to the final essay output. The specific export name can override the saved dashboard name for this essay only.",
                "Output Settings",
            )
            st.text_input("Specific export name (optional)", key="essay_specific_name")
            opt_left, opt_right = st.columns(2, gap="small")
            with opt_left:
                st.checkbox("Include saved or specific name in export", key="profile_output_include_name_input")
                st.checkbox("Include date in export", key="profile_output_include_date_input")
            with opt_right:
                st.checkbox("Include essay heading suggestion", key="profile_essay_include_heading_input")
                st.checkbox("Include self-check tip", key="profile_essay_include_tip_input")
            st.caption("Save Profile from the dashboard if you want these export settings to become the active session defaults.")

    with right:
        with st.container(border=True):
            render_card_header(
                "Essay Preview",
                "The draft appears here once generated. Export it when the structure and tone look right.",
                "Preview",
            )
            st.text_area("Essay output", value=st.session_state.essay_response, height=560, disabled=True)
            if st.button("Clear Result", key="essay_clear_result", use_container_width=True):
                queue_reset("clear_essay_result", "Essay result cleared.")
                st.rerun()
            if st.session_state.essay_response.strip():
                temp_saved = {
                    "output_include_name": st.session_state.output_include_name,
                    "output_include_date": st.session_state.output_include_date,
                    "essay_include_heading": st.session_state.essay_include_heading,
                    "essay_include_tip": st.session_state.essay_include_tip,
                }
                st.session_state.output_include_name = st.session_state.profile_output_include_name_input
                st.session_state.output_include_date = st.session_state.profile_output_include_date_input
                st.session_state.essay_include_heading = st.session_state.profile_essay_include_heading_input
                st.session_state.essay_include_tip = st.session_state.profile_essay_include_tip_input
                try:
                    render_download_button(
                        st.session_state.essay_title.strip() or "Essay Draft",
                        st.session_state.essay_response,
                        sanitize_filename((st.session_state.essay_title.strip() or "essay_draft").lower().replace(" ", "_")),
                        "essay",
                        "clear_essay_workspace",
                        name_override=st.session_state.essay_specific_name.strip(),
                    )
                finally:
                    st.session_state.output_include_name = temp_saved["output_include_name"]
                    st.session_state.output_include_date = temp_saved["output_include_date"]
                    st.session_state.essay_include_heading = temp_saved["essay_include_heading"]
                    st.session_state.essay_include_tip = temp_saved["essay_include_tip"]


def render_activity_page(ai_ready: bool) -> None:
    render_workspace_header("activity", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Activity Builder",
                "Set the topic, activity type, and optional level details, then generate a ready-to-use school activity.",
                "Compose",
            )
            st.text_input("Activity topic or title", key="activity_title")
            st.text_input("Activity type", key="activity_type")
            st.text_input("Level or class (optional)", key="activity_level")
            st.text_area("Specific instructions", key="activity_prompt", height=180)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Activity", key="activity_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.activity_title.strip()
                activity_type = st.session_state.activity_type.strip() or "Activity"
                level = st.session_state.activity_level.strip()
                request = st.session_state.activity_prompt.strip()
                if not title:
                    st.warning("Enter the activity topic or title before generating.")
                else:
                    prompt = (
                        "Create a structured educational activity for offline study use.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Activity title or topic: {title}\n"
                        f"Activity type: {activity_type}\n"
                        f"Level or class: {level or 'Not specified'}\n"
                        f"Specific request: {request or 'No extra request provided'}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Activity Title:\n"
                        "Objective:\n"
                        "Instructions:\n"
                        "Activity Proper:\n"
                        "Answer Guide:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "activity draft")
                    if result:
                        st.session_state.activity_response = result
                        st.success("Activity ready.")
            if action_b.button("Clear Activity Form", key="activity_clear_form", use_container_width=True):
                queue_reset("clear_activity_workspace", "Activity workspace cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Activity Preview",
                "The generated activity appears here in a clean structure you can export directly to Word.",
                "Preview",
            )
            st.text_area("Generated activity", value=st.session_state.activity_response, height=560, disabled=True)
            if st.button("Clear Result", key="activity_clear_result", use_container_width=True):
                queue_reset("clear_activity_result", "Activity result cleared.")
                st.rerun()
            if st.session_state.activity_response.strip():
                render_download_button(
                    st.session_state.activity_title.strip() or "Activity Draft",
                    st.session_state.activity_response,
                    sanitize_filename((st.session_state.activity_title.strip() or "activity_draft").lower().replace(" ", "_")),
                    "activity",
                    "clear_activity_workspace",
                )


def render_document_page(ai_ready: bool) -> None:
    render_workspace_header("document", "academics", "Back To Academics")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Document Builder",
                "Generate a structured school document such as a handout, report, reviewer, or formal academic write-up.",
                "Compose",
            )
            st.text_input("Document title", key="document_title")
            st.text_input("Document type", key="document_type")
            st.text_input("Audience or purpose (optional)", key="document_audience")
            st.text_area("Specific content request", key="document_prompt", height=180)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Document", key="document_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.document_title.strip()
                doc_type = st.session_state.document_type.strip() or "Study Handout"
                audience = st.session_state.document_audience.strip()
                request = st.session_state.document_prompt.strip()
                if not title:
                    st.warning("Enter the document title before generating.")
                elif not request:
                    st.warning("Describe what the document should contain before generating.")
                else:
                    prompt = (
                        "Create a structured educational document.\n"
                        f"Saved dashboard profile:\n{identity_block() or 'No saved profile'}\n\n"
                        f"Document title: {title}\n"
                        f"Document type: {doc_type}\n"
                        f"Audience or purpose: {audience or 'Not specified'}\n"
                        f"Content request: {request}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Document Type:\n"
                        "Purpose:\n"
                        "Main Content:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "document draft")
                    if result:
                        st.session_state.document_response = result
                        st.success("Document ready.")
            if action_b.button("Clear Document Form", key="document_clear_form", use_container_width=True):
                queue_reset("clear_document_workspace", "Document workspace cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Document Preview",
                "The generated document draft appears here and can be exported directly to Word.",
                "Preview",
            )
            st.text_area("Generated document", value=st.session_state.document_response, height=560, disabled=True)
            if st.button("Clear Result", key="document_clear_result", use_container_width=True):
                queue_reset("clear_document_result", "Document result cleared.")
                st.rerun()
            if st.session_state.document_response.strip():
                render_download_button(
                    st.session_state.document_title.strip() or "Document Draft",
                    st.session_state.document_response,
                    sanitize_filename((st.session_state.document_title.strip() or "document_draft").lower().replace(" ", "_")),
                    "document",
                    "clear_document_workspace",
                )


def render_codefix_page(ai_ready: bool) -> None:
    render_workspace_header("codefix", "developer", "Back To Developer")

    left, right = st.columns([1.02, 0.98], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Code Fix Builder",
                "Paste the broken code, add the error or symptom, and explain the expected behavior so the app can generate a cleaner fix.",
                "Compose",
            )
            st.text_input("Issue title (optional)", key="codefix_title")
            st.text_input("Language or stack", key="codefix_language")
            st.text_area("Error message or symptoms", key="codefix_error", height=140)
            st.text_area("Code snippet", key="codefix_source", height=240)
            st.text_area("Expected behavior (optional)", key="codefix_expectation", height=110)
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Generate Code Fix", key="codefix_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                title = st.session_state.codefix_title.strip()
                language = st.session_state.codefix_language.strip() or "Code"
                error_text = st.session_state.codefix_error.strip()
                source = st.session_state.codefix_source.strip()
                expectation = st.session_state.codefix_expectation.strip()
                if not source:
                    st.warning("Paste the code snippet before generating a fix.")
                elif not error_text:
                    st.warning("Add the error message or symptoms before generating a fix.")
                else:
                    prompt = (
                        "Fix the following code issue for study and debugging support.\n\n"
                        f"Language or stack: {language}\n"
                        f"Issue title: {title or 'No title provided'}\n"
                        f"Error message or symptoms:\n{error_text}\n\n"
                        f"Code snippet:\n{source}\n\n"
                        f"Expected behavior:\n{expectation or 'Not provided'}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Issue Summary:\n"
                        "Root Cause:\n"
                        "Fixed Version:\n"
                        "Why It Works:\n"
                        "Next Checks:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "code fix")
                    if result:
                        st.session_state.codefix_response = result
                        st.success("Code fix ready.")
            if action_b.button("Clear Code Fixer", key="codefix_clear_form", use_container_width=True):
                queue_reset("clear_codefix_workspace", "Code fixer workspace cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Code Fix Preview",
                "The cleaned-up fix and explanation appear here. Export it if you want to keep debugging notes in Word format.",
                "Preview",
            )
            st.text_area("Fixed result", value=st.session_state.codefix_response, height=560, disabled=True)
            if st.button("Clear Result", key="codefix_clear_result", use_container_width=True):
                queue_reset("clear_codefix_result", "Code fix result cleared.")
                st.rerun()
            if st.session_state.codefix_response.strip():
                export_title = st.session_state.codefix_title.strip() or f"{st.session_state.codefix_language.strip() or 'Code'} Code Fix"
                render_download_button(
                    export_title,
                    st.session_state.codefix_response,
                    sanitize_filename(export_title.lower().replace(" ", "_")),
                    "codefix",
                    "clear_codefix_workspace",
                )


STATE_DEFAULTS.update(
    {
        "workspaces": {},
        "active_workspace_id": "",
        "workspace_new_name_input": "",
        "workspace_new_description_input": "",
        "workspace_analysis_result": "",
        "hugyoku_task_input": "",
        "hugyoku_attachment_note": "",
        "hugyoku_stage": 1,
        "hugyoku_understanding": "",
        "hugyoku_refinement_prompt": "",
        "hugyoku_refinement_round": 0,
        "hugyoku_output_sections": [],
        "hugyoku_output_title": "",
        "hugyoku_output_format": "docx",
        "hugyoku_output_raw": "",
        "hugyoku_result_prompt": "",
        "hugyoku_generation_note": "",
        "hugyoku_last_bundle": "",
        "hugyoku_last_ocr_status": "",
        "source_lab_question_input": "",
        "source_lab_answer": "",
        "source_lab_analysis": "",
        "reviewer_request_input": "",
        "reviewer_response": "",
        "flashcard_request_input": "",
        "flashcard_response": "",
        "practice_test_request_input": "",
        "practice_test_response": "",
        "answer_checker_question_input": "",
        "answer_checker_reference_input": "",
        "answer_checker_user_answer_input": "",
        "answer_checker_response": "",
        "rubric_title_input": "",
        "rubric_prompt_input": "",
        "rubric_focus_input": ["Clarity", "Structure", "Depth"],
        "rubric_response": "",
        "batch_topic_input": "",
        "batch_request_input": "",
        "batch_response": "",
        "history_entries": [],
        "model_override": "",
        "model_choice_input": "HF Secret Default",
        "model_custom_input": "",
        "is_authenticated": False,
        "auth_user_id": 0,
        "auth_username": "",
        "auth_display_name": "",
        "auth_role": "",
        "auth_permissions": {},
        "login_username_input": "",
        "login_password_input": "",
        "bootstrap_username_input": "",
        "bootstrap_display_name_input": "",
        "bootstrap_password_input": "",
        "bootstrap_confirm_password_input": "",
        "admin_new_username": "",
        "admin_new_display_name": "",
        "admin_new_password": "",
        "admin_new_role": "member",
        "admin_selected_user_id": 0,
        "admin_reset_password": "",
        "admin_filter_role": "All",
        "last_generation_model": "",
        "last_generation_label": "",
        "last_generation_status": "",
        "last_generation_time": "",
        "last_generation_note": "",
        "verification_last_model": "",
        "verification_last_status": "",
        "verification_last_time": "",
        "verification_last_note": "",
        "codegen_title": "",
        "codegen_stack_profile": "Python",
        "codegen_custom_stack": "",
        "codegen_description": "",
        "codegen_expectation": "",
        "codegen_attachment_note": "",
        "codegen_response": "",
        "codefix_stack_profile": "Python",
        "codefix_custom_stack": "",
        "selftest_stack_profile": "Python",
        "selftest_custom_stack": "",
        "selftest_response": "",
        "selftest_last_model": "",
        "compare_model_a": "Active Session Model",
        "compare_model_b": "Qwen/Qwen2.5-Coder-7B-Instruct",
        "compare_custom_model_a": "",
        "compare_custom_model_b": "",
        "compare_prompt": "",
        "compare_output_a": "",
        "compare_output_b": "",
        "compare_used_model_a": "",
        "compare_used_model_b": "",
        "export_template": "Academic Classic",
        "settings_history_limit": 120,
    }
)

TOOL_FOLDERS.update(
    {
        "workspace_analysis": "workspace_analysis",
        "grounded_answer": "source_grounded_answers",
        "documentqa": "document_qna",
        "reviewer": "reviewer_generator",
        "flashcards": "flashcard_generator",
        "practice_test": "practice_test_generator",
        "answer_checker": "answer_checker",
        "rubric": "rubric_writer",
        "batch": "batch_output_generator",
        "codegen": "code_generator",
        "hugyoku": "hugyoku_universal",
    }
)

PAGE_DETAILS.update(
    {
        "hugyoku": {
            "title": "Hugyoku",
            "subtitle": "Universal academic command workflow that reads the task, confirms the AI understanding, loops on corrections, then generates an editable final result and export.",
        },
        "workspaces": {
            "title": "Workspaces",
            "subtitle": "Create project workspaces, collect source files and screenshots, and keep all generated outputs grouped by task.",
        },
        "history": {
            "title": "History",
            "subtitle": "Review recent uploads, generations, and exports across the active session.",
        },
        "settings": {
            "title": "Settings",
            "subtitle": "Control the active AI model, export template, and session-level app behavior.",
        },
        "admin": {
            "title": "Admin",
            "subtitle": "Control user access, roles, and account status for the app before anyone reaches the main workspace.",
        },
    }
)

EXPORT_TEMPLATES = {
    "Academic Classic": {
        "font": "Times New Roman",
        "body_size": 12,
        "title_size": 18,
        "heading_size": 13,
        "title_color": RGBColor(48, 87, 138),
        "heading_color": RGBColor(60, 76, 104),
        "meta_color": RGBColor(88, 88, 88),
        "line_spacing": 1.2,
        "indent_essay": True,
    },
    "Reviewer Notes": {
        "font": "Arial",
        "body_size": 11,
        "title_size": 17,
        "heading_size": 12,
        "title_color": RGBColor(35, 124, 101),
        "heading_color": RGBColor(46, 84, 72),
        "meta_color": RGBColor(98, 98, 98),
        "line_spacing": 1.1,
        "indent_essay": False,
    },
    "Formal Report": {
        "font": "Calibri",
        "body_size": 11,
        "title_size": 19,
        "heading_size": 12,
        "title_color": RGBColor(75, 70, 132),
        "heading_color": RGBColor(67, 67, 67),
        "meta_color": RGBColor(100, 100, 100),
        "line_spacing": 1.15,
        "indent_essay": False,
    },
}


def clone_default_value(value: object) -> object:
    if isinstance(value, dict):
        return value.copy()
    if isinstance(value, list):
        return list(value)
    if isinstance(value, set):
        return set(value)
    return value


def auth_timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def get_auth_connection() -> sqlite3.Connection:
    connection = sqlite3.connect(AUTH_DB_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def default_permissions_for_role(role: str) -> dict[str, bool]:
    return dict(ROLE_PERMISSIONS.get(role, ROLE_PERMISSIONS["member"]))


def encode_permissions(permissions: dict[str, bool]) -> str:
    return json.dumps({key: bool(value) for key, value in permissions.items()}, sort_keys=True)


def decode_permissions(raw_value: str | None, role: str) -> dict[str, bool]:
    base = default_permissions_for_role(role)
    if not raw_value:
        return base
    try:
        loaded = json.loads(raw_value)
    except Exception:
        return base
    if not isinstance(loaded, dict):
        return base
    for key, value in loaded.items():
        base[str(key)] = bool(value)
    return base


def hash_password(password: str, salt: bytes | None = None, iterations: int = 200000) -> str:
    working_salt = salt or os.urandom(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), working_salt, iterations)
    return f"pbkdf2_sha256${iterations}${working_salt.hex()}${digest.hex()}"


def verify_password(password: str, stored_hash: str) -> bool:
    try:
        scheme, iteration_text, salt_hex, digest_hex = stored_hash.split("$", 3)
        if scheme != "pbkdf2_sha256":
            return False
        iterations = int(iteration_text)
        candidate = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), bytes.fromhex(salt_hex), iterations).hex()
        return hmac.compare_digest(candidate, digest_hex)
    except Exception:
        return False


def initialize_auth_storage() -> None:
    with get_auth_connection() as connection:
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                display_name TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL,
                is_active INTEGER NOT NULL DEFAULT 1,
                permissions_json TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                last_login_at TEXT
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS auth_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT,
                event_type TEXT NOT NULL,
                details TEXT NOT NULL,
                created_at TEXT NOT NULL
            )
            """
        )


def log_auth_event(username: str, event_type: str, details: str) -> None:
    with get_auth_connection() as connection:
        connection.execute(
            "INSERT INTO auth_logs (username, event_type, details, created_at) VALUES (?, ?, ?, ?)",
            (username, event_type, details, auth_timestamp()),
        )


def auth_user_count() -> int:
    with get_auth_connection() as connection:
        row = connection.execute("SELECT COUNT(*) AS total FROM users").fetchone()
    return int(row["total"]) if row else 0


def normalize_user_row(row: sqlite3.Row | None) -> dict[str, object] | None:
    if row is None:
        return None
    role = str(row["role"])
    permissions = decode_permissions(row["permissions_json"], role)
    return {
        "id": int(row["id"]),
        "username": str(row["username"]),
        "display_name": str(row["display_name"]),
        "role": role,
        "is_active": bool(row["is_active"]),
        "permissions": permissions,
        "created_at": str(row["created_at"]),
        "updated_at": str(row["updated_at"]),
        "last_login_at": str(row["last_login_at"] or ""),
    }


def get_user_by_username(username: str) -> dict[str, object] | None:
    with get_auth_connection() as connection:
        row = connection.execute("SELECT * FROM users WHERE LOWER(username) = LOWER(?)", (username.strip(),)).fetchone()
    return normalize_user_row(row)


def get_user_with_hash_by_username(username: str) -> tuple[dict[str, object] | None, str]:
    with get_auth_connection() as connection:
        row = connection.execute("SELECT * FROM users WHERE LOWER(username) = LOWER(?)", (username.strip(),)).fetchone()
    if row is None:
        return None, ""
    return normalize_user_row(row), str(row["password_hash"])


def get_user_by_id(user_id: int) -> dict[str, object] | None:
    with get_auth_connection() as connection:
        row = connection.execute("SELECT * FROM users WHERE id = ?", (int(user_id),)).fetchone()
    return normalize_user_row(row)


def list_auth_users() -> list[dict[str, object]]:
    with get_auth_connection() as connection:
        rows = connection.execute("SELECT * FROM users ORDER BY role DESC, username ASC").fetchall()
    return [normalize_user_row(row) for row in rows if normalize_user_row(row) is not None]


def active_admin_count(exclude_user_id: int | None = None) -> int:
    users = list_auth_users()
    count = 0
    for user in users:
        if exclude_user_id is not None and int(user["id"]) == int(exclude_user_id):
            continue
        if bool(user["is_active"]) and str(user["role"]) in {"admin", "super_admin"}:
            count += 1
    return count


def create_auth_user(username: str, display_name: str, password: str, role: str = "member") -> tuple[bool, str]:
    username_clean = username.strip()
    display_clean = display_name.strip() or username_clean
    role_clean = role.strip() if role.strip() in ROLE_PERMISSIONS else "member"
    if not username_clean or not password.strip():
        return False, "Username and password are required."
    if get_user_by_username(username_clean):
        return False, "That username already exists."
    now = auth_timestamp()
    password_hash = hash_password(password.strip())
    permissions_json = encode_permissions(default_permissions_for_role(role_clean))
    with get_auth_connection() as connection:
        connection.execute(
            """
            INSERT INTO users (username, display_name, password_hash, role, is_active, permissions_json, created_at, updated_at)
            VALUES (?, ?, ?, ?, 1, ?, ?, ?)
            """,
            (username_clean, display_clean, password_hash, role_clean, permissions_json, now, now),
        )
    log_auth_event(username_clean, "user_created", f"Created role={role_clean}")
    return True, "User created."


def update_auth_user(user_id: int, display_name: str, role: str, is_active: bool, permissions: dict[str, bool]) -> tuple[bool, str]:
    user = get_user_by_id(user_id)
    if not user:
        return False, "User not found."
    role_clean = role if role in ROLE_PERMISSIONS else user["role"]
    if str(user["role"]) in {"admin", "super_admin"}:
        keeps_admin_access = role_clean in {"admin", "super_admin"} and bool(is_active)
        if not keeps_admin_access and active_admin_count(exclude_user_id=int(user_id)) == 0:
            return False, "At least one active admin account must remain."
    permissions_json = encode_permissions(permissions)
    now = auth_timestamp()
    with get_auth_connection() as connection:
        connection.execute(
            """
            UPDATE users
            SET display_name = ?, role = ?, is_active = ?, permissions_json = ?, updated_at = ?
            WHERE id = ?
            """,
            (display_name.strip() or user["display_name"], role_clean, 1 if is_active else 0, permissions_json, now, int(user_id)),
        )
    log_auth_event(user["username"], "user_updated", f"role={role_clean}, active={is_active}")
    return True, "User updated."


def update_auth_password(user_id: int, new_password: str) -> tuple[bool, str]:
    user = get_user_by_id(user_id)
    if not user:
        return False, "User not found."
    if not new_password.strip():
        return False, "Password cannot be empty."
    with get_auth_connection() as connection:
        connection.execute(
            "UPDATE users SET password_hash = ?, updated_at = ? WHERE id = ?",
            (hash_password(new_password.strip()), auth_timestamp(), int(user_id)),
        )
    log_auth_event(user["username"], "password_reset", "Password updated by admin.")
    return True, "Password updated."


def mark_user_login(username: str) -> None:
    with get_auth_connection() as connection:
        connection.execute(
            "UPDATE users SET last_login_at = ?, updated_at = ? WHERE LOWER(username) = LOWER(?)",
            (auth_timestamp(), auth_timestamp(), username.strip()),
        )


def apply_auth_session(user: dict[str, object]) -> None:
    st.session_state.is_authenticated = True
    st.session_state.auth_user_id = int(user["id"])
    st.session_state.auth_username = str(user["username"])
    st.session_state.auth_display_name = str(user["display_name"])
    st.session_state.auth_role = str(user["role"])
    st.session_state.auth_permissions = dict(user["permissions"])


def logout_current_user() -> None:
    username = st.session_state.get("auth_username", "")
    st.session_state.is_authenticated = False
    st.session_state.auth_user_id = 0
    st.session_state.auth_username = ""
    st.session_state.auth_display_name = ""
    st.session_state.auth_role = ""
    st.session_state.auth_permissions = {}
    st.session_state.active_page = "dashboard"
    if username:
        log_auth_event(str(username), "logout", "User logged out.")


def authenticate_user(username: str, password: str) -> tuple[bool, str]:
    user, stored_hash = get_user_with_hash_by_username(username)
    if not user:
        return False, "Invalid username or password."
    if not user["is_active"]:
        return False, "This account is disabled."
    if not verify_password(password, stored_hash):
        return False, "Invalid username or password."
    apply_auth_session(user)
    mark_user_login(user["username"])
    log_auth_event(user["username"], "login", f"role={user['role']}")
    return True, "Login successful."


def auth_permissions() -> dict[str, bool]:
    raw = st.session_state.get("auth_permissions")
    if isinstance(raw, dict):
        return {str(key): bool(value) for key, value in raw.items()}
    return {}


def can_access_page(page: str) -> bool:
    if not st.session_state.get("is_authenticated", False):
        return False
    page_map = {
        "dashboard": "dashboard",
        "workspaces": "workspaces",
        "academics": "academics",
        "quiz": "academics",
        "assignment": "academics",
        "essay": "academics",
        "activity": "academics",
        "document": "academics",
        "developer": "developer",
        "codefix": "developer",
        "history": "history",
        "settings": "settings",
        "admin": "admin",
    }
    permission_key = page_map.get(page, "dashboard")
    return bool(auth_permissions().get(permission_key, False))


def auth_log_rows(limit: int = 80) -> list[dict[str, str]]:
    with get_auth_connection() as connection:
        rows = connection.execute(
            "SELECT username, event_type, details, created_at FROM auth_logs ORDER BY id DESC LIMIT ?",
            (int(limit),),
        ).fetchall()
    return [
        {
            "username": str(row["username"] or "-"),
            "event_type": str(row["event_type"]),
            "details": str(row["details"]),
            "created_at": str(row["created_at"]),
        }
        for row in rows
    ]


def ensure_state() -> None:
    for key, value in STATE_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = clone_default_value(value)


def current_model_name() -> str:
    override = (st.session_state.get("model_override") or "").strip()
    return override or read_secret("HF_MODEL")


def load_client(requested_model: str | None = None) -> tuple[object | None, str | None, str | None]:
    token = read_secret("HF_TOKEN")
    model = (requested_model or "").strip() or current_model_name()

    if InferenceClient is None:
        return None, None, "Install the packages in requirements.txt to enable AI features."
    if not token or not model:
        return None, None, "Add HF_TOKEN and HF_MODEL to Streamlit secrets or local .env to enable AI features."

    return InferenceClient(api_key=token), model, None


def workspace_timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def short_preview(text: str, limit: int = 220) -> str:
    cleaned = re.sub(r"\s+", " ", text.strip())
    if len(cleaned) <= limit:
        return cleaned
    return f"{cleaned[:limit].rstrip()}..."


def new_workspace_record(name: str, description: str = "") -> dict[str, object]:
    slug = sanitize_filename(name.lower().replace(" ", "_"))
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return {
        "id": f"{slug}_{stamp}",
        "name": name.strip() or "General Workspace",
        "description": description.strip(),
        "created_at": workspace_timestamp(),
        "notes": "",
        "files": [],
        "images": [],
        "outputs": [],
    }


def ensure_workspace_bootstrap() -> None:
    workspaces = st.session_state.get("workspaces")
    if not isinstance(workspaces, dict):
        st.session_state.workspaces = {}
    if not st.session_state.workspaces:
        workspace = new_workspace_record("General Workspace", "Default workspace for new uploads and generated outputs.")
        st.session_state.workspaces = {str(workspace["id"]): workspace}
        st.session_state.active_workspace_id = str(workspace["id"])
        return
    if st.session_state.active_workspace_id not in st.session_state.workspaces:
        st.session_state.active_workspace_id = next(iter(st.session_state.workspaces.keys()))


def active_workspace() -> dict[str, object]:
    ensure_workspace_bootstrap()
    return st.session_state.workspaces[st.session_state.active_workspace_id]


def active_workspace_name() -> str:
    return str(active_workspace().get("name", "General Workspace"))


def workspace_option_label(workspace_id: str) -> str:
    workspace = st.session_state.workspaces[workspace_id]
    file_count = len(workspace.get("files", []))
    image_count = len(workspace.get("images", []))
    return f"{workspace['name']} ({file_count} files, {image_count} images)"


def create_workspace(name: str, description: str) -> bool:
    normalized_name = name.strip()
    if not normalized_name:
        return False
    workspace = new_workspace_record(normalized_name, description)
    st.session_state.workspaces[str(workspace["id"])] = workspace
    st.session_state.active_workspace_id = str(workspace["id"])
    append_history_entry(
        "workspace",
        f"Created workspace: {normalized_name}",
        description.strip() or "No description provided.",
        "workspaces",
        normalized_name,
    )
    return True


def delete_active_workspace() -> bool:
    ensure_workspace_bootstrap()
    if len(st.session_state.workspaces) <= 1:
        return False
    workspace = active_workspace()
    removed_name = str(workspace.get("name", "Workspace"))
    del st.session_state.workspaces[st.session_state.active_workspace_id]
    st.session_state.active_workspace_id = next(iter(st.session_state.workspaces.keys()))
    append_history_entry("workspace", f"Deleted workspace: {removed_name}", "Workspace removed from the current session.", "workspaces", removed_name)
    return True


def workspace_word_count(workspace: dict[str, object]) -> int:
    return sum(int(file_entry.get("words", 0)) for file_entry in workspace.get("files", []))


def workspace_source_bundle(workspace: dict[str, object], include_notes: bool = True) -> str:
    parts: list[str] = []
    for file_entry in workspace.get("files", []):
        name = str(file_entry.get("name", "source.txt"))
        text = str(file_entry.get("text", "")).strip()
        if text:
            parts.append(f"[Source: {name}]\n{text}")
    if include_notes:
        notes = str(workspace.get("notes", "")).strip()
        if notes:
            parts.append(f"[Workspace Notes]\n{notes}")
    for image_entry in workspace.get("images", []):
        name = str(image_entry.get("name", "image"))
        caption = str(image_entry.get("caption", "")).strip()
        if caption:
            parts.append(f"[Screenshot Note: {name}]\n{caption}")
        else:
            parts.append(f"[Screenshot: {name}]\nScreenshot uploaded. OCR is not enabled, so use a manual note if you need this image included in source-grounded answers.")
    return "\n\n".join(parts).strip()


def trim_prompt_source(text: str, limit: int = 14000) -> str:
    cleaned = text.strip()
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rstrip() + "\n\n[Source content truncated for prompt safety.]"


def upsert_named_entry(entries: list[dict[str, object]], entry: dict[str, object]) -> list[dict[str, object]]:
    target_name = str(entry.get("name", "")).strip().lower()
    filtered = [item for item in entries if str(item.get("name", "")).strip().lower() != target_name]
    filtered.insert(0, entry)
    return filtered


def add_source_files_to_active_workspace(uploaded_files: list[object]) -> tuple[int, list[str]]:
    workspace = active_workspace()
    added = 0
    failures: list[str] = []
    for uploaded_file in uploaded_files:
        try:
            text = read_uploaded_document(uploaded_file)
            entry = {
                "name": getattr(uploaded_file, "name", "source.txt"),
                "suffix": Path(getattr(uploaded_file, "name", "source.txt")).suffix.lower(),
                "text": text,
                "words": count_words(text),
                "added_at": workspace_timestamp(),
            }
            workspace["files"] = upsert_named_entry(list(workspace.get("files", [])), entry)
            append_history_entry(
                "upload",
                f"Added source file: {entry['name']}",
                f"{entry['words']} words added to {workspace['name']}.",
                "workspaces",
                str(workspace["name"]),
            )
            added += 1
        except Exception as exc:
            failures.append(f"{getattr(uploaded_file, 'name', 'file')}: {exc}")
    return added, failures


def add_images_to_active_workspace(uploaded_images: list[object], caption: str = "") -> int:
    workspace = active_workspace()
    added = 0
    for uploaded_image in uploaded_images:
        entry = {
            "name": getattr(uploaded_image, "name", "image.png"),
            "bytes": uploaded_image.getvalue(),
            "caption": caption.strip(),
            "added_at": workspace_timestamp(),
        }
        workspace["images"] = upsert_named_entry(list(workspace.get("images", [])), entry)
        append_history_entry(
            "upload",
            f"Added screenshot: {entry['name']}",
            "Screenshot attached to the active workspace.",
            "workspaces",
            str(workspace["name"]),
        )
        added += 1
    return added


def append_history_entry(kind: str, title: str, details: str, page: str, workspace_name: str = "", category: str = "") -> None:
    entries = list(st.session_state.get("history_entries", []))
    entries.insert(
        0,
        {
            "timestamp": workspace_timestamp(),
            "kind": kind,
            "title": title,
            "details": details,
            "page": page,
            "workspace": workspace_name or active_workspace_name(),
            "category": category,
        },
    )
    limit = max(int(st.session_state.get("settings_history_limit", 120) or 120), 20)
    st.session_state.history_entries = entries[:limit]


def remember_output(kind: str, title: str, content: str, page: str, category: str) -> None:
    workspace = active_workspace()
    outputs = list(workspace.get("outputs", []))
    outputs.insert(
        0,
        {
            "timestamp": workspace_timestamp(),
            "kind": kind,
            "title": title,
            "category": category,
            "content": content,
            "preview": short_preview(content, 260),
        },
    )
    workspace["outputs"] = outputs[:24]
    append_history_entry("generation", title, f"{kind} generated in {page}.", page, str(workspace["name"]), category)


def extract_section_value(text: str, labels: list[str]) -> str:
    wanted = {normalize_section_label(label) for label in labels}
    for block in parse_structured_blocks(text):
        if block["type"] != "section":
            continue
        heading = normalize_section_label(str(block.get("heading", "")))
        if heading in wanted:
            return str(block.get("content", "")).strip()
    return ""


def build_code_diff(source: str, fixed_version: str) -> str:
    if not source.strip() or not fixed_version.strip():
        return ""
    diff_lines = difflib.unified_diff(
        source.splitlines(),
        fixed_version.splitlines(),
        fromfile="original",
        tofile="fixed",
        lineterm="",
    )
    return "\n".join(diff_lines).strip()


def developer_stack_profile() -> dict[str, object]:
    model_name = current_model_name().lower()
    is_coder_model = "coder" in model_name or "codeqwen" in model_name
    if is_coder_model:
        best_fit = [
            "Python",
            "JavaScript",
            "TypeScript",
            "HTML/CSS",
            "React",
            "Next.js",
            "Node.js / Express",
            "SQL",
            "JSON / YAML / Config",
            "Bash / PowerShell",
            "Java",
            "C#",
            "PHP",
            "Go",
            "Rust",
            "C / C++",
        ]
        good_fit = ["Kotlin", "Swift", "Flutter / Dart", "Docker / DevOps"]
        low_fit = ["Unity / Game Engine", "Embedded / Microcontroller", "Video / Vision pipelines"]
        profile_name = "Code-specialized profile"
        note = (
            "Active model looks code-specialized. It should be strongest on mainstream backend, frontend, scripting, "
            "and general software engineering tasks."
        )
    else:
        best_fit = [
            "Python",
            "JavaScript",
            "TypeScript",
            "HTML/CSS",
            "React",
            "Next.js",
            "Node.js / Express",
            "SQL",
            "JSON / YAML / Config",
            "Bash / PowerShell",
        ]
        good_fit = ["Java", "C#", "PHP", "Go", "Rust", "C / C++"]
        low_fit = ["Kotlin", "Swift", "Flutter / Dart", "Unity / Game Engine", "Embedded / Microcontroller", "Video / Vision pipelines"]
        profile_name = "General instruct profile"
        note = (
            "Active model is a general instruct model with improved coding, but not a code-only model. "
            "It is safer on common scripting, web, and structured-data tasks than on platform-specific or low-level stacks."
        )
    ordered_options = best_fit + good_fit + low_fit + ["Custom"]
    return {
        "is_coder_model": is_coder_model,
        "profile_name": profile_name,
        "best_fit": best_fit,
        "good_fit": good_fit,
        "low_fit": low_fit,
        "ordered_options": ordered_options,
        "note": note,
    }


MODEL_SELECTION_OPTIONS = [
    "HF Secret Default",
    "Active Session Model",
    "Qwen/Qwen2.5-7B-Instruct",
    "Qwen/Qwen2.5-Coder-7B-Instruct",
    "meta-llama/Llama-3.1-8B-Instruct",
    "mistralai/Mistral-Nemo-Instruct-2407",
    "Custom",
]


def resolve_model_selection(selection: str, custom_value: str = "") -> str:
    choice = (selection or "").strip()
    if choice == "HF Secret Default":
        return read_secret("HF_MODEL")
    if choice == "Active Session Model":
        return current_model_name()
    if choice == "Custom":
        return custom_value.strip()
    return choice


def sync_model_selector_state() -> None:
    override = (st.session_state.get("model_override") or "").strip()
    if override:
        if override in MODEL_SELECTION_OPTIONS:
            st.session_state.model_choice_input = override
            st.session_state.model_custom_input = ""
        else:
            st.session_state.model_choice_input = "Custom"
            st.session_state.model_custom_input = override
    else:
        st.session_state.model_choice_input = "HF Secret Default"
        st.session_state.model_custom_input = ""


def record_last_generation(label: str, model_name: str, status: str, note: str = "") -> None:
    st.session_state.last_generation_label = label
    st.session_state.last_generation_model = model_name
    st.session_state.last_generation_status = status
    st.session_state.last_generation_time = workspace_timestamp()
    st.session_state.last_generation_note = note


def run_generation_with_details(
    prompt: str,
    label: str,
    requested_model: str | None = None,
    track_global: bool = True,
) -> tuple[str | None, str | None]:
    client, model, error = load_client(requested_model)
    if error:
        if track_global:
            record_last_generation(label, requested_model or current_model_name(), "error", error)
        st.error(error)
        return None, None
    with st.spinner(f"Generating {label}..."):
        try:
            result = generate_text(prompt, label, client=client, model=model)
            if track_global:
                record_last_generation(label, model or "", "success", "Generation completed.")
            return result, model
        except Exception as exc:
            if track_global:
                record_last_generation(label, model or requested_model or "", "error", str(exc))
            st.error(str(exc))
            return None, model


def run_model_verification_probe(target_model: str) -> tuple[bool, str]:
    probe_prompt = (
        "Verification probe for coding assistant connectivity.\n"
        "Respond in plain text with exactly these sections:\n"
        "Verification: OK\n"
        "Focus: coding\n"
        "MathCheck: 2+2=4\n"
        "Ready: YES\n"
        "Do not add extra sections."
    )
    result, used_model = run_generation_with_details(
        probe_prompt,
        "model verification probe",
        requested_model=target_model,
        track_global=False,
    )
    if result:
        st.session_state.verification_last_model = used_model or target_model
        st.session_state.verification_last_status = "success"
        st.session_state.verification_last_time = workspace_timestamp()
        st.session_state.verification_last_note = short_preview(result, 180)
        return True, result
    st.session_state.verification_last_model = target_model
    st.session_state.verification_last_status = "error"
    st.session_state.verification_last_time = workspace_timestamp()
    st.session_state.verification_last_note = "Probe failed."
    return False, ""


def developer_selftest_prompt(stack_name: str) -> str:
    prompts = {
        "Python": "Write a small function that validates email input and returns a normalized lowercase email or raises a clear ValueError.",
        "JavaScript": "Write a small async function that fetches JSON, handles HTTP errors, and returns typed-safe friendly data checks in plain JavaScript.",
        "TypeScript": "Write a typed utility that groups a list of records by a key and returns a strongly typed map.",
        "React": "Write a small React component with loading, error, and success states for fetching a list of items.",
        "Next.js": "Write a minimal Next.js route handler plus a page snippet that fetches and renders data safely.",
        "Node.js / Express": "Write an Express route with input validation, async error handling, and JSON response structure.",
        "SQL": "Write a query that returns top customers by total spend with clear grouping and ordering.",
        "HTML/CSS": "Write a responsive card layout with semantic HTML and mobile-first CSS.",
        "Bash / PowerShell": "Write a script that checks whether a file exists, creates a backup, and logs the result safely.",
    }
    task = prompts.get(stack_name, f"Write a small but realistic starter example in {stack_name}.")
    return (
        "Developer self-test.\n\n"
        f"Target stack: {stack_name}\n"
        f"Task: {task}\n\n"
        "Return these plain-text sections in this exact order:\n"
        "Self-Test Summary:\n"
        "Code Sample:\n"
        "Risks Or Review Points:\n"
    )


def resolve_stack_choice(selected: str, custom_value: str = "") -> tuple[str, str]:
    profile = developer_stack_profile()
    custom_clean = custom_value.strip()
    if selected == "Custom":
        if not custom_clean:
            return "Custom stack", "review"
        return custom_clean, "review"
    if selected in profile["best_fit"]:
        return selected, "best"
    if selected in profile["good_fit"]:
        return selected, "good"
    return selected, "review"


def stack_confidence_message(level: str) -> str:
    messages = {
        "best": "Best-fit stack for the active model profile.",
        "good": "Usable stack, but review framework details and environment-specific code carefully.",
        "review": "Manual review strongly advised. This stack is more likely to need corrections or tighter prompting.",
    }
    return messages.get(level, "Manual review strongly advised.")


def read_codegen_reference_bundle(doc_files: list[object], image_files: list[object], attachment_note: str) -> tuple[str, list[str]]:
    parts: list[str] = []
    issues: list[str] = []
    for uploaded_file in doc_files:
        try:
            text = read_uploaded_document(uploaded_file)
            parts.append(f"[Reference File: {getattr(uploaded_file, 'name', 'reference.txt')}]\n{text}")
        except Exception as exc:
            issues.append(f"{getattr(uploaded_file, 'name', 'file')}: {exc}")
    for uploaded_image in image_files:
        parts.append(
            f"[Reference Image: {getattr(uploaded_image, 'name', 'image')}]\n"
            "Image attached. This current setup is text-first, so image content is not parsed automatically."
        )
    if attachment_note.strip():
        parts.append(f"[Attachment Note]\n{attachment_note.strip()}")
    return "\n\n".join(parts).strip(), issues


def active_export_template() -> dict[str, object]:
    template_name = str(st.session_state.get("export_template") or "Academic Classic")
    return EXPORT_TEMPLATES.get(template_name, EXPORT_TEMPLATES["Academic Classic"])


def add_body_paragraph(document: Document, text: str, indent_first_line: bool = False, italic: bool = False) -> None:
    if not text.strip():
        return
    style = active_export_template()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = float(style["line_spacing"])
    paragraph.paragraph_format.space_after = Pt(8)
    if indent_first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
    run = paragraph.add_run(text.strip())
    run.font.name = str(style["font"])
    run.font.size = Pt(int(style["body_size"]))
    run.italic = italic


def render_docx_bytes(
    title: str,
    body: str,
    category: str = "generic",
    metadata_lines: list[str] | None = None,
    output_options: dict[str, bool] | None = None,
) -> bytes:
    style = active_export_template()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    metadata_lines = metadata_lines or []
    output_options = output_options or {}
    export_title = title.strip() or "Untitled Document"

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(export_title)
    title_run.bold = True
    title_run.font.name = str(style["font"])
    title_run.font.size = Pt(int(style["title_size"]))
    title_run.font.color.rgb = style["title_color"]

    if metadata_lines:
        meta_paragraph = document.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_paragraph.paragraph_format.space_after = Pt(10)
        meta_run = meta_paragraph.add_run(" | ".join(metadata_lines))
        meta_run.italic = True
        meta_run.font.name = str(style["font"])
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = style["meta_color"]

    parsed_blocks = parse_structured_blocks(body)
    essay_subtitle: str | None = None
    essay_tip: str | None = None
    body_items: list[dict[str, object]] = []

    for item in parsed_blocks:
        if item["type"] == "section":
            label = normalize_section_label(str(item["heading"]))
            content = str(item.get("content", "")).strip()
            if category == "essay":
                if label in {"heading suggestion", "title suggestion", "suggested heading"}:
                    if content:
                        essay_subtitle = content
                    continue
                if label in {"essay body", "essay"}:
                    if content:
                        body_items.append({"type": "paragraph", "content": content})
                    continue
                if label in {"self-check tip", "quick self-check tip", "closing self-check tip"}:
                    if content:
                        essay_tip = content
                    continue
            body_items.append(item)
        else:
            body_items.append(item)

    if category == "essay" and output_options.get("essay_include_heading", True) and essay_subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.paragraph_format.space_after = Pt(12)
        subtitle_run = subtitle_paragraph.add_run(essay_subtitle)
        subtitle_run.italic = True
        subtitle_run.font.name = str(style["font"])
        subtitle_run.font.size = Pt(max(int(style["body_size"]) - 1, 10))
        subtitle_run.font.color.rgb = style["meta_color"]

    for item in body_items:
        if item["type"] == "section":
            heading_text = str(item["heading"]).strip().rstrip(":")
            if heading_text:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(heading_text)
                run.bold = True
                run.font.name = str(style["font"])
                run.font.size = Pt(int(style["heading_size"]))
                run.font.color.rgb = style["heading_color"]
            content = str(item.get("content", "")).strip()
            if content:
                add_body_paragraph(document, content, indent_first_line=(category == "essay" and bool(style["indent_essay"])))
        else:
            add_body_paragraph(document, str(item["content"]), indent_first_line=(category == "essay" and bool(style["indent_essay"])))

    if category == "essay" and output_options.get("essay_include_tip", True) and essay_tip:
        tip_heading = document.add_paragraph()
        tip_heading.paragraph_format.space_before = Pt(10)
        tip_heading.paragraph_format.space_after = Pt(2)
        tip_run = tip_heading.add_run("Self-Check Tip")
        tip_run.bold = True
        tip_run.font.name = str(style["font"])
        tip_run.font.size = Pt(int(style["heading_size"]))
        tip_run.font.color.rgb = style["heading_color"]
        add_body_paragraph(document, essay_tip, italic=True)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def prepare_download_history(reset_action: str, notice: str, level: str, title: str, category: str) -> None:
    append_history_entry("export", title, f"Prepared browser download for {category}.", st.session_state.active_page, active_workspace_name(), category)
    queue_reset(reset_action, notice, level)


def render_download_button(title: str, body: str, default_name: str, category: str, reset_action: str, name_override: str | None = None) -> None:
    docx_bytes, docx_filename, local_docx_path = build_export_document(
        title,
        body,
        default_name,
        category,
        name_override=name_override,
    )
    if local_docx_path:
        render_route_block("Local save route", local_docx_path)
        local_col, open_col = st.columns(2, gap="small")
        if local_col.button(
            "Save To Selected Folder",
            key=f"save_local_{category}_{sanitize_filename(default_name)}",
            use_container_width=True,
            type="primary",
        ):
            try:
                saved_path = save_local_export_file(docx_bytes, local_docx_path)
                append_history_entry("export", f"Saved export: {docx_filename}", saved_path, st.session_state.active_page, active_workspace_name(), category)
                queue_reset(reset_action, f"Saved to {saved_path}")
                st.rerun()
            except Exception as exc:
                st.error(f"Could not save to the selected folder: {exc}")
        if open_col.button(
            "Open Selected Folder",
            key=f"open_local_{category}_{sanitize_filename(default_name)}",
            use_container_width=True,
        ):
            try:
                target_folder = Path(local_docx_path).parent
                target_folder.mkdir(parents=True, exist_ok=True)
                open_in_file_manager(str(target_folder))
            except Exception as exc:
                st.error(f"Could not open the selected folder: {exc}")
    else:
        render_route_block("Browser download", f"{docx_filename}\nSaved directly by your browser.")
    st.download_button(
        "Download Word File",
        data=docx_bytes,
        file_name=docx_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_{category}_{sanitize_filename(default_name)}",
        use_container_width=True,
        on_click=prepare_download_history,
        args=(reset_action, f"Prepared download: {docx_filename}", "info", title, category),
        type="secondary" if local_docx_path else "primary",
    )


def ocr_supported() -> bool:
    return RapidOCR is not None and Image is not None and np is not None


def get_ocr_engine() -> object:
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        if not ocr_supported():
            raise RuntimeError("OCR support is not installed in this environment.")
        _OCR_ENGINE = RapidOCR()
    return _OCR_ENGINE


def extract_image_text_from_bytes(image_bytes: bytes) -> str:
    if not ocr_supported():
        raise RuntimeError("OCR support is not available in this environment.")
    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    results, _elapsed = get_ocr_engine()(np.array(image))
    if not results:
        return ""
    parts: list[str] = []
    for item in results:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            text_value = item[1]
            if isinstance(text_value, str) and text_value.strip():
                parts.append(text_value.strip())
    return "\n".join(parts).strip()


def read_hugyoku_reference_bundle(doc_files: list[object], image_files: list[object], attachment_note: str) -> tuple[str, list[str], str]:
    parts: list[str] = []
    issues: list[str] = []
    ocr_status = "No images attached."
    for uploaded_file in doc_files:
        try:
            text = read_uploaded_document(uploaded_file)
            parts.append(f"[Reference File: {getattr(uploaded_file, 'name', 'reference.txt')}]\n{text}")
        except Exception as exc:
            issues.append(f"{getattr(uploaded_file, 'name', 'file')}: {exc}")
    if image_files:
        if ocr_supported():
            extracted_count = 0
            for uploaded_image in image_files:
                try:
                    text = extract_image_text_from_bytes(uploaded_image.getvalue())
                    label = getattr(uploaded_image, "name", "image")
                    if text:
                        parts.append(f"[Image OCR: {label}]\n{text}")
                        extracted_count += 1
                    else:
                        parts.append(f"[Image OCR: {label}]\nNo readable text was detected from this image.")
                except Exception as exc:
                    issues.append(f"{getattr(uploaded_image, 'name', 'image')}: OCR failed ({exc})")
            ocr_status = f"OCR processed {extracted_count} of {len(image_files)} image(s)."
        else:
            for uploaded_image in image_files:
                parts.append(
                    f"[Image Attachment: {getattr(uploaded_image, 'name', 'image')}]\n"
                    "Image attached, but OCR is not available in this environment."
                )
            ocr_status = "Images attached, but OCR support is not available here."
    if attachment_note.strip():
        parts.append(f"[Attachment Note]\n{attachment_note.strip()}")
    return "\n\n".join(parts).strip(), issues, ocr_status


def detect_requested_export_format(task_text: str, understanding_text: str = "") -> str:
    explicit = extract_section_value(
        understanding_text,
        ["Requested File Format", "Preferred File Format", "Submission Format", "Output Format"],
    ).strip().lower()
    combined = f"{explicit}\n{task_text}".lower()
    if any(token in combined for token in ["pdf", ".pdf", "portable document"]):
        return "pdf"
    if any(token in combined for token in ["txt", ".txt", "text file", "plain text"]):
        return "txt"
    if any(token in combined for token in ["docx", ".docx", "word file", "microsoft word", "ms word"]):
        return "docx"
    return "docx"


def compile_hugyoku_sections(sections: list[dict[str, object]]) -> str:
    parts: list[str] = []
    for section in sections:
        heading = str(section.get("heading", "")).strip()
        content = str(section.get("content", "")).strip()
        if heading and content:
            parts.append(f"{heading}:\n{content}")
        elif heading:
            parts.append(heading)
        elif content:
            parts.append(content)
    return "\n\n".join(part for part in parts if part.strip()).strip()


def build_hugyoku_sections(text: str) -> list[dict[str, object]]:
    sections: list[dict[str, object]] = []
    for index, block in enumerate(parse_structured_blocks(text), start=1):
        if block["type"] == "section":
            sections.append(
                {
                    "heading": str(block.get("heading", "")).strip() or f"Section {index}",
                    "content": str(block.get("content", "")).strip(),
                }
            )
        else:
            sections.append(
                {
                    "heading": f"Section {index}",
                    "content": str(block.get("content", "")).strip(),
                }
            )
    if not sections and text.strip():
        sections.append({"heading": "Main Output", "content": text.strip()})
    return sections


def sync_hugyoku_sections_from_widgets() -> None:
    updated: list[dict[str, object]] = []
    for index, section in enumerate(list(st.session_state.get("hugyoku_output_sections", []))):
        heading = str(st.session_state.get(f"hugyoku_section_heading_{index}", section.get("heading", ""))).strip()
        content = str(st.session_state.get(f"hugyoku_section_content_{index}", section.get("content", ""))).strip()
        updated.append({"heading": heading, "content": content})
    st.session_state.hugyoku_output_sections = updated
    st.session_state.hugyoku_output_raw = compile_hugyoku_sections(updated)


def prime_hugyoku_section_widgets(sections: list[dict[str, object]]) -> None:
    for index, section in enumerate(sections):
        st.session_state[f"hugyoku_section_heading_{index}"] = str(section.get("heading", "")).strip()
        st.session_state[f"hugyoku_section_content_{index}"] = str(section.get("content", "")).strip()


def guess_hugyoku_title(task_text: str, understanding_text: str, sections: list[dict[str, object]]) -> str:
    explicit = extract_section_value(understanding_text, ["Requested Output", "Task Summary"]).strip()
    for section in sections:
        heading = str(section.get("heading", "")).strip().lower()
        if heading in {"title", "essay title", "document title"}:
            content = str(section.get("content", "")).strip()
            if content:
                return content
    if explicit:
        return explicit[:90]
    cleaned = " ".join(task_text.strip().split())
    return cleaned[:90] if cleaned else "Hugyoku Output"


def render_txt_bytes(title: str, body: str, metadata_lines: list[str] | None = None) -> bytes:
    lines: list[str] = [title.strip() or "Untitled Document"]
    if metadata_lines:
        lines.append(" | ".join(metadata_lines))
    if body.strip():
        lines.extend(["", body.strip()])
    return "\n".join(lines).encode("utf-8")


def render_pdf_bytes(title: str, body: str, metadata_lines: list[str] | None = None) -> bytes:
    if canvas is None or LETTER is None:
        raise RuntimeError("PDF export requires reportlab in this environment.")
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    left_margin = 72
    top = height - 72
    bottom = 72
    line_height = 14
    y = top

    def write_line(text: str, *, font: str = "Helvetica", size: int = 11) -> None:
        nonlocal y
        if y <= bottom:
            pdf.showPage()
            y = top
        pdf.setFont(font, size)
        pdf.drawString(left_margin, y, text)
        y -= line_height if size <= 11 else line_height + 2

    write_line(title.strip() or "Untitled Document", font="Helvetica-Bold", size=18)
    if metadata_lines:
        for meta_line in metadata_lines:
            for wrapped in textwrap.wrap(meta_line, width=86) or [""]:
                write_line(wrapped, font="Helvetica-Oblique", size=10)
        y -= 4

    for paragraph in compile_hugyoku_sections(build_hugyoku_sections(body)).split("\n"):
        wrapped_lines = textwrap.wrap(paragraph, width=90) or [""]
        for wrapped in wrapped_lines:
            write_line(wrapped)

    pdf.save()
    return buffer.getvalue()


def build_hugyoku_export_payload(
    title: str,
    body: str,
    export_format: str,
    *,
    name_override: str | None = None,
) -> tuple[bytes, str, str, str | None]:
    requested = export_format.lower().strip()
    if requested not in {"docx", "pdf", "txt"}:
        requested = "docx"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = sanitize_filename(name_override or title or "hugyoku_output")
    metadata_lines = export_metadata_lines(category="hugyoku", name_override=name_override)
    export_root = (st.session_state.export_root_path or "").strip() if save_destination_mode() == "local" else ""
    internal_folder = TOOL_FOLDERS.get("hugyoku", "hugyoku_universal")
    if requested == "pdf":
        payload = render_pdf_bytes(title, body, metadata_lines)
        filename = f"{stem}_{stamp}.pdf"
        mime = "application/pdf"
    elif requested == "txt":
        payload = render_txt_bytes(title, body, metadata_lines)
        filename = f"{stem}_{stamp}.txt"
        mime = "text/plain"
    else:
        payload = render_docx_bytes(
            title,
            body,
            category="hugyoku",
            metadata_lines=metadata_lines,
            output_options=current_output_settings(),
        )
        filename = f"{stem}_{stamp}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    local_path = str(Path(export_root) / internal_folder / filename) if export_root else None
    return payload, filename, mime, local_path


def prepare_hugyoku_download(file_name: str) -> None:
    append_history_entry(
        "export",
        f"Prepared Hugyoku download: {file_name}",
        "Browser download prepared for the universal Hugyoku workflow.",
        "hugyoku",
        active_workspace_name(),
        "hugyoku",
    )
    st.session_state.flash_message = f"Prepared download: {file_name}"
    st.session_state.flash_level = "info"


def render_hugyoku_export_button(title: str, body: str, export_format: str) -> None:
    file_bytes, file_name, mime_type, local_path = build_hugyoku_export_payload(title, body, export_format, name_override=title)
    if local_path:
        render_route_block("Local save route", local_path)
        local_col, open_col = st.columns(2, gap="small")
        if local_col.button("Save To Selected Folder", key="save_local_hugyoku_export", use_container_width=True, type="primary"):
            try:
                saved_path = save_local_export_file(file_bytes, local_path)
                append_history_entry("export", f"Saved Hugyoku export: {file_name}", saved_path, "hugyoku", active_workspace_name(), "hugyoku")
                st.success(f"Saved to {saved_path}")
            except Exception as exc:
                st.error(f"Could not save to the selected folder: {exc}")
        if open_col.button("Open Selected Folder", key="open_local_hugyoku_export", use_container_width=True):
            try:
                target_folder = Path(local_path).parent
                target_folder.mkdir(parents=True, exist_ok=True)
                open_in_file_manager(str(target_folder))
            except Exception as exc:
                st.error(f"Could not open the selected folder: {exc}")
    else:
        render_route_block("Browser download", f"{file_name}\nSaved directly by your browser.")
    st.download_button(
        f"Download {export_format.upper()} File",
        data=file_bytes,
        file_name=file_name,
        mime=mime_type,
        key=f"download_hugyoku_{sanitize_filename(title)}_{export_format}",
        use_container_width=True,
        on_click=prepare_hugyoku_download,
        args=(file_name,),
        type="secondary" if local_path else "primary",
    )


def build_hugyoku_understanding_prompt(task_text: str, reference_bundle: str) -> str:
    return (
        "Analyze the user's school or academic task before generating anything.\n"
        "Your job is to explain what you understand from the request so the user can confirm or correct it.\n\n"
        f"User task:\n{task_text or 'No direct task text provided. Infer from references.'}\n\n"
        f"Reference bundle:\n{trim_prompt_source(reference_bundle) if reference_bundle else 'No reference files provided.'}\n\n"
        "Return these plain-text sections in this exact order:\n"
        "Task Summary:\n"
        "What You Understand:\n"
        "Requested Output:\n"
        "Requested File Format:\n"
        "Important Requirements:\n"
        "Missing or Uncertain Parts:\n"
        "Next Confirmation Step:\n"
    )


def build_hugyoku_refinement_prompt(task_text: str, reference_bundle: str, current_understanding: str, refinement_prompt: str) -> str:
    return (
        "Revise your understanding of the user's academic task.\n"
        "Keep the task understanding accurate, specific, and ready for confirmation before generation.\n\n"
        f"Original task:\n{task_text or 'No direct task text provided.'}\n\n"
        f"Reference bundle:\n{trim_prompt_source(reference_bundle) if reference_bundle else 'No reference files provided.'}\n\n"
        f"Current understanding:\n{current_understanding}\n\n"
        f"User correction or refinement:\n{refinement_prompt}\n\n"
        "Return these plain-text sections in this exact order:\n"
        "Task Summary:\n"
        "What You Understand:\n"
        "Requested Output:\n"
        "Requested File Format:\n"
        "Important Requirements:\n"
        "Missing or Uncertain Parts:\n"
        "Next Confirmation Step:\n"
    )


def build_hugyoku_generation_prompt(task_text: str, reference_bundle: str, understanding_text: str, export_format: str, revision_prompt: str = "") -> str:
    return (
        "Generate the final academic deliverable based on the confirmed understanding.\n"
        "Write the best possible study-safe output, suitable for review and submission drafting.\n\n"
        f"Original task:\n{task_text or 'No direct task text provided.'}\n\n"
        f"Confirmed understanding:\n{understanding_text}\n\n"
        f"Reference bundle:\n{trim_prompt_source(reference_bundle) if reference_bundle else 'No reference files provided.'}\n\n"
        f"Requested file format: {export_format.upper()}\n"
        f"Additional revision prompt: {revision_prompt or 'None'}\n\n"
        "Return clear plain-text sections with headings and content.\n"
        "Use this order when possible:\n"
        "Title:\n"
        "Final Output:\n"
        "Supporting Notes:\n"
        "Submission Notes:\n\n"
        "Do not use markdown fence blocks."
    )


def clear_workspace_analysis_workspace() -> None:
    st.session_state.workspace_analysis_result = ""


def clear_source_lab_workspace() -> None:
    st.session_state.source_lab_question_input = ""
    st.session_state.source_lab_answer = ""
    st.session_state.source_lab_analysis = ""


def clear_reviewer_workspace() -> None:
    st.session_state.reviewer_request_input = ""
    st.session_state.reviewer_response = ""


def clear_flashcard_workspace() -> None:
    st.session_state.flashcard_request_input = ""
    st.session_state.flashcard_response = ""


def clear_practice_test_workspace() -> None:
    st.session_state.practice_test_request_input = ""
    st.session_state.practice_test_response = ""


def clear_answer_checker_workspace() -> None:
    st.session_state.answer_checker_question_input = ""
    st.session_state.answer_checker_reference_input = ""
    st.session_state.answer_checker_user_answer_input = ""
    st.session_state.answer_checker_response = ""


def clear_rubric_workspace() -> None:
    st.session_state.rubric_title_input = ""
    st.session_state.rubric_prompt_input = ""
    st.session_state.rubric_focus_input = ["Clarity", "Structure", "Depth"]
    st.session_state.rubric_response = ""


def clear_batch_workspace() -> None:
    st.session_state.batch_topic_input = ""
    st.session_state.batch_request_input = ""
    st.session_state.batch_response = ""


def clear_history_entries() -> None:
    st.session_state.history_entries = []


def clear_model_selection() -> None:
    st.session_state.model_choice_input = "HF Secret Default"
    st.session_state.model_override = ""
    st.session_state.model_custom_input = ""


def clear_codegen_workspace() -> None:
    st.session_state.codegen_title = ""
    st.session_state.codegen_stack_profile = "Python"
    st.session_state.codegen_custom_stack = ""
    st.session_state.codegen_description = ""
    st.session_state.codegen_expectation = ""
    st.session_state.codegen_attachment_note = ""
    st.session_state.codegen_response = ""


def clear_selftest_workspace() -> None:
    st.session_state.selftest_stack_profile = "Python"
    st.session_state.selftest_custom_stack = ""
    st.session_state.selftest_response = ""
    st.session_state.selftest_last_model = ""


def clear_compare_workspace() -> None:
    st.session_state.compare_model_a = "Active Session Model"
    st.session_state.compare_model_b = "Qwen/Qwen2.5-Coder-7B-Instruct"
    st.session_state.compare_custom_model_a = ""
    st.session_state.compare_custom_model_b = ""
    st.session_state.compare_prompt = ""
    st.session_state.compare_output_a = ""
    st.session_state.compare_output_b = ""
    st.session_state.compare_used_model_a = ""
    st.session_state.compare_used_model_b = ""


def clear_hugyoku_workspace() -> None:
    st.session_state.hugyoku_task_input = ""
    st.session_state.hugyoku_attachment_note = ""
    st.session_state.hugyoku_stage = 1
    st.session_state.hugyoku_understanding = ""
    st.session_state.hugyoku_refinement_prompt = ""
    st.session_state.hugyoku_refinement_round = 0
    st.session_state.hugyoku_output_sections = []
    st.session_state.hugyoku_output_title = ""
    st.session_state.hugyoku_output_format = "docx"
    st.session_state.hugyoku_output_raw = ""
    st.session_state.hugyoku_result_prompt = ""
    st.session_state.hugyoku_generation_note = ""
    st.session_state.hugyoku_last_bundle = ""
    st.session_state.hugyoku_last_ocr_status = ""


RESET_ACTIONS.update(
    {
        "clear_workspace_analysis_workspace": clear_workspace_analysis_workspace,
        "clear_source_lab_workspace": clear_source_lab_workspace,
        "clear_reviewer_workspace": clear_reviewer_workspace,
        "clear_flashcard_workspace": clear_flashcard_workspace,
        "clear_practice_test_workspace": clear_practice_test_workspace,
        "clear_answer_checker_workspace": clear_answer_checker_workspace,
        "clear_rubric_workspace": clear_rubric_workspace,
        "clear_batch_workspace": clear_batch_workspace,
        "clear_history_entries": clear_history_entries,
        "clear_model_selection": clear_model_selection,
        "clear_codegen_workspace": clear_codegen_workspace,
        "clear_selftest_workspace": clear_selftest_workspace,
        "clear_compare_workspace": clear_compare_workspace,
        "clear_hugyoku_workspace": clear_hugyoku_workspace,
    }
)


def render_history_snippets(limit: int = 5, page_filter: str | None = None) -> None:
    entries = list(st.session_state.get("history_entries", []))
    if page_filter:
        entries = [entry for entry in entries if entry.get("page") == page_filter]
    if not entries:
        st.caption("No history entries yet.")
        return
    for entry in entries[:limit]:
        with st.container(border=True):
            render_card_header(
                str(entry.get("title", "History Entry")),
                str(entry.get("details", "")),
                f"{entry.get('timestamp', '')} • {entry.get('kind', 'event')}",
                anchor="EV",
                tier="tertiary",
                compact=True,
            )
            render_meta_grid(
                [
                    ("Page", str(entry.get("page", "-"))),
                    ("Workspace", str(entry.get("workspace", "-"))),
                    ("Category", str(entry.get("category", "-")) or "-"),
                ]
            )


def render_workspace_outputs(limit: int = 6) -> None:
    workspace = active_workspace()
    outputs = list(workspace.get("outputs", []))
    if not outputs:
        st.caption("No generated outputs saved in this workspace yet.")
        return
    for output in outputs[:limit]:
        with st.container(border=True):
            render_card_header(
                str(output.get("title", "Saved Output")),
                str(output.get("preview", "")),
                f"{output.get('timestamp', '')} • {output.get('kind', 'output')}",
                anchor="OP",
                tier="tertiary",
                compact=True,
            )
            st.caption(f"Category: {output.get('category', '-')}")


def current_identity_name() -> str:
    return st.session_state.saved_name or st.session_state.auth_display_name or st.session_state.auth_username or "Guest"


def render_login_gate() -> None:
    has_users = auth_user_count() > 0
    hero_left, hero_right = st.columns([1.3, 0.95], gap="large")
    with hero_left:
        with st.container(border=True):
            st.markdown(
                """
                <div class="app-kicker">Hugyoku Access Control</div>
                <div class="app-hero-title">Secure Login Before Workspace Access</div>
                <div class="app-card-subtitle">
                  Authentication now protects the full app. Users only reach the main workspace after login,
                  and admins can control roles, active accounts, and menu access from one place.
                </div>
                """,
                unsafe_allow_html=True,
            )
            render_tag_row(["Login First", "Role-Based Access", "Admin Panel", "Local Account Storage"])

    with hero_right:
        with st.container(border=True):
            render_card_header(
                "Access Status",
                "Set up the first admin account or sign in with an existing account.",
                "Security",
            )
            render_meta_grid(
                [
                    ("Users in database", str(auth_user_count())),
                    ("Database", AUTH_DB_PATH.name),
                    ("Auth mode", "Bootstrap required" if not has_users else "Login required"),
                ]
            )

    if not has_users:
        left, right = st.columns([1.1, 0.9], gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Create The First Admin",
                    "No users exist yet. Create the first admin account here. This becomes the account that can manage everyone else.",
                    "Bootstrap",
                )
                with st.form("bootstrap_form", clear_on_submit=False):
                    st.text_input("Admin username", key="bootstrap_username_input")
                    st.text_input("Display name", key="bootstrap_display_name_input")
                    st.text_input("Password", key="bootstrap_password_input", type="password")
                    st.text_input("Confirm password", key="bootstrap_confirm_password_input", type="password")
                    bootstrap_submitted = st.form_submit_button(
                        "Create First Admin",
                        use_container_width=True,
                        type="primary",
                    )
                if bootstrap_submitted:
                    username = st.session_state.bootstrap_username_input.strip()
                    display_name = st.session_state.bootstrap_display_name_input.strip()
                    password = st.session_state.bootstrap_password_input
                    confirm = st.session_state.bootstrap_confirm_password_input
                    if not username or not password:
                        st.warning("Username and password are required.")
                    elif password != confirm:
                        st.warning("Passwords do not match.")
                    else:
                        ok, message = create_auth_user(username, display_name or username, password, role="super_admin")
                        if ok:
                            queue_reset("clear_bootstrap_form", "First admin account created. You can sign in now.")
                            st.rerun()
                        else:
                            st.error(message)
        with right:
            with st.container(border=True):
                render_card_header(
                    "Bootstrap Notes",
                    "This setup stores local users in a SQLite database beside the app, not inside your public GitHub repo secrets.",
                    "Guide",
                )
                st.markdown(
                    """
                    <ol class="app-guide-list">
                      <li>Create the first admin once on the machine where the app runs.</li>
                      <li>After that, login is required before any workspace appears.</li>
                      <li>Admins can create members and control which menus each account can use.</li>
                    </ol>
                    """,
                    unsafe_allow_html=True,
                )
        return

    login_left, login_right = st.columns([1.0, 1.0], gap="large")
    with login_left:
        with st.container(border=True):
            render_card_header(
                "Sign In",
                "Use your Hugyoku account to open the workspace. The app stays locked until login succeeds.",
                "Authentication",
            )
            with st.form("login_form", clear_on_submit=False):
                st.text_input("Username", key="login_username_input")
                st.text_input("Password", key="login_password_input", type="password")
                login_submitted = st.form_submit_button("Login To Hugyoku", use_container_width=True, type="primary")
            if login_submitted:
                success, message = authenticate_user(st.session_state.login_username_input, st.session_state.login_password_input)
                if success:
                    queue_reset(
                        "clear_login_form",
                        f"Welcome back, {st.session_state.auth_display_name or st.session_state.auth_username}.",
                    )
                    st.rerun()
                st.error(message)

    with login_right:
        with st.container(border=True):
            render_card_header(
                "What Happens After Login",
                "The app only shows the menus and pages your role is allowed to use.",
                "Access Model",
            )
            render_route_block("Access rules", "member -> core workspace tools\nadmin -> workspace tools + admin panel\ninactive user -> blocked")
            recent_logs = auth_log_rows(limit=4)
            if recent_logs:
                for log_row in recent_logs[:4]:
                    render_route_block(
                        f"{log_row['created_at']} • {log_row['event_type']}",
                        f"{log_row['username']}\n{log_row['details']}",
                    )


def render_admin_page() -> None:
    users = list_auth_users()
    role_filter = st.session_state.admin_filter_role
    filtered_users = users if role_filter == "All" else [user for user in users if user["role"] == role_filter]

    with st.container(border=True):
        render_card_header(
            PAGE_DETAILS["admin"]["title"],
            PAGE_DETAILS["admin"]["subtitle"],
            "Admin Control",
        )
        render_tag_row(["User Manager", "Role Permissions", "Auth Logs", "Account Status"])

    overview_tab, users_tab, logs_tab = st.tabs(["Overview", "Users", "Activity Logs"])

    with overview_tab:
        total_users = len(users)
        active_users = sum(1 for user in users if user["is_active"])
        admin_users = sum(1 for user in users if user["role"] in {"admin", "super_admin"})
        member_users = sum(1 for user in users if user["role"] == "member")
        viewer_users = sum(1 for user in users if user["role"] == "viewer")
        with st.container(border=True):
            render_card_header(
                "Access Overview",
                "Quick visibility into who can use the app and which accounts currently have elevated privileges.",
                "Overview",
            )
            render_meta_grid(
                [
                    ("Total users", str(total_users)),
                    ("Active users", str(active_users)),
                    ("Admins", str(admin_users)),
                    ("Current operator", st.session_state.auth_username or "-"),
                ]
            )
            render_meta_grid(
                [
                    ("Members", str(member_users)),
                    ("Viewers", str(viewer_users)),
                    ("Inactive", str(max(total_users - active_users, 0))),
                    ("Admin safeguard", "Enabled"),
                ]
            )
            render_route_block(
                "Role model",
                "super_admin/admin -> full access\nmember -> main app access\nviewer -> no developer access unless enabled manually",
            )

    with users_tab:
        create_col, manage_col = st.columns([0.95, 1.05], gap="large")
        with create_col:
            with st.container(border=True):
                render_card_header(
                    "Create User",
                    "Create a new local account and assign its base role in one step.",
                    "User Management",
                )
                with st.form("admin_create_user_form", clear_on_submit=False):
                    st.text_input("New username", key="admin_new_username")
                    st.text_input("Display name", key="admin_new_display_name")
                    st.text_input("Temporary password", key="admin_new_password", type="password")
                    st.selectbox("Role", options=["member", "viewer", "admin"], key="admin_new_role")
                    create_user_submitted = st.form_submit_button(
                        "Create User Account",
                        use_container_width=True,
                        type="primary",
                    )
                if create_user_submitted:
                    ok, message = create_auth_user(
                        st.session_state.admin_new_username,
                        st.session_state.admin_new_display_name or st.session_state.admin_new_username,
                        st.session_state.admin_new_password,
                        st.session_state.admin_new_role,
                    )
                    if ok:
                        queue_reset("clear_admin_create_user_form", message)
                        st.rerun()
                    else:
                        st.error(message)
                if filtered_users:
                    st.markdown("#### Current Directory")
                    for user in filtered_users[:6]:
                        status_note = "active" if user["is_active"] else "disabled"
                        render_route_block(
                            f"{user['username']} • {user['role']}",
                            f"{user['display_name']}\n{status_note}",
                        )

        with manage_col:
            with st.container(border=True):
                render_card_header(
                    "Manage Existing Users",
                    "Edit roles, status, and page-level privileges. Password resets are local and immediate.",
                    "Access Control",
                )
                filter_col, select_col = st.columns([0.45, 0.55], gap="small")
                filter_col.selectbox("Role filter", options=["All", "super_admin", "admin", "member", "viewer"], key="admin_filter_role")
                selected_options = filtered_users or users
                if not selected_options:
                    st.caption("No users available.")
                else:
                    selected_index = 0
                    existing_id = int(st.session_state.admin_selected_user_id or 0)
                    option_ids = [int(user["id"]) for user in selected_options]
                    if existing_id in option_ids:
                        selected_index = option_ids.index(existing_id)
                    selected_user_id = select_col.selectbox(
                        "Select user",
                        options=option_ids,
                        index=selected_index,
                        format_func=lambda value: next(
                            f"{user['username']} ({user['role']})" for user in selected_options if int(user["id"]) == int(value)
                        ),
                        key="admin_selected_user_id",
                    )
                    selected_user = get_user_by_id(int(selected_user_id))
                    if selected_user:
                        default_permissions = selected_user["permissions"]
                        is_super_admin_target = str(selected_user["role"]) == "super_admin"
                        current_is_super_admin = st.session_state.auth_role == "super_admin"
                        can_manage_target = current_is_super_admin or not is_super_admin_target
                        if not can_manage_target:
                            st.warning("Only a super admin can edit another super admin account.")
                        display_name = st.text_input("Display name", value=str(selected_user["display_name"]), key=f"admin_display_name_{selected_user_id}")
                        role_value = st.selectbox(
                            "Role",
                            options=["super_admin", "admin", "member", "viewer"],
                            index=["super_admin", "admin", "member", "viewer"].index(str(selected_user["role"])),
                            key=f"admin_role_{selected_user_id}",
                        )
                        active_value = st.checkbox("Account active", value=bool(selected_user["is_active"]), key=f"admin_active_{selected_user_id}")
                        st.markdown("#### Page Permissions")
                        perm_cols = st.columns(2, gap="small")
                        permissions_keys = ["dashboard", "workspaces", "academics", "developer", "history", "settings", "admin"]
                        permission_values: dict[str, bool] = {}
                        for idx, permission_key in enumerate(permissions_keys):
                            target_col = perm_cols[idx % 2]
                            with target_col:
                                permission_values[permission_key] = st.checkbox(
                                    permission_key.replace("_", " ").title(),
                                    value=bool(default_permissions.get(permission_key, False)),
                                    key=f"perm_{selected_user_id}_{permission_key}",
                                )
                        action_row = st.columns(2, gap="small")
                        if action_row[0].button(
                            "Save User Changes",
                            key=f"save_user_{selected_user_id}",
                            use_container_width=True,
                            type="primary",
                            disabled=not can_manage_target,
                        ):
                            ok, message = update_auth_user(
                                int(selected_user_id),
                                display_name,
                                role_value,
                                active_value,
                                permission_values,
                            )
                            if ok:
                                if int(selected_user_id) == int(st.session_state.auth_user_id):
                                    refreshed = get_user_by_id(int(selected_user_id))
                                    if refreshed:
                                        apply_auth_session(refreshed)
                                st.session_state.flash_message = message
                                st.session_state.flash_level = "success"
                                st.rerun()
                            else:
                                st.error(message)
                        st.text_input("Reset password", key="admin_reset_password", type="password")
                        if action_row[1].button(
                            "Update Password",
                            key=f"reset_user_password_{selected_user_id}",
                            use_container_width=True,
                            disabled=not can_manage_target,
                        ):
                            ok, message = update_auth_password(int(selected_user_id), st.session_state.admin_reset_password)
                            if ok:
                                queue_reset("clear_admin_password_reset", message)
                                st.rerun()
                            else:
                                st.error(message)

    with logs_tab:
        with st.container(border=True):
            render_card_header(
                "Authentication Logs",
                "Recent account events for this local app instance.",
                "Audit Trail",
            )
            for row in auth_log_rows(limit=60):
                render_route_block(
                    f"{row['created_at']} • {row['event_type']}",
                    f"{row['username']}\n{row['details']}",
                )


def render_header(ai_ready: bool, model_label: str, ai_message: str) -> None:
    workspace = active_workspace()
    title = f"{current_identity_name()}'s Premium Study Desk"
    subtitle = (
        "Unified dashboard for workspaces, source-grounded analysis, study generators, "
        "developer support, history, and export controls."
    )
    status_text = "AI ready for academics" if ai_ready else ai_message
    chip_class = "ready" if ai_ready else ("waiting" if "Add HF_TOKEN" in ai_message or "Streamlit secrets" in ai_message else "offline")
    current_page_title = PAGE_DETAILS.get(st.session_state.active_page, {}).get("title", "Dashboard")
    current_mode = "Local folder picker" if save_destination_mode() == "local" else "Browser download"
    current_template = str(st.session_state.get("export_template", "Academic Classic"))

    left, right = st.columns([1.72, 0.98], gap="large")
    with left:
        with st.container(border=True):
            st.markdown(
                f"""
                <div class="app-hero-shell">
                  <div class="app-card-topline">
                    <span class="app-anchor-badge">HQ</span>
                    <span class="app-tier-pill primary">Command Center</span>
                  </div>
                  <div class="app-hero-title">{html_text(title)}</div>
                  <div class="app-hero-lead">{html_text(subtitle)}</div>
                  <div class="app-hero-note">Focused workflows for research, drafting, review, and export.</div>
                  <div class="app-luxury-rule"></div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            render_kpi_row(
                [
                    ("Current page", current_page_title),
                    ("Export mode", current_mode),
                    ("Workspace", str(workspace.get("name", "General Workspace"))),
                ]
            )
            action_a, action_b, action_c = st.columns([1.05, 1, 0.92], gap="small")
            if action_a.button("Open Hugyoku", key="hero_open_hugyoku", type="primary", use_container_width=True):
                go("hugyoku")
            if action_b.button("Open Academics", key="hero_open_academics", use_container_width=True):
                go("academics")
            if action_c.button("Open Developer", key="hero_open_developer", use_container_width=True):
                go("developer")
    with right:
        with st.container(border=True):
            render_card_header(
                "Live Control Panel",
                "Real-time session status, model awareness, and operator context stay visible in one compact panel.",
                "System Status",
                anchor="AI",
                tier="secondary",
                compact=True,
            )
            st.markdown(
                f"<div class='app-status-pill {chip_class}'>{html_text(status_text)}</div>",
                unsafe_allow_html=True,
            )
            render_meta_grid(
                [
                    ("Signed in as", st.session_state.auth_username or "No active user"),
                    ("Active model", model_label),
                ]
            )
            with st.expander("Session details", expanded=False):
                render_meta_grid(
                    [
                        ("Role", st.session_state.auth_role or "None"),
                        ("History entries", str(len(st.session_state.get("history_entries", [])))),
                        ("Template", current_template),
                    ]
                )
            if st.button("Refresh AI Status", key="refresh_ai_status", use_container_width=True):
                st.rerun()


def render_sidebar(ai_ready: bool, model_label: str, ai_message: str) -> None:
    nav_items = [
        ("hugyoku", "01 Hugyoku"),
        ("dashboard", "02 Dashboard"),
        ("workspaces", "03 Workspaces"),
        ("academics", "04 Academics"),
        ("developer", "05 Developer"),
        ("history", "06 History"),
        ("settings", "07 Settings"),
    ]
    if can_access_page("admin"):
        nav_items.append(("admin", "08 Admin"))

    with st.sidebar:
        st.markdown(
            """
            <div class="app-sidebar-brand-row">
              <span class="app-anchor-badge">HY</span>
              <div>
                <div class="app-sidebar-brand">Hugyoku</div>
                <div class="app-sidebar-copy">Private academic command desk for grounded research, drafting, exports, and developer support.</div>
                <div class="app-sidebar-quiet">Compact navigation rail with live session context.</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown("<div class='app-nav-section'>Navigate</div>", unsafe_allow_html=True)

        academic_pages = {"academics", "quiz", "assignment", "essay", "activity", "document"}
        developer_pages = {"developer", "codefix"}
        for page, label in nav_items:
            if not can_access_page(page):
                continue
            active = (
                page == "academics" and st.session_state.active_page in academic_pages
            ) or (
                page == "developer" and st.session_state.active_page in developer_pages
            ) or st.session_state.active_page == page
            if st.button(
                label,
                key=f"nav_{page}",
                type="primary" if active else "secondary",
                use_container_width=True,
            ):
                go(page)

        ensure_workspace_bootstrap()
        workspace_ids = list(st.session_state.workspaces.keys())
        with st.container(border=True):
            render_card_header(
                "Workspace Focus",
                "Switch context before generating, reviewing, or exporting anything.",
                "Control",
                anchor="WS",
                tier="secondary",
                compact=True,
            )
            selected_workspace = st.selectbox(
                "Active workspace",
                options=workspace_ids,
                index=workspace_ids.index(st.session_state.active_workspace_id),
                format_func=workspace_option_label,
                key="sidebar_workspace_selector",
            )
            if selected_workspace != st.session_state.active_workspace_id:
                st.session_state.active_workspace_id = selected_workspace
                st.rerun()
            workspace = active_workspace()
            render_kpi_row(
                [
                    ("Files", str(len(workspace.get("files", [])))),
                    ("Images", str(len(workspace.get("images", [])))),
                ]
            )
            st.caption("Your active workspace controls what gets grouped, cited, and exported.")

        with st.container(border=True):
            render_card_header(
                "Session Control",
                "Your identity, role, and model context are grouped here so the sidebar stays compact.",
                "Access",
                anchor="ID",
                tier="tertiary",
                compact=True,
            )
            render_meta_grid(
                [
                    ("User", st.session_state.auth_display_name or st.session_state.auth_username or "Unknown user"),
                    ("Role", st.session_state.auth_role or "none"),
                ]
            )
            st.caption("Use the app for studying, drafting, and understanding source material before submitting anything.")
            if st.button("Logout", key="sidebar_logout_button", use_container_width=True):
                logout_current_user()
                st.session_state.flash_message = "You have been logged out."
                st.session_state.flash_level = "info"
                st.rerun()

        with st.expander("Session diagnostics", expanded=False):
            st.caption("AI status")
            st.write("Ready" if ai_ready else ai_message)
            st.caption(model_label)
            st.caption(f"Export mode: {'Local folder picker' if save_destination_mode() == 'local' else 'Browser download'}")


def render_hugyoku_page(ai_ready: bool) -> None:
    workspace = active_workspace()
    doc_files = list(st.session_state.get("hugyoku_reference_docs") or [])
    image_files = list(st.session_state.get("hugyoku_reference_images") or [])
    output_sections = list(st.session_state.get("hugyoku_output_sections", []))
    has_understanding = bool(st.session_state.hugyoku_understanding.strip())
    has_result = bool(output_sections)
    stage = int(st.session_state.get("hugyoku_stage", 1) or 1)
    if stage >= 3 and has_result:
        stage = max(stage, 3)
    elif has_understanding:
        stage = max(stage, 2)
    else:
        stage = 1
    st.session_state.hugyoku_stage = stage

    with st.container(border=True):
        render_card_header(
            PAGE_DETAILS["hugyoku"]["title"],
            PAGE_DETAILS["hugyoku"]["subtitle"],
            "Universal Flow",
            anchor="HY",
            tier="primary",
        )
        render_kpi_row(
            [
                ("Workspace", str(workspace.get("name", "General Workspace"))),
                ("Docs", str(len(doc_files))),
                ("Images", str(len(image_files))),
                ("Status", "Result ready" if stage >= 3 else ("Understanding ready" if stage >= 2 else "Awaiting task")),
            ]
        )
        st.caption("Flow: 1) Task intake -> 2) AI understanding and corrections -> 3) Final result, editing, and export.")

    def persist_hugyoku_result(result_text: str, export_format: str, generation_note: str) -> None:
        sections = build_hugyoku_sections(result_text)
        output_title = guess_hugyoku_title(st.session_state.hugyoku_task_input, st.session_state.hugyoku_understanding, sections)
        st.session_state.hugyoku_output_sections = sections
        prime_hugyoku_section_widgets(sections)
        st.session_state.hugyoku_output_raw = compile_hugyoku_sections(sections)
        st.session_state.hugyoku_output_title = output_title
        st.session_state.hugyoku_output_format = export_format
        st.session_state.hugyoku_generation_note = generation_note
        remember_output("Hugyoku", output_title, st.session_state.hugyoku_output_raw, "hugyoku", "hugyoku")

    def run_understanding_analysis() -> None:
        task_text = st.session_state.hugyoku_task_input.strip()
        reference_bundle, issues, ocr_status = read_hugyoku_reference_bundle(doc_files, image_files, st.session_state.hugyoku_attachment_note)
        for issue in issues:
            st.warning(issue)
        if not task_text and not reference_bundle:
            st.warning("Add a task, a document, or an image first so Hugyoku has something to analyze.")
            return
        result, used_model = run_generation_with_details(
            build_hugyoku_understanding_prompt(task_text, reference_bundle),
            "hugyoku task understanding",
        )
        if result:
            st.session_state.hugyoku_understanding = result
            st.session_state.hugyoku_last_bundle = reference_bundle
            st.session_state.hugyoku_last_ocr_status = ocr_status
            st.session_state.hugyoku_stage = 2
            st.session_state.hugyoku_refinement_round = 1
            st.session_state.hugyoku_output_sections = []
            st.session_state.hugyoku_output_raw = ""
            st.session_state.hugyoku_output_title = ""
            st.session_state.hugyoku_result_prompt = ""
            st.session_state.hugyoku_output_format = detect_requested_export_format(task_text, result)
            st.session_state.hugyoku_generation_note = f"Understanding prepared using {used_model or current_model_name() or 'the active model'}."
            append_history_entry(
                "analysis",
                "Hugyoku understanding ready",
                ocr_status or "Task understanding generated.",
                "hugyoku",
                active_workspace_name(),
                "hugyoku",
            )
            st.session_state.flash_message = "Step 1 complete. Review the AI understanding in Step 2."
            st.session_state.flash_level = "success"
            st.rerun()

    def run_understanding_revision() -> None:
        refinement_prompt = st.session_state.hugyoku_refinement_prompt.strip()
        if not refinement_prompt:
            st.warning("Add a correction prompt first, or skip this step if the understanding is already correct.")
            return
        task_text = st.session_state.hugyoku_task_input.strip()
        reference_bundle = st.session_state.hugyoku_last_bundle
        if not reference_bundle:
            reference_bundle, issues, ocr_status = read_hugyoku_reference_bundle(doc_files, image_files, st.session_state.hugyoku_attachment_note)
            for issue in issues:
                st.warning(issue)
            st.session_state.hugyoku_last_bundle = reference_bundle
            st.session_state.hugyoku_last_ocr_status = ocr_status
        result, used_model = run_generation_with_details(
            build_hugyoku_refinement_prompt(task_text, reference_bundle, st.session_state.hugyoku_understanding, refinement_prompt),
            "hugyoku understanding revision",
        )
        if result:
            st.session_state.hugyoku_understanding = result
            st.session_state.hugyoku_stage = 2
            st.session_state.hugyoku_refinement_round = max(int(st.session_state.hugyoku_refinement_round or 0), 0) + 1
            st.session_state.hugyoku_output_sections = []
            st.session_state.hugyoku_output_raw = ""
            st.session_state.hugyoku_output_title = ""
            st.session_state.hugyoku_result_prompt = ""
            st.session_state.hugyoku_output_format = detect_requested_export_format(f"{task_text}\n{refinement_prompt}", result)
            st.session_state.hugyoku_generation_note = f"Understanding refined using {used_model or current_model_name() or 'the active model'}."
            st.session_state.hugyoku_refinement_prompt = ""
            append_history_entry(
                "analysis",
                "Hugyoku understanding revised",
                "The task understanding loop ran again with user corrections.",
                "hugyoku",
                active_workspace_name(),
                "hugyoku",
            )
            st.session_state.flash_message = "Understanding updated. If it looks correct now, continue to the final result."
            st.session_state.flash_level = "success"
            st.rerun()

    def run_final_generation() -> None:
        task_text = st.session_state.hugyoku_task_input.strip()
        reference_bundle = st.session_state.hugyoku_last_bundle
        if not reference_bundle:
            reference_bundle, issues, ocr_status = read_hugyoku_reference_bundle(doc_files, image_files, st.session_state.hugyoku_attachment_note)
            for issue in issues:
                st.warning(issue)
            st.session_state.hugyoku_last_bundle = reference_bundle
            st.session_state.hugyoku_last_ocr_status = ocr_status
        export_format = detect_requested_export_format(
            f"{task_text}\n{st.session_state.hugyoku_refinement_prompt}",
            st.session_state.hugyoku_understanding,
        )
        result, used_model = run_generation_with_details(
            build_hugyoku_generation_prompt(
                task_text,
                reference_bundle,
                st.session_state.hugyoku_understanding,
                export_format,
                st.session_state.hugyoku_refinement_prompt.strip(),
            ),
            "hugyoku final output",
        )
        if result:
            persist_hugyoku_result(
                result,
                export_format,
                f"Final result generated using {used_model or current_model_name() or 'the active model'}.",
            )
            st.session_state.hugyoku_stage = 3
            st.session_state.flash_message = "Final result generated. Step 3 is now ready for editing and export."
            st.session_state.flash_level = "success"
            st.rerun()

    if has_result:
        sync_hugyoku_sections_from_widgets()
        compiled_output = st.session_state.hugyoku_output_raw.strip()
        result_left, result_right = st.columns([1.05, 0.95], gap="large")
        with result_left:
            with st.container(border=True):
                render_card_header(
                    "Final Result",
                    "This is now the main focus. Review the output here first, then export or refine it from the panel on the right.",
                    "Step 4",
                    anchor="RS",
                    tier="primary",
                )
                render_preview_panel(
                    "Output Preview",
                    "The generated result stays pinned here so it is easy to see and review.",
                    "Final output preview",
                    compiled_output,
                    height=500,
                    empty_title="No result generated yet",
                    empty_body="Generate the final result first. It will appear here once the draft is ready.",
                    anchor="OP",
                    tier="secondary",
                )
                if compiled_output:
                    render_hugyoku_export_button(
                        st.session_state.hugyoku_output_title or "Hugyoku Output",
                        compiled_output,
                        st.session_state.hugyoku_output_format,
                    )
        with result_right:
            with st.container(border=True):
                render_card_header(
                    "Edit And Refine",
                    "Adjust each section, change the format, or send one more prompt after you see the result.",
                    "Controls",
                    anchor="ED",
                    tier="secondary",
                )
                st.text_input("Output title", key="hugyoku_output_title")
                st.selectbox(
                    "Export format",
                    options=["docx", "pdf", "txt"],
                    format_func=lambda value: value.upper(),
                    key="hugyoku_output_format",
                )
                for index, section in enumerate(output_sections):
                    heading_key = f"hugyoku_section_heading_{index}"
                    content_key = f"hugyoku_section_content_{index}"
                    if heading_key not in st.session_state:
                        st.session_state[heading_key] = str(section.get("heading", "")).strip()
                    if content_key not in st.session_state:
                        st.session_state[content_key] = str(section.get("content", "")).strip()
                    with st.expander(st.session_state.get(heading_key) or f"Section {index + 1}", expanded=index == 0):
                        st.text_input("Section heading", key=heading_key)
                        st.text_area("Section content", key=content_key, height=150)
                row_a, row_b = st.columns(2, gap="small")
                if row_a.button("Apply Section Edits", key="hugyoku_apply_section_edits", use_container_width=True, type="primary"):
                    sync_hugyoku_sections_from_widgets()
                    st.success("Section edits applied to the result preview.")
                if row_b.button("Start New Hugyoku Task", key="hugyoku_reset_after_result", use_container_width=True):
                    queue_reset("clear_hugyoku_workspace", "Hugyoku workflow reset. You can start a new task now.")
                    st.rerun()
                st.text_area(
                    "Additional prompt after result",
                    key="hugyoku_result_prompt",
                    height=120,
                    placeholder="Example: Make it more formal, shorten the answer, or convert it into a reviewer format.",
                )
                refine_a, refine_b = st.columns(2, gap="small")
                if refine_a.button("Regenerate With Additional Prompt", key="hugyoku_regenerate_output", use_container_width=True, type="secondary", disabled=not ai_ready):
                    additional_prompt = st.session_state.hugyoku_result_prompt.strip()
                    if not additional_prompt:
                        st.warning("Add an additional prompt first so Hugyoku knows what to improve.")
                    else:
                        sync_hugyoku_sections_from_widgets()
                        current_body = st.session_state.hugyoku_output_raw
                        revision_prompt = f"{additional_prompt}\n\nCurrent editable draft:\n{current_body}"
                        result, used_model = run_generation_with_details(
                            build_hugyoku_generation_prompt(
                                st.session_state.hugyoku_task_input.strip(),
                                st.session_state.hugyoku_last_bundle,
                                st.session_state.hugyoku_understanding,
                                st.session_state.hugyoku_output_format,
                                revision_prompt,
                            ),
                            "hugyoku refined output",
                        )
                        if result:
                            persist_hugyoku_result(
                                result,
                                st.session_state.hugyoku_output_format,
                                f"Result refined using {used_model or current_model_name() or 'the active model'}.",
                            )
                            st.session_state.hugyoku_stage = 3
                            st.session_state.hugyoku_result_prompt = ""
                            st.session_state.flash_message = "The final result was refined and updated."
                            st.session_state.flash_level = "success"
                            st.rerun()
                if refine_b.button("Save Current Draft To Workspace", key="hugyoku_save_current_draft", use_container_width=True):
                    sync_hugyoku_sections_from_widgets()
                    remember_output(
                        "Hugyoku Draft",
                        st.session_state.hugyoku_output_title or "Hugyoku Draft",
                        st.session_state.hugyoku_output_raw,
                        "hugyoku",
                        "hugyoku",
                    )
                    st.success("Current Hugyoku draft saved into the active workspace history.")
                with st.expander("Result notes", expanded=False):
                    render_route_block(
                        "Detected export route",
                        f"{st.session_state.hugyoku_output_format.upper()} output ready for {'local save' if save_destination_mode() == 'local' else 'browser download'}.",
                    )
                    render_route_block(
                        "Why this draft exists",
                        st.session_state.hugyoku_generation_note or "No generation note yet.",
                    )

    with st.expander("1. Task intake", expanded=stage == 1):
        with st.container(border=True):
            render_card_header(
                "Task Intake",
                "Describe the task and attach anything the AI should read first. Keep this focused on the assignment itself.",
                "Step 1",
                anchor="IN",
                tier="secondary",
            )
            st.text_area(
                "Input task",
                key="hugyoku_task_input",
                height=160,
                placeholder="Example: Make me an essay about Aristotle's life and submit it as a Word file. Include introduction, key contributions, and conclusion.",
            )
            upload_left, upload_right = st.columns(2, gap="small")
            with upload_left:
                st.file_uploader(
                    "Insert reference files",
                    type=["docx", "pdf", "txt", "md"],
                    accept_multiple_files=True,
                    key="hugyoku_reference_docs",
                    help="Attach assignment sheets, rubrics, notes, or source documents.",
                )
            with upload_right:
                st.file_uploader(
                    "Insert image",
                    type=["png", "jpg", "jpeg", "webp"],
                    accept_multiple_files=True,
                    key="hugyoku_reference_images",
                    help="Attach screenshots, worksheet images, or photo references.",
                )
            st.text_area(
                "Attachment note (optional)",
                key="hugyoku_attachment_note",
                height=85,
                placeholder="Use this only if the attachments need extra context before analysis.",
            )
            action_a, action_b = st.columns(2, gap="small")
            if action_a.button("Analyze Task Understanding", key="hugyoku_analyze_task", use_container_width=True, type="primary", disabled=not ai_ready):
                run_understanding_analysis()
            if action_b.button("Reset Hugyoku Flow", key="hugyoku_clear_workspace", use_container_width=True):
                queue_reset("clear_hugyoku_workspace", "Hugyoku workflow cleared.")
                st.rerun()

    if stage >= 2 and has_understanding:
        with st.expander("2. Review AI understanding", expanded=stage == 2):
            top_left, top_right = st.columns([1.08, 0.92], gap="large")
            with top_left:
                with st.container(border=True):
                    render_preview_panel(
                        "AI Understanding",
                        "Read this first. If the AI got the task wrong, use the correction box below. If it got the task right, continue to the final result.",
                        "Understanding preview",
                        st.session_state.hugyoku_understanding,
                        height=300,
                        empty_title="No understanding yet",
                        empty_body="Start from task intake first.",
                        anchor="UA",
                        tier="secondary",
                    )
            with top_right:
                with st.container(border=True):
                    render_card_header(
                        "Quick Summary",
                        "Use this short summary to confirm whether the AI understood the request correctly.",
                        "Step 2",
                        anchor="CF",
                        tier="tertiary",
                    )
                    render_meta_grid(
                        [
                            ("Requested output", extract_section_value(st.session_state.hugyoku_understanding, ["Requested Output"]) or "Not analyzed yet"),
                            ("Requested format", (extract_section_value(st.session_state.hugyoku_understanding, ["Requested File Format"]) or st.session_state.hugyoku_output_format).upper()),
                        ]
                    )
                    render_route_block(
                        "Task summary",
                        extract_section_value(st.session_state.hugyoku_understanding, ["Task Summary"]) or "No task summary yet.",
                    )
                    with st.expander("More analysis notes", expanded=False):
                        render_route_block(
                            "Important requirements",
                            extract_section_value(st.session_state.hugyoku_understanding, ["Important Requirements"]) or "No requirements captured yet.",
                        )
                        render_route_block(
                            "Missing or uncertain parts",
                            extract_section_value(st.session_state.hugyoku_understanding, ["Missing or Uncertain Parts"]) or "No missing parts were flagged.",
                        )
                        render_route_block("OCR status", st.session_state.hugyoku_last_ocr_status or "No image analysis has been run yet.")

            with st.container(border=True):
                render_card_header(
                    "Correct Or Continue",
                    "If may mali, ilagay mo dito ang correction. Kung tama na, diretso ka na sa final result.",
                    "Next Action",
                    anchor="NX",
                    tier="secondary",
                )
                st.text_area(
                    "Correction or clarification prompt",
                    key="hugyoku_refinement_prompt",
                    height=120,
                    placeholder="Example: Focus only on the significance of the name, keep it concise, and submit it as DOCX.",
                )
                action_a, action_b = st.columns(2, gap="small")
                if action_a.button("Revise Understanding", key="hugyoku_revise_understanding", use_container_width=True, type="secondary", disabled=not ai_ready):
                    run_understanding_revision()
                if action_b.button("Generate Final Result", key="hugyoku_generate_result", use_container_width=True, type="primary", disabled=not ai_ready):
                    run_final_generation()


def render_dashboard() -> None:
    workspace = active_workspace()
    paths = folder_path_lines()
    current_mode = save_destination_mode()
    picker_ready = local_folder_picker_available()

    left, right = st.columns([1.15, 0.85], gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Profile Section",
                "Set your identity, choose how exports should behave, and keep the dashboard disciplined with compact progressive controls.",
                "Dashboard Identity",
                anchor="ID",
                tier="primary",
            )
            render_kpi_row(
                [
                    ("Identity", st.session_state.saved_name or "Session default"),
                    ("Date stamp", "Enabled" if st.session_state.profile_include_date_input else "Off"),
                    ("Destination", "Local" if st.session_state.profile_save_destination_mode_input == "local" else "Browser"),
                ]
            )
            identity_left, identity_right = st.columns([1.35, 0.65], gap="small")
            with identity_left:
                st.text_input("Enter your name", key="profile_name_input", placeholder="Name to attach to exports")
            with identity_right:
                st.checkbox("Add date today", key="profile_include_date_input")

            with st.expander("Save Destination", expanded=True):
                st.radio(
                    "Save destination",
                    options=["browser", "local"],
                    format_func=lambda value: (
                        "Browser Download (Mobile + PC Recommended)"
                        if value == "browser"
                        else "Local Folder Picker (Desktop Only)"
                    ),
                    key="profile_save_destination_mode_input",
                    horizontal=True,
                )

                if st.session_state.profile_save_destination_mode_input == "browser":
                    render_route_block(
                        "Save behavior",
                        "Works on both mobile and PC.\nTap the download button and let your browser save the file to your device.",
                    )
                    render_route_block("Download route", browser_export_label())
                else:
                    folder_col, pick_col, clear_col = st.columns([1.7, 0.9, 0.8], gap="small")
                    folder_col.text_input(
                        "Selected local folder",
                        value=st.session_state.profile_export_root_path_input or "",
                        disabled=True,
                        placeholder="No local folder selected yet",
                    )
                    if pick_col.button(
                        "Choose Folder",
                        key="choose_export_folder",
                        use_container_width=True,
                        disabled=not picker_ready,
                    ):
                        try:
                            selected_folder = pick_local_folder()
                            if selected_folder:
                                queue_export_root_selection(selected_folder)
                                st.rerun()
                        except Exception as exc:
                            st.warning(str(exc))
                    if clear_col.button("Clear Folder", key="clear_export_folder", use_container_width=True):
                        queue_export_root_selection(None)
                        st.rerun()

                    if picker_ready:
                        render_route_block(
                            "Desktop save behavior",
                            "Exports save directly into the selected folder with automatic tool subfolders.",
                        )
                    else:
                        st.info("Local folder picking is not available here, so browser download is the cross-device fallback.")

            with st.expander("Output Options", expanded=False):
                opt_left, opt_right = st.columns(2, gap="small")
                with opt_left:
                    st.checkbox("Include saved name in export", key="profile_output_include_name_input")
                    st.checkbox("Include date in export", key="profile_output_include_date_input")
                with opt_right:
                    st.checkbox("Include essay heading suggestion", key="profile_essay_include_heading_input")
                    st.checkbox("Include self-check tip", key="profile_essay_include_tip_input")

            action_left, action_right = st.columns(2, gap="small")
            if action_left.button("Save Profile", use_container_width=True, type="primary", key="dashboard_save_profile"):
                save_profile()
                st.success("Profile saved for this session.")
            if action_right.button("Clear Saved Profile", use_container_width=True, key="dashboard_clear_profile"):
                queue_reset("clear_profile", "Profile cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Active Workspace Snapshot",
                "Track the current project workspace, uploaded materials, output template, and export routing in one place.",
                "Workspace Overview",
                anchor="WS",
                tier="secondary",
            )
            render_meta_grid(
                [
                    ("Workspace", str(workspace.get("name", "General Workspace"))),
                    ("Source files", str(len(workspace.get("files", [])))),
                    ("Screenshots", str(len(workspace.get("images", [])))),
                    ("Saved outputs", str(len(workspace.get("outputs", [])))),
                ]
            )
            render_route_block(
                "Save destination",
                paths["main"] if current_mode == "local" else "Browser download to the user's device",
            )
            render_route_block(
                "Current mode",
                "Local folder picker" if current_mode == "local" else "Browser download (mobile + PC)",
            )
            render_route_block("Selected template", str(st.session_state.get("export_template", "Academic Classic")))

    with st.expander("Additional dashboard panels", expanded=False):
        extra_a, extra_b, extra_c, extra_d = st.tabs(["Quick Launch", "Recent Activity", "Routing", "Outputs"])
        with extra_a:
            first_col, second_col = st.columns(2, gap="small")
            if first_col.button("Open Workspaces", key="dashboard_open_workspaces", use_container_width=True, type="primary"):
                go("workspaces")
            if second_col.button("Open Academics", key="dashboard_open_academics", use_container_width=True):
                go("academics")
            third_col, fourth_col = st.columns(2, gap="small")
            if third_col.button("Open Developer", key="dashboard_open_developer", use_container_width=True):
                go("developer")
            if fourth_col.button("Open Settings", key="dashboard_open_settings", use_container_width=True):
                go("settings")
        with extra_b:
            render_history_snippets(limit=4)
        with extra_c:
            render_route_block("Academics suite", "\n".join([paths["reviewer"], paths["flashcards"], paths["practice_test"], paths["rubric"], paths["batch"]]))
            render_route_block("Developer suite", paths["codefix"])
        with extra_d:
            render_workspace_outputs(limit=4)


def render_workspaces_page(ai_ready: bool) -> None:
    workspace = active_workspace()
    with st.container(border=True):
        render_card_header(
            PAGE_DETAILS["workspaces"]["title"],
            PAGE_DETAILS["workspaces"]["subtitle"],
            "Workspace Manager",
            anchor="WS",
            tier="primary",
        )
        render_tag_row(["Project Workspaces", "Multi-File Upload", "Screenshot Intake", "Workspace Analysis"])

    top_left, top_right = st.columns(2, gap="large")
    with top_left:
        with st.container(border=True):
            render_card_header(
                "Create Workspace",
                "Use dedicated workspaces to separate subjects, projects, and source libraries.",
                "Create",
            )
            st.text_input("Workspace name", key="workspace_new_name_input")
            st.text_area("Workspace description", key="workspace_new_description_input", height=110)
            if st.button("Create Workspace", key="create_workspace_button", use_container_width=True, type="primary"):
                if create_workspace(st.session_state.workspace_new_name_input, st.session_state.workspace_new_description_input):
                    st.session_state.workspace_new_name_input = ""
                    st.session_state.workspace_new_description_input = ""
                    st.success("Workspace created.")
                    st.rerun()
                else:
                    st.warning("Enter a workspace name first.")

    with top_right:
        with st.container(border=True):
            render_card_header(
                "Active Workspace",
                "Switch workspaces here, inspect current stats, or remove the current workspace.",
                "Current Context",
            )
            workspace_ids = list(st.session_state.workspaces.keys())
            selector = st.selectbox(
                "Choose workspace",
                options=workspace_ids,
                index=workspace_ids.index(st.session_state.active_workspace_id),
                format_func=workspace_option_label,
                key="workspace_page_selector",
            )
            if selector != st.session_state.active_workspace_id:
                st.session_state.active_workspace_id = selector
                st.rerun()
            workspace = active_workspace()
            render_meta_grid(
                [
                    ("Created", str(workspace.get("created_at", "-"))),
                    ("Files", str(len(workspace.get("files", [])))),
                    ("Images", str(len(workspace.get("images", [])))),
                    ("Words", str(workspace_word_count(workspace))),
                ]
            )
            if len(st.session_state.workspaces) > 1 and st.button("Delete Active Workspace", key="delete_workspace_button", use_container_width=True):
                if delete_active_workspace():
                    st.success("Workspace deleted.")
                    st.rerun()

    mid_left, mid_right = st.columns(2, gap="large")
    with mid_left:
        with st.container(border=True):
            render_card_header(
                "Source Library",
                "Upload multiple text documents and add them to the active workspace in one step.",
                "Files",
            )
            uploaded_files = st.file_uploader(
                "Upload .docx, .pdf, .txt, or .md files",
                type=["docx", "pdf", "txt", "md"],
                accept_multiple_files=True,
                key="workspace_file_upload",
            )
            if st.button("Add Files To Workspace", key="add_workspace_files", use_container_width=True, type="primary"):
                if uploaded_files:
                    added, failures = add_source_files_to_active_workspace(uploaded_files)
                    if added:
                        st.success(f"Added {added} file(s) to {active_workspace_name()}.")
                    for failure in failures:
                        st.warning(failure)
                    st.rerun()
                else:
                    st.warning("Upload at least one file first.")
            file_names = [str(item.get("name", "")) for item in workspace.get("files", [])]
            st.caption("Current files: " + (", ".join(file_names[:8]) if file_names else "No files added yet."))

    with mid_right:
        with st.container(border=True):
            render_card_header(
                "Workspace Notes",
                "Keep manual notes, instructions, or OCR fallbacks attached to the current workspace.",
                "Notes",
            )
            notes_value = st.text_area(
                "Workspace notes",
                value=str(workspace.get("notes", "")),
                height=200,
                key=f"workspace_notes_editor_{workspace['id']}",
            )
            if st.button("Save Workspace Notes", key=f"save_workspace_notes_{workspace['id']}", use_container_width=True):
                workspace["notes"] = notes_value.strip()
                append_history_entry("workspace", f"Updated notes for {workspace['name']}", "Workspace notes saved.", "workspaces", str(workspace["name"]))
                st.success("Workspace notes saved.")

    bottom_left, bottom_right = st.columns(2, gap="large")
    with bottom_left:
        with st.container(border=True):
            render_card_header(
                "Screenshot Intake",
                "Attach screenshots or photos to the active workspace. OCR is not enabled yet, but the image record stays with the workspace.",
                "Images",
            )
            image_caption = st.text_input("Optional image note", key="workspace_image_caption")
            uploaded_images = st.file_uploader(
                "Upload screenshots",
                type=["png", "jpg", "jpeg", "webp"],
                accept_multiple_files=True,
                key="workspace_image_upload",
            )
            if st.button("Add Screenshots", key="add_workspace_images", use_container_width=True):
                if uploaded_images:
                    added = add_images_to_active_workspace(uploaded_images, image_caption)
                    st.success(f"Added {added} screenshot(s) to {active_workspace_name()}.")
                    st.rerun()
                else:
                    st.warning("Upload at least one image first.")
            with st.expander("Preview attached screenshots", expanded=False):
                if workspace.get("images"):
                    preview_images = [item.get("bytes") for item in workspace.get("images", [])[:3]]
                    st.image(preview_images, width=180)
                else:
                    st.caption("No screenshots attached yet.")

    with bottom_right:
        with st.container(border=True):
            render_card_header(
                "Workspace Analysis",
                "Generate a quick analysis over all current sources and notes inside the active workspace.",
                "Analysis",
            )
            source_bundle = workspace_source_bundle(workspace)
            st.caption(f"Prompt-ready source length: {len(source_bundle)} characters")
            if st.button(
                "Analyze Workspace Sources",
                key="analyze_workspace_sources",
                use_container_width=True,
                type="primary",
                disabled=not ai_ready,
            ):
                if not source_bundle:
                    st.warning("Add files, screenshots, or notes to the workspace first.")
                else:
                    prompt = (
                        "Analyze the study workspace sources below.\n\n"
                        f"Workspace: {workspace['name']}\n"
                        f"Description: {workspace.get('description', '') or 'No description provided'}\n\n"
                        f"Sources:\n{trim_prompt_source(source_bundle)}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Summary:\n"
                        "Key Concepts:\n"
                        "Likely Focus Areas:\n"
                        "Study Advice:\n\n"
                        "Do not use markdown symbols like ### or ####."
                    )
                    result = run_generation(prompt, "workspace analysis")
                    if result:
                        st.session_state.workspace_analysis_result = result
                        remember_output("Workspace Analysis", f"{workspace['name']} Workspace Analysis", result, "workspaces", "workspace_analysis")
                        st.success("Workspace analysis ready.")
            if st.session_state.workspace_analysis_result.strip():
                st.text_area("Analysis output", value=st.session_state.workspace_analysis_result, height=260, disabled=True)
            else:
                render_route_block(
                    "No analysis yet",
                    "Generate the workspace analysis first. This panel stays compact until there is a result to review.",
                )
            if st.session_state.workspace_analysis_result.strip():
                render_download_button(
                    f"{workspace['name']} Workspace Analysis",
                    st.session_state.workspace_analysis_result,
                    sanitize_filename(f"{workspace['name']}_workspace_analysis".lower().replace(" ", "_")),
                    "workspace_analysis",
                    "clear_workspace_analysis_workspace",
                )
    with st.expander("Additional workspace panels", expanded=False):
        list_left, list_right = st.columns(2, gap="large")
        with list_left:
            with st.container(border=True):
                render_card_header(
                    "Source Inventory",
                    "A quick list of the file assets currently attached to the active workspace.",
                    "Inventory",
                    compact=True,
                )
                if workspace.get("files"):
                    for file_entry in workspace.get("files", [])[:10]:
                        render_route_block(
                            str(file_entry.get("name", "source.txt")),
                            f"{file_entry.get('words', 0)} words • added {file_entry.get('added_at', '-')}",
                        )
                else:
                    st.caption("No source files in this workspace yet.")

        with list_right:
            with st.container(border=True):
                render_card_header(
                    "Workspace Outputs",
                    "Generated outputs remain attached to the active workspace for quick reuse.",
                    "Outputs",
                    compact=True,
                )
                render_workspace_outputs(limit=6)


def render_source_lab_tab(ai_ready: bool, workspace: dict[str, object], source_bundle: str) -> None:
    src_left, src_right = st.columns(2, gap="large")
    with src_left:
        with st.container(border=True):
            render_card_header(
                "Multi-File Source Lab",
                "Add more source files directly from this page or review the current active source bundle before asking grounded questions.",
                "Sources",
            )
            inline_files = st.file_uploader(
                "Add more files to the active workspace",
                type=["docx", "pdf", "txt", "md"],
                accept_multiple_files=True,
                key="source_lab_inline_upload",
            )
            if st.button("Add Files To Active Workspace", key="source_lab_add_files", use_container_width=True):
                if inline_files:
                    added, failures = add_source_files_to_active_workspace(inline_files)
                    if added:
                        st.success(f"Added {added} file(s) to {active_workspace_name()}.")
                    for failure in failures:
                        st.warning(failure)
                    st.rerun()
                else:
                    st.warning("Upload at least one file first.")
            with st.expander("Review merged source preview", expanded=False):
                if source_bundle.strip():
                    st.text_area("Merged source preview", value=source_bundle, height=280, disabled=True)
                else:
                    render_route_block(
                        "No source bundle yet",
                        "Add files, notes, or screenshots to the active workspace first, then the merged source preview will appear here.",
                    )

    with src_right:
        with st.container(border=True):
            render_preview_panel(
                "Source-Grounded Answers",
                "Generate a grounded analysis with explicit file citations based on the active workspace sources.",
                "Grounded analysis output",
                st.session_state.source_lab_analysis,
                height=310,
                empty_title="No grounded analysis yet",
                empty_body="Run the grounded analysis action after adding workspace material. This preview will stay compact until there is a real result to review.",
                anchor="GA",
                tier="secondary",
            )
            if st.button(
                "Generate Grounded Analysis",
                key="source_lab_analyze",
                use_container_width=True,
                type="primary",
                disabled=not ai_ready,
            ):
                if not source_bundle:
                    st.warning("Add source files or notes to the active workspace first.")
                else:
                    prompt = (
                        "Analyze the provided sources and keep the response grounded in them.\n\n"
                        f"Sources:\n{trim_prompt_source(source_bundle)}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Summary:\n"
                        "Important Details:\n"
                        "Source Citations:\n\n"
                        "In Source Citations, cite files using the form [Source: filename]."
                    )
                    result = run_generation(prompt, "grounded analysis")
                    if result:
                        st.session_state.source_lab_analysis = result
                        remember_output("Grounded Analysis", f"{workspace['name']} Grounded Analysis", result, "academics", "grounded_answer")
                        st.success("Grounded analysis ready.")
            if st.session_state.source_lab_analysis.strip():
                render_download_button(
                    f"{workspace['name']} Grounded Analysis",
                    st.session_state.source_lab_analysis,
                    sanitize_filename(f"{workspace['name']}_grounded_analysis".lower().replace(" ", "_")),
                    "grounded_answer",
                    "clear_source_lab_workspace",
                )

    qa_left, qa_right = st.columns(2, gap="large")
    with qa_left:
        with st.container(border=True):
            render_card_header(
                "Document Q&A Mode",
                "Ask questions directly against the active workspace and require file-based citations in the answer.",
                "Q&A",
            )
            st.text_area("Question", key="source_lab_question_input", height=140)
            if st.button(
                "Answer With Citations",
                key="source_lab_answer_question",
                use_container_width=True,
                type="primary",
                disabled=not ai_ready,
            ):
                question = st.session_state.source_lab_question_input.strip()
                if not question:
                    st.warning("Enter a question first.")
                elif not source_bundle:
                    st.warning("Add source material to the active workspace first.")
                else:
                    prompt = (
                        "Answer the user's question using only the provided sources.\n"
                        "If the answer is not supported by the sources, say so clearly.\n\n"
                        f"Question: {question}\n\n"
                        f"Sources:\n{trim_prompt_source(source_bundle)}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Answer:\n"
                        "Cited Sources:\n"
                        "Confidence:\n\n"
                        "Use the form [Source: filename] in Cited Sources."
                    )
                    result = run_generation(prompt, "document Q&A answer")
                    if result:
                        st.session_state.source_lab_answer = result
                        remember_output("Document Q&A", f"{workspace['name']} Q&A", result, "academics", "documentqa")
                        st.success("Document Q&A answer ready.")
            if st.session_state.source_lab_answer.strip():
                st.text_area("Q&A answer", value=st.session_state.source_lab_answer, height=250, disabled=True)
            else:
                render_route_block(
                    "No answer yet",
                    "Enter a source-grounded question first. The answer preview will only expand once there is a result.",
                )
            if st.session_state.source_lab_answer.strip():
                render_download_button(
                    f"{workspace['name']} Document Answer",
                    st.session_state.source_lab_answer,
                    sanitize_filename(f"{workspace['name']}_document_answer".lower().replace(" ", "_")),
                    "documentqa",
                    "clear_source_lab_workspace",
                )

    with qa_right:
        with st.container(border=True):
            render_card_header(
                "Screenshot Input",
                "Screenshots stay attached to the active workspace. Use workspace notes to supply manual OCR text or important context.",
                "Image Input",
            )
            if workspace.get("images"):
                st.image([item.get("bytes") for item in workspace.get("images", [])[:4]], width=180)
                st.caption("Attached screenshots are stored with the active workspace.")
            else:
                st.caption("No screenshots attached yet. Add them from Workspaces.")
            if st.button("Open Workspaces", key="source_lab_open_workspaces", use_container_width=True):
                go("workspaces")


def render_study_tools_tab(ai_ready: bool, workspace: dict[str, object], source_bundle: str) -> None:
    reviewer_tab, flashcard_tab, test_tab, checker_tab = st.tabs(
        ["Reviewer Generator", "Flashcard Generator", "Practice Test", "Answer Checker"]
    )

    with reviewer_tab:
        left, right = st.columns(2, gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Reviewer Generator",
                    "Turn the active workspace or a custom request into structured review notes.",
                    "Study Tool",
                )
                st.text_area("Reviewer focus or instructions", key="reviewer_request_input", height=170)
                if st.button("Generate Reviewer", key="generate_reviewer", use_container_width=True, type="primary", disabled=not ai_ready):
                    prompt = (
                        "Create reviewer notes for the provided study materials.\n\n"
                        f"Reviewer request: {st.session_state.reviewer_request_input.strip() or 'Use the active workspace as-is.'}\n\n"
                        f"Sources:\n{trim_prompt_source(source_bundle or workspace.get('notes', ''))}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Summary:\n"
                        "Key Terms:\n"
                        "Important Details:\n"
                        "Quick Recall Questions:\n"
                    )
                    result = run_generation(prompt, "reviewer notes")
                    if result:
                        st.session_state.reviewer_response = result
                        remember_output("Reviewer Notes", f"{workspace['name']} Reviewer", result, "academics", "reviewer")
                        st.success("Reviewer notes ready.")
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Reviewer Preview",
                    "Review the notes before exporting.",
                    "Reviewer output",
                    st.session_state.reviewer_response,
                    height=340,
                    empty_title="No reviewer notes yet",
                    empty_body="Generate reviewer notes from the left panel. The preview appears only once there is content worth reviewing.",
                    anchor="RV",
                    tier="secondary",
                )
                if st.session_state.reviewer_response.strip():
                    render_download_button(
                        f"{workspace['name']} Reviewer Notes",
                        st.session_state.reviewer_response,
                        sanitize_filename(f"{workspace['name']}_reviewer_notes".lower().replace(" ", "_")),
                        "reviewer",
                        "clear_reviewer_workspace",
                    )

    with flashcard_tab:
        left, right = st.columns(2, gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Flashcard Generator",
                    "Generate concise question-and-answer cards from the active workspace or a custom request.",
                    "Study Tool",
                )
                st.text_area("Flashcard focus or instructions", key="flashcard_request_input", height=170)
                if st.button("Generate Flashcards", key="generate_flashcards", use_container_width=True, type="primary", disabled=not ai_ready):
                    prompt = (
                        "Create flashcards from the provided study materials.\n\n"
                        f"Flashcard request: {st.session_state.flashcard_request_input.strip() or 'Use the active workspace as-is.'}\n\n"
                        f"Sources:\n{trim_prompt_source(source_bundle or workspace.get('notes', ''))}\n\n"
                        "Return a numbered list of flashcards in this pattern:\n"
                        "Q:\n"
                        "A:\n\n"
                        "Keep each answer short and review-friendly."
                    )
                    result = run_generation(prompt, "flashcards")
                    if result:
                        st.session_state.flashcard_response = result
                        remember_output("Flashcards", f"{workspace['name']} Flashcards", result, "academics", "flashcards")
                        st.success("Flashcards ready.")
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Flashcard Preview",
                    "Review the generated cards before export.",
                    "Flashcard output",
                    st.session_state.flashcard_response,
                    height=340,
                    empty_title="No flashcards yet",
                    empty_body="Generate flashcards from the left panel. The preview will appear here once the deck is ready.",
                    anchor="FC",
                    tier="secondary",
                )
                if st.session_state.flashcard_response.strip():
                    render_download_button(
                        f"{workspace['name']} Flashcards",
                        st.session_state.flashcard_response,
                        sanitize_filename(f"{workspace['name']}_flashcards".lower().replace(" ", "_")),
                        "flashcards",
                        "clear_flashcard_workspace",
                    )

    with test_tab:
        left, right = st.columns(2, gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Practice Test Generator",
                    "Generate multiple-choice, true/false, and short-answer practice items from the active workspace.",
                    "Study Tool",
                )
                st.text_area("Practice test focus or instructions", key="practice_test_request_input", height=170)
                if st.button("Generate Practice Test", key="generate_practice_test", use_container_width=True, type="primary", disabled=not ai_ready):
                    prompt = (
                        "Create a study practice test using the provided materials.\n\n"
                        f"Practice test request: {st.session_state.practice_test_request_input.strip() or 'Use the active workspace as-is.'}\n\n"
                        f"Sources:\n{trim_prompt_source(source_bundle or workspace.get('notes', ''))}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Multiple Choice:\n"
                        "True or False:\n"
                        "Short Answer:\n"
                        "Answer Key:\n"
                    )
                    result = run_generation(prompt, "practice test")
                    if result:
                        st.session_state.practice_test_response = result
                        remember_output("Practice Test", f"{workspace['name']} Practice Test", result, "academics", "practice_test")
                        st.success("Practice test ready.")
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Practice Test Preview",
                    "Review the generated practice test before export.",
                    "Practice test output",
                    st.session_state.practice_test_response,
                    height=340,
                    empty_title="No practice test yet",
                    empty_body="Generate the practice test first. The output preview and export controls will appear here after that.",
                    anchor="PT",
                    tier="secondary",
                )
                if st.session_state.practice_test_response.strip():
                    render_download_button(
                        f"{workspace['name']} Practice Test",
                        st.session_state.practice_test_response,
                        sanitize_filename(f"{workspace['name']}_practice_test".lower().replace(" ", "_")),
                        "practice_test",
                        "clear_practice_test_workspace",
                    )

    with checker_tab:
        left, right = st.columns(2, gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Answer Checker",
                    "Compare a student answer against the question, optional reference notes, and the active workspace sources.",
                    "Study Tool",
                )
                st.text_area("Question", key="answer_checker_question_input", height=110)
                st.text_area("Reference notes (optional)", key="answer_checker_reference_input", height=110)
                st.text_area("Your answer", key="answer_checker_user_answer_input", height=140)
                if st.button("Check Answer", key="check_answer_button", use_container_width=True, type="primary", disabled=not ai_ready):
                    question = st.session_state.answer_checker_question_input.strip()
                    user_answer = st.session_state.answer_checker_user_answer_input.strip()
                    if not question or not user_answer:
                        st.warning("Enter both the question and your answer first.")
                    else:
                        prompt = (
                            "Evaluate the student's answer using the provided references and source material.\n\n"
                            f"Question:\n{question}\n\n"
                            f"Reference notes:\n{st.session_state.answer_checker_reference_input.strip() or 'No separate reference notes provided.'}\n\n"
                            f"Active workspace sources:\n{trim_prompt_source(source_bundle or workspace.get('notes', ''))}\n\n"
                            f"Student answer:\n{user_answer}\n\n"
                            "Return these plain-text sections in this exact order:\n"
                            "Score:\n"
                            "What Works:\n"
                            "Missing or Weak Parts:\n"
                            "Improved Answer:\n"
                        )
                        result = run_generation(prompt, "answer check")
                        if result:
                            st.session_state.answer_checker_response = result
                            remember_output("Answer Check", f"{workspace['name']} Answer Check", result, "academics", "answer_checker")
                            st.success("Answer review ready.")
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Answer Checker Preview",
                    "Review the score, gaps, and improved answer.",
                    "Answer checker output",
                    st.session_state.answer_checker_response,
                    height=390,
                    empty_title="No answer review yet",
                    empty_body="Submit a question and a student answer first. The review panel stays lightweight until there is a scored response.",
                    anchor="AN",
                    tier="secondary",
                )
                if st.session_state.answer_checker_response.strip():
                    render_download_button(
                        f"{workspace['name']} Answer Check",
                        st.session_state.answer_checker_response,
                        sanitize_filename(f"{workspace['name']}_answer_check".lower().replace(" ", "_")),
                        "answer_checker",
                        "clear_answer_checker_workspace",
                    )


def render_writing_studio_tab(ai_ready: bool, workspace: dict[str, object], source_bundle: str) -> None:
    rubric_tab, batch_tab, legacy_tab = st.tabs(["Rubric Mode", "Batch Output Generator", "Legacy Builders"])

    with rubric_tab:
        left, right = st.columns(2, gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Rubric Mode",
                    "Generate writing aligned to selected rubric criteria while still using the active workspace as reference.",
                    "Writing Studio",
                )
                st.text_input("Writing title", key="rubric_title_input")
                st.text_area("Writing request", key="rubric_prompt_input", height=170)
                st.multiselect(
                    "Rubric focus",
                    options=["Clarity", "Structure", "Depth", "Grammar", "Critical Thinking", "Evidence"],
                    key="rubric_focus_input",
                )
                if st.button("Generate Rubric-Aligned Draft", key="generate_rubric_draft", use_container_width=True, type="primary", disabled=not ai_ready):
                    title = st.session_state.rubric_title_input.strip()
                    request = st.session_state.rubric_prompt_input.strip()
                    focus = ", ".join(st.session_state.rubric_focus_input) or "Clarity, Structure, Depth"
                    if not title or not request:
                        st.warning("Enter both the title and writing request first.")
                    else:
                        prompt = (
                            "Write a structured academic draft aligned to the selected rubric.\n\n"
                            f"Title: {title}\n"
                            f"Rubric focus: {focus}\n"
                            f"Request: {request}\n\n"
                            f"Source material:\n{trim_prompt_source(source_bundle or workspace.get('notes', ''))}\n\n"
                            "Return these plain-text sections in this exact order:\n"
                            "Title Suggestion:\n"
                            "Draft:\n"
                            "Rubric Alignment Notes:\n"
                        )
                        result = run_generation(prompt, "rubric-aligned draft")
                        if result:
                            st.session_state.rubric_response = result
                            remember_output("Rubric Draft", f"{title} Rubric Draft", result, "academics", "rubric")
                            st.success("Rubric-aligned draft ready.")
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Rubric Draft Preview",
                    "Review the aligned draft and rubric notes before export.",
                    "Rubric mode output",
                    st.session_state.rubric_response,
                    height=390,
                    empty_title="No rubric draft yet",
                    empty_body="Generate a rubric-aligned draft from the left panel. The review and export controls will appear here once the draft is ready.",
                    anchor="RB",
                    tier="secondary",
                )
                if st.session_state.rubric_response.strip():
                    render_download_button(
                        st.session_state.rubric_title_input.strip() or f"{workspace['name']} Rubric Draft",
                        st.session_state.rubric_response,
                        sanitize_filename((st.session_state.rubric_title_input.strip() or f"{workspace['name']}_rubric_draft").lower().replace(" ", "_")),
                        "rubric",
                        "clear_rubric_workspace",
                    )

    with batch_tab:
        left, right = st.columns(2, gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Batch Output Generator",
                    "Generate multiple deliverables from one topic or one workspace source bundle in a single run.",
                    "Writing Studio",
                )
                st.text_input("Batch topic", key="batch_topic_input")
                st.text_area("Batch instructions", key="batch_request_input", height=170)
                if st.button("Generate Batch Outputs", key="generate_batch_outputs", use_container_width=True, type="primary", disabled=not ai_ready):
                    topic = st.session_state.batch_topic_input.strip() or str(workspace.get("name", "Workspace"))
                    request = st.session_state.batch_request_input.strip() or "Generate a study summary, reviewer, flashcards, and essay outline."
                    prompt = (
                        "Create multiple study outputs from the same source material.\n\n"
                        f"Topic: {topic}\n"
                        f"Request: {request}\n\n"
                        f"Source material:\n{trim_prompt_source(source_bundle or workspace.get('notes', ''))}\n\n"
                        "Return these plain-text sections in this exact order:\n"
                        "Summary:\n"
                        "Reviewer Notes:\n"
                        "Flashcards:\n"
                        "Essay Outline:\n"
                    )
                    result = run_generation(prompt, "batch output package")
                    if result:
                        st.session_state.batch_response = result
                        remember_output("Batch Output", f"{topic} Batch Output", result, "academics", "batch")
                        st.success("Batch output ready.")
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Batch Output Preview",
                    "Review the combined output package before export.",
                    "Batch output",
                    st.session_state.batch_response,
                    height=390,
                    empty_title="No batch output yet",
                    empty_body="Generate the batch package from the left panel. The preview stays compact until the combined output is ready.",
                    anchor="BT",
                    tier="secondary",
                )
                if st.session_state.batch_response.strip():
                    render_download_button(
                        st.session_state.batch_topic_input.strip() or f"{workspace['name']} Batch Output",
                        st.session_state.batch_response,
                        sanitize_filename((st.session_state.batch_topic_input.strip() or f"{workspace['name']}_batch_output").lower().replace(" ", "_")),
                        "batch",
                        "clear_batch_workspace",
                    )

    with legacy_tab:
        with st.container(border=True):
            render_card_header(
                "Legacy Builders",
                "The older dedicated pages still exist for users who prefer the original single-tool workflows.",
                "Compatibility",
            )
            row_a = st.columns(3, gap="small")
            if row_a[0].button("Open Quiz Solver", key="legacy_quiz_open", use_container_width=True):
                go("quiz")
            if row_a[1].button("Open Assignment Solver", key="legacy_assignment_open", use_container_width=True):
                go("assignment")
            if row_a[2].button("Open Essay Builder", key="legacy_essay_open", use_container_width=True):
                go("essay")
            row_b = st.columns(2, gap="small")
            if row_b[0].button("Open Activity Builder", key="legacy_activity_open", use_container_width=True):
                go("activity")
            if row_b[1].button("Open Document Builder", key="legacy_document_open", use_container_width=True):
                go("document")


def render_export_center_tab() -> None:
    left, right = st.columns(2, gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Professional Export Templates",
                "Choose the document style used by all Word exports generated in this session.",
                "Export Center",
            )
            st.radio(
                "Export template",
                options=list(EXPORT_TEMPLATES.keys()),
                key="export_template",
                horizontal=False,
            )
            template = EXPORT_TEMPLATES[st.session_state.export_template]
            render_meta_grid(
                [
                    ("Font", str(template["font"])),
                    ("Body size", str(template["body_size"])),
                    ("Title size", str(template["title_size"])),
                    ("Line spacing", str(template["line_spacing"])),
                ]
            )
    with right:
        with st.container(border=True):
            render_card_header(
                "Export Routes",
                "Local saves and browser downloads both use the same current export template.",
                "Routing",
            )
            paths = folder_path_lines()
            with st.expander("Open route list", expanded=False):
                render_route_block("Reviewer", paths["reviewer"])
                render_route_block("Flashcards", paths["flashcards"])
                render_route_block("Practice Test", paths["practice_test"])
                render_route_block("Rubric Draft", paths["rubric"])
                render_route_block("Batch Output", paths["batch"])


def render_academics_hub(ai_ready: bool) -> None:
    workspace = active_workspace()
    source_bundle = workspace_source_bundle(workspace)
    with st.container(border=True):
        render_card_header(
            PAGE_DETAILS["academics"]["title"],
            "Use the active workspace as the source of truth, then branch into one focused study flow at a time.",
            "Academics Suite",
            anchor="AC",
            tier="primary",
        )
        render_kpi_row(
            [
                ("Workspace", str(workspace.get("name", "General Workspace"))),
                ("Source files", str(len(workspace.get("files", [])))),
                ("Images", str(len(workspace.get("images", [])))),
                ("Template", str(st.session_state.get("export_template", "Academic Classic"))),
            ]
        )
        with st.expander("Workspace context", expanded=False):
            render_route_block(
                "Current source bundle",
                f"{len(source_bundle)} characters ready for prompting" if source_bundle else "No files or notes have been attached to this workspace yet.",
            )

    source_tab, study_tab, writing_tab, export_tab = st.tabs(
        ["Source Lab", "Study Tools", "Writing Studio", "Export Center"]
    )

    with source_tab:
        render_source_lab_tab(ai_ready, workspace, source_bundle)
    with study_tab:
        render_study_tools_tab(ai_ready, workspace, source_bundle)
    with writing_tab:
        render_writing_studio_tab(ai_ready, workspace, source_bundle)
    with export_tab:
        render_export_center_tab()


def render_developer_hub(ai_ready: bool) -> None:
    stack_profile = developer_stack_profile()
    with st.container(border=True):
        render_card_header(
            PAGE_DETAILS["developer"]["title"],
            "Keep code generation and debugging separate from the academic tools, with one focused workflow visible at a time.",
            "Developer Suite",
            anchor="DV",
            tier="primary",
        )
        render_kpi_row(
            [
                ("Active model", current_model_name() or "No active model"),
                ("Stack profile", str(stack_profile["profile_name"])),
                ("Code focus", "High" if stack_profile["is_coder_model"] else "Moderate"),
            ]
        )
        with st.expander("Capability context", expanded=False):
            st.caption(stack_profile["note"])

    fix_tab, generator_tab, stack_tab, selftest_tab, compare_tab, diff_tab, notes_tab = st.tabs(
        ["Code Fixer", "Code Generator", "Stack Guide", "Self-Test", "Model Compare", "Code Diff View", "Debug Notes"]
    )

    with fix_tab:
        selected_fix_stack, selected_fix_level = resolve_stack_choice(
            st.session_state.codefix_stack_profile,
            st.session_state.codefix_custom_stack,
        )
        left, right = st.columns([1.02, 0.98], gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Code Fix Builder",
                    "Paste the failing code, choose the target stack, and add the error or symptom. The stack list is filtered for what the current model is more likely to handle well.",
                    "Compose",
                )
                st.text_input("Issue title (optional)", key="codefix_title")
                st.selectbox(
                    "Tech stack",
                    options=stack_profile["ordered_options"],
                    key="codefix_stack_profile",
                )
                if st.session_state.codefix_stack_profile == "Custom":
                    st.text_input("Custom stack", key="codefix_custom_stack")
                render_route_block("Stack fit", f"{selected_fix_stack}\n{stack_confidence_message(selected_fix_level)}")
                st.text_area("Error message or symptoms", key="codefix_error", height=120)
                st.text_area("Code snippet", key="codefix_source", height=220)
                st.text_area("Expected behavior (optional)", key="codefix_expectation", height=95)
                action_a, action_b = st.columns(2, gap="small")
                if action_a.button("Generate Code Fix", key="developer_codefix_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                    title = st.session_state.codefix_title.strip()
                    language = selected_fix_stack or "Code"
                    error_text = st.session_state.codefix_error.strip()
                    source = st.session_state.codefix_source.strip()
                    expectation = st.session_state.codefix_expectation.strip()
                    st.session_state.codefix_language = language
                    if not source:
                        st.warning("Paste the code snippet before generating a fix.")
                    elif not error_text:
                        st.warning("Add the error message or symptoms before generating a fix.")
                    else:
                        prompt = (
                            "Fix the following code issue for study and debugging support.\n\n"
                            f"Language or stack: {language}\n"
                            f"Stack confidence level: {selected_fix_level}\n"
                            f"Issue title: {title or 'No title provided'}\n"
                            f"Error message or symptoms:\n{error_text}\n\n"
                            f"Code snippet:\n{source}\n\n"
                            f"Expected behavior:\n{expectation or 'Not provided'}\n\n"
                            "Return these plain-text sections in this exact order:\n"
                            "Issue Summary:\n"
                            "Root Cause:\n"
                            "Fixed Version:\n"
                            "Why It Works:\n"
                            "Next Checks:\n"
                        )
                        result = run_generation(prompt, "code fix")
                        if result:
                            st.session_state.codefix_response = result
                            remember_output("Code Fix", title or f"{language} Code Fix", result, "developer", "codefix")
                            st.success("Code fix ready.")
                if action_b.button("Clear Code Fixer", key="developer_codefix_clear", use_container_width=True):
                    queue_reset("clear_codefix_workspace", "Code fixer workspace cleared.")
                    st.rerun()
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Code Fix Preview",
                    "The fix, explanation, and next checks appear here.",
                    "Fixed result",
                    st.session_state.codefix_response,
                    height=430,
                    empty_title="No code fix yet",
                    empty_body="Generate the fix from the left panel first. The preview and export controls will appear only when there is a real debugging result to inspect.",
                    anchor="FX",
                    tier="secondary",
                )
                if st.session_state.codefix_response.strip():
                    export_title = st.session_state.codefix_title.strip() or f"{selected_fix_stack or 'Code'} Code Fix"
                    render_download_button(
                        export_title,
                        st.session_state.codefix_response,
                        sanitize_filename(export_title.lower().replace(" ", "_")),
                        "codefix",
                        "clear_codefix_workspace",
                    )

    with generator_tab:
        selected_codegen_stack, selected_codegen_level = resolve_stack_choice(
            st.session_state.codegen_stack_profile,
            st.session_state.codegen_custom_stack,
        )
        left, right = st.columns([1.02, 0.98], gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Code Generator",
                    "Describe what you want built, or upload reference documents and notes so the model can infer the target without a long manual description.",
                    "Compose",
                )
                st.text_input("Project or file title", key="codegen_title")
                st.selectbox(
                    "Target stack",
                    options=stack_profile["ordered_options"],
                    key="codegen_stack_profile",
                )
                if st.session_state.codegen_stack_profile == "Custom":
                    st.text_input("Custom target stack", key="codegen_custom_stack")
                render_route_block("Stack fit", f"{selected_codegen_stack}\n{stack_confidence_message(selected_codegen_level)}")
                st.text_area("Build description (optional if references already explain the task)", key="codegen_description", height=130)
                st.text_area("Expected output or constraints", key="codegen_expectation", height=100)
                st.file_uploader(
                    "Reference documents",
                    type=["docx", "pdf", "txt", "md"],
                    accept_multiple_files=True,
                    key="codegen_reference_docs",
                )
                st.file_uploader(
                    "Reference images",
                    type=["png", "jpg", "jpeg", "webp"],
                    accept_multiple_files=True,
                    key="codegen_reference_images",
                )
                st.text_area(
                    "Attachment note (use this if the images need extra explanation)",
                    key="codegen_attachment_note",
                    height=88,
                )
                action_a, action_b = st.columns(2, gap="small")
                if action_a.button("Generate Code", key="developer_codegen_generate", use_container_width=True, type="primary", disabled=not ai_ready):
                    title = st.session_state.codegen_title.strip() or f"{selected_codegen_stack} Generator"
                    description = st.session_state.codegen_description.strip()
                    expectation = st.session_state.codegen_expectation.strip()
                    doc_files = st.session_state.get("codegen_reference_docs") or []
                    image_files = st.session_state.get("codegen_reference_images") or []
                    reference_bundle, issues = read_codegen_reference_bundle(
                        doc_files,
                        image_files,
                        st.session_state.codegen_attachment_note,
                    )
                    for issue in issues:
                        st.warning(issue)
                    if not description and not reference_bundle:
                        st.warning("Add a build description or upload at least one reference file.")
                    else:
                        prompt = (
                            "Generate code or project scaffolding for the user's request.\n\n"
                            f"Target stack: {selected_codegen_stack}\n"
                            f"Stack confidence level: {selected_codegen_level}\n"
                            f"Project title: {title}\n"
                            f"Build description:\n{description or 'Infer the goal from the attached references.'}\n\n"
                            f"Expected output or constraints:\n{expectation or 'No extra constraints provided.'}\n\n"
                            f"Reference material:\n{trim_prompt_source(reference_bundle) if reference_bundle else 'No attachments provided.'}\n\n"
                            "Return these plain-text sections in this exact order:\n"
                            "Solution Summary:\n"
                            "Generated Code:\n"
                            "Setup Notes:\n"
                            "Next Steps:\n"
                        )
                        result = run_generation(prompt, "code generator output")
                        if result:
                            st.session_state.codegen_response = result
                            remember_output("Code Generator", title, result, "developer", "codegen")
                            st.success("Code output ready.")
                if action_b.button("Clear Code Generator", key="developer_codegen_clear", use_container_width=True):
                    queue_reset("clear_codegen_workspace", "Code generator cleared.")
                    st.rerun()
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Generated Code Preview",
                    "Review the generated code and notes before exporting or re-prompting.",
                    "Generated code output",
                    st.session_state.codegen_response,
                    height=430,
                    empty_title="No generated code yet",
                    empty_body="Describe the build or attach reference material first. The preview area expands when there is generated code to review.",
                    anchor="CG",
                    tier="secondary",
                )
                if st.session_state.codegen_response.strip():
                    export_title = st.session_state.codegen_title.strip() or f"{selected_codegen_stack} Code Output"
                    render_download_button(
                        export_title,
                        st.session_state.codegen_response,
                        sanitize_filename(export_title.lower().replace(" ", "_")),
                        "codegen",
                        "clear_codegen_workspace",
                    )

    with stack_tab:
        with st.container(border=True):
            render_card_header(
                "Stack Fit Guide",
                "This guide keeps the developer tools honest about what the current model is more likely to handle well.",
                "Capability Profile",
            )
            render_meta_grid(
                [
                    ("Active model", current_model_name() or "No active model"),
                    ("Profile", stack_profile["profile_name"]),
                    ("Code focus", "High" if stack_profile["is_coder_model"] else "Moderate"),
                ]
            )
            st.caption(stack_profile["note"])
            render_route_block("Best fit", "\n".join(stack_profile["best_fit"]))
            render_route_block("Good with review", "\n".join(stack_profile["good_fit"]))
            render_route_block("Manual review strongly advised", "\n".join(stack_profile["low_fit"]))
            st.caption("Reference files are parsed for text documents. Image attachments are kept as references, but this current setup does not perform OCR.")

    with selftest_tab:
        selected_selftest_stack, _selftest_level = resolve_stack_choice(
            st.session_state.selftest_stack_profile,
            st.session_state.selftest_custom_stack,
        )
        left, right = st.columns([1.02, 0.98], gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Developer Self-Test",
                    "Run a standard coding prompt against the active model for the selected stack so we can check whether the model behaves well before trusting it on a real task.",
                    "Verification",
                )
                st.selectbox("Self-test stack", options=stack_profile["ordered_options"], key="selftest_stack_profile")
                if st.session_state.selftest_stack_profile == "Custom":
                    st.text_input("Custom self-test stack", key="selftest_custom_stack")
                render_route_block("Current test target", f"{selected_selftest_stack}\nModel: {current_model_name() or 'No active model'}")
                if st.button("Run Self-Test", key="run_developer_selftest", use_container_width=True, type="primary", disabled=not ai_ready):
                    prompt = developer_selftest_prompt(selected_selftest_stack)
                    result, used_model = run_generation_with_details(prompt, f"{selected_selftest_stack} self-test")
                    if result:
                        st.session_state.selftest_response = result
                        st.session_state.selftest_last_model = used_model or current_model_name()
                        remember_output("Developer Self-Test", f"{selected_selftest_stack} Self-Test", result, "developer", "codefix")
                        st.success("Self-test completed.")
                if st.button("Clear Self-Test", key="clear_developer_selftest", use_container_width=True):
                    queue_reset("clear_selftest_workspace", "Developer self-test cleared.")
                    st.rerun()
        with right:
            with st.container(border=True):
                render_preview_panel(
                    "Self-Test Result",
                    "This shows the latest self-test output and the exact requested model id used for that run.",
                    "Self-test output",
                    st.session_state.selftest_response,
                    height=300,
                    empty_title="No self-test yet",
                    empty_body="Run the self-test from the left panel first. The result area remains compact until there is an actual benchmark response.",
                    anchor="ST",
                    tier="secondary",
                )
                render_meta_grid(
                    [
                        ("Last self-test model", st.session_state.selftest_last_model or "No self-test run yet"),
                        ("Last generation model", st.session_state.last_generation_model or "No generation yet"),
                        ("Last generation label", st.session_state.last_generation_label or "None"),
                    ]
                )

    with compare_tab:
        left, right = st.columns([1.02, 0.98], gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Compare Two Models",
                    "Run the same coding prompt against two requested model ids and inspect the outputs side by side.",
                    "Comparison",
                )
                compare_a, compare_b = st.columns(2, gap="small")
                compare_a.selectbox("Model A", options=MODEL_SELECTION_OPTIONS, key="compare_model_a")
                compare_b.selectbox("Model B", options=MODEL_SELECTION_OPTIONS, key="compare_model_b")
                if st.session_state.compare_model_a == "Custom":
                    st.text_input("Custom Model A", key="compare_custom_model_a")
                if st.session_state.compare_model_b == "Custom":
                    st.text_input("Custom Model B", key="compare_custom_model_b")
                st.text_area(
                    "Comparison prompt",
                    key="compare_prompt",
                    height=160,
                    placeholder="Example: Build a FastAPI CRUD endpoint with validation and explain the route structure.",
                )
                action_a, action_b = st.columns(2, gap="small")
                if action_a.button("Run Model Comparison", key="run_model_compare", use_container_width=True, type="primary", disabled=not ai_ready):
                    compare_prompt = st.session_state.compare_prompt.strip()
                    model_a = resolve_model_selection(st.session_state.compare_model_a, st.session_state.compare_custom_model_a)
                    model_b = resolve_model_selection(st.session_state.compare_model_b, st.session_state.compare_custom_model_b)
                    if not compare_prompt:
                        st.warning("Enter a comparison prompt first.")
                    elif not model_a or not model_b:
                        st.warning("Both model slots need a valid model id.")
                    else:
                        result_a, used_a = run_generation_with_details(compare_prompt, "model compare A", requested_model=model_a, track_global=False)
                        result_b, used_b = run_generation_with_details(compare_prompt, "model compare B", requested_model=model_b, track_global=False)
                        if result_a:
                            st.session_state.compare_output_a = result_a
                            st.session_state.compare_used_model_a = used_a or model_a
                        if result_b:
                            st.session_state.compare_output_b = result_b
                            st.session_state.compare_used_model_b = used_b or model_b
                        if result_a or result_b:
                            st.session_state.last_generation_label = "model comparison"
                            st.session_state.last_generation_model = st.session_state.compare_used_model_b or st.session_state.compare_used_model_a
                            st.session_state.last_generation_status = "success"
                            st.session_state.last_generation_time = workspace_timestamp()
                            st.session_state.last_generation_note = "Side-by-side model comparison completed."
                            st.success("Model comparison completed.")
                if action_b.button("Clear Comparison", key="clear_model_compare", use_container_width=True):
                    queue_reset("clear_compare_workspace", "Model comparison cleared.")
                    st.rerun()
        with right:
            model_col_a, model_col_b = st.columns(2, gap="small")
            with model_col_a:
                with st.container(border=True):
                    render_preview_panel(
                        "Model A Output",
                        st.session_state.compare_used_model_a or "No comparison run yet.",
                        "Compare output A",
                        st.session_state.compare_output_a,
                        height=300,
                        empty_title="No Model A output yet",
                        empty_body="Run the model comparison first. This side stays compact until the first output arrives.",
                        anchor="A",
                        tier="secondary",
                    )
            with model_col_b:
                with st.container(border=True):
                    render_preview_panel(
                        "Model B Output",
                        st.session_state.compare_used_model_b or "No comparison run yet.",
                        "Compare output B",
                        st.session_state.compare_output_b,
                        height=300,
                        empty_title="No Model B output yet",
                        empty_body="Run the model comparison first. This side stays compact until the second output arrives.",
                        anchor="B",
                        tier="secondary",
                    )

    with diff_tab:
        fixed_version = extract_section_value(st.session_state.codefix_response, ["Fixed Version"])
        diff_text = build_code_diff(st.session_state.codefix_source, fixed_version)
        left, right = st.columns(2, gap="large")
        with left:
            with st.container(border=True):
                render_card_header(
                    "Before / After Diff",
                    "The diff is built from the original snippet and the Fixed Version section returned by the AI.",
                    "Diff",
                )
                if diff_text:
                    st.code(diff_text, language="diff")
                else:
                    st.caption("Generate a code fix first to build a diff view.")
        with right:
            with st.container(border=True):
                render_card_header(
                    "Parsed Fix Sections",
                    "Useful when you want the root cause and next checks separated from the raw output.",
                    "Sections",
                )
                st.text_area("Root cause", value=extract_section_value(st.session_state.codefix_response, ["Root Cause"]), height=120, disabled=True)
                st.text_area("Why it works", value=extract_section_value(st.session_state.codefix_response, ["Why It Works"]), height=120, disabled=True)
                st.text_area("Next checks", value=extract_section_value(st.session_state.codefix_response, ["Next Checks"]), height=120, disabled=True)

    with notes_tab:
        with st.container(border=True):
            render_card_header(
                "Debug Notes",
                "Recent debugging outputs, last-used model details, and comparison context remain visible here for continuity.",
                "Developer Notes",
            )
            render_meta_grid(
                [
                    ("Last generation model", st.session_state.last_generation_model or "No generation yet"),
                    ("Last generation label", st.session_state.last_generation_label or "None"),
                    ("Last generation time", st.session_state.last_generation_time or "None"),
                ]
            )
            render_history_snippets(limit=8, page_filter="developer")
            row = st.columns(2, gap="small")
            if row[0].button("Open Standalone Code Fixer", key="open_standalone_codefix", use_container_width=True):
                go("codefix")
            if row[1].button("Clear Developer History", key="clear_dev_history_hint", use_container_width=True):
                go("history")


def render_history_page() -> None:
    entries = list(st.session_state.get("history_entries", []))
    with st.container(border=True):
        render_card_header(
            PAGE_DETAILS["history"]["title"],
            PAGE_DETAILS["history"]["subtitle"],
            "History",
            anchor="HS",
            tier="primary",
        )
        render_tag_row(["Uploads", "Generations", "Exports", "Workspace Events"])

    filter_left, filter_right = st.columns(2, gap="large")
    with filter_left:
        kind_filter = st.selectbox(
            "Filter by kind",
            options=["All"] + sorted({str(entry.get("kind", "")) for entry in entries if entry.get("kind")}),
            key="history_kind_filter",
        )
    with filter_right:
        page_filter = st.selectbox(
            "Filter by page",
            options=["All"] + sorted({str(entry.get("page", "")) for entry in entries if entry.get("page")}),
            key="history_page_filter",
        )

    filtered_entries = entries
    if kind_filter != "All":
        filtered_entries = [entry for entry in filtered_entries if entry.get("kind") == kind_filter]
    if page_filter != "All":
        filtered_entries = [entry for entry in filtered_entries if entry.get("page") == page_filter]

    top_left, top_right = st.columns([1.05, 0.95], gap="large")
    with top_left:
        with st.container(border=True):
            render_card_header(
                "Recent Session Timeline",
                "The most recent events appear first so it is easy to reconstruct what happened in the session.",
                "Timeline",
            )
            if filtered_entries:
                for entry in filtered_entries[:20]:
                    render_route_block(
                        f"{entry.get('timestamp', '')} • {entry.get('title', '')}",
                        f"{entry.get('details', '')}\nPage: {entry.get('page', '-')}\nWorkspace: {entry.get('workspace', '-')}",
                    )
            else:
                st.caption("No history entries match the current filters.")
    with top_right:
        with st.container(border=True):
            render_card_header(
                "Workspace Output Archive",
                "Saved outputs are grouped under the current workspace and remain reusable during this session.",
                "Archive",
            )
            render_workspace_outputs(limit=10)
            if st.button("Clear History", key="clear_history_button", use_container_width=True):
                queue_reset("clear_history_entries", "History cleared.", "info")
                st.rerun()


def render_settings_page(ai_ready: bool) -> None:
    secret_model = read_secret("HF_MODEL") or "No secret model configured"
    sync_model_selector_state()

    left, right = st.columns(2, gap="large")
    with left:
        with st.container(border=True):
            render_card_header(
                "Model Switcher",
                "Override the model coming from secrets for this Streamlit session.",
                "Settings",
                anchor="MD",
                tier="primary",
            )
            st.selectbox(
                "Model source",
                options=[item for item in MODEL_SELECTION_OPTIONS if item != "Active Session Model"],
                key="model_choice_input",
            )
            if st.session_state.model_choice_input == "Custom":
                st.text_input("Custom model id", key="model_custom_input")
            render_meta_grid(
                [
                    ("Secret default", secret_model),
                    ("Current override", st.session_state.model_override or "None"),
                    ("AI status", "Ready" if ai_ready else "Not ready"),
                ]
            )
            save_model_col, reset_model_col = st.columns(2, gap="small")
            if save_model_col.button("Apply Model Choice", key="apply_model_choice", use_container_width=True, type="primary"):
                choice = st.session_state.model_choice_input
                if choice == "HF Secret Default":
                    st.session_state.model_override = ""
                elif choice == "Custom":
                    st.session_state.model_override = st.session_state.model_custom_input.strip()
                else:
                    st.session_state.model_override = choice
                append_history_entry("settings", "Updated model selection", st.session_state.model_override or "Using HF secret default model.", "settings", active_workspace_name(), "model")
                st.session_state.flash_message = "Model selection updated."
                st.session_state.flash_level = "success"
                st.rerun()
            if reset_model_col.button("Reset To Secret Model", key="reset_model_choice", use_container_width=True):
                queue_reset("clear_model_selection", "Model override cleared.")
                st.rerun()

    with right:
        with st.container(border=True):
            render_card_header(
                "Session Controls",
                "Adjust session-level export defaults and history retention.",
                "Settings",
                anchor="SC",
                tier="secondary",
            )
            st.radio("Default export template", options=list(EXPORT_TEMPLATES.keys()), key="export_template", horizontal=False)
            st.number_input("History entry limit", min_value=20, max_value=300, step=10, key="settings_history_limit")
            render_route_block("Current package root", current_package_root_name())
            if st.button("Open History", key="settings_open_history", use_container_width=True):
                go("history")

    verify_left, verify_right = st.columns(2, gap="large")
    with verify_left:
        with st.container(border=True):
            render_card_header(
                "Real Model Verification Panel",
                "This runs a live probe against the currently active requested model id. It verifies the app request path and the provider response, not the hidden physical backend identity behind the provider.",
                "Verification",
            )
            render_meta_grid(
                [
                    ("Requested active model", current_model_name() or "No active model"),
                    ("Last verified model", st.session_state.verification_last_model or "No probe yet"),
                    ("Last verify status", st.session_state.verification_last_status or "Not run"),
                ]
            )
            if st.button("Run Live Verification Probe", key="run_live_model_probe", use_container_width=True, type="primary", disabled=not ai_ready):
                target_model = current_model_name()
                if not target_model:
                    st.warning("No active model is configured.")
                else:
                    ok, _probe = run_model_verification_probe(target_model)
                    if ok:
                        st.success("Live model probe completed.")
                    else:
                        st.error("Live model probe failed.")
            render_route_block(
                "Verification note",
                st.session_state.verification_last_note or "No verification probe has been run yet.",
            )

    with verify_right:
        with st.container(border=True):
            render_card_header(
                "Last Generation Used Model",
                "This records the exact requested model id used by the last generation call handled by the app.",
                "Generation Trace",
            )
            render_meta_grid(
                [
                    ("Model", st.session_state.last_generation_model or "No generation yet"),
                    ("Label", st.session_state.last_generation_label or "None"),
                    ("Status", st.session_state.last_generation_status or "None"),
                    ("Time", st.session_state.last_generation_time or "None"),
                ]
            )
            render_route_block(
                "Last generation note",
                st.session_state.last_generation_note or "No generation has been recorded yet.",
            )


def main() -> None:
    # Wide layout plus auto sidebar collapse gives better behavior on narrow screens.
    st.set_page_config(page_title="Hugyoku | Premium Academics Suite", layout="wide", initial_sidebar_state="auto")
    st.markdown(THEME_CSS, unsafe_allow_html=True)
    initialize_auth_storage()
    ensure_state()
    apply_pending_state_actions()
    render_flash_message()

    if not st.session_state.is_authenticated:
        render_login_gate()
        return

    ensure_workspace_bootstrap()

    if not can_access_page(st.session_state.active_page):
        fallback_page = "dashboard"
        for candidate in ["hugyoku", "dashboard", "workspaces", "academics", "developer", "history", "settings"]:
            if can_access_page(candidate):
                fallback_page = candidate
                break
        st.session_state.active_page = fallback_page
        st.session_state.flash_message = "That page is not allowed for your account."
        st.session_state.flash_level = "warning"
        st.rerun()

    _client, model, error = load_client()
    ai_ready = error is None
    model_label = model or "No model configured"
    ai_message = "AI ready for academics" if ai_ready else error

    render_sidebar(ai_ready, model_label, ai_message)
    render_header(ai_ready, model_label, ai_message)
    render_flash_message()

    page = st.session_state.active_page
    if page == "hugyoku":
        render_hugyoku_page(ai_ready)
    elif page == "dashboard":
        render_dashboard()
    elif page == "workspaces":
        render_workspaces_page(ai_ready)
    elif page == "academics":
        render_academics_hub(ai_ready)
    elif page == "developer":
        render_developer_hub(ai_ready)
    elif page == "history":
        render_history_page()
    elif page == "settings":
        render_settings_page(ai_ready)
    elif page == "admin":
        render_admin_page()
    elif page == "quiz":
        render_quiz_page(ai_ready)
    elif page == "assignment":
        render_assignment_page(ai_ready)
    elif page == "essay":
        render_essay_page(ai_ready)
    elif page == "activity":
        render_activity_page(ai_ready)
    elif page == "document":
        render_document_page(ai_ready)
    elif page == "codefix":
        render_codefix_page(ai_ready)
    else:
        go("dashboard")


if __name__ == "__main__":
    main()
